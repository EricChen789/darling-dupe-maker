"""
Clean ROD DOCX template: replace sample data with {{PLACEHOLDER}} markers.
The ROD template is paragraph-based, not table-based.
Pattern: header (P0-P12) → officers data blocks (P13-P36) → page number (P37)
"""
from docx import Document
import os

path = r"D:\myproject\darling-dupe-maker\_templates_cleaned\ROD_register_template.docx"
doc = Document(path)

# ── Header replacements ──
def replace_in_para(para, old, new):
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

replace_in_para(doc.paragraphs[0], "Testing Company Limited", "{{CO_NAME}}")
replace_in_para(doc.paragraphs[1], "0101234", "{{CO_BR}}")
replace_in_para(doc.paragraphs[2], "05 APRIL 2024", "{{REPORT_DATE}}")
replace_in_para(doc.paragraphs[3], "2", "{{QUORUM}}")

# ── Remove all sample officer data paragraphs (P13 through P36) ──
# Keep column headers (P4-P12) and page number (P37)
# We'll mark the data region with special markers
# First, add a marker paragraph before and after data
body = doc.element.body

# Find the paragraphs that contain sample data (P13-P36)
# Strategy: mark P13 with a unique ID and P37 with a marker
from docx.oxml import OxmlElement

# P13: "ABC TESTING" → mark as start of data block
for run in doc.paragraphs[13].runs:
    run.text = "{{ROD_DATA_START}}"

# P37: "- 1 -" → mark as end of data block
for run in doc.paragraphs[37].runs:
    if "- 1 -" in run.text:
        run.text = "{{ROD_DATA_END}}"

# Now remove paragraphs between start and end markers (P14-P36)
# But keep the marker paragraphs
paras_to_clear = list(range(14, 25)) + list(range(25, 37))
# Actually, let's keep the sample data block as a template for the XML structure
# Just replace all text in data paragraphs with placeholders

# For director block: P13-P24
dir_placeholders = [
    "{{DIR1_NAME}}",          # P13: Name
    "{{DIR1_ADDR1}}",         # P14: Address line 1
    "{{DIR1_ADDR2}}",         # P15: Address line 2
    "{{DIR1_ADDR3}}",         # P16: Address line 3
    "{{DIR1_DOB}}",           # P17: Date of Birth
    "{{DIR1_POB}}",           # P18: Place of Birth
    "{{DIR1_NATION}}",        # P19: Nationality
    "{{DIR1_ID}}",            # P20: ID No
    "{{DIR1_POSITION}}",      # P21: Position
    "{{DIR1_DATE_APP}}",      # P22: Date Appointed
    "{{DIR1_REASON}}",        # P23: Reason
    "{{DIR1_DATE_CEA}}",      # P24: Date Ceased
]

for i, ph in enumerate(dir_placeholders):
    p = doc.paragraphs[13 + i]
    for run in p.runs:
        if run.text.strip():
            run.text = ph
            break
    else:
        # No existing runs with text — add one
        if p.runs:
            p.runs[0].text = ph

# P25: "Secretary" label + P26: Secretary date — keep as template for secretary section
# Keep these as-is for now; we'll handle in cloud function

# For second director block: P27-P36  (same structure, different placeholders)
dir2_placeholders = [
    "{{DIR2_NAME}}",
    "{{DIR2_ADDR1}}", "{{DIR2_ADDR2}}", "{{DIR2_ADDR3}}",
    "{{DIR2_DOB}}", "{{DIR2_POB}}", "{{DIR2_NATION}}", "{{DIR2_ID}}",
    "{{DIR2_POSITION}}", "{{DIR2_DATE_APP}}",
]

for i, ph in enumerate(dir2_placeholders):
    p = doc.paragraphs[27 + i]
    for run in p.runs:
        if run.text.strip():
            run.text = ph
            break
    else:
        if p.runs:
            p.runs[0].text = ph

# Save cleaned template
out_path = r"D:\myproject\darling-dupe-maker\_templates_cleaned\ROD_register_template.docx"
doc.save(out_path)
print(f"Cleaned ROD template saved: {out_path}")
for i, p in enumerate(doc.paragraphs):
    t = p.text[:120]
    print(f'  P[{i}]: "{t}"')
