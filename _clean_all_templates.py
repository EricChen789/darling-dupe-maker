"""
Fill placeholder data into ROM and SCR templates from .doc conversions.
Each template has empty data rows — we fill them with placeholders for cloud cloning.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree
import zipfile, os, shutil

DATA_FONT_SIZE = Pt(7.5)  # Smaller font for injected data

def set_cell_text(cell, text, font_size=None):
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run(text)
    if font_size:
        run.font.size = font_size

# ── Header XML manipulation (for templates where python-docx can't easily access header paragraphs) ──
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def add_run_to_para(para_el, text, sz_val='16', szCs_val='16'):
    """Add a <w:r> with <w:t>text</w:t> and font size to a paragraph element."""
    r = etree.SubElement(para_el, f'{{{W_NS}}}r')
    rPr = etree.SubElement(r, f'{{{W_NS}}}rPr')
    etree.SubElement(rPr, f'{{{W_NS}}}sz').set(f'{{{W_NS}}}val', sz_val)
    etree.SubElement(rPr, f'{{{W_NS}}}szCs').set(f'{{{W_NS}}}val', szCs_val)
    t = etree.SubElement(r, f'{{{W_NS}}}t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def set_header_placeholders(docx_path, placeholder_map):
    """
    Open DOCX, parse header1.xml, find empty paragraphs by index and fill with
    placeholder text. placeholder_map: {para_index: placeholder_text}
    Preserves existing paragraph properties (font size, style).
    """
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/header1.xml':
                    tree = etree.fromstring(data)
                    paras = tree.findall(f'.//{{{W_NS}}}p')
                    for idx, text in placeholder_map.items():
                        if idx < len(paras):
                            p = paras[idx]
                            # Remove existing runs in this paragraph
                            for r in p.findall(f'{{{W_NS}}}r'):
                                p.remove(r)
                            # Get existing paragraph props for font size reference
                            pPr = p.find(f'{{{W_NS}}}pPr')
                            sz_val = '15'  # 7.5pt default for data
                            if pPr is not None:
                                rPr = pPr.find(f'{{{W_NS}}}rPr')
                                if rPr is not None:
                                    sz_el = rPr.find(f'{{{W_NS}}}sz')
                                    if sz_el is not None:
                                        sz_val = sz_el.get(f'{{{W_NS}}}val', '15')
                            add_run_to_para(p, text, sz_val=sz_val, szCs_val='16')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone='yes')
                zout.writestr(item, data)
    shutil.move(tmp, docx_path)

def replace_in_para(para, old, new):
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

OUT = r"D:\myproject\darling-dupe-maker\_templates_cleaned"

# ═══════════════════════════ ROM ═══════════════════════════
print("=== Cleaning ROM (Register of members.doc) ===")
rom = Document(os.path.join(OUT, "ROM_register_template.docx"))
table = rom.tables[0]

# Header
for p in table.rows[0].cells[0].paragraphs:
    for run in p.runs:
        if "Name of Company" in run.text:
            run.text = "{{CO_NAME}}"
for p in table.rows[1].cells[0].paragraphs:
    for run in p.runs:
        if "Company Number" in run.text:
            run.text = "Company Number: {{CO_BR}}"
for cell in table.rows[1].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            if "REGISTER OF MEMBERS" in run.text:
                run.text = "REGISTER OF MEMBERS as at {{REPORT_DATE}}"

# Fill data Row 9 (first shareholder) — 19 columns with placeholders
# Based on header analysis, these columns are:
# Col    Content
# 0      Full Name
# 1-2    Certificate Number (Acq)
# 3      Distinctive From (Acq)
# 4      Distinctive To (Acq)
# 5      No. of Shares (Acq)
# 6-7    Occupation (in header) / Consideration Paid (Acq) in data
# 8      Transfer Deed No (or Date Entered label)
# 9      Certificate Number (Xfer)
# 10     From (Xfer)
# 11     To (Xfer)
# 12-13  No. of Shares (Xfer)
# 14     Consideration Paid (Xfer)
# 15     Total Shares Held
# 16-17  Remarks
# 18     Entry Made By
data_phs = {
    0: "{{SH_NAME}}",
    1: "{{SH_CERT_ACQ}}",
    3: "{{SH_DIST_FM}}",
    4: "{{SH_DIST_TO}}",
    5: "{{SH_SHARES_ACQ}}",
    # Note: cells[6] and [7] are the same merged cell (Consideration Paid + Occupation share this area)
    7: "{{SH_CONS_ACQ}}",
    8: "{{SH_DATE_APP}}",
    9: "{{SH_CERT_XFER}}",
    12: "{{SH_DATE_CEA}}",
    15: "{{SH_TOTAL}}",
    16: "{{SH_REMARKS}}",
    18: "{{SH_ENTRY_BY}}",
}
for ci, ph in data_phs.items():
    if ph:
        set_cell_text(table.rows[9].cells[ci], ph, DATA_FONT_SIZE)

# Row 10: Address + Occupation (separate placeholders in the same merged cell)
set_cell_text(table.rows[10].cells[0], "{{SH_ADDR}}  |  {{SH_OCCUPATION}}", DATA_FONT_SIZE)

rom.save(os.path.join(OUT, "ROM_register_template.docx"))
print("ROM done")

# ═══════════════════════════ SCR ═══════════════════════════
print("\n=== Cleaning SCR (重要控制人.doc) ===")
scr = Document(os.path.join(OUT, "SCR_register_template.docx"))

# The SCR has 6 tables:
# Table[0]: Main SCR register (7 cols, 3 rows: header + 2 empty)
# Table[1]: Additional Matters (2x2)
# Table[2]: Designated Rep table (2x7)
# Table[3]: Another SCR table (3x7)
# Table[4]: Additional Matters (2x2)
# Table[5]: Another Designated Rep table (2x7)

# Fill data cells in Table[0] (main SCR — first controller)
t0 = scr.tables[0]
scr_phs = {
    0: "{{SCR_ENTRY_DATE}}",
    1: "{{SCR_NAME}}",
    2: "{{SCR_ADDR}}",
    3: "{{SCR_ID}}",
    4: "{{SCR_NATURE}}",
    5: "{{SCR_DATES}}",
    6: "{{SCR_REMARKS}}",
}
for ci, ph in scr_phs.items():
    set_cell_text(t0.rows[1].cells[ci], ph, DATA_FONT_SIZE)

# Also fill row 2 with second set of placeholders (for overflow)
scr2_phs = {
    0: "{{SCR2_ENTRY_DATE}}",
    1: "{{SCR2_NAME}}",
    2: "{{SCR2_ADDR}}",
    3: "{{SCR2_ID}}",
    4: "{{SCR2_NATURE}}",
    5: "{{SCR2_DATES}}",
    6: "{{SCR2_REMARKS}}",
}
for ci, ph in scr2_phs.items():
    set_cell_text(t0.rows[2].cells[ci], ph, DATA_FONT_SIZE)

# Table[2]: Designated Rep (page 1 — max 1 rep)
t2 = scr.tables[2]
rep_phs = {
    0: "{{REP1_ENTRY_DATE}}",
    1: "{{REP1_NAME}}",
    2: "{{REP1_ADDR}}",
    3: "{{REP1_CAPACITY}}",
    4: "{{REP1_TEL}}",
    5: "{{REP1_DATES}}",
    6: "{{REP1_REMARKS}}",
}
for ci, ph in rep_phs.items():
    set_cell_text(t2.rows[1].cells[ci], ph, DATA_FONT_SIZE)

# Table[5]: Designated Rep overflow page (page 2 — for 2nd+ reps)
t5 = scr.tables[5]
rep2_phs = {
    0: "{{REP2_ENTRY_DATE}}",
    1: "{{REP2_NAME}}",
    2: "{{REP2_ADDR}}",
    3: "{{REP2_CAPACITY}}",
    4: "{{REP2_TEL}}",
    5: "{{REP2_DATES}}",
    6: "{{REP2_REMARKS}}",
}
for ci, ph in rep2_phs.items():
    set_cell_text(t5.rows[1].cells[ci], ph, DATA_FONT_SIZE)

scr.save(os.path.join(OUT, "SCR_register_template.docx"))
print("SCR tables done")

# ── Header: fill company name / number / jurisdiction placeholders ──
scr_path = os.path.join(OUT, "SCR_register_template.docx")
set_header_placeholders(scr_path, {
    2: "{{CO_NAME}}",          # Empty para after "NAME OF COMPANY"
    14: "{{CO_BR}}",           # Empty para after "COMPANY NUMBER"
    17: "{{CO_JURISDICTION}}", # Empty para after "JURISDICTION"
})
print("SCR header done")

# ═══════════════════════════ Transfer Resolutions ═══════════════════════════
print("\n=== Cleaning Transfer Resolutions ===")
tr = Document(os.path.join(OUT, "Transfer_resolutions_template.docx"))

# Paragraph-based template — replace key text with placeholders
for pi, p in enumerate(tr.paragraphs):
    text = p.text.strip()
    if not text:
        continue

    # P0: Company name
    if text == "Testing Company Limited":
        p.clear()
        p.add_run("{{CO_NAME}}")

    # P1: Company number
    elif "Company Number:" in text:
        p.clear()
        p.add_run("Company Number: {{CO_BR}}")

    # P13: Seller name
    elif text.startswith("Seller:"):
        p.clear()
        p.add_run("Seller: {{SELLER_NAME}}")

    # P14: Seller HKID
    elif "Hong Kong ID No" in text and "Y000000" in text:
        p.clear()
        p.add_run("Hong Kong ID No {{SELLER_ID}}")

    # P15: Buyer name
    elif text.startswith("Buyer:"):
        p.clear()
        p.add_run("Buyer: {{BUYER_NAME}}")

    # P16: Buyer HKID
    elif "Hong Kong ID No" in text and "Y231456" in text:
        p.clear()
        p.add_run("Hong Kong ID No {{BUYER_ID}}")

    # P17: No. of Shares
    elif "No. of Shares:" in text:
        p.clear()
        p.add_run("No. of Shares: {{SHARES}}")

    # P18: Share description
    elif "HK$" in text and "Ordinary" in text:
        p.clear()
        p.add_run("{{SHARE_DESC}}")

    # P24: Authorised persons (gap between "that" and "be authorised")
    elif "be authorised to issue" in text:
        # Runs: ["It was resolved that", "and", "\tbe authorised..."]
        # Replace the middle run(s) with {{AUTH_PERSONS}}
        for run in p.runs:
            if "and" in run.text and len(run.text.strip()) <= 5:
                run.text = run.text.replace("and", " {{AUTH_PERSONS}} ")

    # P33: Signature date line
    elif "..../..../...." in text:
        p.clear()
        p.add_run(".....................................................      {{SIG_DATE}}")

    # P34: Director signatory name (the last "Bcd Testing" occurrence)
    elif text == "Bcd Testing":
        p.clear()
        p.add_run("{{DIRECTOR_NAME}}")

tr.save(os.path.join(OUT, "Transfer_resolutions_template.docx"))
print("Transfer Resolutions done")

print("\n=== All templates cleaned ===")
