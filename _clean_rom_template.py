"""
Clean ROM template from Register of members.doc:
- Replace company name/BR in header
- Add a sample shareholder row with placeholders for data cloning
"""
from docx import Document
from docx.oxml import OxmlElement
import os

path = r"D:\myproject\darling-dupe-maker\_templates_cleaned\ROM_register_template.docx"
doc = Document(path)
table = doc.tables[0]

# ── Header replacements ──
def set_cell_text(cell, text):
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run(text)

# Row 0, Cell 0: "Name of Company" → {{CO_NAME}}
for p in table.rows[0].cells[0].paragraphs:
    for run in p.runs:
        if "Name of Company" in run.text:
            run.text = "{{CO_NAME}}"

# Row 1, Cell 0: "Company Number" → change to show BR
for p in table.rows[1].cells[0].paragraphs:
    for run in p.runs:
        if "Company Number" in run.text:
            run.text = "Company Number: {{CO_BR}}"

# Row 1, Cell 13: "REGISTER OF MEMBERS" → add date
for p in table.rows[1].cells[13].paragraphs:
    for run in p.runs:
        if "REGISTER OF MEMBERS" in run.text:
            run.text = "REGISTER OF MEMBERS as at {{REPORT_DATE}}"

# ── Add a sample data row (Row 9) with placeholders ──
# The data section: each shareholder uses 5 rows (rows 9-13 for block 1)
# Each row corresponds to: Name, Address, Subscription, [empty], [empty]
# But the actual share data goes into specific columns

# Since the template has no sample data, let's fill Row 9 with the key placeholders
# The 19 columns map (from header analysis):
# 0=Name, 1-5=SharesAcq sub-cols, 6-7=Occupation, 8-14=SharesXfer/DateEntered, 15=Total, 16-17=Remarks, 18=EntryBy

# Actually, let me just fill a few cells with recognizable placeholders
data_cells_ph = {
    0: "{{SH_NAME}}",       # Full Name
    6: "{{SH_OCCUPATION}}",  # Occupation
    8: "{{SH_DATE_APP}}",    # Date Entered (in Shares Transferred section header)
    15: "{{SH_TOTAL}}",      # Total Shares
    16: "{{SH_REMARKS}}",     # Remarks
    18: "{{SH_ENTRY_BY}}",    # Entry Made By
}
for col_idx, ph in data_cells_ph.items():
    set_cell_text(table.rows[9].cells[col_idx], ph)

# Save
out_path = r"D:\myproject\darling-dupe-maker\_templates_cleaned\ROM_register_template.docx"
doc.save(out_path)
print(f"ROM cleaned: {out_path}")

# Print key cells
print("\nHeader check:")
print(f"  Row[0] Cell[0]: '{table.rows[0].cells[0].text[:80]}'")
print(f"  Row[1] Cell[0]: '{table.rows[1].cells[0].text[:80]}'")
print(f"  Row[1] Cell[13]: '{table.rows[1].cells[13].text[:80]}'")
print(f"  Row[9] Cell[0]: '{table.rows[9].cells[0].text[:80]}'")
print(f"  Row[9] Cell[6]: '{table.rows[9].cells[6].text[:80]}'")
