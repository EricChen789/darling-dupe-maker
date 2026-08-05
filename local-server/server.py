#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Local API server for darling-dupe-maker
Mimics Cloudflare Pages Functions API with SQLite
Dev mode: no auth required for any operation
"""
import sqlite3
import uuid
import hashlib
import hmac
import json
import os
import re
import time
import base64
import smtplib
import threading
import urllib.request
import tempfile
import requests
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.utils import formataddr
from datetime import datetime, timedelta, timezone
from flask import Flask, request, jsonify, g, Response
from fpdf import FPDF, XPos, YPos

# Word COM for RTF→PDF conversion (Windows only)
try:
    import pythoncom
    import win32com.client
    _HAS_WORD_COM = True
except ImportError:
    _HAS_WORD_COM = False

# Thread-safe lock for Word COM (Word is single-instance)
_word_lock = threading.Lock()

# RTF template paths (Paul Tang reference files)
_RTF_DIR = os.path.join(os.path.dirname(__file__), '..', '..', '秘书系统文件', '汇出股权转让文件')
_RTF_BS_NOTE = os.path.join(_RTF_DIR, 'Testing Bought Sold note_fixed.rtf')
_RTF_INSTRUMENT = os.path.join(_RTF_DIR, 'Testing Instrument of transfer.rtf')
_RTF_CERTIFICATE = os.path.join(_RTF_DIR, 'Testing Share Certificate.rtf')

# Register RTF templates (Paul Tang reference samples)
_RTF_REGISTER_DIR = os.path.join(os.path.dirname(__file__), '..', '..', '秘书系统文件', 'rod rom')
_RTF_ROM = os.path.join(_RTF_REGISTER_DIR, 'Register of members.doc')
_RTF_ROD = os.path.join(_RTF_REGISTER_DIR, 'Testing ROD.rtf')

# DOCX register templates (Paul Tang format)
_ROM_DOCX_TEMPLATE = os.path.join(
    os.path.dirname(__file__), '..', '..', '秘书系统文件', '登记册', '股東登記冊_PaulTang格式.docx'
)
_SCR_DOCX_TEMPLATE = os.path.join(
    os.path.dirname(__file__), '..', '..', '秘书系统文件', '登记册', '重要控制人登記冊_PaulTang格式.docx'
)

# python-docx for register DOCX generation
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

app = Flask(__name__)
DB_PATH = os.path.join(os.path.dirname(__file__), 'local.db')
JWT_SECRET = 'local-dev-secret-do-not-use-in-production'
AI_MOCK = os.environ.get('AI_MOCK', 'true').lower() == 'true'  # 默认 Mock 模式，CI 设 false 调真实 API

# ─── Database ───
def get_db():
    if 'db' not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA journal_mode=WAL")
        g.db.execute("PRAGMA foreign_keys=ON")
    return g.db

@app.teardown_appcontext
def close_db(e):
    db = g.pop('db', None)
    if db:
        db.close()

# ─── Cloud Write-Through Sync ────────────────────────────────────────────
# Best-effort: every local write pushes to cloud in a background thread.
# Set SYNC_ENABLED=false to disable.

CLOUD_SYNC_BASE = "https://secretary-system-9cl.pages.dev"
CLOUD_SYNC_SECRET = "d54891cbb3df4705dec96bef87297447df29ee64bebf7413d1186786f777ecb8"
CLOUD_SYNC_SUB = "f0e8d7c6-b5a4-4932-8180-abcdef123456"
CLOUD_SYNC_EMAIL = "admin@localhost"
SYNC_ENABLED = os.environ.get("SYNC_ENABLED", "true").lower() == "true"

SYNC_TABLES = {
    "companies", "persons", "presenters", "significant_controllers",
    "person_company_roles", "shareholders", "officers", "share_transactions",
    "resolutions", "reminders", "nar1_filings", "change_events",
    "form_linkages", "company_versions", "form_history",
}

_sync_token = None
_sync_token_lock = threading.Lock()

def _get_sync_token():
    """Generate self-signed JWT for cloud API writes (matches sync_tool.py)."""
    global _sync_token
    with _sync_token_lock:
        if _sync_token is not None:
            return _sync_token
        header = {"alg": "HS256", "typ": "JWT"}
        now = int(time.time())
        payload = {"sub": CLOUD_SYNC_SUB, "email": CLOUD_SYNC_EMAIL,
                   "iat": now, "exp": now + 3600}
        header_b64 = base64.urlsafe_b64encode(
            json.dumps(header).encode()).rstrip(b"=").decode()
        payload_b64 = base64.urlsafe_b64encode(
            json.dumps(payload).encode()).rstrip(b"=").decode()
        signing_input = f"{header_b64}.{payload_b64}"
        sig = hmac.new(CLOUD_SYNC_SECRET.encode(),
                       signing_input.encode(), hashlib.sha256).digest()
        sig_b64 = base64.urlsafe_b64encode(sig).rstrip(b"=").decode()
        _sync_token = f"{signing_input}.{sig_b64}"
        return _sync_token

def _sync_to_cloud(method, path, body=None):
    """Fire-and-forget sync to cloud. Silent failure — best-effort only."""
    if not SYNC_ENABLED:
        return
    def _do():
        try:
            token = _get_sync_token()
            url = f"{CLOUD_SYNC_BASE}{path}"
            data = json.dumps(body).encode("utf-8") if body else None
            req = urllib.request.Request(url, data=data, method=method)
            req.add_header("Content-Type", "application/json")
            req.add_header("Authorization", f"Bearer {token}")
            req.add_header("User-Agent",
                           "Mozilla/5.0 (Windows NT 10.0; Win64; x64) LocalDev/1.0")
            urllib.request.urlopen(req, timeout=10)
        except Exception:
            pass
    threading.Thread(target=_do, daemon=True).start()

# ─── CORS ───
@app.after_request
def cors(resp):
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
    resp.headers['Access-Control-Allow-Headers'] = 'Authorization, Content-Type, apikey'
    return resp

# ─── JWT ───
def base64url(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).rstrip(b'=').decode()

JWT_TTL_SECONDS = 7 * 24 * 60 * 60  # token 有效期 7 天（配合 verify_jwt 的 exp 檢查）

def sign_jwt(payload: dict) -> str:
    header = base64url(json.dumps({"alg": "HS256", "typ": "JWT"}).encode())
    now = int(time.time())
    p = {"exp": now + JWT_TTL_SECONDS, **payload, "iat": now}
    payload_b64 = base64url(json.dumps(p).encode())
    sig = hmac.new(JWT_SECRET.encode(), f"{header}.{payload_b64}".encode(), hashlib.sha256).digest()
    return f"{header}.{payload_b64}.{base64url(sig)}"

def verify_jwt(token: str) -> dict | None:
    try:
        parts = token.split('.')
        if len(parts) != 3:
            return None
        header_b64, payload_b64, sig_b64 = parts
        expected_sig = base64url(hmac.new(JWT_SECRET.encode(), f"{header_b64}.{payload_b64}".encode(), hashlib.sha256).digest())
        if sig_b64 != expected_sig:
            return None
        # Fix padding for base64 decode
        payload_b64 = payload_b64 + '=' * (4 - len(payload_b64) % 4)
        payload = json.loads(base64.urlsafe_b64decode(payload_b64))
        if payload.get('exp') and payload['exp'] < time.time():
            return None
        return payload
    except Exception as e:
        print(f"[JWT] Verification failed: {e}")
        return None

# ─── Password ───
def hash_password(password: str) -> str:
    salt = os.urandom(16)
    key = hashlib.pbkdf2_hmac('sha256', password.encode(), salt, 100000, dklen=32)
    return f"{base64url(salt)}:{base64url(key)}"

def verify_password(password: str, stored: str) -> bool:
    try:
        salt_b64, hash_b64 = stored.split(':')
        salt = base64.urlsafe_b64decode(salt_b64 + '=' * (4 - len(salt_b64) % 4))
        key = hashlib.pbkdf2_hmac('sha256', password.encode(), salt, 100000, dklen=32)
        return base64url(key) == hash_b64
    except (ValueError, TypeError) as e:
        print(f"[AUTH] Password verification error: {e}")
        return False

# ─── Auth (dev mode: always returns admin) ───
def get_user():
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    if not token:
        return None
    payload = verify_jwt(token)
    if not payload:
        return None
    return {
        'id': payload['sub'],
        'email': payload['email'],
        'display_name': payload.get('display_name', ''),
        'role': payload.get('role', 'user'),
    }

# ─── Init DB ───
def init_db():
    db = sqlite3.connect(DB_PATH)
    db.executescript("""
        CREATE TABLE IF NOT EXISTS auth_users (
            id TEXT PRIMARY KEY,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            display_name TEXT DEFAULT '',
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS user_roles (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            role TEXT NOT NULL,
            created_at TEXT NOT NULL DEFAULT '',
            updated_at TEXT NOT NULL DEFAULT (datetime('now')),
            UNIQUE (user_id, role)
        );

        CREATE TABLE IF NOT EXISTS company_versions (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL,
            version_no INTEGER NOT NULL DEFAULT 1,
            snapshot TEXT NOT NULL DEFAULT '{}',
            changed_fields TEXT NOT NULL DEFAULT '[]',
            change_summary TEXT DEFAULT '',
            changed_by TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS companies (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            chinese_name TEXT DEFAULT '',
            company_number TEXT DEFAULT '',
            trading_name TEXT DEFAULT '',
            business_nature TEXT DEFAULT '',
            company_type TEXT DEFAULT '私人公司 Private company',
            business_code TEXT DEFAULT '',
            company_group TEXT DEFAULT '',
            quorum TEXT DEFAULT '',
            register_date TEXT DEFAULT '',
            reg_flat TEXT DEFAULT '',
            reg_building TEXT DEFAULT '',
            reg_street TEXT DEFAULT '',
            reg_district TEXT DEFAULT '',
            reg_region TEXT DEFAULT '香港 Hong Kong',
            incorporation_date TEXT DEFAULT '',
            jurisdiction TEXT DEFAULT 'Hong Kong',
            ci_file_path TEXT DEFAULT '',
            br_file_path TEXT DEFAULT '',
            preferred_presenter_id TEXT,
            presenter_reference TEXT DEFAULT '',
            status TEXT DEFAULT 'active',
            ci_number TEXT DEFAULT '',
            email TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            signer_role_id TEXT,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS officers (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL REFERENCES companies(id) ON DELETE CASCADE,
            name_english TEXT NOT NULL DEFAULT '',
            name_chinese TEXT DEFAULT '',
            identity TEXT NOT NULL DEFAULT 'natural',
            role TEXT NOT NULL CHECK (role IN ('director', 'secretary')),
            id_number TEXT DEFAULT '',
            address TEXT DEFAULT '',
            date_appointed TEXT DEFAULT '',
            date_ceased TEXT DEFAULT '',
            place_incorporated TEXT DEFAULT '',
            company_number_ref TEXT DEFAULT '',
            service_address TEXT DEFAULT '',
            passport_number TEXT DEFAULT '',
            passport_expiry TEXT DEFAULT '',
            whatsapp TEXT DEFAULT '',
            email TEXT DEFAULT '',
            passport_file_path TEXT DEFAULT '',
            id_card_file_path TEXT DEFAULT '',
            address_proof_file_path TEXT DEFAULT '',
            tcsp_number TEXT DEFAULT '',
            previous_name_chinese TEXT DEFAULT '',
            previous_name_english TEXT DEFAULT '',
            alias_chinese TEXT DEFAULT '',
            alias_english TEXT DEFAULT '',
            date_of_birth TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS shareholders (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL REFERENCES companies(id) ON DELETE CASCADE,
            name TEXT NOT NULL DEFAULT '',
            shares INTEGER NOT NULL DEFAULT 0,
            identity TEXT NOT NULL DEFAULT 'natural',
            id_number TEXT DEFAULT '',
            name_chinese TEXT DEFAULT '',
            name_english TEXT DEFAULT '',
            address TEXT DEFAULT '',
            email TEXT DEFAULT '',
            share_type TEXT DEFAULT '',
            service_address TEXT DEFAULT '',
            issue_price TEXT DEFAULT '',
            currency TEXT DEFAULT 'HKD',
            paid_up TEXT DEFAULT '',
            unpaid TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS persons (
            id TEXT PRIMARY KEY,
            identity TEXT NOT NULL DEFAULT 'natural',
            name_english TEXT NOT NULL DEFAULT '',
            name_chinese TEXT DEFAULT '',
            previous_name_english TEXT DEFAULT '',
            previous_name_chinese TEXT DEFAULT '',
            alias_english TEXT DEFAULT '',
            alias_chinese TEXT DEFAULT '',
            id_number TEXT DEFAULT '',
            passport_number TEXT DEFAULT '',
            passport_expiry TEXT DEFAULT '',
            passport_country TEXT DEFAULT '',
            address TEXT DEFAULT '',
            service_address TEXT DEFAULT '',
            addr_flat TEXT DEFAULT '',
            addr_building TEXT DEFAULT '',
            addr_street TEXT DEFAULT '',
            addr_district TEXT DEFAULT '',
            addr_region TEXT DEFAULT '',
            svc_addr_flat TEXT DEFAULT '',
            svc_addr_building TEXT DEFAULT '',
            svc_addr_street TEXT DEFAULT '',
            svc_addr_district TEXT DEFAULT '',
            svc_addr_region TEXT DEFAULT '',
            email TEXT DEFAULT '',
            whatsapp TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            place_incorporated TEXT DEFAULT '',
            company_number_ref TEXT DEFAULT '',
            tcsp_number TEXT DEFAULT '',
            passport_file_path TEXT DEFAULT '',
            id_card_file_path TEXT DEFAULT '',
            address_proof_file_path TEXT DEFAULT '',
            normalized_key TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            date_of_birth TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS person_company_roles (
            id TEXT PRIMARY KEY,
            person_id TEXT NOT NULL,
            company_id TEXT NOT NULL,
            role TEXT NOT NULL,
            date_appointed TEXT DEFAULT '',
            date_ceased TEXT DEFAULT '',
            service_address_override TEXT DEFAULT '',
            shares INTEGER DEFAULT 0,
            share_type TEXT DEFAULT '',
            currency TEXT DEFAULT 'HKD',
            issue_price TEXT DEFAULT '',
            paid_up TEXT DEFAULT '',
            unpaid TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            is_reserve INTEGER DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS presenters (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            address TEXT DEFAULT '',
            contact TEXT DEFAULT '',
            type TEXT NOT NULL DEFAULT 'individual',
            phone TEXT DEFAULT '',
            fax TEXT DEFAULT '',
            email TEXT DEFAULT '',
            reference TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS significant_controllers (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL REFERENCES companies(id) ON DELETE CASCADE,
            identity TEXT NOT NULL DEFAULT 'natural',
            name_english TEXT NOT NULL DEFAULT '',
            name_chinese TEXT DEFAULT '',
            id_number TEXT DEFAULT '',
            address TEXT DEFAULT '',
            service_address TEXT DEFAULT '',
            date_became TEXT DEFAULT '',
            date_ceased TEXT DEFAULT '',
            nature_shares INTEGER DEFAULT 0,
            nature_voting INTEGER DEFAULT 0,
            nature_appoint INTEGER DEFAULT 0,
            nature_influence INTEGER DEFAULT 0,
            nature_trust INTEGER DEFAULT 0,
            nature_other TEXT DEFAULT '',
            is_designated_rep INTEGER DEFAULT 0,
            designated_rep_name TEXT DEFAULT '',
            designated_rep_contact TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS company_logs (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL,
            company_name_hint TEXT DEFAULT '',
            source_folder TEXT DEFAULT '',
            doc_type TEXT DEFAULT '',
            original_filename TEXT DEFAULT '',
            storage_path TEXT DEFAULT '',
            html_content TEXT DEFAULT '',
            text_content TEXT DEFAULT '',
            doc_date TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS reminders (
            id TEXT PRIMARY KEY,
            company_id TEXT,
            reminder_type TEXT DEFAULT 'NAR1',
            title TEXT DEFAULT '',
            due_date TEXT DEFAULT '',
            status TEXT DEFAULT 'pending',
            notes TEXT DEFAULT '',
            notified_at TEXT DEFAULT NULL,
            completed_at TEXT DEFAULT NULL,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS resolutions (
            id TEXT PRIMARY KEY,
            company_id TEXT,
            title TEXT DEFAULT '',
            content TEXT DEFAULT '',
            resolution_type TEXT DEFAULT '',
            resolution_date TEXT DEFAULT '',
            signers TEXT DEFAULT '',
            is_ai_generated INTEGER DEFAULT 0,
            notes TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS secretary_templates (
            id TEXT PRIMARY KEY,
            label TEXT DEFAULT '',
            identity TEXT DEFAULT 'corporate',
            name_english TEXT DEFAULT '',
            name_chinese TEXT DEFAULT '',
            id_number TEXT DEFAULT '',
            br_number TEXT DEFAULT '',
            tcsp_number TEXT DEFAULT '',
            place_incorporated TEXT DEFAULT '',
            address TEXT DEFAULT '',
            service_address TEXT DEFAULT '',
            email TEXT DEFAULT '',
            phone TEXT DEFAULT '',
            is_default INTEGER DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS share_transactions (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL REFERENCES companies(id) ON DELETE CASCADE,
            transaction_date TEXT NOT NULL DEFAULT '',
            transaction_type TEXT NOT NULL DEFAULT 'transfer',
            from_person_id TEXT,
            from_name TEXT DEFAULT '',
            to_person_id TEXT,
            to_name TEXT DEFAULT '',
            shares INTEGER NOT NULL DEFAULT 0,
            share_type TEXT DEFAULT '',
            currency TEXT DEFAULT 'HKD',
            price_per_share TEXT DEFAULT '',
            total_consideration TEXT DEFAULT '',
            instrument_number TEXT DEFAULT '',
            notes TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS invoices (
            id TEXT PRIMARY KEY,
            company_id TEXT NOT NULL REFERENCES companies(id) ON DELETE CASCADE,
            invoice_number TEXT DEFAULT '',
            description TEXT DEFAULT '',
            amount REAL NOT NULL DEFAULT 0,
            currency TEXT DEFAULT 'HKD',
            status TEXT DEFAULT 'pending',
            issue_date TEXT DEFAULT '',
            due_date TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS email_templates (
            id TEXT PRIMARY KEY,
            name TEXT DEFAULT '',
            template_type TEXT DEFAULT 'general',
            subject TEXT DEFAULT '',
            body TEXT DEFAULT '',
            is_default INTEGER DEFAULT 0,
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );

        CREATE TABLE IF NOT EXISTS email_logs (
            id TEXT PRIMARY KEY,
            company_id TEXT,
            template_id TEXT,
            to_email TEXT DEFAULT '',
            cc_email TEXT DEFAULT '',
            subject TEXT DEFAULT '',
            body TEXT DEFAULT '',
            status TEXT DEFAULT 'sent',
            scheduled_at TEXT DEFAULT NULL,
            sent_at TEXT DEFAULT NULL,
            error TEXT DEFAULT '',
            created_at TEXT NOT NULL DEFAULT (datetime('now')),
            updated_at TEXT NOT NULL DEFAULT (datetime('now'))
        );
    """)
    db.commit()

    # Seed admin user if not exists
    admin = db.execute("SELECT id FROM auth_users WHERE email = 'admin@localhost'").fetchone()
    if not admin:
        uid = str(uuid.uuid4())
        pwd_hash = hash_password('admin123')
        db.execute("INSERT INTO auth_users (id, email, password_hash, display_name) VALUES (?, ?, ?, ?)",
                   (uid, 'admin@localhost', pwd_hash, 'Admin'))
        db.execute("INSERT INTO user_roles (id, user_id, role) VALUES (?, ?, ?)",
                   (str(uuid.uuid4()), uid, 'admin'))
        db.commit()
        print(f"[INIT] Created admin user: admin@localhost / admin123")
    db.close()

# ─── Font & PDF helpers ───
FONT_PATH = os.path.join(os.path.dirname(__file__), 'NotoSansTC-Regular.otf')
FONT_URL = "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/TraditionalChinese/NotoSansTC-Regular.otf"
SYSTEM_FONT_FALLBACKS = [
    "C:/Windows/Fonts/simhei.ttf",
    "C:/Windows/Fonts/STXIHEI.TTF",
    "C:/Windows/Fonts/STKAITI.TTF",
    "C:/Windows/Fonts/STSONG.TTF",
    "C:/Windows/Fonts/simkai.ttf",
]
_cached_font_path = None

def find_font():
    """Find an available CJK font. Returns path or None. Caches result."""
    global _cached_font_path
    if _cached_font_path:
        return _cached_font_path

    # 1. Downloaded font
    if os.path.exists(FONT_PATH) and os.path.getsize(FONT_PATH) > 10000:
        _cached_font_path = FONT_PATH
        return FONT_PATH

    # 2. Try to download
    try:
        req = urllib.request.Request(FONT_URL, headers={'User-Agent': 'Mozilla/5.0'})
        urllib.request.urlretrieve(FONT_URL, FONT_PATH)
        if os.path.exists(FONT_PATH) and os.path.getsize(FONT_PATH) > 10000:
            print(f"[FONT] Downloaded Noto Sans TC ({os.path.getsize(FONT_PATH)} bytes)")
            _cached_font_path = FONT_PATH
            return FONT_PATH
    except Exception as e:
        print(f"[FONT] Download failed: {e}")

    # 3. Fall back to Windows system CJK fonts
    for f in SYSTEM_FONT_FALLBACKS:
        if os.path.exists(f):
            print(f"[FONT] Using system font: {f}")
            _cached_font_path = f
            return f

    print("[FONT] WARNING: No CJK font found, PDF will have no Chinese text!")
    return None

def create_pdf(landscape=False):
    """Create an fpdf2 FPDF instance with Chinese font (cached) + Times New Roman."""
    orient = 'L' if landscape else 'P'
    font_path = find_font()
    pdf = FPDF(orientation=orient, unit='pt', format='A4')
    if font_path:
        try:
            pdf.add_font('TC', style='', fname=font_path)
            pdf.add_font('TC', style='B', fname=font_path)
        except Exception as e:
            print(f"[PDF] Failed to load font {font_path}: {e}")
    # Register Times New Roman for English text (Paul Tang reference style)
    tnr_dir = 'C:/Windows/Fonts'
    for style_name, fname in [('', 'times.ttf'), ('B', 'timesbd.ttf'), ('I', 'timesi.ttf'), ('BI', 'timesbi.ttf')]:
        tnr_path = os.path.join(tnr_dir, fname)
        if os.path.exists(tnr_path):
            try:
                pdf.add_font('TNR', style=style_name, fname=tnr_path)
            except Exception:
                pass
    if not font_path:
        pdf.set_font('Helvetica', size=10)
    pdf.set_auto_page_break(auto=True, margin=60)
    return pdf

def pdf_draw(pdf, text, x, y, size=10, gray=0, color=None, bold=False):
    """Draw text at absolute position. color=(r,g,b) or None=black; gray=int 0-255 fallback."""
    pdf.set_xy(x, y)
    style = 'B' if bold else ''
    pdf.set_font('TC', style, size=size)
    if color:
        pdf.set_text_color(*color)
    elif gray != 0:
        pdf.set_text_color(gray, gray, gray)
    else:
        pdf.set_text_color(0, 0, 0)
    pdf.cell(0, size + 2, text or '', new_x="LMARGIN", new_y="NEXT")


def pdf_draw_field(pdf, label, value, x_label, x_value, y, value_width, size=9, label_gray=100, line_height=None):
    """Draw a label-value pair with manual multi-line wrapping for long values.

    Returns the new y position after drawing.
    Uses cell() for all text — no multi_cell, to avoid font/page-break issues.
    """
    if line_height is None:
        line_height = size + 3

    pdf.set_font('TC', size=size)
    text = value or '-'

    # Measure how many lines this value needs
    tw = pdf.get_string_width(text)
    lines = max(1, -(-tw // value_width))  # ceiling division

    # Draw label with cell() at (x_label, y)
    pdf.set_xy(x_label, y)
    pdf.set_font('TC', size=size)
    pdf.set_text_color(label_gray, label_gray, label_gray)
    pdf.cell(0, line_height, label)

    # Draw value — manual line-wrapping via multiple cell() calls
    pdf.set_font('TC', size=size)
    pdf.set_text_color(0, 0, 0)

    if lines == 1:
        pdf.set_xy(x_value, y)
        pdf.cell(value_width, line_height, text)
    else:
        # Split text into lines that fit within value_width.
        # Walk character by character, measuring cumulative width.
        wrapped_lines = []
        current = ''
        for ch in text:
            trial = current + ch
            if pdf.get_string_width(trial) > value_width and current:
                wrapped_lines.append(current)
                current = ch
            else:
                current = trial
        if current:
            wrapped_lines.append(current)
        # Ensure we have at least `lines` entries (safety)
        while len(wrapped_lines) < lines:
            wrapped_lines.append('')

        for i, line_text in enumerate(wrapped_lines):
            if i >= lines:
                break
            pdf.set_xy(x_value, y + i * line_height)
            pdf.cell(value_width, line_height, line_text)

    return y + lines * line_height

def seed_gray(val):
    """Convert 0-1 float to 0-255 int for fpdf2."""
    return int(round(val * 255))

def rget(row, key, default=''):
    """Safely get value from sqlite3.Row or dict."""
    try:
        val = row[key]
        return val if val is not None else default
    except (KeyError, IndexError):
        return default

# ─── Auto-migration ───
def auto_migrate():
    """Add missing columns to existing tables (handles schema evolution)."""
    db = sqlite3.connect(DB_PATH)

    def ensure_columns(table, columns):
        """Ensure all columns exist in table. columns is {col_name: col_type}."""
        try:
            rows = db.execute(f"PRAGMA table_info({table})").fetchall()
            existing = {row[1] for row in rows}
        except sqlite3.OperationalError:
            return  # table doesn't exist yet
        for col, col_type in columns.items():
            if col not in existing:
                try:
                    db.execute(f"ALTER TABLE {table} ADD COLUMN {col} {col_type}")
                    print(f"[MIGRATE] Added {table}.{col}")
                except sqlite3.OperationalError as e:
                    print(f"[MIGRATE] Skip {table}.{col}: {e}")

    # officers: add updated_at (P3-1)
    ensure_columns('officers', {
        'updated_at': "TEXT NOT NULL DEFAULT (datetime('now'))",
    })

    # persons: split address fields (NP-05 地址分拆欄位) — 通訊 addr_* + 送達 svc_addr_*
    ensure_columns('persons', {
        'addr_flat': "TEXT DEFAULT ''",
        'addr_building': "TEXT DEFAULT ''",
        'addr_street': "TEXT DEFAULT ''",
        'addr_district': "TEXT DEFAULT ''",
        'addr_region': "TEXT DEFAULT ''",
        'svc_addr_flat': "TEXT DEFAULT ''",
        'svc_addr_building': "TEXT DEFAULT ''",
        'svc_addr_street': "TEXT DEFAULT ''",
        'svc_addr_district': "TEXT DEFAULT ''",
        'svc_addr_region': "TEXT DEFAULT ''",
    })

    # shareholders: add updated_at (P3-1)
    ensure_columns('shareholders', {
        'updated_at': "TEXT NOT NULL DEFAULT (datetime('now'))",
    })

    # user_roles: add updated_at (P3-1) — use simple default for SQLite compat
    ensure_columns('user_roles', {
        'updated_at': "TEXT DEFAULT ''",
    })

    # auth_users: add is_active (P1 — 停用/啟用用户，10.3)
    ensure_columns('auth_users', {
        'is_active': "INTEGER NOT NULL DEFAULT 1",
    })

    # resolutions: add 4 missing columns (P0-2)
    ensure_columns('resolutions', {
        'resolution_date': "TEXT DEFAULT ''",
        'signers': "TEXT DEFAULT ''",
        'is_ai_generated': "INTEGER DEFAULT 0",
        'notes': "TEXT DEFAULT ''",
    })

    # secretary_templates: replace old schema with new one (P0-1)
    try:
        rows = db.execute("PRAGMA table_info(secretary_templates)").fetchall()
        existing = {row[1] for row in rows}
    except sqlite3.OperationalError:
        existing = set()

    if 'name' in existing and 'label' not in existing:
        # Old schema detected — add all new columns
        new_columns = {
            'label': "TEXT DEFAULT ''",
            'identity': "TEXT DEFAULT 'corporate'",
            'name_english': "TEXT DEFAULT ''",
            'name_chinese': "TEXT DEFAULT ''",
            'id_number': "TEXT DEFAULT ''",
            'br_number': "TEXT DEFAULT ''",
            'tcsp_number': "TEXT DEFAULT ''",
            'place_incorporated': "TEXT DEFAULT ''",
            'address': "TEXT DEFAULT ''",
            'service_address': "TEXT DEFAULT ''",
            'email': "TEXT DEFAULT ''",
            'phone': "TEXT DEFAULT ''",
            'is_default': "INTEGER DEFAULT 0",
        }
        for col, col_type in new_columns.items():
            if col not in existing:
                try:
                    db.execute(f"ALTER TABLE secretary_templates ADD COLUMN {col} {col_type}")
                    print(f"[MIGRATE] Added secretary_templates.{col}")
                except sqlite3.OperationalError as e:
                    print(f"[MIGRATE] Skip secretary_templates.{col}: {e}")
        # Migrate data: copy name→label, content→(unused, keep as-is)
        try:
            count = db.execute("UPDATE secretary_templates SET label = COALESCE(name, '') WHERE label = '' OR label IS NULL").rowcount
            if count:
                print(f"[MIGRATE] Migrated {count} secretary_templates rows: name→label")
        except sqlite3.OperationalError as e:
            print(f"[MIGRATE] Data migration skip: {e}")

    # share_transactions: add columns that the frontend (RegistersTab) expects
    tx_columns = {
        'from_person_id': 'TEXT',
        'from_name': "TEXT DEFAULT ''",
        'to_person_id': 'TEXT',
        'to_name': "TEXT DEFAULT ''",
        'share_type': "TEXT DEFAULT ''",
        'currency': "TEXT DEFAULT 'HKD'",
        'price_per_share': "TEXT DEFAULT ''",
        'total_consideration': "TEXT DEFAULT ''",
        'instrument_number': "TEXT DEFAULT ''",
        'updated_at': "TEXT NOT NULL DEFAULT (datetime('now'))",
    }
    ensure_columns('share_transactions', tx_columns)

    # email_logs: add email_type column (EM-04/05)
    ensure_columns('email_logs', {'email_type': "TEXT DEFAULT 'general'"})

    db.commit()

    # Seed default email templates if none exist (Email Module 8.1)
    try:
        cnt = db.execute("SELECT COUNT(*) FROM email_templates").fetchone()[0]
    except sqlite3.OperationalError:
        cnt = None
    if cnt == 0:
        seeds = [
            ('invoice', '發票通知', '【{company_name}】服務發票 {invoice_number}',
             '致 {client_name}：\n\n茲附上貴公司 {company_name}（商業登記號碼：{br_number}）的服務發票。\n\n發票編號：{invoice_number}\n金額：{currency} {amount}\n到期日：{due_date}\n\n請於到期日前安排付款，如有查詢歡迎與我們聯絡。\n\n此致\n{sender_name}\n{today}'),
            ('collection', '客户資料收集', '【{company_name}】周年申報所需資料',
             '致 {client_name}：\n\n為辦理貴公司 {company_name} 的周年申報（NAR1），現需向  閣下收集以下資料：\n\n1. 各董事及股東之身份證明文件副本\n2. 最新之註冊辦事處地址證明\n3. 股本結構如有變動之詳情\n\n煩請於 {due_date} 前回覆，以便我們準時辦理。\n\n此致\n{sender_name}\n{today}'),
            ('reminder', '周年申報到期提醒', '【提醒】{company_name} 周年申報將於 {due_date} 到期',
             '致 {client_name}：\n\n謹此提醒，貴公司 {company_name}（商業登記號碼：{br_number}）的周年申報表（NAR1）將於 {due_date} 到期。\n\n請盡快與我們聯絡以安排辦理，避免逾期罰款。\n\n此致\n{sender_name}\n{today}'),
        ]
        for ttype, name, subject, body in seeds:
            db.execute(
                "INSERT INTO email_templates (id, name, template_type, subject, body, is_default) VALUES (?, ?, ?, ?, ?, 1)",
                (str(uuid.uuid4()), name, ttype, subject, body))
        db.commit()
        print(f"[MIGRATE] Seeded {len(seeds)} default email templates")

    # Seed sample email logs for EM-04/05/06 screenshots (if none exist)
    try:
        log_cnt = db.execute("SELECT COUNT(*) FROM email_logs").fetchone()[0]
    except sqlite3.OperationalError:
        log_cnt = None
    if log_cnt == 0:
        # Get first company & template ids
        company = db.execute("SELECT id, name, email FROM companies LIMIT 1").fetchone()
        cid = company[0] if company else None
        templates = {r[0]: r[1] for r in db.execute("SELECT template_type, id FROM email_templates").fetchall()}
        now = datetime.now(timezone.utc).isoformat()
        # EM-04: sent invoice email
        if cid and 'invoice' in templates:
            db.execute(
                "INSERT INTO email_logs (id, company_id, template_id, to_email, subject, body, status, sent_at, email_type) "
                "VALUES (?, ?, ?, ?, ?, ?, 'sent', ?, 'invoice')",
                (str(uuid.uuid4()), cid, templates['invoice'], 'client@example.com',
                 '【PAUL TANG AND COMPANY LIMITED】服務發票 INV-2026-0042',
                 '致 Tang Siu Fan：\n\n茲附上貴公司 PAUL TANG AND COMPANY LIMITED（商業登記號碼：07281051）的服務發票。\n\n發票編號：INV-2026-0042\n金額：HKD 8,500.00\n到期日：2026/07/15\n\n請於到期日前安排付款，如有查詢歡迎與我們聯絡。\n\n此致\nMuse Labs 公司秘書\n2026/07/05',
                 now))
        # EM-05: sent collection email
        if cid and 'collection' in templates:
            db.execute(
                "INSERT INTO email_logs (id, company_id, template_id, to_email, subject, body, status, sent_at, email_type) "
                "VALUES (?, ?, ?, ?, ?, ?, 'sent', ?, 'collection')",
                (str(uuid.uuid4()), cid, templates['collection'], 'client@example.com',
                 '【PAUL TANG AND COMPANY LIMITED】周年申報所需資料',
                 '致 Tang Siu Fan：\n\n為辦理貴公司 PAUL TANG AND COMPANY LIMITED 的周年申報（NAR1），現需向 閣下收集以下資料：\n\n1. 各董事及股東之身份證明文件副本\n2. 最新之註冊辦事處地址證明\n3. 股本結構如有變動之詳情\n\n煩請於 2026/07/20 前回覆，以便我們準時辦理。\n\n此致\nMuse Labs 公司秘書\n2026/07/03',
                 now))
        # EM-06: scheduled reminder (future)
        if cid and 'reminder' in templates:
            future = (datetime.now(timezone.utc) + timedelta(days=5)).isoformat()
            db.execute(
                "INSERT INTO email_logs (id, company_id, template_id, to_email, subject, body, status, scheduled_at, email_type) "
                "VALUES (?, ?, ?, ?, ?, ?, 'scheduled', ?, 'reminder')",
                (str(uuid.uuid4()), cid, templates['reminder'], 'client@example.com',
                 '【提醒】PAUL TANG AND COMPANY LIMITED 周年申報將於 2026/07/31 到期',
                 '致 Tang Siu Fan：\n\n謹此提醒，貴公司 PAUL TANG AND COMPANY LIMITED（商業登記號碼：07281051）的周年申報表（NAR1）將於 2026/07/31 到期。\n\n請盡快與我們聯絡以安排辦理，避免逾期罰款。\n\n此致\nMuse Labs 公司秘書\n2026/07/09',
                 future))
        db.commit()
        after_cnt = db.execute("SELECT COUNT(*) FROM email_logs").fetchone()[0]
        print(f"[MIGRATE] Seeded {after_cnt} sample email logs (invoice/collection/reminder)")

    # Seed baseline (v1) version snapshot for any company that has none (VE-01)
    try:
        company_ids = [r[0] for r in db.execute("SELECT id FROM companies").fetchall()]
        seeded = 0
        for cid in company_ids:
            has = db.execute(
                "SELECT COUNT(*) FROM company_versions WHERE company_id = ?", (cid,)
            ).fetchone()[0]
            if has == 0:
                cols = [c[1] for c in db.execute("PRAGMA table_info(companies)").fetchall()]
                row = db.execute("SELECT * FROM companies WHERE id = ?", (cid,)).fetchone()
                d = dict(zip(cols, row))
                snap = {k: (d.get(k) if d.get(k) is not None else '') for k in VERSION_FIELDS}
                db.execute(
                    "INSERT INTO company_versions (id, company_id, version_no, snapshot, changed_fields, change_summary) "
                    "VALUES (?, ?, 1, ?, '[]', '建立初始版本')",
                    (str(uuid.uuid4()), cid, json.dumps(snap, ensure_ascii=False)))
                seeded += 1
        if seeded:
            db.commit()
            print(f"[MIGRATE] Seeded baseline versions for {seeded} companies")
    except sqlite3.OperationalError as e:
        print(f"[MIGRATE] Baseline version seed skip: {e}")

    db.close()

# ─── Email module ───
# Priority: Resend API > SMTP > simulated
# - RESEND_API_KEY: use Resend API (free 100/day, onboarding@resend.dev)
# - SMTP_HOST: use SMTP (legacy)
# - neither: simulated (log only, no real send)
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
# Fallback: read from local file (Windows env var propagation can be flaky with Flask reloader)
if not RESEND_API_KEY:
    _key_file = os.path.join(os.path.dirname(__file__), '.resend_key')
    if os.path.exists(_key_file):
        with open(_key_file, 'r') as f:
            RESEND_API_KEY = f.read().strip()
SMTP_HOST = os.environ.get('SMTP_HOST', '')
SMTP_PORT = int(os.environ.get('SMTP_PORT', '587'))
SMTP_USER = os.environ.get('SMTP_USER', '')
SMTP_PASS = os.environ.get('SMTP_PASS', '')
SMTP_FROM = os.environ.get('SMTP_FROM', SMTP_USER or 'no-reply@muselabs.local')
SMTP_FROM_NAME = os.environ.get('SMTP_FROM_NAME', 'Muse Labs 公司秘書')


def substitute_vars(text, variables):
    """Replace {key} placeholders with values (unknown placeholders left intact)."""
    if not text:
        return ''
    def repl(m):
        key = m.group(1)
        return str(variables.get(key, m.group(0)))
    return re.sub(r'\{([a-zA-Z_][a-zA-Z0-9_]*)\}', repl, text)


def deliver_email(to_email, subject, body, cc_email=''):
    """Send an email via Resend API > SMTP > simulated.
    Returns (ok: bool, error: str, simulated: bool)."""
    if not to_email:
        return False, 'Missing recipient', False

    # ── Resend API (preferred) ──
    if RESEND_API_KEY:
        try:
            payload = {
                'from': f'{SMTP_FROM_NAME} <onboarding@resend.dev>',
                'to': [to_email],
                'subject': subject,
                'html': body.replace('\n', '<br>'),
            }
            if cc_email:
                payload['cc'] = [cc_email]
            resp = requests.post(
                'https://api.resend.com/emails',
                json=payload,
                headers={
                    'Authorization': f'Bearer {RESEND_API_KEY}',
                    'Content-Type': 'application/json',
                },
                timeout=15,
            )
            if resp.ok:
                print(f"[EMAIL:RESEND] to={to_email} subject={subject!r} id={resp.json().get('id','?')}")
                return True, '', False
            else:
                err = f'Resend HTTP {resp.status_code}: {resp.text[:300]}'
                print(f"[EMAIL:FAILED] to={to_email}: {err}")
                return False, err, False
        except Exception as e:
            print(f"[EMAIL:FAILED] to={to_email}: {e}")
            return False, str(e), False

    # ── SMTP (legacy) ──
    if SMTP_HOST:
        try:
            msg = MIMEMultipart()
            msg['From'] = formataddr((SMTP_FROM_NAME, SMTP_FROM))
            msg['To'] = to_email
            if cc_email:
                msg['Cc'] = cc_email
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            recipients = [e.strip() for e in (to_email + ',' + cc_email).split(',') if e.strip()]
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=20) as s:
                s.starttls()
                if SMTP_USER:
                    s.login(SMTP_USER, SMTP_PASS)
                s.sendmail(SMTP_FROM, recipients, msg.as_string())
            print(f"[EMAIL:SMTP] to={to_email} subject={subject!r}")
            return True, '', False
        except Exception as e:
            print(f"[EMAIL:FAILED] to={to_email}: {e}")
            return False, str(e), False

    # ── Simulated (no backend configured) ──
    print(f"[EMAIL:SIMULATED] to={to_email} cc={cc_email} subject={subject!r}")
    return True, '', True


def process_scheduled_emails():
    """Send any scheduled emails whose scheduled_at has passed. Called by the
    background scheduler thread. Uses its own connection (runs off-request)."""
    db = sqlite3.connect(DB_PATH)
    db.row_factory = sqlite3.Row
    try:
        now = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%S')
        due = db.execute(
            "SELECT * FROM email_logs WHERE status = 'scheduled' AND scheduled_at IS NOT NULL "
            "AND scheduled_at <= ? ORDER BY scheduled_at ASC LIMIT 50",
            (now,)
        ).fetchall()
        for row in due:
            ok, err, _sim = deliver_email(row['to_email'], row['subject'], row['body'], row['cc_email'] or '')
            db.execute(
                "UPDATE email_logs SET status = ?, sent_at = ?, error = ?, updated_at = datetime('now') WHERE id = ?",
                ('sent' if ok else 'failed', datetime.now(timezone.utc).isoformat() if ok else None, err, row['id'])
            )
            db.commit()
            print(f"[SCHEDULER] Processed scheduled email {row['id']} -> {'sent' if ok else 'failed'}")
    except sqlite3.OperationalError as e:
        print(f"[SCHEDULER] skip: {e}")
    finally:
        db.close()


def scheduler_loop():
    while True:
        try:
            process_scheduled_emails()
        except Exception as e:
            print(f"[SCHEDULER] error: {e}")
        time.sleep(60)


# ─── Routes ───
TABLES = ['companies', 'officers', 'shareholders', 'persons', 'person_company_roles',
          'presenters', 'significant_controllers', 'company_logs', 'reminders', 'invoices',
          'resolutions', 'secretary_templates', 'share_transactions', 'user_roles',
          'email_templates', 'email_logs', 'company_versions',
          'change_events', 'nar1_filings', 'form_linkages']

@app.route('/api/send-email', methods=['POST', 'OPTIONS'])
def send_email():
    """Send (or schedule) an email. Body: {to, cc, subject, body, company_id,
    template_id, scheduled_at, variables}. If `variables` is provided, {key}
    placeholders in subject/body are substituted server-side. If `scheduled_at`
    is a future ISO timestamp, the email is queued (status='scheduled') and sent
    by the background scheduler; otherwise it is delivered immediately."""
    if request.method == 'OPTIONS':
        return ('', 204)
    data = request.json or {}
    to_email = (data.get('to') or '').strip()
    cc_email = (data.get('cc') or '').strip()
    subject = data.get('subject') or ''
    body = data.get('body') or ''
    variables = data.get('variables') or {}
    if variables:
        subject = substitute_vars(subject, variables)
        body = substitute_vars(body, variables)
    if not to_email:
        return jsonify({'error': '缺少收件人 (to)'}), 400

    scheduled_at = (data.get('scheduled_at') or '').strip()
    log_id = str(uuid.uuid4())
    db = get_db()

    # Lookup template_type for the log record (EM-04/05)
    email_type = 'general'
    template_id = data.get('template_id')
    if template_id:
        trow = db.execute("SELECT template_type FROM email_templates WHERE id = ?", (template_id,)).fetchone()
        if trow:
            email_type = trow[0]

    # Determine if this is a future-dated (scheduled) send.
    is_scheduled = False
    if scheduled_at:
        try:
            when = datetime.fromisoformat(scheduled_at.replace('Z', ''))
            if when.tzinfo is None:
                when = when.replace(tzinfo=timezone.utc)
            is_scheduled = when > datetime.now(timezone.utc)
        except (ValueError, TypeError):
            is_scheduled = False

    if is_scheduled:
        db.execute(
            "INSERT INTO email_logs (id, company_id, template_id, to_email, cc_email, subject, body, status, scheduled_at, email_type) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, 'scheduled', ?, ?)",
            (log_id, data.get('company_id'), data.get('template_id'), to_email, cc_email, subject, body, scheduled_at, email_type)
        )
        db.commit()
        return jsonify({'success': True, 'id': log_id, 'status': 'scheduled', 'scheduled_at': scheduled_at})

    ok, err, simulated = deliver_email(to_email, subject, body, cc_email)
    status = 'sent' if ok else 'failed'
    db.execute(
        "INSERT INTO email_logs (id, company_id, template_id, to_email, cc_email, subject, body, status, sent_at, error, email_type) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (log_id, data.get('company_id'), data.get('template_id'), to_email, cc_email, subject, body,
         status, datetime.now(timezone.utc).isoformat() if ok else None, err, email_type)
    )
    db.commit()
    return jsonify({'success': ok, 'id': log_id, 'status': status, 'simulated': simulated, 'error': err}), (200 if ok else 502)

@app.route('/api/auth/register', methods=['POST'])
def auth_register():
    data = request.json
    email = (data.get('email', '')).lower().strip()
    password = data.get('password', '')
    display_name = data.get('display_name', email)
    role = data.get('role', 'user')

    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400

    db = get_db()
    existing = db.execute("SELECT id FROM auth_users WHERE email = ?", (email,)).fetchone()
    if existing:
        return jsonify({'error': 'Email already exists'}), 409

    uid = str(uuid.uuid4())
    pwd_hash = hash_password(password)
    db.execute("INSERT INTO auth_users (id, email, password_hash, display_name) VALUES (?, ?, ?, ?)",
               (uid, email, pwd_hash, display_name))
    if role == 'admin':
        db.execute("INSERT INTO user_roles (id, user_id, role) VALUES (?, ?, ?)",
                   (str(uuid.uuid4()), uid, 'admin'))
    db.commit()
    return jsonify({'id': uid, 'email': email, 'display_name': display_name}), 201

@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    data = request.json
    email = (data.get('email', '')).lower().strip()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400

    db = get_db()
    user = db.execute("SELECT * FROM auth_users WHERE email = ?", (email,)).fetchone()
    if not user or not verify_password(password, user['password_hash']):
        return jsonify({'error': 'Invalid email or password'}), 401

    if 'is_active' in user.keys() and user['is_active'] == 0:
        return jsonify({'error': 'Account is deactivated'}), 403

    roles = [r['role'] for r in db.execute(
        "SELECT role FROM user_roles WHERE user_id = ?", (user['id'],)).fetchall()]
    role = 'admin' if 'admin' in roles else ('moderator' if 'moderator' in roles else 'user')

    token = sign_jwt({'sub': user['id'], 'email': user['email'],
                      'display_name': user['display_name'], 'role': role})
    return jsonify({
        'token': token,
        'user': {'id': user['id'], 'email': user['email'],
                 'display_name': user['display_name'], 'role': role}
    })

@app.route('/api/auth/me', methods=['GET'])
def auth_me():
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401
    return jsonify(u)


@app.route('/api/auth/change-password', methods=['POST'])
def auth_change_password():
    """修改密码 — 与云端 /api/auth/change-password 100% 对齐"""
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401
    data = request.json or {}
    current_password = data.get('current_password', '')
    new_password = data.get('new_password', '')
    if not current_password or not new_password:
        return jsonify({'error': 'Current and new password required'}), 400
    db = get_db()
    row = db.execute("SELECT password_hash FROM auth_users WHERE id = ?", (u['id'],)).fetchone()
    if not row or not verify_password(current_password, row['password_hash']):
        return jsonify({'error': 'Current password is incorrect'}), 401
    db.execute("UPDATE auth_users SET password_hash = ? WHERE id = ?",
               (hash_password(new_password), u['id']))
    db.commit()
    return jsonify({'success': True})


# ─── 用户管理（admin，10.1–10.3；本地 dev 模式不強制鑑權）───
VALID_ROLES = ('admin', 'moderator', 'user')


@app.route('/api/admin/users', methods=['GET', 'OPTIONS'])
def admin_users_list():
    if request.method == 'OPTIONS':
        return ('', 204)
    db = get_db()
    users = db.execute(
        "SELECT id, email, display_name, is_active, created_at FROM auth_users ORDER BY created_at"
    ).fetchall()
    roles = db.execute("SELECT user_id, role FROM user_roles").fetchall()
    role_map = {}
    for r in roles:
        role_map.setdefault(r['user_id'], []).append(r['role'])
    out = []
    for u in users:
        d = dict(u)
        d['roles'] = role_map.get(u['id'], [])
        out.append(d)
    return jsonify(out)


@app.route('/api/admin/users', methods=['POST'])
def admin_users_create():
    data = request.json or {}
    email = (data.get('email', '')).lower().strip()
    password = data.get('password', '')
    display_name = data.get('display_name') or email
    role = data.get('role', 'user')
    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400
    db = get_db()
    if db.execute("SELECT id FROM auth_users WHERE email = ?", (email,)).fetchone():
        return jsonify({'error': 'Email already exists'}), 409
    uid = str(uuid.uuid4())
    db.execute(
        "INSERT INTO auth_users (id, email, password_hash, display_name, is_active, created_at) VALUES (?, ?, ?, ?, 1, ?)",
        (uid, email, hash_password(password), display_name, datetime.now().isoformat()))
    if role in VALID_ROLES:
        db.execute("INSERT INTO user_roles (id, user_id, role) VALUES (?, ?, ?)",
                   (str(uuid.uuid4()), uid, role))
    db.commit()
    return jsonify({'id': uid, 'email': email, 'display_name': display_name,
                    'roles': [role] if role in VALID_ROLES else [], 'is_active': 1}), 201


@app.route('/api/admin/users/<uid>', methods=['PUT'])
def admin_users_update(uid):
    data = request.json or {}
    db = get_db()
    if not db.execute("SELECT id FROM auth_users WHERE id = ?", (uid,)).fetchone():
        return jsonify({'error': 'User not found'}), 404
    if 'is_active' in data:
        db.execute("UPDATE auth_users SET is_active = ? WHERE id = ?",
                   (1 if data['is_active'] else 0, uid))
    if 'display_name' in data:
        db.execute("UPDATE auth_users SET display_name = ? WHERE id = ?", (data['display_name'], uid))
    if 'password' in data and data['password']:
        db.execute("UPDATE auth_users SET password_hash = ? WHERE id = ?",
                   (hash_password(data['password']), uid))
    if 'role' in data:
        db.execute("DELETE FROM user_roles WHERE user_id = ?", (uid,))
        if data['role'] in VALID_ROLES:
            db.execute("INSERT INTO user_roles (id, user_id, role) VALUES (?, ?, ?)",
                       (str(uuid.uuid4()), uid, data['role']))
    db.commit()
    return jsonify({'success': True})


@app.route('/api/admin/users/<uid>', methods=['DELETE'])
def admin_users_delete(uid):
    db = get_db()
    db.execute("DELETE FROM user_roles WHERE user_id = ?", (uid,))
    db.execute("DELETE FROM auth_users WHERE id = ?", (uid,))
    db.commit()
    return jsonify({'success': True})

# ─── Register PDF helpers (matching Paul Tang RTF samples) ───
MARGIN = 28  # ~1cm, matching RTF sample margins
PAGE_H = 842
PAGE_W = 595
CONTENT_W = PAGE_W - MARGIN * 2
BLUE = (0, 0, 255)
GREY_HDR = (227, 227, 227)


def pdf_draw_right(pdf, text, x, y, size=10, color=None):
    """Draw text right-aligned at given right-edge x position."""
    pdf.set_font('TC', size=size)
    if color:
        pdf.set_text_color(*color)
    else:
        pdf.set_text_color(0, 0, 0)
    tw = pdf.get_string_width(str(text or ''))
    pdf.set_xy(x - tw, y)
    pdf.cell(tw, size + 2, str(text or ''))


def pdf_line_h(pdf, x1, x2, y, color=(180, 180, 180), width=0.2):
    """Draw a horizontal line."""
    pdf.set_draw_color(*color)
    pdf.set_line_width(width)
    pdf.line(x1, y, x2, y)
    pdf.set_line_width(0.2)


def pdf_rect_fill(pdf, x, y, w, h, color):
    """Draw a filled rectangle."""
    pdf.set_fill_color(*color)
    pdf.rect(x, y, w, h, style='F')


def pdf_wrap_text(pdf, text, width, size=7):
    """Split text into lines that fit within width. Handles \\n and CJK."""
    if not text:
        return ['']
    pdf.set_font('TC', size=size)
    result = []
    for paragraph in str(text).split('\n'):
        if not paragraph:
            result.append('')
            continue
        current = ''
        for ch in paragraph:
            trial = current + ch
            if pdf.get_string_width(trial) > width and current:
                result.append(current)
                current = ch
            else:
                current = trial
        if current:
            result.append(current)
    if not result:
        result.append('')
    return result


def draw_register_header(pdf, company, title_en, quorum=None):
    """Draw the standard register page header — Paul Tang black & white style.
    Returns y position after header + separator line."""
    y = 45

    # Company name — TNR bold, black
    co_name = rget(company, 'name') or ''
    pdf.set_font('TNR', 'B', 14)
    pdf.set_text_color(0, 0, 0)
    tw = pdf.get_string_width(co_name)
    pdf.set_xy(PAGE_W / 2 - tw / 2, y)
    pdf.cell(tw, 16, co_name)
    y += 19

    br = rget(company, 'company_number')
    cn_line = f"Company Number:  {br}" if br else "Company Number:"
    pdf.set_font('TNR', '', 10)
    tw2 = pdf.get_string_width(cn_line)
    pdf.set_xy(PAGE_W / 2 - tw2 / 2, y)
    pdf.cell(tw2, 12, cn_line)

    if quorum is not None:
        pdf.set_font('TNR', '', 9)
        q_text = f"Quorum:  {quorum}"
        qw = pdf.get_string_width(q_text)
        pdf.set_xy(PAGE_W - MARGIN - qw, y)
        pdf.cell(qw, 12, q_text)
    y += 20

    today = datetime.now().strftime('%d/%m/%Y')
    title_full = f"{title_en} AT {today}"
    pdf.set_font('TNR', 'B', 12)
    pdf.set_text_color(0, 0, 0)
    pdf.set_xy(MARGIN, y)
    pdf.cell(0, 14, title_full)
    y += 22

    pdf_line_h(pdf, MARGIN, PAGE_W - MARGIN, y, color=(0, 0, 0), width=0.8)
    return y + 12


def draw_form_label(pdf, label, x_label, y, size=9):
    """Draw a bold form label (like 'Name', 'Address', 'Security')."""
    pdf_draw(pdf, label, x_label, y, size=size, bold=True)


def draw_form_value(pdf, value, x_val, y, size=9):
    """Draw a form value after a label."""
    pdf_draw(pdf, value, x_val, y, size=size)


def draw_grey_header_row(pdf, cols, y):
    """Draw a white-background header row (Paul Tang style). cols = [(text, x, w), ...].
    Returns y below row. Text wraps within column width. No grey fill, no vertical lines."""
    hdr_size = 7
    label_lines = []
    max_lines = 1
    for text, x, w in cols:
        lines = pdf_wrap_text(pdf, text, w - 6, hdr_size)
        label_lines.append(lines)
        max_lines = max(max_lines, len(lines))
    row_h = max(max_lines * 11 + 6, 22)

    # White background — no fill (transparent)
    # Draw header text in black
    pdf.set_text_color(0, 0, 0)
    for i, (text, x, w) in enumerate(cols):
        for li, line_text in enumerate(label_lines[i]):
            pdf.set_xy(x, y + 3 + li * 10)
            # Use TNR for ASCII, TC for CJK
            has_cjk = any('一' <= ch <= '鿿' or '　' <= ch <= '〿' for ch in line_text)
            if has_cjk:
                pdf.set_font('TC', size=hdr_size)
            else:
                pdf.set_font('TNR', size=hdr_size)
            pdf.cell(w, 10, line_text)

    # Top + bottom borders only (black, no vertical dividers)
    pdf.set_draw_color(0, 0, 0)
    pdf.line(MARGIN, y, PAGE_W - MARGIN, y)
    pdf.line(MARGIN, y + row_h, PAGE_W - MARGIN, y + row_h)
    pdf.set_draw_color(0, 0, 0)

    return y + row_h


def draw_data_row(pdf, cols, y, alt=False):
    """Draw a data row (Paul Tang style — no alternating colour, no vertical lines).
    cols = [(text, x, w), ...]. Returns y below row."""
    size = 7
    wrapped = []
    max_lines = 1
    for text, x, w in cols:
        lines = pdf_wrap_text(pdf, str(text or ''), w - 6, size)
        wrapped.append(lines)
        max_lines = max(max_lines, len(lines))
    row_h = max(max_lines * 11 + 6, 20)

    # No alternating background
    pdf.set_text_color(0, 0, 0)
    for i, (text, x, w) in enumerate(cols):
        for li, line_text in enumerate(wrapped[i]):
            pdf.set_xy(x, y + 3 + li * 10)
            has_cjk = any('一' <= ch <= '鿿' or '　' <= ch <= '〿' for ch in line_text)
            if has_cjk:
                pdf.set_font('TC', size=size)
            else:
                pdf.set_font('TNR', size=size)
            pdf.cell(w, 10, line_text)

    # Thin bottom border only (no vertical lines)
    pdf.set_draw_color(180, 180, 180)
    pdf.line(MARGIN, y + row_h, PAGE_W - MARGIN, y + row_h)
    pdf.set_draw_color(0, 0, 0)

    return y + row_h


# ─── PDF Generation ───

# ROM Share Transaction sub-column layout:
# This matches the 11 sub-columns in the RTF sample's grey header
# Positions calculated proportionally from the RTF twip coordinates
def _rom_txn_cols():
    """Return [(label, x, w), ...] for the ROM share transaction table."""
    x0 = MARGIN + 5
    return [
        ("Date Entered\n/ Ceased",      x0,        57),
        ("Transaction\nType",            x0 + 58,   60),
        ("Units",                        x0 + 119,  75),
        ("Par Value\nPer Share",         x0 + 195,  75),
        ("Paid Up Value\nPer Share",    x0 + 271,  75),
        ("Certificate\nNo",             x0 + 347,  50),
        ("Distinctive\nNumbers",        x0 + 398,  115),
        ("Balance",                     x0 + 514,  55),
        ("Transferred To/From,\nRedeemed, Reissued", x0 + 570, 140),
    ]


# ═══════════════════════════════════════════════════════════════
# ROM DOCX template fill (Paul Tang format)
# ═══════════════════════════════════════════════════════════════

def _fmt_date_rom(val):
    """Normalize a date value to DD/MM/YYYY string for ROM template."""
    if not val or val == '-':
        return '-'
    val_str = str(val).strip()
    # Already DD/MM/YYYY
    if re.match(r'^\d{1,2}/\d{1,2}/\d{4}$', val_str):
        return val_str
    # ISO format YYYY-MM-DD
    if re.match(r'^\d{4}-\d{2}-\d{2}$', val_str):
        try:
            dt = datetime.strptime(val_str, '%Y-%m-%d')
            return dt.strftime('%d/%m/%Y')
        except:
            pass
    # Raw number DDMMYYYY
    if re.match(r'^\d{8}$', val_str):
        try:
            dt = datetime.strptime(val_str, '%d%m%Y')
            return dt.strftime('%d/%m/%Y')
        except:
            pass
        try:
            dt = datetime.strptime(val_str, '%Y%m%d')
            return dt.strftime('%d/%m/%Y')
        except:
            pass
    return val_str


def _fill_rom_docx_template(db, company_id):
    """Fill the Paul Tang ROM DOCX template with company & shareholder data.

    Opens 股東登記冊_PaulTang格式.docx, replaces header placeholders
    with actual data, fills the 18-column table with shareholder rows,
    and saves to a temp .docx file.

    Returns path to temp .docx, or None on failure.
    """
    if not _HAS_DOCX:
        return None
    if not os.path.exists(_ROM_DOCX_TEMPLATE):
        print(f"[ROM] DOCX template not found: {_ROM_DOCX_TEMPLATE}")
        return None

    from copy import deepcopy
    from docx.oxml import OxmlElement

    # ── Load data ──
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return None

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role = 'shareholder'",
        (company_id,)).fetchall()

    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        ph = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({ph})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    txs = db.execute(
        "SELECT * FROM share_transactions WHERE company_id = ? ORDER BY transaction_date",
        (company_id,)).fetchall()

    co_name = rget(company, 'name') or ''
    co_br = rget(company, 'company_number') or ''
    today = datetime.now()
    report_date = today.strftime('%d %B %Y').upper()

    # ── Build tx_by_person map ──
    tx_by_person = {}
    for t in txs:
        key = (rget(t, 'from_name') or rget(t, 'to_name') or '').strip().upper()
        if key:
            tx_by_person.setdefault(key, []).append(t)

    # ── Build shareholder data list ──
    shareholders = []
    for role in roles:
        p = person_map.get(role['person_id'], {})
        name_en = (rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)')[:80]

        addr, region = _get_person_address(db, p['id'])
        if not addr:
            addr = (rget(p, 'address') or '')[:100]
        if region and region not in (addr or ''):
            addr = f"{addr}, {region}".strip(', ')
        addr = (addr or '')[:120]

        occupation = rget(p, 'occupation') or ''
        date_app = _fmt_date_rom(rget(role, 'date_appointed') or '-')
        date_cea = _fmt_date_rom(rget(role, 'date_ceased') or '-')
        shares_held = int(rget(role, 'shares') or 0)
        cert_no = rget(role, 'certificate_number') or '-'
        currency = rget(role, 'currency') or 'HKD'
        issue_price = rget(role, 'issue_price') or '1.00'

        # Calculate total shares including transactions
        person_name_key = name_en.strip().upper()
        person_txs = tx_by_person.get(person_name_key, [])
        total_shares = shares_held
        for tx in person_txs:
            tx_shares = int(rget(tx, 'shares') or 0)
            is_in = (rget(tx, 'to_name') or '').strip().upper() == person_name_key
            is_out = (rget(tx, 'from_name') or '').strip().upper() == person_name_key
            if is_in:
                total_shares += tx_shares
            elif is_out:
                total_shares -= tx_shares

        shareholders.append({
            'full_name': name_en,
            'addr': addr,
            'occupation': occupation,
            'date_app': date_app,
            'date_cea': date_cea if date_cea else '-',
            'shares_held': shares_held,
            'cert_no': cert_no,
            'currency': currency,
            'issue_price': issue_price,
            'total_shares': total_shares,
            'txs': person_txs,
        })

    # ── Open template ──
    doc = DocxDocument(_ROM_DOCX_TEMPLATE)

    # ── Fill header paragraphs ──
    # P[0]: "Name of Company" → actual company name
    p0 = doc.paragraphs[0]
    for run in p0.runs:
        if 'Name of Company' in run.text:
            run.text = run.text.replace('Name of Company', co_name)
            break

    # P[2]: "Company Number  ...  REGISTER OF MEMBERS"
    p2 = doc.paragraphs[2]
    for run in p2.runs:
        if 'Company Number' in run.text:
            run.text = f'Company Number: {co_br}'
            break
    for run in p2.runs:
        if 'REGISTER OF MEMBERS' in run.text:
            run.text = f'REGISTER OF MEMBERS as at {report_date}'
            break

    # ── Fill main table (Table[3] = 18-col data table) ──
    # Tables 0-2 are placeholder/dummy tables in the template
    table = doc.tables[3]

    # Save the first data row's XML before deleting it (for cloning)
    template_tr_xml = deepcopy(table.rows[2]._tr) if len(table.rows) > 2 else None

    # Remove all sample data rows (rows 2, 3, 4… from end to preserve indices)
    for i in range(len(table.rows) - 1, 1, -1):
        table._tbl.remove(table.rows[i]._tr)

    if not shareholders:
        # Add a "no shareholders" row
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run('(No shareholders / 尚無股東記錄)')
        _set_docx_run_font(row.cells[0].paragraphs[0].runs[0], Pt(6))
    else:
        for si, sh in enumerate(shareholders):
            # Clone the template row or reuse for first shareholder
            if template_tr_xml is not None:
                new_tr = template_tr_xml if si == 0 else deepcopy(template_tr_xml)
                table._tbl.append(new_tr)
            else:
                table.add_row()

            row = table.rows[-1]
            cells = row.cells

            # 18-column data map: (col_index, value)
            cell_data = [
                (0,  sh['full_name']),                              # Full Name
                (1,  sh['addr']),                                    # Address
                (2,  sh['occupation']),                              # Occupation
                (3,  ''),                                            # Merchant
                (4,  sh['date_app']),                                # Date Entered as Member
                (5,  sh['date_cea']),                                # Date Ceasing to be Member
                (6,  sh['cert_no']),                                 # Cert No (Acq)
                (7,  '-'),                                           # Distinctive Nos (Acq)
                (8,  str(sh['shares_held'])),                        # No. of Shares (Acq)
                (9,  f"{sh['currency']}${sh['issue_price']}"),       # Consideration Paid (Acq)
                (10, '-'),                                           # Transfer Deed No (Acq)
                (11, '-'),                                           # Cert No (Xfer)
                (12, '-'),                                           # Distinctive Nos (Xfer)
                (13, '-'),                                           # No. of Shares (Xfer)
                (14, '-'),                                           # Consideration Paid (Xfer)
                (15, str(sh['total_shares'])),                       # Total Shares Held
                (16, ''),                                            # Remarks
                (17, ''),                                            # Entry Made By
            ]

            for ci, val in cell_data:
                if ci >= len(cells):
                    continue
                cell = cells[ci]
                # Clear existing content
                for p in cell.paragraphs:
                    p.clear()
                # Add new text
                run = cell.paragraphs[0].add_run(str(val) if val else '')
                run.font.size = Pt(6)
                run.font.name = 'Arial'
                # Set East-Asian font for CJK rendering
                _set_cjk_fallback(run, 'DengXian')

    # ── Save to temp file ──
    tmp_docx = tempfile.mktemp(suffix='.docx')
    doc.save(tmp_docx)
    return tmp_docx


def _set_cjk_fallback(run, font_name):
    """Set East-Asian font fallback on a run element for CJK text."""
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


@app.route('/api/generate-shareholders-register-pdf', methods=['POST'])
def generate_shareholders_register_pdf():
    """Register of Members (ROM) — matching Paul Tang & Co reference format.
    Portrait A4, 18-column table, white headers, Times New Roman, thin black lines."""
    data = request.get_json(silent=True) or {}
    company_id = data.get('companyId', '')
    if not company_id:
        return jsonify({'error': 'companyId required'}), 400

    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return jsonify({'error': 'Company not found'}), 404

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role = 'shareholder'",
        (company_id,)).fetchall()
    txs = db.execute(
        "SELECT * FROM share_transactions WHERE company_id = ? ORDER BY transaction_date",
        (company_id,)).fetchall()

    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        placeholders = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({placeholders})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    tx_by_person = {}
    for t in txs:
        key = (rget(t, 'from_name') or rget(t, 'to_name') or '').strip().upper()
        if key:
            tx_by_person.setdefault(key, []).append(t)

    # ── Helper: cover unused SH2 for single shareholder ──
    def _cover_sh2_if_single(pdf_bytes, n_shareholders):
        if n_shareholders != 1:
            return pdf_bytes
        import fitz as _fitz
        _d = _fitz.open(stream=pdf_bytes, filetype="pdf")
        areas = _d[0].search_for("Full Name")
        if len(areas) >= 2:
            _sh2_y = areas[1].y0 - 8
            _d[0].draw_rect(
                _fitz.Rect(0, _sh2_y, _d[0].rect.width, _d[0].rect.height),
                color=None, fill=(1, 1, 1), width=0)
            # Save to temp file then read back (stream write() unreliable)
            import tempfile as _tf
            _tmp = _tf.mktemp(suffix='.pdf')
            _d.save(_tmp, deflate=True)
            _d.close()
            with open(_tmp, 'rb') as _f:
                pdf_bytes = _f.read()
            try:
                os.unlink(_tmp)
            except:
                pass
        else:
            _d.close()
        return pdf_bytes

    # ── Helper: draw missing bottom border for transaction sub-table ──
    def _fix_tx_table_borders(pdf_bytes):
        """Add missing bottom border and column dividers to transaction sub-tables.
        Uses PyMuPDF to search for the grey-header transaction table area
        and draws proper borders below the data rows."""
        import fitz as _f2
        _d = _f2.open(stream=pdf_bytes, filetype="pdf")
        for pi in range(_d.page_count):
            page = _d[pi]
            # Search for the tx table grey-header text "CertNo" to locate table area
            areas = page.search_for("CertNo")
            if not areas:
                continue
            # For each tx sub-table found
            for area in areas:
                tx_top = area.y0  # top of tx sub-table header
                # Find the bottom of the data area: search for the data row below header
                # Look for text below tx_top+10 in the same x-range
                data_bottom = tx_top + 38  # default: 3 rows × ~13pt
                # Find all text blocks below the header to determine actual bottom
                blocks = page.get_text("dict")["blocks"]
                for b in blocks:
                    if "lines" not in b:
                        continue
                    by0 = b["bbox"][1]
                    if by0 > tx_top + 10 and by0 < tx_top + 80:
                        # This is a data row — update bottom
                        data_bottom = max(data_bottom, b["bbox"][3] + 6)
                # Draw bottom border line
                page.draw_line(
                    _f2.Point(28, data_bottom),
                    _f2.Point(567, data_bottom),
                    color=(0, 0, 0), width=0.5)
                # Draw column divider lines from header bottom to bottom border
                # Find header column positions from grey-header text
                hdr_text = page.get_textbox(area)
                # Get all text spans in the header row for column positions
                hdr_blocks = page.get_text("dict", clip=_f2.Rect(28, tx_top - 2, 567, tx_top + 20))["blocks"]
                col_xs = set()
                for b2 in hdr_blocks:
                    if "lines" not in b2:
                        continue
                    for line in b2["lines"]:
                        for span in line["spans"]:
                            # Record left edge of each text span as column start
                            col_xs.add(round(span["bbox"][0]))
                # Add right edge too
                col_xs.add(567)
                # Draw vertical lines at column boundaries
                prev_x = 28
                for cx in sorted(col_xs):
                    if cx > prev_x + 15:  # min column width
                        page.draw_line(
                            _f2.Point(prev_x, tx_top),
                            _f2.Point(prev_x, data_bottom),
                            color=(0.5, 0.5, 0.5), width=0.2)
                        prev_x = cx
                    elif cx > prev_x:
                        prev_x = cx
        # Save to temp file then read back
        import tempfile as _tf
        _tmp = _tf.mktemp(suffix=".pdf")
        _d.save(_tmp, deflate=True)
        _d.close()
        with open(_tmp, "rb") as _f:
            pdf_bytes = _f.read()
        try:
            os.unlink(_tmp)
        except:
            pass
        return pdf_bytes

    # ── Path 1 (NEW): DOCX template fill → Word COM → PDF (Paul Tang format) ──
    if _HAS_WORD_COM and _HAS_DOCX:
        docx_path = _fill_rom_docx_template(db, company_id)
        if docx_path:
            pdf_bytes = _docx_to_pdf_via_word(docx_path)
            try:
                os.unlink(docx_path)
            except:
                pass
            if pdf_bytes:
                return jsonify({'pdf': base64.b64encode(pdf_bytes).decode('ascii')})

    # ── Path 2: Old DOC template → Word COM Find & Replace → PDF ──
    if _HAS_WORD_COM:
        pdf_bytes = _rtf_rom_to_pdf(db, company_id)
        if pdf_bytes:
            pdf_bytes = _cover_sh2_if_single(pdf_bytes, len(roles))
            pdf_bytes = _fix_tx_table_borders(pdf_bytes)
            return jsonify({'pdf': base64.b64encode(pdf_bytes).decode('ascii')})

    # ── Path 3: DOCX built from scratch → Word COM → PDF ──
    if _HAS_WORD_COM and _HAS_DOCX:
        docx_path = _build_rom_register_docx(db, company_id)
        if docx_path:
            pdf_bytes = _docx_to_pdf_via_word(docx_path)
            try:
                os.unlink(docx_path)
            except:
                pass
            if pdf_bytes:
                pdf_bytes = _cover_sh2_if_single(pdf_bytes, len(roles))
                pdf_bytes = _fix_tx_table_borders(pdf_bytes)
                return jsonify({'pdf': base64.b64encode(pdf_bytes).decode('ascii')})

    # ── Path 4: fpdf2 (no external dependencies) ──
    M = 28  # margin
    PW, PH = 595, 842  # A4 portrait in pt
    CW = PW - 2 * M  # content width

    pdf = create_pdf()
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    # ── Helper: draw text with TNR (English) or TC (CJK) ──
    def tnr(text, x, y, size=9, bold=False, align='L', color=(0, 0, 0)):
        """Draw text with Times New Roman font."""
        style = 'B' if bold else ''
        pdf.set_font('TNR', style, size)
        pdf.set_text_color(*color)
        if align == 'C':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw / 2
        elif align == 'R':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw
        pdf.set_xy(x, y)
        pdf.cell(0, size + 2, str(text or ''))

    def tc(text, x, y, size=9, bold=False, align='L', color=(0, 0, 0)):
        """Draw text with CJK font."""
        style = 'B' if bold else ''
        pdf.set_font('TC', style, size)
        pdf.set_text_color(*color)
        if align == 'C':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw / 2
        elif align == 'R':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw
        pdf.set_xy(x, y)
        pdf.cell(0, size + 2, str(text or ''))

    def line_h(x1, x2, y, w=0.3, color=(0, 0, 0)):
        pdf.set_draw_color(*color)
        pdf.set_line_width(w)
        pdf.line(x1, y, x2, y)

    def line_v(x, y1, y2, w=0.3, color=(0, 0, 0)):
        pdf.set_draw_color(*color)
        pdf.set_line_width(w)
        pdf.line(x, y1, x, y2)

    # ── Page Header ──
    y = 35
    co_name = rget(company, 'name') or ''
    tnr(co_name, PW / 2, y, size=12, bold=True, align='C')
    y += 18
    br = rget(company, 'company_number') or ''
    tnr(f"Company Number: {br}", PW / 2, y, size=9, align='C')
    y += 2

    # "REGISTER OF MEMBERS" right-aligned
    tnr("REGISTER OF MEMBERS", PW - M, 35, size=13, bold=True, align='R')

    y += 14
    line_h(M, PW - M, y, w=0.4)

    # ── Table columns (18 cols, based on Paul Tang reference) ──
    # Adjust widths for portrait A4 (CW ≈ 539pt)
    col_w = [
        54,  # 1. Full Name
        60,  # 2. Address
        37,  # 3. Occupation
        37,  # 4. Merchant
        40,  # 5. Date Entered
        40,  # 6. Date Ceasing
        # Shares Acquired (5 sub-cols):
        28,  # 7. Cert No
        46,  # 8. Distinctive Nos
        28,  # 9. No. of Shares
        30,  # 10. Consideration
        30,  # 11. Transfer Deed No
        # Shares Transferred (4 sub-cols):
        28,  # 12. Cert No
        46,  # 13. Distinctive Nos
        28,  # 14. No. of Shares
        30,  # 15. Consideration
        30,  # 16. Total Shares Held
        32,  # 17. Remarks
        32,  # 18. Entry Made By
    ]
    # Build x positions
    col_x = [M]
    for w in col_w[:-1]:
        col_x.append(col_x[-1] + w)

    def draw_table_border(top_y, bottom_y):
        """Draw outer border and vertical lines for the table."""
        line_h(M, PW - M, top_y, w=0.4)
        line_h(M, PW - M, bottom_y, w=0.4)
        line_v(M, top_y, bottom_y, w=0.4)
        line_v(PW - M, top_y, bottom_y, w=0.4)

    def draw_row_sep(row_y, start_col=0):
        """Draw horizontal line across row, and vertical lines between cols."""
        end_x = PW - M
        line_h(col_x[start_col], end_x, row_y, w=0.25, color=(100, 100, 100))
        for i in range(start_col, len(col_x)):
            line_v(col_x[i], row_y - 50, row_y, w=0.2, color=(150, 150, 150))

    # ── Table Header Row 1: Main headers ──
    y += 8
    hdr_y0 = y
    hdr1_labels = [
        ("Full\nName", 0), ("Address", 1), ("Occupation", 2), ("Merchant", 3),
        ("Date Entered\nas Member", 4), ("Date Ceasing\nto be Member", 5),
    ]
    # Draw hdr row 1 labels
    hdr_size = 6.5
    row_h1 = 26
    for label, ci in hdr1_labels:
        cx = col_x[ci]
        cw_val = col_w[ci]
        tnr(label, cx + 1, y + 1, size=hdr_size, bold=True)
    # "Shares Acquired" spanning cols 6-10
    sacq_x1 = col_x[6]
    sacq_x2 = col_x[10] + col_w[10]
    sacq_cx = (sacq_x1 + sacq_x2) / 2
    tnr("Shares Acquired", sacq_cx, y + 1, size=hdr_size, bold=True, align='C')
    # "Shares Transferred" spanning cols 11-14
    strf_x1 = col_x[11]
    strf_x2 = col_x[14] + col_w[14]
    strf_cx = (strf_x1 + strf_x2) / 2
    tnr("Shares Transferred", strf_cx, y + 1, size=hdr_size, bold=True, align='C')
    # Remaining headers
    for label, ci in [("Total\nShares Held", 15), ("Remarks", 16), ("Entry\nMade By", 17)]:
        cx = col_x[ci]
        tnr(label, cx + 1, y + 1, size=hdr_size, bold=True)

    y += row_h1
    line_h(M, PW - M, y, w=0.3)
    # Vertical separators for header row 1
    for ci in [0, 1, 2, 3, 4, 5, 6, 11, 15, 16, 17]:
        line_v(col_x[ci], hdr_y0, y, w=0.2, color=(150, 150, 150))
    line_v(col_x[10] + col_w[10], hdr_y0, y, w=0.2, color=(150, 150, 150))
    line_v(col_x[14] + col_w[14], hdr_y0, y, w=0.2, color=(150, 150, 150))

    # ── Table Header Row 2: Sub-columns under Shares Acquired / Transferred ──
    hdr_y1 = y
    row_h2 = 22
    sub_labels = [
        ("Cert\nNo", 6), ("Distinctive Nos.\n(From → To)", 7),
        ("No. of\nShares", 8), ("Consideration\nPaid", 9), ("Transfer\nDeed No", 10),
        ("Cert\nNo", 11), ("Distinctive Nos.\n(From → To)", 12),
        ("No. of\nShares", 13), ("Consideration\nPaid", 14),
    ]
    for label, ci in sub_labels:
        tnr(label, col_x[ci] + 1, y + 1, size=6, bold=True)
    y += row_h2
    line_h(M, PW - M, y, w=0.4)

    table_top_y = hdr_y0
    table_bottom_y = y

    # ── Data Rows ──
    row_num = 0
    data_size = 6.5
    data_row_h = 16

    if not roles:
        tc("(No shareholders / 尚無股東記錄)", M + 3, y + 8, size=9)
        y += 24
        table_bottom_y = y
    else:
        for r in roles:
            p = person_map.get(r['person_id'], {})
            name_en = rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)'
            is_nat = rget(p, 'identity') != 'corporate'
            hkid = rget(p, 'id_number') or rget(p, 'passport_number') or ''
            addr = rget(p, 'address') or ''
            occupation = rget(p, 'occupation') or ''
            merchant = rget(p, 'merchant') or ''

            date_app = rget(r, 'date_appointed') or '-'
            date_cea = rget(r, 'date_ceased') or ''

            share_type = rget(r, 'share_type') or 'ORD'
            currency = rget(r, 'currency') or 'HKD'
            issue_price = rget(r, 'issue_price') or '1.00'
            shares = rget(r, 'shares') or 0
            cert_no = rget(r, 'certificate_number') or '-'

            person_name_key = name_en.strip().upper()
            person_txs = tx_by_person.get(person_name_key, [])

            # Page break if needed (each member ~3-5 data rows)
            if y + 80 > PH - 50:
                pdf.add_page()
                y = 35
                tnr(co_name, PW / 2, y, size=12, bold=True, align='C')
                y += 18
                tnr(f"Company Number: {br}", PW / 2, y, size=9, align='C')
                tnr("REGISTER OF MEMBERS (Cont'd)", PW - M, 35, size=13, bold=True, align='R')
                y += 14
                line_h(M, PW - M, y, w=0.4)
                y += 8
                # Redraw headers (simplified)
                hdr_y0 = y
                for label, ci in hdr1_labels:
                    tnr(label.replace('\n', ' '), col_x[ci] + 1, y + 1, size=hdr_size, bold=True)
                tnr("Shares Acquired", sacq_cx, y + 1, size=hdr_size, bold=True, align='C')
                tnr("Shares Transferred", strf_cx, y + 1, size=hdr_size, bold=True, align='C')
                for label, ci in [("Total\nShares Held", 15), ("Remarks", 16), ("Entry\nMade By", 17)]:
                    tnr(label, col_x[ci] + 1, y + 1, size=hdr_size, bold=True)
                y += row_h1
                line_h(M, PW - M, y, w=0.3)
                for label, ci in sub_labels:
                    tnr(label.replace('\n', ' '), col_x[ci] + 1, y + 1, size=6, bold=True)
                y += row_h2
                line_h(M, PW - M, y, w=0.4)
                table_top_y = hdr_y0

            row_num += 1
            row_start_y = y

            if not person_txs:
                # Single row with all data
                row_data = [
                    name_en, addr, occupation, merchant, date_app, date_cea,
                    cert_no, '-', str(shares), f"{currency}${issue_price}", '-',
                    '-', '-', '-', '-',
                    str(shares), '', ''
                ]
                for i, val in enumerate(row_data):
                    tnr(str(val)[:30], col_x[i] + 1, y + 1, size=data_size)
                y += data_row_h
            else:
                # First row: member info + first transaction
                running_balance = 0
                first_tx = person_txs[0]
                tx_type = rget(first_tx, 'transaction_type') or 'Transfer'
                tx_date = rget(first_tx, 'transaction_date') or '-'
                tx_shares = int(first_tx['shares'] or 0)
                tx_price = rget(first_tx, 'price_per_share') or issue_price
                tx_currency = rget(first_tx, 'currency') or currency
                tx_from = rget(first_tx, 'from_name') or ''
                tx_to = rget(first_tx, 'to_name') or ''
                tx_inst = rget(first_tx, 'instrument_number') or '-'
                tx_cert = rget(first_tx, 'certificate_number') or cert_no

                is_in = tx_to.strip().upper() == person_name_key
                is_out = tx_from.strip().upper() == person_name_key
                if is_in:
                    running_balance = tx_shares
                elif is_out:
                    running_balance = -tx_shares
                else:
                    running_balance = tx_shares

                first_row = [
                    name_en, addr, occupation, merchant, date_app, date_cea,
                    tx_cert, f"{tx_inst}", str(tx_shares), f"{tx_currency}${tx_price}", tx_inst,
                    '-', '-', '-', '-',
                    str(running_balance), '', ''
                ]
                for i, val in enumerate(first_row):
                    tnr(str(val)[:30], col_x[i] + 1, y + 1, size=data_size)
                y += data_row_h

                # Subsequent transaction rows
                for tx in person_txs[1:]:
                    if y + data_row_h > PH - 50:
                        pdf.add_page()
                        y = 35
                        tnr(co_name, PW / 2, y, size=12, bold=True, align='C')
                        tnr("REGISTER OF MEMBERS (Cont'd)", PW - M, 35, size=13, bold=True, align='R')
                        y += 32
                        line_h(M, PW - M, y, w=0.4)
                        y += 8
                        hdr_y0 = y
                        for label, ci in hdr1_labels:
                            tnr(label.replace('\n', ' '), col_x[ci] + 1, y + 1, size=hdr_size, bold=True)
                        tnr("Shares Acquired", sacq_cx, y + 1, size=hdr_size, bold=True, align='C')
                        tnr("Shares Transferred", strf_cx, y + 1, size=hdr_size, bold=True, align='C')
                        for label, ci in [("Total\nShares Held", 15), ("Remarks", 16), ("Entry\nMade By", 17)]:
                            tnr(label, col_x[ci] + 1, y + 1, size=hdr_size, bold=True)
                        y += row_h1
                        line_h(M, PW - M, y, w=0.3)
                        for label, ci in sub_labels:
                            tnr(label.replace('\n', ' '), col_x[ci] + 1, y + 1, size=6, bold=True)
                        y += row_h2
                        line_h(M, PW - M, y, w=0.4)

                    tx_type = rget(tx, 'transaction_type') or 'Transfer'
                    tx_date = rget(tx, 'transaction_date') or '-'
                    tx_shares = int(tx['shares'] or 0)
                    tx_price = rget(tx, 'price_per_share') or issue_price
                    tx_currency = rget(tx, 'currency') or currency
                    tx_from = rget(tx, 'from_name') or ''
                    tx_to = rget(tx, 'to_name') or ''
                    tx_inst = rget(tx, 'instrument_number') or '-'
                    tx_cert = rget(tx, 'certificate_number') or '-'

                    is_in = tx_to.strip().upper() == person_name_key
                    is_out = tx_from.strip().upper() == person_name_key
                    if is_in:
                        running_balance += tx_shares
                    elif is_out:
                        running_balance -= tx_shares
                    else:
                        running_balance += tx_shares

                    tx_row = [
                        '', '', '', '', tx_date, '',
                        tx_cert, f"{tx_inst}", str(tx_shares), f"{tx_currency}${tx_price}", tx_inst,
                        '-', '-', '-', '-',
                        str(running_balance), '', ''
                    ]
                    for i, val in enumerate(tx_row):
                        if val:
                            tnr(str(val)[:30], col_x[i] + 1, y + 1, size=data_size)
                    y += data_row_h

            # Row separator line
            line_h(M, PW - M, y, w=0.2, color=(180, 180, 180))
            table_bottom_y = y

    # Draw outer table border
    line_h(M, PW - M, table_top_y, w=0.5)
    line_h(M, PW - M, table_bottom_y, w=0.5)
    line_v(M, table_top_y, table_bottom_y, w=0.5)
    line_v(PW - M, table_top_y, table_bottom_y, w=0.5)

    pdf_bytes = bytes(pdf.output())
    pdf_bytes = _fix_tx_table_borders(pdf_bytes)
    import base64 as b64
    return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})

@app.route('/api/generate-directors-register-pdf', methods=['POST'])
def generate_directors_register_pdf():
    """Register of Directors (ROD) — matching RTF sample layout."""
    data = request.get_json(silent=True) or {}
    company_id = data.get('companyId', '')
    if not company_id:
        return jsonify({'error': 'companyId required'}), 400

    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return jsonify({'error': 'Company not found'}), 404

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role IN ('director', 'secretary')",
        (company_id,)).fetchall()
    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        placeholders = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({placeholders})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    directors = [r for r in roles if r['role'] == 'director']
    secretaries = [r for r in roles if r['role'] == 'secretary']

    # ── Try DOCX → Word COM → PDF first ──
    if _HAS_WORD_COM and _HAS_DOCX:
        docx_path = _build_rod_register_docx(db, company_id)
        if docx_path:
            pdf_bytes = _docx_to_pdf_via_word(docx_path)
            try:
                os.unlink(docx_path)
            except:
                pass
            if pdf_bytes:
                return jsonify({'pdf': base64.b64encode(pdf_bytes).decode('ascii')})

    # ── Fallback: fpdf2 ──
    pdf = create_pdf()
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    quorum = len(directors) if directors else None
    y = draw_register_header(pdf, company, "REGISTER OF OFFICERS", quorum)

    # ROD columns matching RTF sample — 6 columns
    x0 = MARGIN + 3
    rod_cols = [
        ("Name / Service /\nResidential Address",   x0,         130),
        ("Date / Place Birth /\nPlace Incorporated /\nOccupation /", x0 + 131,  75),
        ("ID No / Passport\nDetails",               x0 + 207,   67),
        ("Position",                                x0 + 275,   48),
        ("Date(s) Appointed\n/Meeting",             x0 + 324,   78),
        ("Reason / Date(s)\nCeased",                x0 + 403,   78),
    ]

    y = draw_grey_header_row(pdf, rod_cols, y)

    row_num = 0

    def render_section(items, is_secretary=False):
        nonlocal y, row_num
        for r in items:
            p = person_map.get(r['person_id'], {})
            name_en = rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)'
            name_ch = rget(p, 'name_chinese') if rget(p, 'name_english') else ''
            is_nat = rget(p, 'identity') != 'corporate'

            # Build Name/Address block
            addr = (rget(p, 'address') or '') if is_nat else (rget(p, 'registered_office') or rget(p, 'address') or '')
            name_block = name_en
            if name_ch:
                name_block += f"\n{name_ch}"
            if addr:
                name_block += f"\n{addr[:120]}"

            # DOB/Place/Nation block
            if is_nat:
                dob = rget(p, 'date_of_birth') or '-'
                pob = rget(p, 'place_of_birth') or '-'
                nat = rget(p, 'nationality') or '-'
                dob_block = f"{dob}\n{pob}\n{nat}"
            else:
                poi = rget(p, 'place_incorporated') or '-'
                dob_block = f"{poi}\n-\n-"

            # ID block
            id_info = (rget(p, 'id_number') or rget(p, 'passport_number') or '-') if is_nat else (rget(p, 'company_number_ref') or '-')

            # Position
            if is_secretary:
                position = "Secretary"
            else:
                position = "Reserve Director" if rget(r, 'is_reserve') else "Director"

            # Date Appointed
            date_app = rget(r, 'date_appointed') or '-'

            # Date Ceased / Reason
            date_cea = rget(r, 'date_ceased')
            if date_cea:
                reason_block = f"Resigned\n{date_cea}"
            else:
                reason_block = "Current\n現任"

            row_num += 1
            if y + 50 > PAGE_H - 50:
                pdf.add_page()
                y = draw_register_header(pdf, company, "REGISTER OF OFFICERS (Cont'd)", quorum)
                y = draw_grey_header_row(pdf, rod_cols, y)

            row_data = [
                (name_block,       rod_cols[0][1], rod_cols[0][2]),
                (dob_block,        rod_cols[1][1], rod_cols[1][2]),
                (id_info,          rod_cols[2][1], rod_cols[2][2]),
                (position,         rod_cols[3][1], rod_cols[3][2]),
                (date_app,         rod_cols[4][1], rod_cols[4][2]),
                (reason_block,     rod_cols[5][1], rod_cols[5][2]),
            ]
            y = draw_data_row(pdf, row_data, y, alt=(row_num % 2 == 0))

    # Directors first
    if directors:
        render_section(directors)
    else:
        pdf_draw(pdf, "(No directors / 尚無董事記錄)", MARGIN + 5, y + 12, size=8)
        y += 22

    # Secretaries below directors with same column layout
    if secretaries:
        if y + 40 > PAGE_H - 50:
            pdf.add_page()
            y = draw_register_header(pdf, company, "REGISTER OF OFFICERS (Cont'd)", quorum)
        y += 6
        render_section(secretaries, is_secretary=True)

    pdf_bytes = bytes(pdf.output())
    import base64 as b64
    return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})


@app.route('/api/generate-secretaries-register-pdf', methods=['POST'])
def generate_secretaries_register_pdf():
    """Standalone Register of Secretaries — matching ROD style."""
    data = request.get_json(silent=True) or {}
    company_id = data.get('companyId', '')
    if not company_id:
        return jsonify({'error': 'companyId required'}), 400

    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return jsonify({'error': 'Company not found'}), 404

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role = 'secretary'",
        (company_id,)).fetchall()
    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        placeholders = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({placeholders})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    pdf = create_pdf()
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    y = draw_register_header(pdf, company, "REGISTER OF COMPANY SECRETARIES")

    x0 = MARGIN + 3
    sec_cols = [
        ("Name / Service /\nResidential Address",   x0,         140),
        ("ID No / Passport\nDetails",               x0 + 141,   80),
        ("TCSP Licence\nNo.",                       x0 + 222,   60),
        ("Position",                                x0 + 283,   55),
        ("Date(s) Appointed",                       x0 + 339,   75),
        ("Reason / Date(s)\nCeased",                x0 + 415,   75),
    ]

    y = draw_grey_header_row(pdf, sec_cols, y)
    row_num = 0

    if not roles:
        pdf_draw(pdf, "(No secretaries / 尚無公司秘書記錄)", MARGIN + 5, y + 12, size=8)
    else:
        for r in roles:
            p = person_map.get(r['person_id'], {})
            name_en = rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)'
            name_ch = rget(p, 'name_chinese') if rget(p, 'name_english') else ''
            is_nat = rget(p, 'identity') != 'corporate'

            addr = (rget(p, 'address') or '') if is_nat else (rget(p, 'registered_office') or rget(p, 'address') or '')
            name_block = name_en
            if name_ch:
                name_block += f"\n{name_ch}"
            if addr:
                name_block += f"\n{addr[:100]}"

            id_info = (rget(p, 'id_number') or rget(p, 'passport_number') or '-') if is_nat else (rget(p, 'company_number_ref') or '-')
            tcsp = rget(p, 'tcsp_number') or '-'
            position = "Company Secretary"
            date_app = rget(r, 'date_appointed') or '-'
            date_cea = rget(r, 'date_ceased')
            reason_block = f"Resigned\n{date_cea}" if date_cea else "Current\n現任"

            row_num += 1
            if y + 40 > PAGE_H - 50:
                pdf.add_page()
                y = draw_register_header(pdf, company, "REGISTER OF COMPANY SECRETARIES (Cont'd)")
                y = draw_grey_header_row(pdf, sec_cols, y)

            row_data = [
                (name_block,       sec_cols[0][1], sec_cols[0][2]),
                (id_info,          sec_cols[1][1], sec_cols[1][2]),
                (tcsp,             sec_cols[2][1], sec_cols[2][2]),
                (position,         sec_cols[3][1], sec_cols[3][2]),
                (date_app,         sec_cols[4][1], sec_cols[4][2]),
                (reason_block,     sec_cols[5][1], sec_cols[5][2]),
            ]
            y = draw_data_row(pdf, row_data, y, alt=(row_num % 2 == 0))

    pdf_bytes = bytes(pdf.output())
    import base64 as b64
    return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})


# ═══════════════════════════════════════════════════════════════
# SCR DOCX template fill (Paul Tang format)
# ═══════════════════════════════════════════════════════════════

def _build_nature_of_control(sc):
    """Build a human-readable 'Nature of Control' string from SCR flags."""
    parts = []
    if rget(sc, 'nature_shares') or 0:
        parts.append('>25% shares')
    if rget(sc, 'nature_voting') or 0:
        parts.append('>25% voting rights')
    if rget(sc, 'nature_appoint') or 0:
        parts.append('>25% appointment rights')
    if rget(sc, 'nature_influence') or 0:
        parts.append('Sig. influence')
    if rget(sc, 'nature_trust') or 0:
        parts.append('Trust')
    other = rget(sc, 'nature_other') or ''
    if other:
        parts.append(other)
    return ', '.join(parts) if parts else ''


def _fill_scr_docx_template(db, company_id):
    """Fill the Paul Tang SCR DOCX template with company & controller data.

    Opens 重要控制人登記冊_PaulTang格式.docx, replaces header placeholders
    with actual data, fills the 7-column controller table, and saves to
    a temp .docx file.

    Returns path to temp .docx, or None on failure.
    """
    if not _HAS_DOCX:
        return None
    if not os.path.exists(_SCR_DOCX_TEMPLATE):
        print(f"[SCR] DOCX template not found: {_SCR_DOCX_TEMPLATE}")
        return None

    from copy import deepcopy
    from docx.oxml import OxmlElement

    # ── Load data ──
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return None

    controllers = db.execute(
        "SELECT * FROM significant_controllers WHERE company_id = ? ORDER BY created_at",
        (company_id,)).fetchall()

    co_name = rget(company, 'name') or ''
    co_br = rget(company, 'company_number') or ''

    # Jurisdiction — default to Hong Kong
    jurisdiction = (rget(company, 'jurisdiction') or 'Hong Kong').strip()
    jurisdiction_zh = '香港' if jurisdiction.upper() in ('HONG KONG', 'HK', 'HONGKONG') else jurisdiction

    # ── Open template ──
    doc = DocxDocument(_SCR_DOCX_TEMPLATE)

    # ── Fill header table (Table[0]) ──
    t0 = doc.tables[0]
    # Cell [0,0]: "NAME OF COMPANY: <co_name> | 公司名稱: | COMPANY NUMBER: <br>"
    cell_header = t0.rows[0].cells[0]
    for p in cell_header.paragraphs:
        for run in p.runs:
            if 'PAUL TANG AND COMPANY LIMITED' in run.text:
                run.text = run.text.replace('PAUL TANG AND COMPANY LIMITED', co_name)
            # Replace the template BR number (07281051) with actual BR
            if '07281051' in run.text:
                run.text = run.text.replace('07281051', co_br)

    # ── Fill jurisdiction paragraphs ──
    # P[1]: "JURISDICTION:  HONG KONG"
    if len(doc.paragraphs) > 1:
        p1 = doc.paragraphs[1]
        for run in p1.runs:
            if 'HONG KONG' in run.text.upper():
                run.text = run.text.replace(
                    run.text[run.text.upper().index('HONG KONG'):run.text.upper().index('HONG KONG')+9],
                    jurisdiction.upper()
                )
    # P[2]: Chinese jurisdiction line
    if len(doc.paragraphs) > 2:
        p2 = doc.paragraphs[2]
        for run in p2.runs:
            if 'HONG KONG' in run.text.upper():
                run.text = run.text.replace(
                    run.text[run.text.upper().index('HONG KONG'):run.text.upper().index('HONG KONG')+9],
                    jurisdiction.upper()
                )

    # ── Fill data table (Table[1]) ──
    t1 = doc.tables[1]

    # Save template row XML (row 1 = sample data row)
    template_tr_xml = deepcopy(t1.rows[1]._tr) if len(t1.rows) > 1 else None

    # Remove all data rows (keep only header row 0)
    for i in range(len(t1.rows) - 1, 0, -1):
        t1._tbl.remove(t1.rows[i]._tr)

    if not controllers:
        # Add empty row
        row = t1.add_row()
        row.cells[1].paragraphs[0].add_run('(No significant controllers / 無重要控制人)')
    else:
        for si, sc in enumerate(controllers):
            if template_tr_xml is not None:
                new_tr = template_tr_xml if si == 0 else deepcopy(template_tr_xml)
                t1._tbl.append(new_tr)
            else:
                t1.add_row()

            row = t1.rows[-1]
            cells = row.cells

            identity = rget(sc, 'identity') or 'natural'
            name_en = rget(sc, 'name_english') or rget(sc, 'name_chinese') or '(unnamed)'
            id_no = rget(sc, 'id_number') or '-'
            addr = rget(sc, 'address') or rget(sc, 'service_address') or ''
            date_became = _fmt_date_rom(rget(sc, 'date_became') or '-')
            created_date = (rget(sc, 'created_at') or '')[:10]
            entry_date = _fmt_date_rom(created_date) if created_date else '-'
            nature = _build_nature_of_control(sc)
            designated = ''
            if rget(sc, 'is_designated_rep') or 0:
                rep_name = rget(sc, 'designated_rep_name') or ''
                rep_contact = rget(sc, 'designated_rep_contact') or ''
                designated = f'Designated Rep: {rep_name}'
                if rep_contact:
                    designated += f' / {rep_contact}'

            # Build ID/PPT string
            if identity == 'natural':
                id_str = f'ID/PPT: {id_no} | Natural Person'
            else:
                id_str = f'Company No: {id_no} | Legal Entity'

            # Build remarks
            remarks_parts = []
            if rget(sc, 'date_ceased'):
                remarks_parts.append(f'Ceased: {_fmt_date_rom(rget(sc, "date_ceased"))}')
            if designated:
                remarks_parts.append(designated)
            remarks = '; '.join(remarks_parts) if remarks_parts else ('Current / 現任' if identity == 'natural' else 'Current / 現任')

            # 7-column data mapping
            cell_data = [
                entry_date,          # [0] Entry Date
                name_en,             # [1] Name
                addr,                # [2] Address
                id_str,              # [3] ID/PPT No.
                nature,              # [4] Nature of Control
                f'{date_became} / ', # [5] Becoming Date
                remarks,             # [6] Remarks
            ]

            for ci, val in enumerate(cell_data):
                if ci >= len(cells):
                    continue
                cell = cells[ci]
                for p in cell.paragraphs:
                    p.clear()
                run = cell.paragraphs[0].add_run(str(val) if val else '')
                run.font.size = Pt(7)
                run.font.name = 'Arial'
                _set_cjk_fallback(run, 'DengXian')

    # ── Save to temp file ──
    tmp_docx = tempfile.mktemp(suffix='.docx')
    doc.save(tmp_docx)
    return tmp_docx


@app.route('/api/generate-scr-pdf', methods=['POST'])
def generate_scr_pdf():
    """Significant Controllers Register — matching Paul Tang & Co reference format.
    Landscape A4, 7-column table, bilingual headers, grid borders, Times New Roman + TC."""
    data = request.get_json(silent=True) or {}
    company_id = data.get('companyId', '')
    if not company_id:
        return jsonify({'error': 'companyId required'}), 400

    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return jsonify({'error': 'Company not found'}), 404

    scrs = db.execute(
        "SELECT * FROM significant_controllers WHERE company_id = ? ORDER BY created_at",
        (company_id,)).fetchall()

    # ── Path 1 (NEW): DOCX template fill → Word COM → PDF (Paul Tang format) ──
    if _HAS_WORD_COM and _HAS_DOCX:
        docx_path = _fill_scr_docx_template(db, company_id)
        if docx_path:
            pdf_bytes = _docx_to_pdf_via_word(docx_path)
            try:
                os.unlink(docx_path)
            except:
                pass
            if pdf_bytes:
                return jsonify({'pdf': base64.b64encode(pdf_bytes).decode('ascii')})

    # ── Fallback: fpdf2 ──
    M = 28
    PW, PH = 842, 595  # Landscape A4
    CW = PW - 2 * M  # 786pt content width

    pdf = create_pdf(landscape=True)
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    # ── Helpers ──
    def has_cjk(text):
        s = str(text or '')
        return any('一' <= ch <= '鿿' or '　' <= ch <= '〿' or '＀' <= ch <= '￯' for ch in s)

    def tnr(text, x, y, size=8, bold=False, align='L'):
        style = 'B' if bold else ''
        pdf.set_font('TNR', style, size)
        pdf.set_text_color(0, 0, 0)
        tw = pdf.get_string_width(str(text or ''))
        if align == 'C': x = x - tw / 2
        elif align == 'R': x = x - tw
        pdf.set_xy(x, y)
        pdf.cell(tw + 2, size + 2, str(text or ''))

    def tc(text, x, y, size=8, bold=False, align='L'):
        style = 'B' if bold else ''
        pdf.set_font('TC', style, size)
        pdf.set_text_color(0, 0, 0)
        tw = pdf.get_string_width(str(text or ''))
        if align == 'C': x = x - tw / 2
        elif align == 'R': x = x - tw
        pdf.set_xy(x, y)
        pdf.cell(tw + 2, size + 2, str(text or ''))

    def line_h(x1, x2, y, w=0.3, color=(0, 0, 0)):
        pdf.set_draw_color(*color)
        pdf.set_line_width(w)
        pdf.line(x1, y, x2, y)

    def line_v(x, y1, y2, w=0.3, color=(0, 0, 0)):
        pdf.set_draw_color(*color)
        pdf.set_line_width(w)
        pdf.line(x, y1, x, y2)

    def draw_cell_border(x0, y0, w, h):
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.3)
        pdf.rect(x0, y0, w, h)

    def wrap_text_to_lines(text, font_name, style, size, max_w):
        if not text:
            return ['']
        s = str(text)
        pdf.set_font(font_name, style, size)
        all_lines = []
        # First split by explicit newlines, then wrap each chunk
        for chunk in s.split('\n'):
            if not chunk:
                all_lines.append('')
                continue
            current = ''
            for ch in chunk:
                trial = current + ch
                if pdf.get_string_width(trial) > max_w and current:
                    all_lines.append(current)
                    current = ch
                else:
                    current = trial
            if current:
                all_lines.append(current)
        return all_lines if all_lines else ['']

    # ── Company Data ──
    co_name = rget(company, 'name') or ''
    co_name_ch = rget(company, 'chinese_name') or ''
    br = rget(company, 'company_number') or ''

    # ═══════════════════════════════════════════════════════
    # HEADER BLOCK (borderless table: company info left, title right)
    # Per docx: header table first, then JURISDICTION below
    # ═══════════════════════════════════════════════════════

    y = 22
    hdr_size = 8

    # ── Title on right side ──
    title_x = PW - M
    tnr("SIGNIFICANT CONTROLLERS REGISTER", title_x, y, size=13, bold=True, align='R')
    tc("重要控制人登記冊", title_x, y + 18, size=11, bold=True, align='R')

    # ── Header: NAME OF COMPANY full width, then COMPANY NUMBER || JURISDICTION side by side ──
    col_a_x = M

    # Row 1 (English): NAME OF COMPANY — label only, no value/underline
    tnr("NAME OF COMPANY:  ", col_a_x, y, size=hdr_size, bold=True)

    y += 12

    # Row 2 (Chinese): 公司名稱 — with value + underline
    tc("公司名稱:  ", col_a_x, y, size=hdr_size)
    cn_label_w = pdf.get_string_width("公司名稱:  ")
    tc(co_name_ch if co_name_ch else co_name, col_a_x + cn_label_w, y, size=hdr_size)
    cn_val_w = pdf.get_string_width(co_name_ch or co_name)
    line_h(col_a_x + cn_label_w, col_a_x + cn_label_w + max(cn_val_w, 150), y + 11)

    y += 16

    # Row 3 (English): COMPANY NUMBER  |  JURISDICTION — labels only, no values/underlines
    tnr("COMPANY NUMBER:  ", col_a_x, y, size=hdr_size, bold=True)
    num_label_w = pdf.get_string_width("COMPANY NUMBER:  ")
    br_underline_end = col_a_x + num_label_w + 100

    jur_x = br_underline_end + 24
    tnr("JURISDICTION:  ", jur_x, y, size=hdr_size, bold=True)

    y += 14

    # Row 4 (Chinese): 公司編號 _______  司法管轄區: HONG KONG — with underlines
    tc("公司編號:  ", col_a_x, y, size=hdr_size)
    num2_label_w = pdf.get_string_width("公司編號:  ")
    tc(br, col_a_x + num2_label_w, y, size=hdr_size)
    br2_val_w = pdf.get_string_width(br)
    br2_underline_end = col_a_x + num2_label_w + max(br2_val_w, 100)
    line_h(col_a_x + num2_label_w, br2_underline_end, y + 11)

    tc("司法管轄區:  HONG KONG", jur_x, y, size=hdr_size)
    line_h(jur_x, jur_x + 120, y + 11)

    y += 16

    # Separator
    line_h(M, PW - M, y, w=0.5)
    y += 12

    # ═══════════════════════════════════════════════════════
    # DATA TABLE (7 columns, grid borders, bilingual headers)
    # Column widths per docx gridCol proportions with ID column widened
    # ═══════════════════════════════════════════════════════

    # Docx gridCol DXA: merged[1526], 2154, 2835, 2551, 2551, 1701, 1814
    # Total=15132 DXA. Scale to CW=786pt, then widen ID col (+8pt) and narrow Nature (-8pt)
    col_ratios = [1526, 2154, 2835, 2551, 2551, 1701, 1814]
    total_dxa = sum(col_ratios)
    col_w = [r * CW / total_dxa for r in col_ratios]
    # Widen ID column (index 3) and narrow Nature column (index 4)
    col_w[3] += 8   # ID: ~140pt
    col_w[4] -= 8   # Nature: ~124pt

    col_x = [M]
    for w in col_w[:-1]:
        col_x.append(col_x[-1] + w)
    end_x = PW - M

    # ── Bilingual Multi-line Table Headers ──
    hdr_labels = [
        ("Entry Date", 0),
        ("Name", 1),
        ("Correspondence Address\n (for Registrable Person)\n通訊地址（自然人）\nRegistered Office Address (for Legal Entity)\n註冊／主要營業地址\n（法律實體）", 2),
        ("ID / PPT No. (Issuing Country)\n(for Registrable Person)\n身份證／護照號碼\n（簽發國家）（自然人）\nCompany No. (Place of Incorp.)\nLegal Form (for Legal Entity)\n公司編號（成立地方）\n法律形式（法律實體）", 3),
        ("Nature of Control\n控制性質", 4),
        ("Becoming Date\n(Cessation Date)\n起始日期\n（終止日期）", 5),
        ("Remarks\n備註", 6),
    ]

    hdr_lines_counts = [len(lbl.split('\n')) for lbl, _ in hdr_labels]
    hdr_line_h = 10
    max_hdr_lines = max(hdr_lines_counts)
    hdr_h = max_hdr_lines * hdr_line_h + 6

    hdr_y0 = y

    for label, ci in hdr_labels:
        x0 = col_x[ci]
        cw_val = col_w[ci] if ci < len(col_w) else (end_x - col_x[ci])
        draw_cell_border(x0, y, cw_val, hdr_h)
        lines = label.split('\n')
        text_block_h = len(lines) * hdr_line_h
        text_start_y = y + (hdr_h - text_block_h) / 2
        for li, line_text in enumerate(lines):
            font_name = 'TC' if has_cjk(line_text) else 'TNR'
            pdf.set_font(font_name, 'B', 7)
            pdf.set_text_color(0, 0, 0)
            lw = pdf.get_string_width(line_text)
            lx = x0 + (cw_val - lw) / 2
            pdf.set_xy(lx, text_start_y + li * hdr_line_h)
            pdf.cell(lw, hdr_line_h, line_text)

    y += hdr_h
    table_top_y = hdr_y0

    # ── Data Rows ──
    data_size = 8
    min_row_h = 20

    if not scrs:
        empty_h = min_row_h
        for ci in range(len(col_x)):
            cw_val = col_w[ci] if ci < len(col_w) else (end_x - col_x[ci])
            draw_cell_border(col_x[ci], y, cw_val, empty_h)
        tc("(No SCR records / 尚無重要控制人記錄)", M + 4, y + 4, size=8)
        y += empty_h
    else:
        for s in scrs:
            # Build nature of control
            natures = []
            if rget(s, 'nature_shares'): natures.append('>25% shares')
            if rget(s, 'nature_voting'): natures.append('>25% voting')
            if rget(s, 'nature_appoint'): natures.append('Appoint/remove directors')
            if rget(s, 'nature_influence'): natures.append('Sig. influence')
            if rget(s, 'nature_trust'): natures.append('Trust control')
            if rget(s, 'nature_other'): natures.append(rget(s, 'nature_other'))

            is_nat = rget(s, 'identity') != 'corporate'
            name_en = rget(s, 'name_english') or ''
            name_ch = rget(s, 'name_chinese') or ''
            name_display = f"{name_ch}  {name_en}".strip() if name_ch else (name_en or '(unnamed)')

            # ID / Company info
            if is_nat:
                id_no = rget(s, 'id_number') or rget(s, 'passport_number') or '-'
                passport_country = rget(s, 'passport_country') or ''
                id_block = f"ID/PPT: {id_no}"
                if passport_country:
                    id_block += f" ({passport_country})"
                id_block += " | Natural Person"
            else:
                comp_no = rget(s, 'company_number_ref') or '-'
                place_incorp = rget(s, 'place_of_incorporation') or ''
                legal_form = rget(s, 'legal_form') or ''
                id_block = f"Co No: {comp_no}"
                if place_incorp:
                    id_block += f" ({place_incorp})"
                if legal_form:
                    id_block += f" | {legal_form}"
                id_block += " | Body Corporate"

            addr = (rget(s, 'address') or '')[:200]
            nature_text = ', '.join(natures) if natures else '-'
            date_became = rget(s, 'date_became') or '-'
            date_cea = rget(s, 'date_ceased') or ''
            # Paul Tang format: date column has "YYYY-MM-DD  /" (current) or "YYYY-MM-DD  /  YYYY-MM-DD" (ceased)
            date_display = f"{date_became}  /  {date_cea}" if date_cea else f"{date_became}  /"

            entry_date = rget(s, 'created_at') or ''
            if entry_date and len(entry_date) > 10:
                entry_date = entry_date[:10]

            # Remarks: "Current / 現任" goes here per Paul Tang format, + designated rep + user remarks
            remarks_parts = []
            if not date_cea:
                remarks_parts.append("Current / 現任")
            if rget(s, 'is_designated_rep') and rget(s, 'designated_rep_name'):
                remarks_parts.append(f"Rep: {rget(s, 'designated_rep_name')}")
            user_remarks = rget(s, 'remarks') or ''
            if user_remarks:
                remarks_parts.append(user_remarks)
            remarks = '\n'.join(remarks_parts) if remarks_parts else ''
            row_data = [entry_date, name_display, addr, id_block, nature_text, date_display, remarks]

            # Calculate row height from longest wrapped cell
            row_h = min_row_h
            cell_lines_list = []
            for ci, txt in enumerate(row_data):
                if not txt:
                    cell_lines_list.append(1)
                    continue
                font_name = 'TC' if has_cjk(str(txt)) else 'TNR'
                pdf.set_font(font_name, '', data_size)
                cell_pad = 4
                cw_avail = (col_w[ci] if ci < len(col_w) else (end_x - col_x[ci])) - cell_pad
                if cw_avail < 20:
                    cw_avail = 20
                lines = wrap_text_to_lines(str(txt), font_name, '', data_size, cw_avail)
                cell_lines_list.append(len(lines))
                row_h = max(row_h, len(lines) * (data_size + 4) + 4)

            # Page break
            if y + row_h > PH - 70:
                line_h(M, end_x, y, w=0.5)
                pdf.add_page()
                y = 22
                tnr("SIGNIFICANT CONTROLLERS REGISTER (Cont'd)", PW / 2, y, size=13, bold=True, align='C')
                y += 16
                tc("重要控制人登記冊（續）", PW / 2, y, size=11, bold=True, align='C')
                y += 14
                line_h(M, PW - M, y, w=0.5)
                y += 8
                hdr_y0 = y
                for label, ci in hdr_labels:
                    x0 = col_x[ci]
                    cw_val = col_w[ci] if ci < len(col_w) else (end_x - col_x[ci])
                    draw_cell_border(x0, y, cw_val, hdr_h)
                    lines = label.split('\n')
                    text_block_h = len(lines) * hdr_line_h
                    text_start_y = y + (hdr_h - text_block_h) / 2
                    for li, line_text in enumerate(lines):
                        font_name = 'TC' if has_cjk(line_text) else 'TNR'
                        pdf.set_font(font_name, 'B', 7)
                        pdf.set_text_color(0, 0, 0)
                        lw = pdf.get_string_width(line_text)
                        lx = x0 + (cw_val - lw) / 2
                        pdf.set_xy(lx, text_start_y + li * hdr_line_h)
                        pdf.cell(lw, hdr_line_h, line_text)
                y += hdr_h
                table_top_y = hdr_y0

            # Draw row cells
            for ci, txt in enumerate(row_data):
                x0 = col_x[ci]
                cw_val = col_w[ci] if ci < len(col_w) else (end_x - col_x[ci])
                draw_cell_border(x0, y, cw_val, row_h)
                if txt:
                    font_name = 'TC' if has_cjk(str(txt)) else 'TNR'
                    pdf.set_font(font_name, '', data_size)
                    pdf.set_text_color(0, 0, 0)
                    cell_pad = 3
                    cw_avail = cw_val - cell_pad * 2
                    if cw_avail < 20:
                        cw_avail = 20
                    lines = wrap_text_to_lines(str(txt), font_name, '', data_size, cw_avail)
                    is_remarks_col = (ci == 6)  # Remarks column — center per Paul Tang format
                    for li, line_text in enumerate(lines):
                        if is_remarks_col:
                            lw = pdf.get_string_width(line_text)
                            lx = x0 + (cw_val - lw) / 2
                        else:
                            lx = x0 + cell_pad
                        pdf.set_xy(lx, y + 2 + li * (data_size + 4))
                        pdf.cell(cw_avail, data_size + 4, line_text)
            y += row_h

    # Table bottom border
    line_h(M, end_x, y, w=0.5)
    y += 14

    # ═══════════════════════════════════════════════════════
    # ADDITIONAL MATTERS — 2×2 table (header row + content row)
    # ═══════════════════════════════════════════════════════
    add_h_hdr = 26  # header row height
    add_h_content = 48  # content row height
    add_w = CW * 0.5

    # Row 0: Headers — vertically centered with more bottom padding
    add_y0 = y
    draw_cell_border(M, y, add_w, add_h_hdr)
    draw_cell_border(M + add_w, y, add_w, add_h_hdr)
    tnr("Additional Matterse", M + 3, y + 4, size=7, bold=True)
    tc("额外事項", M + 3, y + 14, size=7)
    # Remarks — centered horizontally in right cell
    rmk_text = "Remarks"
    pdf.set_font('TNR', 'B', 7)
    rmk_w = pdf.get_string_width(rmk_text)
    tnr(rmk_text, M + add_w + (add_w - rmk_w) / 2, y + 4, size=7, bold=True)
    rmk_ch = "備註"
    pdf.set_font('TC', '', 7)
    rmk_ch_w = pdf.get_string_width(rmk_ch)
    tc(rmk_ch, M + add_w + (add_w - rmk_ch_w) / 2, y + 14, size=7)
    y += add_h_hdr

    # Row 1: Empty content cells
    draw_cell_border(M, y, add_w, add_h_content)
    draw_cell_border(M + add_w, y, add_w, add_h_content)
    y += add_h_content

    pdf_bytes = bytes(pdf.output())
    import base64 as b64
    return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})

# ─── R2 存儲本地替身（文件系統，鏡像生產 /api/storage/<bucket>/<file...>）───
STORAGE_ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'storage')
STORAGE_BUCKETS = {'pdf-templates', 'company-documents', 'company-logs', 'backups'}


def _storage_resolve(bucket, filepath):
    """把 <bucket>/<filepath> 解析成安全的本地絕對路徑，阻止路徑穿越。非法返回 None。"""
    if bucket not in STORAGE_BUCKETS:
        return None
    bucket_dir = os.path.abspath(os.path.join(STORAGE_ROOT, bucket))
    target = os.path.abspath(os.path.join(bucket_dir, filepath))
    # 確保 target 仍在 bucket_dir 之內（防 ../ 穿越）
    if target != bucket_dir and not target.startswith(bucket_dir + os.sep):
        return None
    return target


@app.route('/api/storage/<bucket>/<path:filepath>', methods=['GET', 'OPTIONS'])
def storage_get(bucket, filepath):
    if request.method == 'OPTIONS':
        return ('', 204)
    target = _storage_resolve(bucket, filepath)
    if not target:
        return jsonify({'error': 'Invalid bucket or path'}), 400
    if not os.path.isfile(target):
        return jsonify({'error': 'File not found'}), 404
    import mimetypes
    ctype = mimetypes.guess_type(target)[0] or 'application/octet-stream'
    with open(target, 'rb') as f:
        data = f.read()
    resp = Response(data, mimetype=ctype)
    dl = request.args.get('download')
    if dl:
        resp.headers['Content-Disposition'] = f'attachment; filename="{dl}"'
    resp.headers['Cache-Control'] = 'public, max-age=3600'
    return resp


@app.route('/api/storage/<bucket>/<path:filepath>', methods=['POST'])
def storage_put(bucket, filepath):
    target = _storage_resolve(bucket, filepath)
    if not target:
        return jsonify({'error': 'Invalid bucket or path'}), 400
    os.makedirs(os.path.dirname(target), exist_ok=True)
    with open(target, 'wb') as f:
        f.write(request.get_data())
    return jsonify({'success': True, 'path': filepath}), 201


@app.route('/api/storage/<bucket>/<path:filepath>', methods=['DELETE'])
def storage_delete(bucket, filepath):
    target = _storage_resolve(bucket, filepath)
    if not target:
        return jsonify({'error': 'Invalid bucket or path'}), 400
    if os.path.isfile(target):
        os.remove(target)
    return jsonify({'success': True})


# ─── Share Transfer PDFs (Instrument / Bought&Sold Note / Share Certificate) ───
# ─── RTF → PDF via Word COM ───

def _rtf_to_pdf_via_word(rtf_path, replacements, address_replacements=None,
                         cert_replacements=None):
    """Use Word COM to open RTF template, find & replace text, save as PDF.

    Args:
        rtf_path: Absolute path to the .rtf template file
        replacements: dict of {find_text: replace_text} — simple single-line replacements.
                      All occurrences are replaced (searches from start each time).
                      Sorted by key length descending to prevent substring conflicts.
        address_replacements: list of (old_lines, new_lines, occurrence) tuples for
                      multi-line address blocks. Each entry replaces the Nth occurrence
                      of old_lines[0] and its continuation lines.
        cert_replacements: dict {placeholder, old_value, new_value} for certificate
                      number replacement — finds placeholder text, then finds old_value
                      after it and replaces only that occurrence (not globally).

    Returns:
        bytes: PDF file content, or None on failure
    """
    if not _HAS_WORD_COM:
        return None
    if not os.path.exists(rtf_path):
        return None

    with _word_lock:
        pythoncom.CoInitialize()
        word = None
        tmp_pdf = None
        try:
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0  # wdAlertsNone

            doc = word.Documents.Open(rtf_path)
            sel = word.Selection

            # ── Simple replacements: find all occurrences, replace each ──
            # Sort by key length DESC to prevent substring conflicts
            # (e.g. "HK$10,000.00" must be processed before "10,000")
            sorted_items = sorted(
                replacements.items(),
                key=lambda kv: len(kv[0]),
                reverse=True,
            )
            for find_text, replace_text in sorted_items:
                if not find_text:
                    continue
                replace_text = str(replace_text or '')
                sel.HomeKey(Unit=6)  # wdStory = go to doc start
                while sel.Find.Execute(
                    FindText=find_text,
                    MatchCase=True,
                    Forward=True,
                    Wrap=0,  # wdFindStop — don't wrap, search start→end
                ):
                    sel.Text = replace_text
                    sel.Collapse(Direction=0)  # wdCollapseEnd — move past replacement for next find

            # ── Multi-line address replacements ──
            if address_replacements:
                for old_lines, new_lines, occurrence in address_replacements:
                    if not old_lines or not new_lines:
                        continue
                    _rtf_replace_address_block(doc, sel, old_lines, new_lines, occurrence)

            # ── Certificate number (positional, not global) ──
            if cert_replacements:
                placeholder = cert_replacements.get('placeholder', '')
                old_val = cert_replacements.get('old_value', '')
                new_val = cert_replacements.get('new_value', '')
                if placeholder and old_val and new_val:
                    _rtf_replace_after_anchor(sel, placeholder, old_val, new_val)

            # Save as PDF (FileFormat 17 = wdFormatPDF)
            tmp_pdf = tempfile.mktemp(suffix='.pdf')
            doc.SaveAs2(tmp_pdf, FileFormat=17)
            doc.Close(False)

            with open(tmp_pdf, 'rb') as f:
                pdf_bytes = f.read()
            return pdf_bytes
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[RTF→PDF] Word COM error: {e}")
            return None
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
            if tmp_pdf and os.path.exists(tmp_pdf):
                try:
                    os.unlink(tmp_pdf)
                except:
                    pass
            pythoncom.CoUninitialize()


def _rtf_replace_address_block(doc, sel, old_lines, new_lines, occurrence=1):
    """Replace the Nth occurrence of a multi-line address block.

    Finds old_lines[0] (occurrence N), then replaces each subsequent line
    with the corresponding new_lines entry. Works positionally — after replacing
    line N, searches forward for the next line to replace.

    Args:
        doc: Word Document object
        sel: Word Selection object
        old_lines: list of strings — the template address lines to search for
        new_lines: list of strings — the replacement address lines
        occurrence: which occurrence of old_lines[0] to target (1-based)
    """
    # Find the Nth occurrence of the first address line
    sel.HomeKey(Unit=6)  # wdStory
    for _ in range(occurrence):
        found = sel.Find.Execute(
            FindText=old_lines[0],
            MatchCase=True,
            Forward=True,
            Wrap=0,
        )
        if not found:
            return
        if _ < occurrence - 1:
            sel.Collapse(Direction=0)  # wdCollapseEnd — move past this occurrence

    # Replace first line
    sel.Text = str(new_lines[0] or '')
    sel.Collapse(Direction=0)  # wdCollapseEnd

    # Replace continuation lines in order (they follow right after)
    for i in range(1, len(old_lines)):
        found = sel.Find.Execute(
            FindText=old_lines[i],
            MatchCase=True,
            Forward=True,
            Wrap=0,
        )
        if not found:
            break
        sel.Text = str(new_lines[i] if i < len(new_lines) else '')
        sel.Collapse(Direction=0)


def _rtf_replace_after_anchor(sel, anchor_text, old_value, new_value):
    """Find anchor_text, then find the next occurrence of old_value and replace it.

    Used for positional replacements where a global find/replace would be
    too broad (e.g. replacing "3" would match ALL "3" characters).
    """
    sel.HomeKey(Unit=6)  # wdStory
    found = sel.Find.Execute(
        FindText=anchor_text,
        MatchCase=True,
        Forward=True,
        Wrap=0,
    )
    if not found:
        return
    sel.Collapse(Direction=0)  # wdCollapseEnd
    found = sel.Find.Execute(
        FindText=old_value,
        MatchCase=True,
        Forward=True,
        Wrap=0,
    )
    if not found:
        return
    sel.Text = str(new_value or '')


# ─────────────────────────────────────────────
# DOCX → PDF via Word COM
# ─────────────────────────────────────────────

def _docx_to_pdf_via_word(docx_path):
    """Open a DOCX file in Word, save as PDF, return pdf_bytes or None.

    This mirrors _rtf_to_pdf_via_word but for DOCX input — no Find & Replace,
    just open, save-as-PDF, close.  Useful when the DOCX was already built by
    python-docx with the correct content.
    """
    if not _HAS_WORD_COM:
        return None
    if not os.path.exists(docx_path):
        return None

    with _word_lock:
        pythoncom.CoInitialize()
        word = None
        tmp_pdf = None
        try:
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0  # wdAlertsNone

            doc = word.Documents.Open(docx_path)
            # Save as PDF (FileFormat 17 = wdFormatPDF)
            tmp_pdf = tempfile.mktemp(suffix='.pdf')
            doc.SaveAs2(tmp_pdf, FileFormat=17)
            doc.Close(False)

            with open(tmp_pdf, 'rb') as f:
                pdf_bytes = f.read()
            return pdf_bytes
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[DOCX→PDF] Word COM error: {e}")
            return None
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
            if tmp_pdf and os.path.exists(tmp_pdf):
                try:
                    os.unlink(tmp_pdf)
                except:
                    pass
            pythoncom.CoUninitialize()


def _build_company_address(company):
    """Build a full registered office address string from company record."""
    parts = []
    for k in ['reg_flat', 'reg_building', 'reg_street', 'reg_district']:
        v = (company.get(k) or '').strip()
        if v:
            parts.append(v)
    region = (company.get('reg_region') or '').strip()
    jurisdiction = (company.get('jurisdiction') or '').strip()
    return ', '.join(parts), region, jurisdiction


def _get_person_address(db, person_id):
    """Get a person's full address from the database."""
    if not person_id:
        return None, None
    p = db.execute(
        "SELECT addr_flat, addr_building, addr_street, addr_district, addr_region, address FROM persons WHERE id = ?",
        (person_id,)).fetchone()
    if not p:
        return None, None
    parts = [p[k] for k in ['addr_flat', 'addr_building', 'addr_street', 'addr_district'] if p[k]]
    addr = ', '.join(parts) if parts else ''
    # Fallback to unstructured address column if structured fields are empty
    if not addr and p['address']:
        addr = p['address']
    region = p['addr_region'] or ''
    return addr, region


def _split_address_lines(addr, region, max_line_chars=55, max_lines=4):
    """Split a full address string into multiple lines for RTF template replacement.

    Splits at comma boundaries to keep lines under max_line_chars.
    Region (e.g. 'HONG KONG', 'KOWLOON') is appended as the last line.
    If the result exceeds max_lines, lines are merged from the end.
    """
    lines = []
    if not addr and not region:
        return lines

    # Split address by commas and rebuild into lines
    if addr:
        parts = [p.strip() for p in addr.split(',')]
        current = ''
        for part in parts:
            test = (current + ', ' + part).strip(', ') if current else part
            if len(test) <= max_line_chars:
                current = test
            else:
                if current:
                    lines.append(current)
                current = part
        if current:
            lines.append(current)

    if region:
        lines.append(region.strip())

    # Merge from end to fit max_lines
    while len(lines) > max_lines and len(lines) > 1:
        # Merge last two lines
        lines[-2] = lines[-2] + ', ' + lines[-1]
        lines.pop()

    return lines


def _build_registered_office_lines(co_addr, co_region, co_juris):
    """Build registered office address lines for Share Certificate template.

    Returns a list of strings (1-4 lines).
    """
    lines = []
    if co_addr:
        # Split into parts
        parts = [p.strip() for p in co_addr.split(',')]
        # Line 1: first part (flat/unit + building)
        lines.append(parts[0] if len(parts) > 0 else co_addr)
        # Line 2: rest of address
        if len(parts) > 1:
            lines.append(', '.join(parts[1:]))
    if co_region:
        lines.append(co_region.strip())
    if co_juris and co_juris != 'Hong Kong':
        lines.append(co_juris.strip())
    # Pad to at least 1 line
    if not lines:
        lines.append('')
    return lines


def _fmt_price(val, currency='HK$'):
    """Format a price value for display."""
    if val is None or val == '':
        return ''
    try:
        return f"{currency}{float(val):,.2f}"
    except (ValueError, TypeError):
        return str(val)


def _fmt_shares(val):
    """Format share count for display."""
    if val is None or val == '' or val == 0:
        return ''
    try:
        return f"{int(val):,}"
    except (ValueError, TypeError):
        return str(val)


# ─── Share Transfer PDF Generation ───

@app.route('/api/generate-share-transfer-pdf', methods=['POST'])
def generate_share_transfer_pdf():
    data = request.get_json(silent=True) or {}
    company_id = data.get('companyId', '')
    if not company_id:
        return jsonify({'error': 'companyId required'}), 400

    doc_type = data.get('documentType', 'instrument_of_transfer')

    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return jsonify({'error': 'Company not found'}), 404

    tx = db.execute(
        "SELECT * FROM share_transactions WHERE company_id = ? ORDER BY transaction_date DESC",
        (company_id,)).fetchone()

    company_dict = dict(company)
    tx_dict = dict(tx) if tx else {}

    # If no explicit transaction, auto-build from company shareholders
    if not tx:
        shareholders = db.execute(
            """SELECT p.id as person_id, p.name_english, p.address,
                      p.addr_flat, p.addr_building, p.addr_street, p.addr_district, p.addr_region,
                      pcr.shares, pcr.share_type, pcr.issue_price
               FROM person_company_roles pcr
               JOIN persons p ON p.id = pcr.person_id
               WHERE pcr.company_id = ? AND pcr.role = 'shareholder'
               ORDER BY p.identity = 'natural' DESC, p.name_english
            """, (company_id,)).fetchall()
        if len(shareholders) >= 2:
            s1, s2 = shareholders[0], shareholders[1]
            tx_dict = {
                'from_person_id': s1['person_id'],
                'from_name': s1['name_english'],
                'to_person_id': s2['person_id'],
                'to_name': s2['name_english'],
                'shares': s1['shares'] or 0,
                'share_type': s1['share_type'] or 'Ordinary',
                'price_per_share': s1['issue_price'] or '1.00',
                'total_consideration': (s1['shares'] or 0) * float(s1['issue_price'] or 1.00) if s1['issue_price'] else 0,
                'transaction_date': datetime.now().strftime('%Y-%m-%d'),
            }

    # Try Word COM RTF→PDF first, fall back to fpdf2
    pdf_bytes = None

    if _HAS_WORD_COM:
        if doc_type == 'share_certificate':
            pdf_bytes = _build_share_certificate_rtf(db, company_dict, tx_dict)
        elif doc_type == 'bought_sold_note':
            pdf_bytes = _build_bought_sold_note_rtf(db, company_dict, tx_dict)
        else:
            pdf_bytes = _build_instrument_of_transfer_rtf(db, company_dict, tx_dict)

    # Fallback to fpdf2 if Word COM failed or not available
    if pdf_bytes is None:
        pdf = create_pdf()
        if doc_type == 'share_certificate':
            _build_share_certificate_fpdf(pdf, company_dict, tx_dict)
        elif doc_type == 'bought_sold_note':
            _build_bought_sold_note_fpdf(pdf, company_dict, tx_dict)
        else:
            _build_instrument_of_transfer_fpdf(pdf, company_dict, tx_dict)
        pdf_bytes = pdf.output()

    return jsonify({'pdf': base64.b64encode(pdf_bytes).decode('utf-8')})


# ═══════════════════════════════════════════
# RTF-based PDF generation (Paul Tang templates)
# ═══════════════════════════════════════════

def _build_instrument_of_transfer_rtf(db, company, tx):
    """Generate Instrument of Transfer PDF from RTF template."""
    if not os.path.exists(_RTF_INSTRUMENT):
        return None

    co_name = company.get('name', '')
    from_name = tx.get('from_name', '')
    to_name = tx.get('to_name', '')
    shares = tx.get('shares', 0) or 0
    par_val = tx.get('price_per_share') or ''
    consideration = tx.get('total_consideration') or ''
    if not consideration and shares and par_val:
        try:
            consideration = f"{shares * float(par_val):,.2f}"
        except (ValueError, TypeError):
            pass

    cons_str = _fmt_price(consideration, 'HK$')
    shares_str = _fmt_shares(shares)

    replacements = {}

    # Company name — "Testing Company Limited"
    if co_name:
        replacements['Testing Company Limited'] = co_name

    # Transferor (from) — "ABC TESTING"
    if from_name:
        replacements['ABC TESTING'] = from_name

    # Transferee (to) — "BCD TESTING"
    if to_name:
        replacements['BCD TESTING'] = to_name

    # Consideration — "HK$10,000.00"
    if cons_str:
        replacements['HK$10,000.00'] = cons_str

    # Shares — "10,000"
    if shares_str:
        replacements['10,000'] = shares_str

    # ── Address replacements ──
    # The Instrument template has TWO address blocks with DIFFERENT line breaks:
    #   Block 1 (Transferor):
    #     "ROOM 405 TUNG NING BUILDING, 249-253 DES VOEUX ROAD CENTRAL, SHEUNG WAN, HONG "
    #     "KONG"
    #   Block 2 (Transferee):
    #     "ROOM 405 TUNG NING BUILDING, 249-253 DES VOEUX ROAD "
    #     "CENTRAL, SHEUNG WAN, HONG KONG"
    addr_replacements = []

    # Transferor address — first occurrence of Block 1 format
    from_addr, from_region = _get_person_address(db, tx.get('from_person_id'))
    if from_addr or from_region:
        xfer_lines = _split_address_lines(from_addr, from_region, max_lines=2)
        addr_replacements.append(([
            'ROOM 405 TUNG NING BUILDING, 249-253 DES VOEUX ROAD CENTRAL, SHEUNG WAN, HONG ',
            'KONG',
        ], xfer_lines, 1))

    # Transferee address — first occurrence of Block 2 format
    to_addr, to_region = _get_person_address(db, tx.get('to_person_id'))
    if to_addr or to_region:
        tfee_lines = _split_address_lines(to_addr, to_region, max_lines=2)
        addr_replacements.append(([
            'ROOM 405 TUNG NING BUILDING, 249-253 DES VOEUX ROAD ',
            'CENTRAL, SHEUNG WAN, HONG KONG',
        ], tfee_lines, 1))

    return _rtf_to_pdf_via_word(_RTF_INSTRUMENT, replacements, addr_replacements or None)


def _build_bought_sold_note_rtf(db, company, tx):
    """Generate Bought & Sold Note PDF from RTF template."""
    if not os.path.exists(_RTF_BS_NOTE):
        return None

    co_name = company.get('name', '')
    from_name = tx.get('from_name', '')
    to_name = tx.get('to_name', '')
    shares = tx.get('shares', 0) or 0
    par_val = tx.get('price_per_share') or ''
    consideration = tx.get('total_consideration') or ''
    tx_date = tx.get('transaction_date', '')

    # Build consideration if not provided
    if not consideration and shares and par_val:
        try:
            consideration = f"{shares * float(par_val):,.2f}"
        except (ValueError, TypeError):
            pass

    # Format values
    shares_str = _fmt_shares(shares)
    par_val_str = _fmt_price(par_val, 'HK$')
    cons_str = _fmt_price(consideration, 'HK$')

    # Format transaction date: DD/MM/YYYY
    tx_date_str = ''
    if tx_date:
        try:
            dt = datetime.strptime(tx_date, '%Y-%m-%d')
            tx_date_str = dt.strftime('%d/%m/%Y')
        except ValueError:
            tx_date_str = tx_date

    # Query occupations for both parties (column may not exist in all DBs)
    seller_occ = ''
    buyer_occ = ''
    from_pid = tx.get('from_person_id')
    to_pid = tx.get('to_person_id')
    try:
        if from_pid:
            p = db.execute("SELECT occupation FROM persons WHERE id = ?", (from_pid,)).fetchone()
            if p and p['occupation']:
                seller_occ = p['occupation']
        if to_pid:
            p = db.execute("SELECT occupation FROM persons WHERE id = ?", (to_pid,)).fetchone()
            if p and p['occupation']:
                buyer_occ = p['occupation']
    except sqlite3.OperationalError:
        pass  # occupation column doesn't exist

    replacements = {}

    # Seller (Transferor) = from_name → "ABC TESTING" in template
    if from_name:
        replacements['ABC TESTING'] = from_name

    # Purchaser (Transferee) = to_name → "BCD TESTING" in template
    if to_name:
        replacements['BCD TESTING'] = to_name

    # Company name → "Testing Company Limited" in template
    if co_name:
        replacements['Testing Company Limited'] = co_name

    # Number of shares → "10,000" in template
    if shares_str:
        replacements['10,000'] = shares_str

    # Price per share → "HK$1.00" in template
    if par_val_str:
        replacements['HK$1.00'] = par_val_str

    # Total consideration → "HK$10,000.00" in template
    if cons_str:
        replacements['HK$10,000.00'] = cons_str

    # Transaction date → "DD/MM/YYYY" in template (both Sold & Bought)
    if tx_date_str:
        replacements['DD/MM/YYYY'] = tx_date_str

    # Occupation → "SELLER_OCCUP" in Bought Note, "BUYER_OCCUP" in Sold Note
    # Always replace — clear placeholder if no occupation available
    replacements['SELLER_OCCUP'] = seller_occ
    replacements['BUYER_OCCUP'] = buyer_occ

    # ── Address replacements ──
    # Template addresses (2-line blocks, 2 occurrences: Seller then Purchaser)
    addr_replacements = []
    tpl_lines = [
        'ROOM 405 TUNG NING BUILDING, 249-253 DES VOEUX ROAD ',
        'CENTRAL, SHEUNG WAN, HONG KONG',
    ]

    # Seller address (first in document order)
    from_addr, from_region = _get_person_address(db, tx.get('from_person_id'))
    if from_addr or from_region:
        seller_lines = _split_address_lines(from_addr, from_region, max_lines=2)
        addr_replacements.append((tpl_lines, seller_lines, 1))

    # Purchaser address (now first remaining, since Seller was just replaced)
    to_addr, to_region = _get_person_address(db, tx.get('to_person_id'))
    if to_addr or to_region:
        buyer_lines = _split_address_lines(to_addr, to_region, max_lines=2)
        addr_replacements.append((tpl_lines, buyer_lines, 1))

    return _rtf_to_pdf_via_word(_RTF_BS_NOTE, replacements, addr_replacements or None)


def _build_share_certificate_rtf(db, company, tx):
    """Generate Share Certificate PDF from RTF template."""
    if not os.path.exists(_RTF_CERTIFICATE):
        return None

    co_name = company.get('name', '')
    co_number = company.get('company_number', '')
    co_ci = company.get('ci_number', '')
    inc_date = company.get('incorporation_date', '')
    jurisdiction = company.get('jurisdiction', 'Hong Kong')
    co_addr, co_region, co_juris = _build_company_address(company)

    to_name = tx.get('to_name', '')
    shares = tx.get('shares', 0) or 0
    cert_no = tx.get('certificate_number') or tx.get('instrument_number', '')
    tx_date = tx.get('transaction_date', '')

    shares_str = _fmt_shares(shares)

    replacements = {}

    # Company name
    if co_name:
        replacements['Testing Company Limited'] = co_name

    # Company number
    if co_number:
        replacements['0101234'] = co_number

    # Shareholder name
    if to_name:
        replacements['BCD TESTING'] = to_name

    # Shares
    if shares_str:
        replacements['10,000'] = shares_str

    # HKID
    if tx.get('to_person_id'):
        pid = tx['to_person_id']
        p = db.execute("SELECT id_number FROM persons WHERE id = ?", (pid,)).fetchone()
        if p and p['id_number']:
            replacements['Y231456(1)'] = p['id_number']

    # Incorporation date
    if inc_date:
        try:
            dt = datetime.strptime(inc_date, '%Y-%m-%d')
            formatted_date = dt.strftime('%d/%m/%Y')
        except ValueError:
            formatted_date = inc_date
        replacements['08/04/2022'] = formatted_date

    # Certificate number — handled via positional replacement (not global,
    # because "3" would match ALL "3" characters in the document).
    cert_replacements = None
    if cert_no:
        cert_replacements = {
            'placeholder': 'Certificate Number:',
            'old_value': '3',
            'new_value': str(cert_no),
        }

    # ── Address replacements ──
    addr_replacements = []

    # 1. Registered office address
    if co_addr or co_region or co_juris:
        reg_lines = _build_registered_office_lines(co_addr, co_region, co_juris)
        tpl_reg_lines = [
            'ROOM 408, JINDIXINGYUAN JINGUANGE,',
            'NO.110, ELING SOUTH ROAD, HUICHENG DISTRICT,',
            'HUIZHOU, GUANGDONG',
            'CHINA',
        ]
        addr_replacements.append((tpl_reg_lines, reg_lines, 1))

    # 2. Shareholder address
    to_addr, to_region = _get_person_address(db, tx.get('to_person_id'))
    if to_addr or to_region:
        holder_lines = _split_address_lines(to_addr, to_region)
        tpl_holder_lines = [
            'ROOM 405 TUNG NING BUILDING',
            '249-253 DES VOEUX ROAD CENTRAL',
            'SHEUNG WAN',
            'HONG KONG',
        ]
        addr_replacements.append((tpl_holder_lines, holder_lines, 1))

    return _rtf_to_pdf_via_word(_RTF_CERTIFICATE, replacements,
                                addr_replacements or None,
                                cert_replacements=cert_replacements)


# ═══════════════════════════════════════════
# FPDF2 fallback PDF generation (original implementations)
# ═══════════════════════════════════════════

def _build_instrument_of_transfer_fpdf(pdf, company, tx):
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)
    y, left = 45, 45
    page_w = 595

    pdf_draw(pdf, "股份轉讓文書 / Instrument of Transfer", left, y, size=15)
    y += 20
    pdf.line(left, y, page_w - left, y)
    y += 14

    def draw_line(label, value):
        nonlocal y
        pdf_draw(pdf, label, left, y, size=10, gray=80)
        pdf_draw(pdf, value or "__________________________", left + 140, y, size=10)
        y += 22

    draw_line("公司名稱 Company:", company.get('name', ''))
    draw_line("BR 號碼:", company.get('company_number', ''))
    if company.get('chinese_name'):
        draw_line("中文名稱:", company.get('chinese_name', ''))

    y += 10
    pdf_draw(pdf, "轉讓詳情 / Transfer Details", left, y, size=12)
    y += 6
    pdf.line(left, y, page_w - left, y)
    y += 16

    itx_shares = tx.get('shares', 0) or 0
    itx_par_val = float(tx.get('price_per_share') or 1.00)
    itx_cons = tx.get('total_consideration') or (itx_shares * itx_par_val)

    draw_line("轉讓人 Transferor:", tx.get('from_name', ''))
    draw_line("受讓人 Transferee:", tx.get('to_name', ''))
    draw_line("股份數目 No. of Shares:",
              f"{itx_shares:,}  of  HK${itx_par_val:,.2f}  each" if tx.get('shares') else "________________")
    draw_line("股份類別 Share Class:", tx.get('share_type', 'Ordinary'))
    draw_line("每股代價 Price per Share:",
              f"{tx.get('currency', '')} {tx.get('price_per_share', '')}" if tx.get('price_per_share') else "________________")
    draw_line("總代價 Total Consideration:",
              f"HK${itx_cons:,.2f}" if (tx.get('total_consideration') or tx.get('shares')) else "________________")
    draw_line("轉讓日期 Transfer Date:", tx.get('transaction_date', ''))
    draw_line("文書編號 Instrument No:", tx.get('instrument_number', ''))

    y += 20
    pdf_draw(pdf, "轉讓人簽署 / Signed by Transferor:", left, y, size=10)
    pdf_draw(pdf, "____________________________", left + 200, y, size=10)
    y += 25
    pdf_draw(pdf, "受讓人簽署 / Signed by Transferee:", left, y, size=10)
    pdf_draw(pdf, "____________________________", left + 200, y, size=10)
    y += 25
    pdf_draw(pdf, "日期 Date: ____/____/________", left, y, size=10)

    pdf_draw(pdf, f"由 Muse Labs Engineering Limited 秘書系統生成 | {datetime.now().strftime('%Y-%m-%d')}",
             left, 30, size=7, gray=150)


def _build_bought_sold_note_fpdf(pdf, company, tx):
    """Bought & Sold Note — matching Paul Tang & Co reference format.
    Free-form layout (no table), tab-stop alignment, Times New Roman, thick title lines."""
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)
    M = 40  # margin
    PW, PH = 595, 842
    half_h = PH / 2
    label_x = M + 5
    value_x = PW / 3 + 10  # ~30% for label, 70% for value

    from_name = tx.get('from_name', '')
    to_name = tx.get('to_name', '')
    shares = tx.get('shares', 0) or 0
    par_val = tx.get('price_per_share', '1.00')
    if not par_val or par_val == '':
        par_val = '1.00'
    consideration = tx.get('total_consideration') or (shares * float(par_val))
    tx_date = tx.get('transaction_date', '')
    co_name = company.get('name', '')

    def tnr(text, x, y, size=12, bold=False, align='L'):
        style = 'B' if bold else ''
        pdf.set_font('TNR', style, size)
        pdf.set_text_color(0, 0, 0)
        if align == 'C':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw / 2
        pdf.set_xy(x, y)
        pdf.cell(0, size + 3, str(text or ''))

    def tc(text, x, y, size=12, bold=False):
        style = 'B' if bold else ''
        pdf.set_font('TC', style, size)
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(x, y)
        pdf.cell(0, size + 3, str(text or ''))

    def draw_row(label, value, y_pos, size=12):
        """Draw label (left) + value (right) pair using tab-stop alignment."""
        tnr(label, label_x, y_pos, size=size)
        tnr(str(value or ''), value_x, y_pos, size=size)

    # ══════════════════════════════════════
    # SOLD NOTE — TOP half
    # ══════════════════════════════════════
    y = PH - 50

    # Title: centered, 16pt bold TNR
    tnr("Sold Note", PW / 2, y, size=16, bold=True, align='C')
    y -= 4
    tc("賣出票據", PW / 2, y, size=11, bold=False)
    y -= 22

    # Thick black line under title (1.5pt, ~70% page width)
    line_w = (PW - 2 * M) * 0.7
    line_start = (PW - line_w) / 2
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(1.5)
    pdf.line(line_start, y, line_start + line_w, y)
    pdf.set_line_width(0.3)
    y -= 18

    # Content rows — 12pt TNR
    draw_row("Name of Purchaser (Transferee):", to_name, y); y -= 18
    draw_row("Address:", "", y); y -= 18
    draw_row("Occupation:", "", y); y -= 18
    draw_row("Name of Company:", co_name, y); y -= 18
    draw_row("Number of Shares:", f"{shares:,}  of  HK${par_val}  each", y); y -= 18
    if consideration:
        draw_row("Consideration Received:", f"HK${consideration:,.2f}", y); y -= 18
    else:
        draw_row("Consideration Received:", "", y); y -= 18

    y -= 10
    # Transferor signature line
    tnr(f"(Transferor)  {from_name}", label_x, y, size=12)
    sig_tf_w = pdf.get_string_width(f"(Transferor)  {from_name}")
    sig_end = label_x + sig_tf_w + 10
    # Signature underline from end of text to right margin
    pdf.set_line_width(0.6)
    pdf.line(sig_end, y + 3, PW - M, y + 3)
    pdf.set_line_width(0.3)
    y -= 16
    tnr(co_name, sig_end, y, size=8)
    y -= 14

    # Date line
    tnr(f"Hong Kong, Dated  {tx_date}", label_x, y, size=12)
    y -= 24

    # ══════════════════════════════════════
    # DIVIDER
    # ══════════════════════════════════════
    pdf.set_draw_color(100, 100, 100)
    pdf.set_line_width(0.5)
    pdf.line(M, half_h, PW - M, half_h)
    pdf.set_line_width(0.3)
    pdf.set_draw_color(0, 0, 0)

    y = half_h - 18

    # ══════════════════════════════════════
    # BOUGHT NOTE — BOTTOM half
    # ══════════════════════════════════════
    # Title: centered, 16pt bold TNR
    tnr("Bought Note", PW / 2, y, size=16, bold=True, align='C')
    y -= 4
    tc("買入票據", PW / 2, y, size=11)
    y -= 22

    # Thick black line under title
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(1.5)
    pdf.line(line_start, y, line_start + line_w, y)
    pdf.set_line_width(0.3)
    y -= 18

    # Content rows
    draw_row("Name of Seller (Transferor):", from_name, y); y -= 18
    draw_row("Address:", "", y); y -= 18
    draw_row("Occupation:", "", y); y -= 18
    draw_row("Name of Company:", co_name, y); y -= 18
    draw_row("Number of Shares:", f"{shares:,}  of  HK${par_val}  each", y); y -= 18
    if consideration:
        draw_row("Consideration Received:", f"HK${consideration:,.2f}", y); y -= 18
    else:
        draw_row("Consideration Received:", "", y); y -= 18

    y -= 10
    # Transferee signature line
    tnr(f"(Transferee)  {to_name}", label_x, y, size=12)
    sig_tee_w = pdf.get_string_width(f"(Transferee)  {to_name}")
    sig_tee_end = label_x + sig_tee_w + 10
    pdf.set_line_width(0.6)
    pdf.line(sig_tee_end, y + 3, PW - M, y + 3)  # horizontal line to right
    pdf.set_line_width(0.3)
    y -= 18
    tnr(f"Hong Kong, Dated  {tx_date}", label_x, y, size=12)

    # Footer
    tnr(f"Generated by Muse Labs | {datetime.now().strftime('%Y-%m-%d')}",
        M, 20, size=7)


def _build_share_certificate_fpdf(pdf, company, tx):
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)
    y, left, page_w = 45, 45, 595
    page_h = 842

    # Ornate border
    pdf.set_draw_color(25, 76, 25)
    pdf.set_line_width(3)
    pdf.rect(15, 15, page_w - 30, page_h - 30)
    pdf.set_draw_color(51, 127, 51)
    pdf.set_line_width(0.5)
    pdf.rect(22, 22, page_w - 44, page_h - 44)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.2)

    y = page_h - 70
    pdf_draw(pdf, "股票證書 / SHARE CERTIFICATE", page_w / 2 - 150, y, size=16)
    y -= 28
    pdf.line(80, y, page_w - 80, y)
    y -= 24

    pdf_draw(pdf, f"公司名稱: {company.get('name', '________________________________')}", 50, y, size=10)
    y -= 18
    if company.get('chinese_name'):
        pdf_draw(pdf, f"中文名稱: {company.get('chinese_name', '')}", 50, y, size=10)
        y -= 18
    pdf_draw(pdf, f"商業登記號碼: {company.get('company_number', '________________')}", 50, y, size=10)
    y -= 18
    reg_office = company.get('address') or company.get('registered_office', '')
    pdf_draw(pdf, f"註冊辦事處地址: {reg_office or '________________________________'}", 50, y, size=10)

    y -= 30
    pdf_draw(pdf, "茲證明 / THIS IS TO CERTIFY that", 50, y, size=10, gray=80)
    y -= 24

    holder_name = tx.get('to_name', '________________________________')
    share_class = tx.get('share_type', 'Ordinary')
    shares_val = tx.get('shares', '________')
    price = tx.get('price_per_share', '____')

    pdf_draw(pdf, holder_name, page_w / 2 - 80, y, size=13)
    y -= 22
    pdf_draw(pdf, "is/are the registered holder(s) of", 50, y, size=10, gray=80)
    y -= 22
    pdf_draw(pdf, f"{shares_val} {share_class} Share(s)", page_w / 2 - 60, y, size=13)
    y -= 22
    pdf_draw(pdf, f"of HK$ {price} each fully paid", 50, y, size=10, gray=80)
    y -= 22
    pdf_draw(pdf, "in the above-named Company", 50, y, size=10, gray=80)

    y -= 30
    pdf_draw(pdf, f"證書編號 Certificate No: {tx.get('instrument_number', '______________')}",
             50, y, size=9, gray=100)

    y -= 50
    pdf.line(50, y, 200, y)
    pdf.line(page_w - 200, y, page_w - 50, y)
    y -= 14
    pdf_draw(pdf, "董事 Director", 50, y, size=8, gray=100)
    pdf_draw(pdf, "公司秘書 Secretary", page_w - 200, y, size=8, gray=100)

    y -= 24
    pdf_draw(pdf, f"簽發日期 Issue Date: {tx.get('transaction_date', '________________')}",
             50, y, size=9)

    pdf_draw(pdf, "由 Muse Labs Engineering Limited 秘書系統生成",
             page_w / 2 - 80, 30, size=7, gray=150)


# ─── Generic CRUD ───
# SQL-safe identifier pattern: alphanumeric + underscore, must start with letter or underscore
_SAFE_ID_RE = re.compile(r'^[a-zA-Z_][a-zA-Z0-9_]*$')

def _safe_ident(name: str) -> bool:
    """Return True if `name` is a safe SQL identifier (no injection risk)."""
    return bool(_SAFE_ID_RE.match(name))


@app.route('/api/<table_name>', methods=['GET'])
def table_list(table_name):
    if table_name not in TABLES:
        return jsonify({'error': 'Not found'}), 404
    db = get_db()
    # Build query from search params — supports operator suffixes: __neq, __gt, __lt,
    # __gte, __lte, __like, __ilike, __in, __is.  Plain keys → eq.
    reserved = {'search', 'limit', 'offset', '_order', '_order_dir'}
    where = []
    bindings = []
    for key in request.args:
        if key in reserved:
            continue
        val = request.args.get(key)
        if val is None or val == '':
            continue
        # Operator suffix: column__op
        if '__' in key:
            col, op = key.rsplit('__', 1)
            if not _safe_ident(col):
                continue  # skip unsafe column names
            if op == 'neq':
                where.append(f"{col} != ?")
                bindings.append(val)
            elif op == 'gt':
                where.append(f"{col} > ?")
                bindings.append(val)
            elif op == 'lt':
                where.append(f"{col} < ?")
                bindings.append(val)
            elif op == 'gte':
                where.append(f"{col} >= ?")
                bindings.append(val)
            elif op == 'lte':
                where.append(f"{col} <= ?")
                bindings.append(val)
            elif op in ('like', 'ilike'):
                where.append(f"{col} LIKE ?")
                bindings.append(val)
            elif op == 'in':
                vals = [v.strip() for v in val.split(',') if v.strip()]
                if vals:
                    placeholders = ','.join(['?'] * len(vals))
                    where.append(f"{col} IN ({placeholders})")
                    bindings.extend(vals)
            elif op == 'is':
                where.append(f"{col} IS {val}")  # val is SQL literal e.g. NULL, NOT NULL
            else:
                where.append(f"{key} = ?")
                bindings.append(val)
        else:
            if not _safe_ident(key):
                continue  # skip unsafe column names
            where.append(f"{key} = ?")
            bindings.append(val)
    search = request.args.get('search')
    if search:
        s = f"%{search}%"
        where.append("(name LIKE ? OR name_english LIKE ? OR name_chinese LIKE ?)")
        bindings.extend([s, s, s])

    # ── Multi-tenant: filter companies by user's accessible company groups ──
    if table_name == 'companies':
        u = get_user()
        if u and u.get('role') != 'admin':
            db2 = get_db()
            user_row = db2.execute(
                "SELECT accessible_company_groups FROM auth_users WHERE id = ?", (u['id'],)
            ).fetchone()
            if user_row:
                groups = (user_row[0] or '').strip()
                if groups and groups != '*':
                    group_list = [g.strip() for g in groups.split(',') if g.strip()]
                    if group_list:
                        placeholders = ','.join(['?'] * len(group_list))
                        where.append(f"(company_group IN ({placeholders}) OR company_group IS NULL OR company_group = '')")
                        bindings.extend(group_list)
    sql = f"SELECT * FROM {table_name}"
    if where:
        sql += " WHERE " + " AND ".join(where)
    # Order: use _order param if provided, otherwise default to created_at DESC
    order_col = request.args.get('_order')
    if order_col and _safe_ident(order_col):
        order_dir = request.args.get('_order_dir', 'asc')
        sql += f" ORDER BY {order_col} {'DESC' if order_dir == 'desc' else 'ASC'}"
    else:
        try:
            db.execute(f"SELECT created_at FROM {table_name} LIMIT 1")
            sql += " ORDER BY created_at DESC"
        except sqlite3.OperationalError:
            pass
    limit = min(int(request.args.get( 'limit', '100')), 1000)
    offset = int(request.args.get( 'offset', '0'))
    sql += f" LIMIT {limit} OFFSET {offset}"
    rows = db.execute(sql, bindings).fetchall()
    return jsonify([dict(r) for r in rows])

@app.route('/api/<table_name>/<item_id>', methods=['GET'])
def table_get(table_name, item_id):
    if table_name not in TABLES:
        return jsonify({'error': 'Not found'}), 404
    db = get_db()
    row = db.execute(f"SELECT * FROM {table_name} WHERE id = ?", (item_id,)).fetchone()
    if not row:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(dict(row))

# ─── Company version snapshots (VE-01/02/03/08) ───
# 需納入版本快照 / 差異比對的公司欄位 → 繁體標籤
VERSION_FIELDS = {
    'name': '英文名稱', 'chinese_name': '中文名稱', 'company_number': '商業登記號碼',
    'ci_number': '公司註冊編號', 'trading_name': '商業名稱', 'business_nature': '業務性質',
    'company_type': '公司類型', 'business_code': '業務代碼', 'status': '狀態',
    'incorporation_date': '成立日期', 'jurisdiction': '司法管轄區',
    'reg_flat': '註冊地址-室/樓/座', 'reg_building': '註冊地址-大廈',
    'reg_street': '註冊地址-街道', 'reg_district': '註冊地址-區', 'reg_region': '註冊地址-地區',
    'email': '電郵地址', 'phone': '電話', 'signer_role_id': '簽署人',
}


def _company_snapshot(db, company_id):
    """回傳公司當前資料的 dict 快照（僅 VERSION_FIELDS 欄位）。"""
    row = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not row:
        return None
    d = dict(row)
    return {k: (d.get(k) if d.get(k) is not None else '') for k in VERSION_FIELDS}


def _latest_version(db, company_id):
    """回傳該公司最新版本 row(dict)，無則 None。"""
    r = db.execute(
        "SELECT * FROM company_versions WHERE company_id = ? ORDER BY version_no DESC LIMIT 1",
        (company_id,)
    ).fetchone()
    return dict(r) if r else None


def record_company_version(db, company_id, changed_by=''):
    """在公司資料變更後寫入一筆版本快照。僅當相對上一版有差異（或首版）時寫入。
    回傳新版本號，或 None（無變化未寫入）。"""
    snap = _company_snapshot(db, company_id)
    if snap is None:
        return None
    prev = _latest_version(db, company_id)
    changed = []
    if prev:
        try:
            prev_snap = json.loads(prev.get('snapshot') or '{}')
        except (ValueError, TypeError):
            prev_snap = {}
        for k in VERSION_FIELDS:
            if str(prev_snap.get(k, '')) != str(snap.get(k, '')):
                changed.append(k)
        if not changed:
            return None  # 無實質變化，不製造重複版本
        version_no = int(prev.get('version_no') or 0) + 1
    else:
        version_no = 1  # 首版（基線），changed 留空
    labels = [VERSION_FIELDS[k] for k in changed]
    summary = ('建立初始版本' if version_no == 1
               else '更新：' + '、'.join(labels) if labels else '更新')
    db.execute(
        "INSERT INTO company_versions (id, company_id, version_no, snapshot, changed_fields, change_summary, changed_by) "
        "VALUES (?, ?, ?, ?, ?, ?, ?)",
        (str(uuid.uuid4()), company_id, version_no, json.dumps(snap, ensure_ascii=False),
         json.dumps(changed, ensure_ascii=False), summary, changed_by)
    )
    return version_no


@app.route('/api/companies/<company_id>/versions', methods=['GET'])
def company_versions_list(company_id):
    """列出公司所有版本（新→舊），含解析後的 snapshot / changed_fields。"""
    db = get_db()
    rows = db.execute(
        "SELECT * FROM company_versions WHERE company_id = ? ORDER BY version_no DESC",
        (company_id,)
    ).fetchall()
    out = []
    for r in rows:
        d = dict(r)
        try:
            d['snapshot'] = json.loads(d.get('snapshot') or '{}')
        except (ValueError, TypeError):
            d['snapshot'] = {}
        try:
            d['changed_fields'] = json.loads(d.get('changed_fields') or '[]')
        except (ValueError, TypeError):
            d['changed_fields'] = []
        out.append(d)
    return jsonify(out)


@app.route('/api/companies/<company_id>/versions/snapshot', methods=['POST'])
def company_version_snapshot(company_id):
    """手動建立一個當前資料的版本快照（VE 手動存檔）。"""
    db = get_db()
    changed_by = (request.json or {}).get('changed_by', '') if request.is_json else ''
    v = record_company_version(db, company_id, changed_by)
    db.commit()
    # ── Write-through: sync to cloud ──
    if v is not None:
        row = db.execute(
            "SELECT * FROM company_versions WHERE company_id = ? AND version_no = ?",
            (company_id, v)).fetchone()
        if row:
            _sync_to_cloud("POST", "/api/company_versions", dict(row))
    return jsonify({'success': True, 'version_no': v, 'created': v is not None})


@app.route('/api/<table_name>', methods=['POST'])
def table_create(table_name):
    if table_name not in TABLES:
        return jsonify({'error': 'Not found'}), 404
    body = request.json
    # Support both single object and array (batch) inserts — matches CF Functions behavior
    rows = body if isinstance(body, list) else [body]
    if not rows:
        return jsonify({'error': 'Empty data'}), 400
    db = get_db()
    ids = []
    for data in rows:
        if not isinstance(data, dict):
            continue
        if 'id' not in data or not data['id']:
            data['id'] = str(uuid.uuid4())
        keys = list(data.keys())
        vals = list(data.values())
        placeholders = ', '.join(['?'] * len(keys))
        db.execute(f"INSERT INTO {table_name} ({', '.join(keys)}) VALUES ({placeholders})", vals)
        ids.append(data['id'])
    db.commit()
    # ── Write-through: sync to cloud ──
    if table_name in SYNC_TABLES:
        for data in rows:
            if isinstance(data, dict):
                _sync_to_cloud("POST", f"/api/{table_name}", data)
    if isinstance(body, list):
        return jsonify({'success': True, 'ids': ids, 'count': len(ids)}), 201
    return jsonify({'success': True, 'id': ids[0] if ids else None}), 201

@app.route('/api/<table_name>/<item_id>', methods=['PUT'])
def table_update(table_name, item_id):
    if table_name not in TABLES:
        return jsonify({'error': 'Not found'}), 404
    data = request.json
    db = get_db()
    sets = [f"{k} = ?" for k in data.keys()]
    vals = list(data.values()) + [item_id]
    db.execute(f"UPDATE {table_name} SET {', '.join(sets)}, updated_at = datetime('now') WHERE id = ?", vals)
    # 公司資料變更後，自動記錄版本快照（VE-01/02/03）
    if table_name == 'companies':
        try:
            record_company_version(db, item_id)
        except Exception as e:
            print(f"[VERSION] snapshot failed: {e}")
    db.commit()
    # ── Write-through: sync to cloud ──
    if table_name in SYNC_TABLES:
        _sync_to_cloud("PUT", f"/api/{table_name}/{item_id}", data)
    return jsonify({'success': True})

@app.route('/api/<table_name>/<item_id>', methods=['DELETE'])
def table_delete(table_name, item_id):
    if table_name not in TABLES:
        return jsonify({'error': 'Not found'}), 404
    db = get_db()

    if table_name == 'companies':
        # ─── Cascade delete: clean up all related records ───
        # 1. Find persons that will become orphaned (only in this company)
        orphan_person_ids = [
            row[0] for row in db.execute(
                "SELECT person_id FROM person_company_roles WHERE company_id = ?", (item_id,)
            ).fetchall()
        ]
        # 2. Delete all related records across tables
        for tbl in ['person_company_roles', 'reminders', 'company_logs',
                     'resolutions', 'significant_controllers', 'share_transactions', 'invoices']:
            db.execute(f"DELETE FROM {tbl} WHERE company_id = ?", (item_id,))
        # 3. Delete orphaned persons (no remaining roles in any company)
        for pid in orphan_person_ids:
            remaining = db.execute(
                "SELECT COUNT(*) FROM person_company_roles WHERE person_id = ?", (pid,)
            ).fetchone()[0]
            if remaining == 0:
                db.execute("DELETE FROM persons WHERE id = ?", (pid,))
        # 4. officers & shareholders have ON DELETE CASCADE, auto-cleaned
        db.execute("DELETE FROM companies WHERE id = ?", (item_id,))
    elif table_name == 'persons':
        # Cascade: clean up person_company_roles before deleting person
        db.execute("DELETE FROM person_company_roles WHERE person_id = ?", (item_id,))
        db.execute("DELETE FROM persons WHERE id = ?", (item_id,))
    else:
        db.execute(f"DELETE FROM {table_name} WHERE id = ?", (item_id,))

    db.commit()
    # ── Write-through: sync to cloud ──
    if table_name in SYNC_TABLES:
        _sync_to_cloud("DELETE", f"/api/{table_name}/{item_id}")
    return jsonify({'success': True})

@app.route('/api/<table_name>', methods=['DELETE'])
def table_delete_filtered(table_name):
    if table_name not in TABLES:
        return jsonify({'error': 'Not found'}), 404
    db = get_db()
    where = []
    bindings = []
    for key in request.args:
        where.append(f"{key} = ?")
        bindings.append(request.args.get(key))
    if where:
        db.execute(f"DELETE FROM {table_name} WHERE {' AND '.join(where)}", bindings)
    db.commit()
    return jsonify({'success': True})

# ─── Form → Company data auto-update (Phase 3.2) ───

def _apply_form_changes_to_company(data, form_code):
    """After generating a form PDF, check if the form data implies changes
    to the company record and apply them automatically.

    Also records a change_event for NAR1 tracking."""
    company_id = data.get('company_id')
    if not company_id:
        return

    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return

    changes = {}
    event_type = None

    if form_code == 'nr1':
        # Address change: update company registered address
        new_flat = data.get('flat', data.get('regFlat', ''))
        new_building = data.get('building', data.get('regBuilding', ''))
        new_street = data.get('street', data.get('regStreet', ''))
        new_district = data.get('district', data.get('regDistrict', ''))
        new_region = data.get('region', data.get('regRegion', ''))

        if new_flat: changes['reg_flat'] = new_flat
        if new_building: changes['reg_building'] = new_building
        if new_street: changes['reg_street'] = new_street
        if new_district: changes['reg_district'] = new_district
        if new_region: changes['reg_region'] = new_region

        if changes:
            event_type = 'address_change'

    elif form_code == 'nnc2':
        # Name change: update company name
        new_name = data.get('newName', data.get('newCompanyName', ''))
        new_chinese_name = data.get('newChineseName', '')

        if new_name and new_name != company['name']:
            changes['name'] = new_name
        if new_chinese_name and new_chinese_name != (company['chinese_name'] or ''):
            changes['chinese_name'] = new_chinese_name

        if changes:
            event_type = 'name_change'

    elif form_code == 'nd4':
        # Director/Secretary resignation — the date_ceased is already set
        # by useUpdateOfficer. No additional company update needed here.
        pass

    elif form_code == 'nd2a':
        # New director/secretary appointment — already handled by useAddOfficer.
        pass

    # Apply changes to companies table
    if changes:
        sets = ', '.join([f"{k} = ?" for k in changes.keys()])
        vals = list(changes.values()) + [company_id]
        db.execute(f"UPDATE companies SET {sets} WHERE id = ?", vals)
        db.commit()

        # Record change event
        if event_type:
            today = datetime.now().strftime('%d/%m/%Y')
            evt_id = str(uuid.uuid4())
            related_form = {'address_change': 'NR1', 'name_change': 'NNC2'}.get(event_type, '')
            db.execute(
                "INSERT INTO change_events (id, company_id, event_type, new_value, change_date, related_form_type) "
                "VALUES (?, ?, ?, ?, ?, ?)",
                (evt_id, company_id, event_type, json.dumps(changes, ensure_ascii=False), today, related_form)
            )
            db.commit()
            print(f"[FORM SYNC] {form_code} → company {company_id}: {event_type} {list(changes.keys())}")


def _apply_nd2b_changes_to_person(data):
    """After generating ND2B, apply person-level changes: address, name, ID, contact.
    Updates the persons table and records change_events for NAR1 tracking."""
    company_id = data.get('companyId') or data.get('company_id') or data.get('selectedCompanyId', '')
    person_id = data.get('personId') or data.get('person_id') or data.get('selectedPersonId', '')
    if not company_id or not person_id:
        print(f"[ND2B SYNC] Skipping — missing company_id({company_id}) or person_id({person_id})")
        return

    db = get_db()
    person = db.execute("SELECT * FROM persons WHERE id = ?", (person_id,)).fetchone()
    if not person:
        print(f"[ND2B SYNC] Person not found: {person_id}")
        return

    change_types = data.get('changeTypes', [])
    if isinstance(change_types, str):
        change_types = [change_types] if change_types else []
    # backward compat
    old_ct = data.get('changeType', '')
    if old_ct and old_ct not in change_types:
        change_types.append(old_ct)

    if not change_types:
        return

    changes = {}
    events = []
    effective_date = data.get('effectiveDate', '')

    for ct in change_types:
        if ct == 'address':
            fmap = {
                'addr_flat': data.get('newFlat', ''),
                'addr_building': data.get('newBuilding', ''),
                'addr_street': data.get('newStreet', ''),
                'addr_district': data.get('newDistrict', ''),
                'addr_region': data.get('newRegion', ''),
            }
            addr_changes = {}
            for col, val in fmap.items():
                if val and val != (person[col] or ''):
                    changes[col] = val
                    addr_changes[col] = val
            if addr_changes:
                events.append(('person_address_change', addr_changes))

        elif ct == 'name':
            new_chinese = data.get('newNameChinese', '')
            new_surname = data.get('newNameSurname', '')
            new_other = data.get('newNameOtherNames', '')
            new_english = f"{new_surname} {new_other}".strip() if new_surname or new_other else ''

            name_changes = {}
            # Move old name to previous_name if new name differs
            if new_chinese and new_chinese != (person['name_chinese'] or ''):
                if person['name_chinese'] and not person['previous_name_chinese']:
                    changes['previous_name_chinese'] = person['name_chinese']
                    name_changes['previous_name_chinese'] = person['name_chinese']
                elif person['name_chinese'] and not person['alias_chinese']:
                    changes['alias_chinese'] = person['name_chinese']
                    name_changes['alias_chinese'] = person['name_chinese']
                changes['name_chinese'] = new_chinese
                name_changes['name_chinese'] = new_chinese

            if new_english and new_english != (person['name_english'] or ''):
                if person['name_english'] and not person['previous_name_english']:
                    changes['previous_name_english'] = person['name_english']
                    name_changes['previous_name_english'] = person['name_english']
                elif person['name_english'] and not person['alias_english']:
                    changes['alias_english'] = person['name_english']
                    name_changes['alias_english'] = person['name_english']
                changes['name_english'] = new_english
                name_changes['name_english'] = new_english

            new_alias_eng = data.get('newAliasEnglish', '')
            new_alias_cn = data.get('newAliasChinese', '')
            if new_alias_eng:
                changes['alias_english'] = new_alias_eng
                name_changes['alias_english'] = new_alias_eng
            if new_alias_cn:
                changes['alias_chinese'] = new_alias_cn
                name_changes['alias_chinese'] = new_alias_cn

            if name_changes:
                events.append(('person_name_change', name_changes))

        elif ct == 'id':
            id_changes = {}
            new_id = data.get('newIdNumber', '')
            if new_id and new_id != (person['id_number'] or ''):
                changes['id_number'] = new_id
                id_changes['id_number'] = new_id
            new_passport = data.get('passportNumber', '')
            if new_passport and new_passport != (person['passport_number'] or ''):
                changes['passport_number'] = new_passport
                id_changes['passport_number'] = new_passport
            if id_changes:
                events.append(('person_id_change', id_changes))

        elif ct == 'contact':
            new_email = data.get('newEmail', '')
            if new_email and new_email != (person['email'] or ''):
                changes['email'] = new_email
                events.append(('person_contact_change', {'email': new_email}))

    # Apply changes to persons table
    if changes:
        sets = ', '.join([f"{k} = ?" for k in changes.keys()])
        vals = list(changes.values()) + [person_id]
        db.execute(f"UPDATE persons SET {sets} WHERE id = ?", vals)
        db.commit()
        print(f"[ND2B SYNC] person {person_id}: updated {list(changes.keys())}")

        # Record change events
        today = datetime.now().strftime('%d/%m/%Y')
        for evt_type, evt_changes in events:
            evt_id = str(uuid.uuid4())
            db.execute(
                "INSERT INTO change_events (id, company_id, event_type, person_id, role, "
                "new_value, change_date, related_form_type) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (evt_id, company_id, evt_type, person_id,
                 data.get('role', ''),
                 json.dumps(evt_changes, ensure_ascii=False),
                 effective_date or today, 'ND2B')
            )
            db.commit()

        # Create company log entry
        log_id = str(uuid.uuid4())
        person_name = person['name_english'] or person['name_chinese'] or 'Unknown'
        company_name_hint = data.get('companyName', '')
        db.execute(
            "INSERT INTO company_logs (id, company_id, company_name_hint, doc_type, doc_date, notes, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, datetime('now'))",
            (log_id, company_id, company_name_hint, 'ND2B', effective_date or today,
             f"ND2B 更改詳情 — {person_name}: {', '.join(change_types)}")
        )
        db.commit()
        print(f"[ND2B SYNC] change_events + log created for {person_name}")


# ─── Special routes ───
@app.route('/api/persons/cleanup-orphans', methods=['POST'])
def cleanup_orphan_persons():
    """Delete all persons that have no person_company_roles (no company binding)."""
    db = get_db()
    result = db.execute('''
        DELETE FROM persons WHERE id IN (
            SELECT p.id FROM persons p
            LEFT JOIN person_company_roles r ON p.id = r.person_id
            WHERE r.person_id IS NULL
        )
    ''')
    db.commit()
    count = result.rowcount
    return jsonify({'success': True, 'deleted': count})

@app.route('/api/search', methods=['GET'])
def search():
    q = request.args.get( 'q', '')
    if not q:
        return jsonify([])
    s = f"%{q}%"
    db = get_db()
    companies = db.execute(
        "SELECT id, name, chinese_name, company_number, ci_number, company_type, status, 'company' as type "
        "FROM companies WHERE name LIKE ? OR chinese_name LIKE ? OR company_number LIKE ? OR ci_number LIKE ? "
        "ORDER BY name LIMIT 30",
        (s, s, s, s)).fetchall()
    persons = db.execute(
        "SELECT id, name_english, name_chinese, identity, id_number, passport_number, 'person' as type "
        "FROM persons WHERE name_english LIKE ? OR name_chinese LIKE ? OR id_number LIKE ? OR passport_number LIKE ? "
        "ORDER BY name_english LIMIT 30",
        (s, s, s, s)).fetchall()
    out = [dict(r) for r in companies]
    for r in persons:
        p = dict(r)
        # 附上關聯公司+角色，讓前端點擊可定位公司登記冊
        p['roles'] = [dict(x) for x in db.execute(
            "SELECT pcr.role, pcr.date_ceased, c.id AS company_id, c.name AS company_name "
            "FROM person_company_roles pcr JOIN companies c ON c.id = pcr.company_id "
            "WHERE pcr.person_id = ?", (p['id'],)).fetchall()]
        out.append(p)
    return jsonify(out)

@app.route('/api/companies/<item_id>/full', methods=['GET'])
def company_full(item_id):
    db = get_db()
    company = db.execute("SELECT * FROM companies WHERE id = ?", (item_id,)).fetchone()
    if not company:
        return jsonify({'error': 'Not found'}), 404
    officers = db.execute("SELECT * FROM officers WHERE company_id = ?", (item_id,)).fetchall()
    # Read shareholders from person_company_roles (same source as frontend hooks)
    shareholder_roles = db.execute(
        "SELECT pcr.*, p.name_english AS person_name_english, p.name_chinese AS person_name_chinese, "
        "p.identity AS person_identity, p.id_number AS person_id_number, p.address AS person_address, "
        "p.email AS person_email, p.service_address AS person_service_address "
        "FROM person_company_roles pcr "
        "LEFT JOIN persons p ON p.id = pcr.person_id "
        "WHERE pcr.company_id = ? AND pcr.role = 'shareholder'",
        (item_id,)).fetchall()
    shareholders = [dict(r) for r in shareholder_roles]
    scrs = db.execute("SELECT * FROM significant_controllers WHERE company_id = ?", (item_id,)).fetchall()
    logs = db.execute("SELECT * FROM company_logs WHERE company_id = ?", (item_id,)).fetchall()
    return jsonify({**dict(company),
                    'officers': [dict(r) for r in officers],
                    'shareholders': [dict(r) for r in shareholders],
                    'significant_controllers': [dict(r) for r in scrs],
                    'logs': [dict(r) for r in logs]})

@app.route('/api/backup', methods=['POST'])
def backup():
    return jsonify({'success': True, 'message': 'Local backup skipped (dev mode)'})

# ─── NAR1 PDF 生成（本地 Python + PyMuPDF） ───
import fitz  # PyMuPDF

_CJK_RE = re.compile(r'[㐀-鿿豈-﫿]')
_PURE_NUMBER_RE = re.compile(r'^[\d,.\s]+$')
_ADDR_FLAT_RE = re.compile(
    r'^(?:flat|room|rm|unit|shop|suite|ste|workshop|portion|floor|fl|level|lvl|\d+/f|g/f|gf|lg/f|ug/f|m/f|b\d*/f)'
    r'|\b(?:tower|twr|block|blk)\b'
    r'|\ball\s+(?:that|those)\b',
    re.IGNORECASE
)
_ADDR_BUILDING_RE = re.compile(
    r'\b(?:building|bldg|mansion|centre|center|tower|house|plaza|estate|court|gardens|industrial|commercial|factory|hotel|complex|arcade|heights|square|place|lane|exchange|finance|financial|plaza|villa)\b',
    re.IGNORECASE
)
_ADDR_STREET_RE = re.compile(
    r'\b(?:road|street|avenue|drive|lane|path|way|boulevard|terrace|row|close|crescent|highway|bridge|pier|ferry|circus|central|plaza)\b|^\d',
    re.IGNORECASE
)
# Broad HK regions that should always merge into country (too large to be a district).
# "Kowloon" is intentionally excluded — it can be a valid district when no more-specific
# area is present (e.g. "22 Nathan Road, Kowloon, Hong Kong").
_ADDR_REGION_RE = re.compile(
    r'^(?:new\s+territories|n\.?\s*t\.?|新界)$',
    re.IGNORECASE
)
_ADDR_COUNTRY_RE = re.compile(
    r'(hong\s*kong|hk\b|china|prc|macau|macao|singapore|taiwan|united\s+\w+|\busa\b|\buk\b|canada|australia|japan|korea|h\.?k\.?\s*sar|kowloon|kln\b|new\s+territories|n\.?\s*t\.?\b|british\s+virgin\s+islands|bvi\b|香港|中國|澳門|台灣|新加坡|日本|韓國|英國|美國|加拿大|澳洲|九龍|新界)',
    re.IGNORECASE
)

def _is_ascii(s):
    return all(ord(c) < 128 for c in s)

def _parse_english_name(full_name):
    """返回 (surname, otherNames)"""
    cleaned = (full_name or '').replace(r'\s+', ' ').strip()
    if not cleaned or not any(c.isascii() and c.isalpha() for c in cleaned):
        return '', ''
    if _CJK_RE.search(cleaned):
        cleaned = _CJK_RE.sub(' ', cleaned).replace(r'\s+', ' ').strip()
        if not cleaned:
            return '', ''
    if ',' in cleaned:
        segs = [s.strip() for s in cleaned.split(',') if s.strip()]
        if len(segs) >= 2:
            return segs[0], ' '.join(segs[1:])
        if len(segs) == 1:
            return segs[0], ''
    parts = cleaned.split()
    if not parts:
        return '', ''
    surname = parts[0].rstrip(',')
    other = ' '.join(parts[1:]).lstrip(', ')
    return surname, other

def _parse_address(addr):
    """Parse a free-text address (typically comma-separated HK address) into
    {flat, building, street, district, country} for NAR1/ND2A/ND2B form fields.

    HK address convention: Flat → Floor → Building → Street No + Name → District → Region/Country
    """
    if not addr:
        return {'flat': '', 'building': '', 'street': '', 'district': '', 'country': ''}
    parts = [s.strip() for s in addr.split(',') if s.strip() and not _PURE_NUMBER_RE.match(s.strip())]
    if not parts:
        return {'flat': '', 'building': '', 'street': '', 'district': '', 'country': ''}

    # Single-part address: if it's a country/region → country, else → street
    if len(parts) == 1:
        if _ADDR_COUNTRY_RE.search(parts[0]):
            return {'flat': '', 'building': '', 'street': '', 'district': '', 'country': parts[0]}
        return {'flat': '', 'building': '', 'street': parts[0], 'district': '', 'country': ''}

    # ── Extract country (last part) ──
    country = ''
    if _ADDR_COUNTRY_RE.search(parts[-1]):
        country = parts.pop()

    # ── Extract district from the end ──
    district = ''
    if len(parts) >= 2:
        candidate = parts[-1]
        is_postal = len(candidate) <= 10 and re.match(r'^[A-Za-z]{0,3}\d[\dA-Za-z]*$', candidate)
        is_street = bool(_ADDR_STREET_RE.search(candidate))
        is_region = bool(_ADDR_REGION_RE.search(candidate))

        if is_region:
            # Broad region (New Territories / NT / 新界) → always merge into country
            region = parts.pop()
            country = region + ', ' + country
            # Re-extract: is there a real district before the region?
            if len(parts) >= 2:
                c2 = parts[-1]
                is_postal2 = len(c2) <= 10 and re.match(r'^[A-Za-z]{0,3}\d[\dA-Za-z]*$', c2)
                if not _ADDR_STREET_RE.search(c2) and not is_postal2 and not _ADDR_REGION_RE.search(c2) and not _ADDR_BUILDING_RE.search(c2) and not _ADDR_FLAT_RE.search(c2):
                    district = parts.pop()
        elif not is_street and not is_postal and not _ADDR_BUILDING_RE.search(candidate) and not _ADDR_FLAT_RE.search(candidate):
            district = parts.pop()

    # ── Extract flat/floor/unit/tower from the front ──
    flat_parts = []
    while len(parts) > 1 and _ADDR_FLAT_RE.search(parts[0]):
        flat_parts.append(parts.pop(0))
    flat = ', '.join(flat_parts)

    # ── Assign building and street ──
    if len(parts) == 1:
        part = parts[0]
        if _ADDR_BUILDING_RE.search(part) and not _ADDR_STREET_RE.search(part):
            building, street = part, ''
        elif _ADDR_FLAT_RE.search(part):
            flat = (flat + ', ' + part).strip(', ')
            building, street = '', ''
        else:
            building, street = '', part
    elif len(parts) >= 2:
        building = parts.pop(0)
        street = ', '.join(parts)
    else:
        building, street = '', ''

    return {'flat': flat, 'building': building, 'street': street, 'district': district, 'country': country}

def _parse_hkid_partial(id_number):
    if not id_number:
        return ''
    return re.sub(r'[()\-\s]', '', id_number).upper()[:4]

def _parse_passport_partial(passport_number):
    if not passport_number:
        return ''
    cleaned = re.sub(r'[^A-Za-z0-9]', '', passport_number).upper()
    return cleaned[:len(cleaned) // 2 + len(cleaned) % 2]

def _fmt_amount(n):
    return f'{n:,.2f}'

def _fmt_int(n):
    return f'{n:,}'

def _build_field_page_map(doc):
    """遍歷所有頁面，建立 field_name → page_index 映射（不存儲 widget 引用，避免 Annot not bound to page）"""
    fmap = {}
    for pi in range(doc.page_count):
        for w in doc[pi].widgets():
            name = w.field_name
            if name:
                fmap[name] = pi
    return fmap

def _set_text(doc, fmap, name, value, align='left'):
    """在指定頁面上查找 widget 並設置值（必須在迭代內完成 update，widget 引用不能外傳）
    align: 'left' (default), 'center', 'right'"""
    if name not in fmap:
        return False
    pi = fmap[name]
    for w in doc[pi].widgets():
        if w.field_name == name:
            try:
                w.field_value = value if value else ''
                if align == 'right':
                    doc.xref_set_key(w._annot.xref, 'Q', '2')
                elif align == 'center':
                    doc.xref_set_key(w._annot.xref, 'Q', '1')
                w.update()
                return True
            except Exception:
                pass
            break
    return False

def _set_text_size(doc, fmap, name, value, font_size):
    """在指定頁面上查找 widget，設定值並縮小字號"""
    if name not in fmap or not value:
        return False
    pi = fmap[name]
    for w in doc[pi].widgets():
        if w.field_name == name:
            try:
                w.text_fontsize = float(font_size)
                w.field_value = str(value)
                w.update()
                return True
            except Exception:
                pass
            break
    return False

def _check(doc, fmap, name, should_check):
    """在指定頁面上查找 checkbox 並勾選"""
    if not should_check or name not in fmap:
        return False
    pi = fmap[name]
    for w in doc[pi].widgets():
        if w.field_name == name:
            try:
                # PyMuPDF 1.28+: w._annot is an Annot object (not dict),
                # so .get('AP') raises AttributeError. Use w.field_value = True
                # which automatically selects the "On" appearance state.
                w.field_value = True
                w.update()
                return True
            except Exception:
                pass
            break
    return False


def _select_dropdown(doc, fmap, name, value):
    """在指定頁面上查找 dropdown 並設置選中值"""
    if name not in fmap:
        return False
    pi = fmap[name]
    for w in doc[pi].widgets():
        if w.field_name == name:
            try:
                w.field_value = value
                w.update()
                return True
            except Exception:
                pass
            break
    return False

# ─── Phase 4: NAR1 Smart Filing Engine ───

def _calculate_nar1_dates(incorporation_date_str):
    """Calculate NAR1 period and due dates from incorporation date.

    Due date = incorporation anniversary each year (recurring annually).
    Only returns the next upcoming filing deadline.

    Returns dict with:
      - period_start: DD/MM/YYYY (incorporation anniversary this year)
      - period_end: DD/MM/YYYY (next anniversary = due date)
      - due_date: DD/MM/YYYY (same as period_end, the next filing deadline)
      - days_remaining: int (days until due_date, negative if overdue)
      - status: 'ok' | 'grace' | 'due_soon' | 'late'
    """
    from datetime import datetime, timedelta
    today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)

    try:
        inc_date = datetime.strptime(incorporation_date_str, '%Y-%m-%d')
    except (ValueError, TypeError):
        # Try DD/MM/YYYY
        try:
            inc_date = datetime.strptime(incorporation_date_str, '%d/%m/%Y')
        except (ValueError, TypeError):
            return None

    # Current year's anniversary
    this_year_anniv = inc_date.replace(year=today.year)
    if this_year_anniv > today:
        # This year's anniversary hasn't happened yet → period started last year
        period_start = this_year_anniv.replace(year=today.year - 1)
        period_end = this_year_anniv
    else:
        # This year's anniversary has passed → period started this year
        period_start = this_year_anniv
        period_end = this_year_anniv.replace(year=today.year + 1)

    # Due date = period_end (incorporation anniversary)
    due_date = period_end

    days_remaining = (due_date - today).days

    # Determine status
    if days_remaining < 0:
        status = 'late'
    elif days_remaining <= 30:
        status = 'due_soon'
    elif days_remaining <= 90:
        status = 'grace'
    else:
        status = 'ok'

    return {
        'period_start': period_start.strftime('%d/%m/%Y'),
        'period_end': period_end.strftime('%d/%m/%Y'),
        'due_date': due_date.strftime('%d/%m/%Y'),
        'days_remaining': days_remaining,
        'status': status,
        'today': today.strftime('%d/%m/%Y'),
    }


def _ensure_nar1_periods(company_id):
    """Create nar1_filings periods for a company if they don't exist.
    Creates all periods from incorporation year to current year.
    Returns list of period dicts."""
    db = get_db()
    company = db.execute(
        "SELECT incorporation_date, nar1_due_date FROM companies WHERE id = ?",
        (company_id,)
    ).fetchone()
    if not company or not company['incorporation_date']:
        return []

    dates = _calculate_nar1_dates(company['incorporation_date'])
    if not dates:
        return []

    # Update company's nar1_due_date if changed
    if company['nar1_due_date'] != dates['due_date']:
        db.execute(
            "UPDATE companies SET nar1_due_date = ? WHERE id = ?",
            (dates['due_date'], company_id)
        )
        db.commit()

    # Check existing periods
    existing = db.execute(
        "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start",
        (company_id,)
    ).fetchall()

    existing_starts = {r['period_start'] for r in existing}

    # Create current period if missing
    if dates['period_start'] not in existing_starts:
        period_id = str(uuid.uuid4())
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        db.execute(
            "INSERT INTO nar1_filings (id, company_id, period_start, period_end, status, created_at) "
            "VALUES (?, ?, ?, ?, 'pending', ?)",
            (period_id, company_id, dates['period_start'], dates['period_end'], now)
        )
        db.commit()

    # Sync NAR1 reminders (30, 7, 1 days before due)
    _sync_nar1_reminders(company_id, dates['due_date'])

    # Refresh and return all periods
    return [dict(r) for r in db.execute(
        "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start DESC",
        (company_id,)
    ).fetchall()]


def _sync_nar1_reminders(company_id, due_date_str):
    """Create/update NAR1 reminders for a company.
    Creates reminders at 30, 7, and 1 days before the due date.
    If reminders already exist, updates their due_date. Deletes old ones
    if the due date has moved past a reminder date."""
    db = get_db()
    from datetime import datetime as dt, timedelta

    try:
        due_date = dt.strptime(due_date_str, '%d/%m/%Y')
    except (ValueError, TypeError):
        return

    # Reminder offsets: (days_before, label_suffix)
    reminders_config = [
        (30, ' - 30天前提醒'),
        (7, ' - 7天前提醒'),
        (1, ' - 到期日提醒'),
    ]

    # Get company name
    company = db.execute("SELECT name FROM companies WHERE id = ?", (company_id,)).fetchone()
    company_name = company['name'] if company else 'Unknown'

    # Get existing NAR1 reminders for this company
    existing = db.execute(
        "SELECT * FROM reminders WHERE company_id = ? AND reminder_type = 'NAR1' "
        "ORDER BY due_date",
        (company_id,)
    ).fetchall()

    existing_dates = {r['due_date']: r for r in existing}
    now = dt.now().strftime('%Y-%m-%d %H:%M:%S')

    for days_before, label_suffix in reminders_config:
        reminder_date = due_date - timedelta(days=days_before)
        reminder_date_str = reminder_date.strftime('%d/%m/%Y')

        if reminder_date_str in existing_dates:
            # Already exists, skip
            del existing_dates[reminder_date_str]
            continue

        # Check if we already have a reminder for this offset
        found = False
        for r in existing:
            # Check if existing reminder date matches this offset
            try:
                ed = dt.strptime(r['due_date'], '%d/%m/%Y')
                expected = due_date - timedelta(days=days_before)
                if ed == expected:
                    found = True
                    del existing_dates[r['due_date']]
                    break
            except (ValueError, TypeError):
                pass

        if not found and reminder_date > dt.now():
            reminder_id = str(uuid.uuid4())
            title = f'NAR1 週年申報{label_suffix} - {company_name}'
            db.execute(
                "INSERT INTO reminders (id, company_id, reminder_type, title, due_date, status, notes, created_at, updated_at) "
                "VALUES (?, ?, 'NAR1', ?, ?, 'pending', ?, ?, ?)",
                (reminder_id, company_id, title, reminder_date_str,
                 f'到期日: {due_date_str}', now, now)
            )
            db.commit()

    # Delete any remaining existing reminders that no longer match
    # (due date has changed and old reminders are stale)
    for stale_date, stale_reminder in existing_dates.items():
        # Only delete if not already past (keep completed reminders)
        if stale_reminder['status'] == 'pending':
            db.execute("DELETE FROM reminders WHERE id = ?", (stale_reminder['id'],))
            db.commit()


def _assign_changes_to_nar1_periods(company_id):
    """Auto-assign unassigned change_events to the correct nar1_filings period
    based on change_date falling within period_start <= change_date < period_end.
    Changes after period_end up to filing_date belong to next period.

    Returns: {assigned_count: int, periods: [...]}
    """
    db = get_db()

    # Get unassigned events
    events = db.execute(
        "SELECT * FROM change_events WHERE company_id = ? AND (nar1_period_id = '' OR nar1_period_id IS NULL) "
        "ORDER BY change_date",
        (company_id,)
    ).fetchall()

    # Get all periods for this company
    periods = db.execute(
        "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start",
        (company_id,)
    ).fetchall()

    if not events or not periods:
        return {'assigned_count': 0, 'periods': [dict(p) for p in periods]}

    assigned = 0
    for event in events:
        evt_date = event['change_date']  # DD/MM/YYYY
        try:
            d = datetime.strptime(evt_date, '%d/%m/%Y')
        except (ValueError, TypeError):
            continue

        # Find matching period
        for period in periods:
            try:
                ps = datetime.strptime(period['period_start'], '%d/%m/%Y')
                pe = datetime.strptime(period['period_end'], '%d/%m/%Y')
            except (ValueError, TypeError):
                continue
            if ps <= d < pe:
                db.execute(
                    "UPDATE change_events SET nar1_period_id = ? WHERE id = ?",
                    (period['id'], event['id'])
                )
                assigned += 1
                break
            # If event date >= period_end, it falls into the NEXT period
            # (handled by the next iteration in the period loop)

    db.commit()

    # Refresh periods
    periods = db.execute(
        "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start DESC",
        (company_id,)
    ).fetchall()

    return {'assigned_count': assigned, 'periods': [dict(p) for p in periods]}


def _get_nar1_changes_summary(company_id, period_id=None):
    """Get a summary of changes for a specific NAR1 period.
    If period_id is None, returns changes for the current (most recent) period.

    Returns: {period: {...}, changes: [...], summary: {director_changes, secretary_changes, ...}}
    """
    db = get_db()

    # Get periods
    periods = db.execute(
        "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start DESC",
        (company_id,)
    ).fetchall()

    if not periods:
        # Ensure periods exist
        _ensure_nar1_periods(company_id)
        periods = db.execute(
            "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start DESC",
            (company_id,)
        ).fetchall()

    if not periods:
        return {'period': None, 'changes': [], 'summary': {}}

    target = None
    if period_id:
        for p in periods:
            if p['id'] == period_id:
                target = p
                break
    if not target:
        target = periods[0]  # most recent period

    # Assign unassigned changes first
    _assign_changes_to_nar1_periods(company_id)

    # Get changes for this period
    changes = db.execute(
        "SELECT * FROM change_events WHERE company_id = ? AND nar1_period_id = ? "
        "ORDER BY change_date DESC",
        (company_id, target['id'])
    ).fetchall()

    # Build summary counts
    summary = {
        'total_changes': len(changes),
        'director_appointments': 0,
        'director_cessations': 0,
        'secretary_appointments': 0,
        'secretary_cessations': 0,
        'shareholder_changes': 0,
        'share_transfers': 0,
        'share_allotments': 0,
        'address_changes': 0,
        'name_changes': 0,
        'other_changes': 0,
    }

    for c in changes:
        et = c['event_type']
        if et == 'director_appoint': summary['director_appointments'] += 1
        elif et == 'director_cease': summary['director_cessations'] += 1
        elif et == 'secretary_appoint': summary['secretary_appointments'] += 1
        elif et == 'secretary_cease': summary['secretary_cessations'] += 1
        elif et in ('shareholder_add', 'shareholder_remove'): summary['shareholder_changes'] += 1
        elif et == 'share_transfer': summary['share_transfers'] += 1
        elif et == 'share_allotment': summary['share_allotments'] += 1
        elif et == 'address_change': summary['address_changes'] += 1
        elif et == 'name_change': summary['name_changes'] += 1
        else: summary['other_changes'] += 1

    return {
        'period': {
            'id': target['id'],
            'period_start': target['period_start'],
            'period_end': target['period_end'],
            'filing_date': target['filing_date'],
            'status': target['status'],
        },
        'changes': [dict(c) for c in changes],
        'summary': summary,
    }


def _fill_nar1_pdf(data):
    """填充 NAR1 PDF 模板，返回 bytes"""
    import os
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NAR1-template-new.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    return_date = data.get('returnDate', '')
    if return_date:
        parts = return_date.split('-')
        year, month, day = parts[0], parts[1], parts[2]
    else:
        from datetime import date
        today = date.today()
        year, month, day = str(today.year), f'{today.month:02d}', f'{today.day:02d}'

    office = data.get('registeredOffice') or {}
    # NAR1 is a HK Companies Registry form — if registered office has
    # address fields but no country/region, default to 'Hong Kong'
    if not office.get('region') and not office.get('country'):
        has_addr = any(office.get(k) for k in ['flat', 'building', 'street', 'district'])
        if has_addr:
            office = dict(office)
            office['region'] = 'Hong Kong'
    br8 = (data.get('brNumber') or '').replace(r'[^0-9A-Za-z]', '')[:8]
    company_type = data.get('companyType') or ''
    ct_lower = company_type.lower()

    # ── P.1 公司資料 ──
    _set_text(doc, fmap, 'fill_1_P.1', br8)
    full_name = '\n'.join(filter(None, [data.get('name', ''), data.get('chineseName', '')]))
    _set_text(doc, fmap, 'fill_2_P.1', full_name)
    _set_text(doc, fmap, 'fill_3_P.1', data.get('tradingName', ''))
    _check(doc, fmap, 'cb_1_P.1', '私人' in company_type or 'private' in ct_lower)
    _check(doc, fmap, 'cb_2_P.1', '公眾' in company_type or 'public' in ct_lower)
    _check(doc, fmap, 'cb_3_P.1', '擔保' in company_type)
    _set_text(doc, fmap, 'fill_4_P.1', data.get('businessCode', ''))
    _set_text(doc, fmap, 'fill_5_P.1', data.get('businessNature', ''))
    _set_text(doc, fmap, 'fill_6_P.1', day)
    _set_text(doc, fmap, 'fill_7_P.1', month)
    _set_text(doc, fmap, 'fill_8_P.1', year)
    _set_text(doc, fmap, 'fill_15_P.1', office.get('flat', ''))
    _set_text(doc, fmap, 'fill_16_P.1', office.get('building', ''))
    _set_text(doc, fmap, 'fill_17_P.1', office.get('street', ''))
    _set_text(doc, fmap, 'fill_18_P.1', office.get('district', ''))
    # 區域下拉
    p1_region = office.get('region') or office.get('country', '')
    if p1_region:
        for name in ['Dropdown1_P.1', 'Dropdown_1_P.1']:
            if name in fmap:
                pi = fmap[name]
                for w in doc[pi].widgets():
                    if w.field_name == name:
                        try:
                            w.field_value = p1_region
                            w.update()
                        except Exception:
                            pass
                        break
                break

    presenter = data.get('presenter') or {}
    if presenter.get('name'):
        _set_text(doc, fmap, 'fill_19_P.1', presenter['name'])
    if presenter.get('address'):
        _set_text(doc, fmap, 'fill_20_P.1', presenter['address'])
    if presenter.get('phone'):
        _set_text_size(doc, fmap, 'fill_21_P.1', presenter['phone'], 10)
    if presenter.get('fax'):
        _set_text_size(doc, fmap, 'fill_22_P.1', presenter['fax'], 10)
    if presenter.get('email'):
        _set_text_size(doc, fmap, 'fill_23_P.1', presenter['email'], 10)
    if presenter.get('reference'):
        _set_text_size(doc, fmap, 'fill_24_P.1', presenter['reference'], 10)

    # ── P.2 股本 ──
    _set_text(doc, fmap, 'fill_1_P.2', br8)
    # 公司電郵（fill_2_P.2）和電話（fill_3_P.2）
    company_email = data.get('companyEmail', '')
    company_phone = data.get('companyPhone', '')
    if company_email:
        _set_text(doc, fmap, 'fill_2_P.2', company_email)
    if company_phone:
        _set_text(doc, fmap, 'fill_3_P.2', company_phone)
    shareholders = data.get('shareholders') or []

    def _norm_class(raw):
        t = (raw or '').strip()
        if not t or 'ord' in t.lower() or '普通' in t:
            return 'ORDINARY SHARES'
        if 'pref' in t.lower() or '優先' in t:
            return 'PREFERENCE SHARES'
        return t.upper()

    def _norm_currency(raw):
        c = (raw or 'HKD').strip().upper()
        return 'HK$' if c in ('HKD', 'HK$') else ('US$' if c in ('USD', 'US$') else c)

    share_type_map = {}
    for sh in shareholders:
        cls = _norm_class(sh.get('shareType', ''))
        cur = _norm_currency(sh.get('currency', ''))
        ip = float(sh.get('issuePrice', 0) or 0)
        key = f'{cls}||{cur}||{ip}'
        if key not in share_type_map:
            share_type_map[key] = {'className': cls, 'currency': cur, 'issuePrice': ip, 'shares': 0, 'paidUp': 0.0, 'unpaid': 0.0}
        info = share_type_map[key]
        info['shares'] += int(sh.get('shares', 0) or 0)
        info['paidUp'] += float(sh.get('paidUp', 0) or 0)
        info['unpaid'] += float(sh.get('unpaid', 0) or 0)

    share_infos = list(share_type_map.values())
    total_shares, total_amount, total_paid, first_currency = 0, 0.0, 0.0, ''
    for i, info in enumerate(share_infos[:4]):
        base = 6 + i * 5
        issued = (info['paidUp'] + info['unpaid']) or (info['issuePrice'] * info['shares']) or (info['shares'] * 1.0)
        _set_text(doc, fmap, f'fill_{base}_P.2', info['className'])
        _set_text(doc, fmap, f'fill_{base+1}_P.2', info['currency'])
        _set_text(doc, fmap, f'fill_{base+2}_P.2', _fmt_int(info['shares']))
        _set_text(doc, fmap, f'fill_{base+3}_P.2', _fmt_amount(issued))
        _set_text(doc, fmap, f'fill_{base+4}_P.2', _fmt_amount(info['paidUp'] or issued))
        total_shares += info['shares']
        total_amount += issued
        total_paid += info['paidUp'] or issued
        if not first_currency:
            first_currency = info['currency']

    if share_infos:
        _set_text(doc, fmap, 'fill_26_P.2', first_currency)
        _set_text(doc, fmap, 'fill_27_P.2', _fmt_int(total_shares))
        _set_text(doc, fmap, 'fill_28_P.2', _fmt_amount(total_amount))
        _set_text(doc, fmap, 'fill_29_P.2', _fmt_amount(total_paid))

    secretaries = data.get('secretaries') or []
    directors = data.get('directors') or []
    nat_secs = [s for s in secretaries if s.get('identity') == 'natural']
    corp_secs = [s for s in secretaries if s.get('identity') == 'corporate']
    nat_dirs = [d for d in directors if d.get('identity') == 'natural']
    corp_dirs = [d for d in directors if d.get('identity') == 'corporate']

    # ── P.3 自然人秘書 ──
    _set_text(doc, fmap, 'fill_1_P.3', br8)
    if nat_secs:
        s = nat_secs[0]
        surname, other = _parse_english_name(s.get('nameEnglish', ''))
        _set_text(doc, fmap, 'fill_2_P.3', s.get('nameChinese', ''))
        _set_text(doc, fmap, 'fill_3_P.3', surname)
        _set_text(doc, fmap, 'fill_4_P.3', other)
        addr = _parse_address(s.get('address', ''))
        _set_text(doc, fmap, 'fill_9_P.3', addr['flat'])
        _set_text(doc, fmap, 'fill_10_P.3', addr['building'])
        _set_text(doc, fmap, 'fill_11_P.3', addr['street'])
        _set_text(doc, fmap, 'fill_12_P.3', addr['district'])
        _set_text(doc, fmap, 'fill_13_P.3', s.get('email', ''))
        hkid = _parse_hkid_partial(s.get('idNumber', ''))
        if hkid:
            _set_text(doc, fmap, 'fill_14_P.3', hkid, align='right')
        if s.get('passportCountry'):
            _set_text(doc, fmap, 'fill_15_P.3', s.get('passportCountry', ''))
        if s.get('passportNumber'):
            _set_text(doc, fmap, 'fill_16_P.3', _parse_passport_partial(s['passportNumber']))

    # ── P.4 法人秘書 ──
    _set_text(doc, fmap, 'fill_1_P.4', br8)
    if corp_secs:
        s = corp_secs[0]
        _set_text(doc, fmap, 'fill_2_P.4', s.get('nameChinese', ''))
        _set_text(doc, fmap, 'fill_3_P.4', s.get('nameEnglish', ''))
        addr = _parse_address(s.get('serviceAddress') or s.get('address', ''))
        _set_text(doc, fmap, 'fill_4_P.4', addr['flat'])
        _set_text(doc, fmap, 'fill_5_P.4', addr['building'])
        _set_text(doc, fmap, 'fill_6_P.4', addr['street'])
        _set_text(doc, fmap, 'fill_7_P.4', addr['district'])
        _set_text(doc, fmap, 'fill_8_P.4', s.get('email', ''))
        _set_text(doc, fmap, 'fill_9_P.4', s.get('companyNumberRef') or s.get('brNumber', '') or s.get('idNumber', ''))
        tcsp = s.get('tcspNumber', '') or s.get('licenceNumber', '')
        if tcsp:
            _set_text(doc, fmap, 'fill_10_P.4', tcsp)

    # ── P.5 自然人董事 ──
    _set_text(doc, fmap, 'fill_1_P.5', br8)
    if nat_dirs:
        d = nat_dirs[0]
        _check(doc, fmap, 'cb_1_P.5', True)
        surname, other = _parse_english_name(d.get('nameEnglish', ''))
        _set_text(doc, fmap, 'fill_3_P.5', d.get('nameChinese', ''))
        _set_text(doc, fmap, 'fill_4_P.5', surname)
        _set_text(doc, fmap, 'fill_5_P.5', other)
        # Directors use registered office as address
        _set_text(doc, fmap, 'fill_10_P.5', office.get('flat', ''))
        _set_text(doc, fmap, 'fill_11_P.5', office.get('building', ''))
        _set_text(doc, fmap, 'fill_12_P.5', office.get('street', ''))
        _set_text(doc, fmap, 'fill_13_P.5', office.get('district', ''))
        _set_text(doc, fmap, 'fill_14_P.5', office.get('region') or office.get('country', ''))
        _set_text(doc, fmap, 'fill_15_P.5', d.get('email', ''))
        hkid = _parse_hkid_partial(d.get('idNumber', ''))
        if hkid:
            _set_text(doc, fmap, 'fill_16_P.5', hkid, align='right')
        if d.get('passportCountry'):
            _set_text(doc, fmap, 'fill_17_P.5', d.get('passportCountry', ''))
        if d.get('passportNumber'):
            _set_text(doc, fmap, 'fill_18_P.5', _parse_passport_partial(d['passportNumber']))

    # ── P.6 法人董事 ──
    _set_text(doc, fmap, 'fill_1_P.6', br8)
    if corp_dirs:
        d = corp_dirs[0]
        _check(doc, fmap, 'cb_1_P.6', True)
        _set_text(doc, fmap, 'fill_3_P.6', d.get('nameChinese', ''))
        _set_text(doc, fmap, 'fill_4_P.6', d.get('nameEnglish', ''))
        _set_text(doc, fmap, 'fill_5_P.6', office.get('flat', ''))
        _set_text(doc, fmap, 'fill_6_P.6', office.get('building', ''))
        _set_text(doc, fmap, 'fill_7_P.6', office.get('street', ''))
        _set_text(doc, fmap, 'fill_8_P.6', office.get('district', ''))
        _set_text(doc, fmap, 'fill_9_P.6', office.get('region') or office.get('country', ''))
        _set_text(doc, fmap, 'fill_10_P.6', d.get('email', ''))
        _set_text(doc, fmap, 'fill_11_P.6', d.get('companyNumberRef') or d.get('brNumber', '') or d.get('idNumber', ''))

    # ── P.7 ──
    _set_text(doc, fmap, 'fill_1_P.7', br8)

    # ── P.8 總結 + 簽署 ──
    _set_text(doc, fmap, 'fill_1_P.8', br8)
    # 14. 股本/公司成員詳情 — 預設勾選「非上市公司的成員詳情列於附表一」
    _check(doc, fmap, 'cb_1_P.8', True)
    valid_members = [sh for sh in shareholders if (int(sh.get('shares', 0) or 0)) > 0]
    is_listed = '上市' in company_type or 'listed' in ct_lower
    if not is_listed:
        _check(doc, fmap, 'cb_4_P.8', True)

    sheet_a = max(0, len(nat_secs) - 1)
    sheet_b = max(0, len(corp_secs) - 1)
    sheet_c = max(0, len(nat_dirs) - 1)
    sheet_d_pages = max(0, (len(corp_dirs) - 1 + 1) // 2) if len(corp_dirs) > 1 else 0
    sched1_pages = 0 if is_listed else ((len(valid_members) + 1) // 2 if valid_members else 0)
    sched2_pages = 1 if is_listed else 0

    if sheet_a > 0:
        _set_text(doc, fmap, 'fill_4_P.8', str(sheet_a))
    if sheet_b > 0:
        _set_text(doc, fmap, 'fill_5_P.8', str(sheet_b))
    if sheet_c > 0:
        _set_text(doc, fmap, 'fill_6_P.8', str(sheet_c))
    if sheet_d_pages > 0:
        _set_text(doc, fmap, 'fill_7_P.8', str(sheet_d_pages))
    if sched1_pages > 0:
        _set_text(doc, fmap, 'fill_9_P.8', str(sched1_pages))
    if sched2_pages > 0:
        _set_text(doc, fmap, 'fill_10_P.8', str(sched2_pages))

    # 簽署人
    signer = data.get('signer')
    signer_name = (signer or {}).get('name') or presenter.get('name', '')
    signer_role = (signer or {}).get('role', '')
    if signer_name:
        _set_text(doc, fmap, 'fill_11_P.8', signer_name)
    if day and month and year:
        _set_text(doc, fmap, 'fill_12_P.8', f'{day}/{month}/{year}')

    # Handle Director / Company Secretary cross-out dropdowns on P.8
    # Dropdown_1_P.8 → strikethrough "Director"  (widget at x=143-205, y=745-756)
    # Dropdown_2_P.8 → strikethrough "Company Secretary" (widget at x=209-343, y=745-756)
    # Each dropdown has two options in /Opt: ' ' (blank) and hex bytes <8484...>
    # that render as a horizontal line with the original font.  We set /V to the
    # hex option + /I index, then draw a PDF line to guarantee visible rendering.
    if signer_role in ('director', 'secretary'):
        # signer_role='secretary' → cross out "Director"     → Dropdown_1_P.8
        # signer_role='director'  → cross out "Company Secretary" → Dropdown_2_P.8
        cross_out_widget = 'Dropdown_1_P.8' if signer_role == 'secretary' else 'Dropdown_2_P.8'

        for widget_name in ('Dropdown_1_P.8', 'Dropdown_2_P.8'):
            if widget_name in fmap:
                pi = fmap[widget_name]
                for w in doc[pi].widgets():
                    if w.field_name == widget_name:
                        try:
                            use_dashes = (widget_name == cross_out_widget)
                            # Set the dropdown to the correct option via its /Opt index
                            opt_idx = 1 if use_dashes else 0
                            doc.xref_set_key(w._annot.xref, 'I', f'[{opt_idx}]')
                            # Set /V to the chosen option value
                            val = w.choice_values[opt_idx]
                            doc.xref_set_key(w._annot.xref, 'V', fitz.get_pdf_str(val))
                            # Visible + printable
                            doc.xref_set_key(w._annot.xref, 'F', '4')

                            # Draw a visible line through the text area
                            # (the font-based AP from the original template is
                            # hard to replicate without its exact subset font)
                            if use_dashes:
                                doc[pi].draw_line(
                                    fitz.Point(w.rect.x0 + 2, w.rect.y0 + w.rect.height / 2),
                                    fitz.Point(w.rect.x1 - 2, w.rect.y0 + w.rect.height / 2),
                                    color=(0, 0, 0), width=1.0
                                )
                        except Exception:
                            pass
                        break

    # ── P.9 附表一（股東，前2人） ──
    if valid_members and not is_listed:
        _set_text(doc, fmap, 'fill_1_P.9', day)
        _set_text(doc, fmap, 'fill_2_P.9', month)
        _set_text(doc, fmap, 'fill_3_P.9', year)
        _set_text(doc, fmap, 'fill_4_P.9', br8)
        if share_infos:
            _set_text(doc, fmap, 'fill_5_P.9', share_infos[0]['className'])
            _set_text(doc, fmap, 'fill_6_P.9', _fmt_int(share_infos[0]['shares']))

        # P.9 slot layout (verified against template field positions 2026-07-24):
        #   pos0=name(7/18)  pos1=shares(16/27)  pos2=surname(8/19)
        #   pos3=other(9/20)  pos4=joint(10/21)
        #   pos5-9=address(11-15/22-26)  pos10=full_addr(17/28)
        slots_ts = [
            {'name': 7, 'surname': 8, 'other': 9, 'shares': 16, 'flat': 11, 'building': 12, 'street': 13, 'district': 14, 'country': 15},
            {'name': 18, 'surname': 19, 'other': 20, 'shares': 27, 'flat': 22, 'building': 23, 'street': 24, 'district': 25, 'country': 26},
        ]
        for idx, sh in enumerate(valid_members[:2]):
            F = slots_ts[idx]
            is_corp = sh.get('identity') == 'corporate'
            full = sh.get('nameEnglish') or sh.get('name', '')
            surname, other = _parse_english_name(full)
            addr = _parse_address(sh.get('address', ''))
            def _safe(v):
                return '' if (v and _PURE_NUMBER_RE.match(v)) else v
            country = _safe(addr['country']) or 'Hong Kong'

            _set_text(doc, fmap, f'fill_{F["name"]}_P.9', sh.get('nameChinese', ''))
            if is_corp:
                _set_text(doc, fmap, f'fill_{F["surname"]}_P.9', full)
            else:
                _set_text(doc, fmap, f'fill_{F["surname"]}_P.9', surname)
                _set_text(doc, fmap, f'fill_{F["other"]}_P.9', other)
            shares_num = int(sh.get('shares', 0) or 0)
            _set_text(doc, fmap, f'fill_{F["shares"]}_P.9', _fmt_int(shares_num) if shares_num > 0 else '0')
            _set_text(doc, fmap, f'fill_{F["flat"]}_P.9', _safe(addr['flat']))
            _set_text(doc, fmap, f'fill_{F["building"]}_P.9', _safe(addr['building']))
            _set_text(doc, fmap, f'fill_{F["street"]}_P.9', _safe(addr['street']))
            _set_text(doc, fmap, f'fill_{F["district"]}_P.9', _safe(addr['district']))
            _set_text(doc, fmap, f'fill_{F["country"]}_P.9', country)

        total_sch1 = (len(valid_members) + 1) // 2
        _set_text(doc, fmap, 'fill_29_P.9', '1')  # current page
        _set_text(doc, fmap, 'fill_30_P.9', str(total_sch1))

    # ── P.10 附表一續（股東 #3+#4）──
    if len(valid_members) > 2 and not is_listed:
        _set_text(doc, fmap, 'fill_1_P.10', day)
        _set_text(doc, fmap, 'fill_2_P.10', month)
        _set_text(doc, fmap, 'fill_3_P.10', year)
        _set_text(doc, fmap, 'fill_4_P.10', br8)
        if share_infos:
            _set_text(doc, fmap, 'fill_5_P.10', share_infos[0]['className'])
            _set_text(doc, fmap, 'fill_6_P.10', _fmt_int(share_infos[0]['shares']))
        # P.10 uses same field structure as P.9 (same x positions, different y)
        #   Slot1: name(7), shares(16), surname(8), other(9), joint(10)
        #          addr: flat(11), building(12), street(13), district(14), country(15)
        #   Slot2: name(19), shares(28), surname(20), other(21), joint(22)
        #          addr: flat(23), building(24), street(25), district(26), country(27)
        slots_p10 = [
            {'name': 7, 'surname': 8, 'other': 9, 'shares': 16, 'flat': 11, 'building': 12, 'street': 13, 'district': 14, 'country': 15},
            {'name': 19, 'surname': 20, 'other': 21, 'shares': 28, 'flat': 23, 'building': 24, 'street': 25, 'district': 26, 'country': 27},
        ]
        for idx, sh in enumerate(valid_members[2:4]):
            F = slots_p10[idx]
            is_corp = sh.get('identity') == 'corporate'
            full = sh.get('nameEnglish') or sh.get('name', '')
            surname, other = _parse_english_name(full)
            addr = _parse_address(sh.get('address', ''))
            def _safe(v):
                return '' if (v and _PURE_NUMBER_RE.match(v)) else v
            country = _safe(addr['country']) or 'Hong Kong'
            _set_text(doc, fmap, f'fill_{F["name"]}_P.10', sh.get('nameChinese', ''))
            if is_corp:
                _set_text(doc, fmap, f'fill_{F["surname"]}_P.10', full)
            else:
                _set_text(doc, fmap, f'fill_{F["surname"]}_P.10', surname)
                _set_text(doc, fmap, f'fill_{F["other"]}_P.10', other)
            shares_num = int(sh.get('shares', 0) or 0)
            _set_text(doc, fmap, f'fill_{F["shares"]}_P.10', _fmt_int(shares_num) if shares_num > 0 else '0')
            _set_text(doc, fmap, f'fill_{F["flat"]}_P.10', _safe(addr['flat']))
            _set_text(doc, fmap, f'fill_{F["building"]}_P.10', _safe(addr['building']))
            _set_text(doc, fmap, f'fill_{F["street"]}_P.10', _safe(addr['street']))
            _set_text(doc, fmap, f'fill_{F["district"]}_P.10', _safe(addr['district']))
            _set_text(doc, fmap, f'fill_{F["country"]}_P.10', country)
        _set_text(doc, fmap, 'fill_31_P.10', '2')
        _set_text(doc, fmap, 'fill_32_P.10', str(total_sch1))

    # ── P.10 附表二（上市公司）──
    if is_listed:
        _set_text(doc, fmap, 'fill_1_P.10', day)
        _set_text(doc, fmap, 'fill_2_P.10', month)
        _set_text(doc, fmap, 'fill_3_P.10', year)
        _set_text(doc, fmap, 'fill_4_P.10', br8)
        if share_infos:
            _set_text(doc, fmap, 'fill_5_P.10', share_infos[0]['className'])
            _set_text(doc, fmap, 'fill_6_P.10', _fmt_int(share_infos[0]['shares']))

    # ── P.11 續頁A：自然人秘書 #2 ──
    if len(nat_secs) > 1:
        s = nat_secs[1]
        _set_text(doc, fmap, 'fill_1_P.11', day)
        _set_text(doc, fmap, 'fill_2_P.11', month)
        _set_text(doc, fmap, 'fill_3_P.11', year)
        _set_text(doc, fmap, 'fill_4_P.11', br8)
        _set_text(doc, fmap, 'fill_5_P.11', s.get('nameChinese', ''))
        surname, other = _parse_english_name(s.get('nameEnglish', ''))
        _set_text(doc, fmap, 'fill_6_P.11', surname)
        _set_text(doc, fmap, 'fill_7_P.11', other)
        addr = _parse_address(s.get('address', ''))
        _set_text(doc, fmap, 'fill_8_P.11', addr['flat'])
        _set_text(doc, fmap, 'fill_9_P.11', addr['building'])  # P.11 has split flat/building row
        _set_text(doc, fmap, 'fill_10_P.11', addr['street'])
        _set_text(doc, fmap, 'fill_11_P.11', addr['district'])
        _set_text(doc, fmap, 'fill_12_P.11', addr.get('country') or addr.get('region', '香港 Hong Kong'))
        _set_text(doc, fmap, 'fill_13_P.11', addr['country'] or '香港 Hong Kong')
        _set_text(doc, fmap, 'fill_15_P.11', s.get('serviceAddress', '') or s.get('address', ''))
        _set_text(doc, fmap, 'fill_16_P.11', s.get('email', ''))
        hkid = _parse_hkid_partial(s.get('idNumber', ''))
        if hkid:
            _set_text(doc, fmap, 'fill_17_P.11', hkid, align='right')
        if s.get('passportCountry'):
            _set_text(doc, fmap, 'fill_18_P.11', s.get('passportCountry', ''))
        if s.get('passportNumber'):
            _set_text(doc, fmap, 'fill_19_P.11', _parse_passport_partial(s['passportNumber']))
        tcsp = s.get('tcspNumber', '')
        if tcsp:
            _set_text(doc, fmap, 'fill_20_P.11', tcsp)

    # ── P.12 續頁B：法人秘書 #2 ──
    if len(corp_secs) > 1:
        s = corp_secs[1]
        _set_text(doc, fmap, 'fill_1_P.12', day)
        _set_text(doc, fmap, 'fill_2_P.12', month)
        _set_text(doc, fmap, 'fill_3_P.12', year)
        _set_text(doc, fmap, 'fill_4_P.12', br8)
        _set_text(doc, fmap, 'fill_5_P.12', s.get('nameChinese', ''))
        _set_text(doc, fmap, 'fill_6_P.12', s.get('nameEnglish', ''))
        addr = _parse_address(s.get('serviceAddress') or s.get('address', ''))
        _set_text(doc, fmap, 'fill_7_P.12', addr['flat'])
        _set_text(doc, fmap, 'fill_8_P.12', addr['building'])
        _set_text(doc, fmap, 'fill_9_P.12', addr['street'])
        _set_text(doc, fmap, 'fill_10_P.12', addr['district'])
        _set_text(doc, fmap, 'fill_11_P.12', s.get('email', ''))
        _set_text(doc, fmap, 'fill_12_P.12', s.get('companyNumberRef') or s.get('brNumber', '') or s.get('idNumber', ''))
        tcsp = s.get('tcspNumber', '')
        if tcsp:
            _set_text(doc, fmap, 'fill_13_P.12', tcsp)

    # ── P.13 續頁C：自然人董事 #2 ──
    if len(nat_dirs) > 1:
        d = nat_dirs[1]
        _set_text(doc, fmap, 'fill_1_P.13', day)
        _set_text(doc, fmap, 'fill_2_P.13', month)
        _set_text(doc, fmap, 'fill_3_P.13', year)
        _set_text(doc, fmap, 'fill_4_P.13', br8)
        _check(doc, fmap, 'cb_1_P.13', True)
        _set_text(doc, fmap, 'fill_5_P.13', d.get('nameChinese', ''))
        surname, other = _parse_english_name(d.get('nameEnglish', ''))
        _set_text(doc, fmap, 'fill_6_P.13', surname)
        _set_text(doc, fmap, 'fill_7_P.13', other)
        _set_text(doc, fmap, 'fill_8_P.13', office.get('flat', ''))
        _set_text(doc, fmap, 'fill_9_P.13', office.get('building', ''))
        _set_text(doc, fmap, 'fill_10_P.13', office.get('street', ''))
        _set_text(doc, fmap, 'fill_11_P.13', office.get('district', ''))
        _set_text(doc, fmap, 'fill_12_P.13', office.get('region') or office.get('country', ''))
        _set_text(doc, fmap, 'fill_15_P.13', office.get('region') or office.get('country', '香港 Hong Kong'))
        _set_text(doc, fmap, 'fill_16_P.13', d.get('email', ''))
        hkid = _parse_hkid_partial(d.get('idNumber', ''))
        if hkid:
            _set_text(doc, fmap, 'fill_17_P.13', hkid, align='right')
        if d.get('passportCountry'):
            _set_text(doc, fmap, 'fill_18_P.13', d.get('passportCountry', ''))
        if d.get('passportNumber'):
            _set_text(doc, fmap, 'fill_19_P.13', _parse_passport_partial(d['passportNumber']))

    # ── P.14 續頁D：法人董事 #2+#3 ──
    extra_corp_dirs = corp_dirs[1:]
    if extra_corp_dirs:
        _set_text(doc, fmap, 'fill_1_P.14', day)
        _set_text(doc, fmap, 'fill_2_P.14', month)
        _set_text(doc, fmap, 'fill_3_P.14', year)
        _set_text(doc, fmap, 'fill_4_P.14', br8)
        # Slot 1: fields 5-14, checkboxes cb_1, cb_2
        # Slot 2: fields 15-24, checkboxes cb_3, cb_4
        slots_p14 = [
            {'cb_dir': 'cb_1_P.14', 'cb_reserve': 'cb_2_P.14', 'name_cn': 5, 'name_en': 6,
             'flat': 7, 'building': 8, 'street': 9, 'district': 10, 'region': 11,
             'country': 12, 'service_addr': 13, 'email': 14},
            {'cb_dir': 'cb_3_P.14', 'cb_reserve': 'cb_4_P.14', 'name_cn': 15, 'name_en': 16,
             'flat': 17, 'building': 18, 'street': 19, 'district': 20, 'region': 21,
             'country': 22, 'service_addr': 23, 'email': 24},
        ]
        for idx, d in enumerate(extra_corp_dirs[:2]):
            F = slots_p14[idx]
            _check(doc, fmap, F['cb_dir'], True)
            _set_text(doc, fmap, f'fill_{F["name_cn"]}_P.14', d.get('nameChinese', ''))
            _set_text(doc, fmap, f'fill_{F["name_en"]}_P.14', d.get('nameEnglish', ''))
            _set_text(doc, fmap, f'fill_{F["flat"]}_P.14', office.get('flat', ''))
            _set_text(doc, fmap, f'fill_{F["building"]}_P.14', office.get('building', ''))
            _set_text(doc, fmap, f'fill_{F["street"]}_P.14', office.get('street', ''))
            _set_text(doc, fmap, f'fill_{F["district"]}_P.14', office.get('district', ''))
            _set_text(doc, fmap, f'fill_{F["region"]}_P.14', office.get('region') or office.get('country', ''))
            _set_text(doc, fmap, f'fill_{F["country"]}_P.14', office.get('region') or office.get('country', '香港 Hong Kong'))
            _set_text(doc, fmap, f'fill_{F["email"]}_P.14', d.get('email', ''))

    # ── P.15 續頁E：公司紀錄保存地點 ──
    company_records = data.get('companyRecords') or []
    valid_records = [r for r in company_records if (r.get('records', '') or '').strip() or (r.get('address', '') or '').strip()]
    if valid_records:
        _set_text(doc, fmap, 'fill_1_P.15', day)
        _set_text(doc, fmap, 'fill_2_P.15', month)
        _set_text(doc, fmap, 'fill_3_P.15', year)
        _set_text(doc, fmap, 'fill_4_P.15', br8)
        records_text = '\n\n'.join([r.get('records', '') or '' for r in valid_records])
        address_text = '\n\n'.join([r.get('address', '') or '' for r in valid_records])
        _set_text(doc, fmap, 'fill_5_P.15', records_text)
        _set_text(doc, fmap, 'fill_6_P.15', address_text)

    # ── 刪除無數據的附頁（P.9-P.15）──
    pages_to_delete = []

    # P.9 (idx 8): 附表一 page 1 — 股東 #1+#2
    if not (valid_members and not is_listed):
        pages_to_delete.append(8)

    # P.10 (idx 9): 附表一 page 2 OR 附表二
    has_sch1_p2 = len(valid_members) > 2 and not is_listed
    has_sch2 = is_listed
    if not (has_sch1_p2 or has_sch2):
        pages_to_delete.append(9)

    # P.11 (idx 10): 續頁A — 第2位自然人秘書
    if len(nat_secs) <= 1:
        pages_to_delete.append(10)

    # P.12 (idx 11): 續頁B — 第2位法人秘書
    if len(corp_secs) <= 1:
        pages_to_delete.append(11)

    # P.13 (idx 12): 續頁C — 第2位自然人董事
    if len(nat_dirs) <= 1:
        pages_to_delete.append(12)

    # P.14 (idx 13): 續頁D — 法人董事 #2+#3
    if len(corp_dirs) <= 1:
        pages_to_delete.append(13)

    # P.15 (idx 14): 續頁E — 公司紀錄保存地點
    if not valid_records:
        pages_to_delete.append(14)

    for pi in reversed(pages_to_delete):
        doc.delete_page(pi)

    # ── 刪除空白頁（原 P.16-P.27，索引因附頁刪除已位移）──
    blank_pages = []
    scan_start = 15 - len(pages_to_delete)  # 原 P.16 的新索引
    for pi in range(scan_start, doc.page_count):
        if not list(doc[pi].widgets()):
            blank_pages.append(pi)
    for pi in reversed(blank_pages):
        doc.delete_page(pi)

    # ── 保存 ──
    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes

@app.route('/api/generate-nar1-pdf', methods=['POST'])
def generate_nar1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nar1_pdf(data)
        import base64 as b64
        result = {'pdf': b64.b64encode(pdf_bytes).decode('ascii')}

        # Phase 4: Auto-assign change events to NAR1 periods after PDF generation
        company_id = data.get('company_id') or data.get('selectedCompanyId')
        if company_id:
            try:
                # Ensure NAR1 periods exist
                _ensure_nar1_periods(company_id)
                # Assign unassigned changes to current period
                assign_result = _assign_changes_to_nar1_periods(company_id)
                if assign_result.get('assigned_count', 0) > 0:
                    result['nar1_assigned_changes'] = assign_result['assigned_count']
                # Get the current period info for the response
                periods = assign_result.get('periods', [])
                if periods:
                    result['nar1_current_period'] = {
                        'id': periods[0]['id'],
                        'period_start': periods[0]['period_start'],
                        'period_end': periods[0]['period_end'],
                    }
            except Exception as e:
                print(f"[NAR1] Warning: Failed to assign changes: {e}")
                # Don't fail the PDF generation for this

        return jsonify(result)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# ─── Phase 4: NAR1 Smart Filing API Endpoints ───

@app.route('/api/companies/<company_id>/nar1-status', methods=['GET'])
def get_nar1_status(company_id):
    """Get NAR1 filing status for a company: dates, grace period, and pending changes count."""
    try:
        db = get_db()
        company = db.execute(
            "SELECT id, name, incorporation_date, nar1_due_date FROM companies WHERE id = ?",
            (company_id,)
        ).fetchone()
        if not company:
            return jsonify({'error': 'Company not found'}), 404

        result = {
            'company_id': company_id,
            'company_name': company['name'],
            'incorporation_date': company['incorporation_date'],
        }

        # Calculate dates
        if company['incorporation_date']:
            dates = _calculate_nar1_dates(company['incorporation_date'])
            if dates:
                result.update({
                    'period_start': dates['period_start'],
                    'period_end': dates['period_end'],
                    'due_date': dates['due_date'],
                    'days_remaining': dates['days_remaining'],
                    'status': dates['status'],
                    'today': dates['today'],
                })

                # Update stored due_date if changed
                if company['nar1_due_date'] != dates['due_date']:
                    db.execute(
                        "UPDATE companies SET nar1_due_date = ? WHERE id = ?",
                        (dates['due_date'], company_id)
                    )
                    db.commit()

        # Get current period
        periods = _ensure_nar1_periods(company_id)
        if periods:
            current = periods[0]
            # Get changes summary for current period
            summary = _get_nar1_changes_summary(company_id)
            result['current_period'] = summary['period']
            result['changes_summary'] = summary['summary']
            result['changes'] = summary['changes']

        return jsonify({'success': True, **result})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/companies/<company_id>/nar1-changes', methods=['GET'])
def get_nar1_changes(company_id):
    """Get changes for a specific NAR1 period. Query params: period_id (optional)."""
    try:
        period_id = request.args.get('period_id')
        summary = _get_nar1_changes_summary(company_id, period_id)
        return jsonify({'success': True, **summary})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/companies/<company_id>/nar1-assign-changes', methods=['POST'])
def assign_nar1_changes(company_id):
    """Manually trigger assignment of unassigned change events to NAR1 periods."""
    try:
        result = _assign_changes_to_nar1_periods(company_id)
        return jsonify({'success': True, **result})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/companies/<company_id>/nar1-file', methods=['POST'])
def file_nar1(company_id):
    """Mark a NAR1 period as filed. Body: {period_id, filing_date (optional)}.
    After filing, auto-assigns post-period changes to the next period."""
    try:
        db = get_db()
        data = request.json or {}
        period_id = data.get('period_id')

        if not period_id:
            # Use current period
            periods = db.execute(
                "SELECT * FROM nar1_filings WHERE company_id = ? ORDER BY period_start DESC LIMIT 1",
                (company_id,)
            ).fetchall()
            if not periods:
                return jsonify({'error': 'No NAR1 periods found'}), 404
            period_id = periods[0]['id']

        filing_date = data.get('filing_date') or datetime.now().strftime('%d/%m/%Y')

        db.execute(
            "UPDATE nar1_filings SET status = 'filed', filing_date = ? WHERE id = ? AND company_id = ?",
            (filing_date, period_id, company_id)
        )
        db.commit()

        # Re-assign changes: events after period_end go to next period
        _assign_changes_to_nar1_periods(company_id)

        return jsonify({
            'success': True,
            'period_id': period_id,
            'filing_date': filing_date,
            'message': 'NAR1 period marked as filed'
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/companies/<company_id>/calculate-nar1-dates', methods=['POST'])
def calculate_nar1_dates(company_id):
    """Calculate NAR1 dates for a company and update the stored due_date.
    Also ensures nar1_filings periods exist."""
    try:
        db = get_db()
        company = db.execute(
            "SELECT incorporation_date FROM companies WHERE id = ?",
            (company_id,)
        ).fetchone()
        if not company:
            return jsonify({'error': 'Company not found'}), 404

        dates = _calculate_nar1_dates(company['incorporation_date'])
        if not dates:
            return jsonify({'error': 'Invalid incorporation date'}), 400

        # Update company
        db.execute(
            "UPDATE companies SET nar1_due_date = ? WHERE id = ?",
            (dates['due_date'], company_id)
        )
        db.commit()

        # Ensure periods
        periods = _ensure_nar1_periods(company_id)

        return jsonify({
            'success': True,
            'dates': dates,
            'periods': periods,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/nar1-due-companies', methods=['GET'])
def list_nar1_due_companies():
    """List companies with upcoming or overdue NAR1 filings.
    Query params: status (ok/grace/due_soon/late), days (int, default 42)."""
    try:
        db = get_db()
        status_filter = request.args.get('status', '')
        days_ahead = int(request.args.get('days', '42'))

        companies = db.execute(
            "SELECT id, name, chinese_name, company_number as br_number, "
            "incorporation_date, nar1_due_date, status as company_status "
            "FROM companies WHERE incorporation_date != '' AND status != 'deleted' "
            "ORDER BY nar1_due_date"
        ).fetchall()

        result = []
        for c in companies:
            dates = _calculate_nar1_dates(c['incorporation_date'])
            if not dates:
                continue
            entry = {
                'company_id': c['id'],
                'company_name': c['name'],
                'chinese_name': c['chinese_name'],
                'br_number': c['br_number'],
                'incorporation_date': c['incorporation_date'],
                **dates,
            }
            if status_filter:
                if dates['status'] == status_filter:
                    result.append(entry)
            elif dates['days_remaining'] <= days_ahead:
                result.append(entry)

        return jsonify({'success': True, 'companies': result, 'total': len(result)})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NR1 PDF 生成（本地 Python + PyMuPDF） ───

def _set_dropdown_value(doc, fmap, field_name, display_text):
    """Set a dropdown/combo field to a specific display-text option."""
    if field_name not in fmap:
        return False
    pi = fmap[field_name]
    for w in doc[pi].widgets():
        if w.field_name == field_name:
            try:
                w.field_value = display_text
                w.update()
                return True
            except Exception:
                return False
    return False


def _fill_nr1_pdf(data):
    """填充 NR1 PDF 模板，返回 bytes"""
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NR1-template.pdf')
    doc = fitz.open(template_path)
    fmap = {}
    for pi in range(doc.page_count):
        for w in doc[pi].widgets():
            if w.field_name:
                fmap[w.field_name] = pi

    br8 = (data.get('brNumber', '') or '').replace(r'[^0-9A-Za-z]', '')[:8]

    field_map = {
        'fill_1_P.1': br8,
        'fill_2_P.1': data.get('companyName', ''),
        'fill_3_P.1': data.get('flat', ''),
        'fill_4_P.1': data.get('building', ''),
        'fill_5_P.1': data.get('street', ''),
        'fill_6_P.1': data.get('district', ''),
        'fill_7_P.1': data.get('addressEffectiveDay', ''),
        'fill_8_P.1': data.get('addressEffectiveMonth', ''),
        'fill_9_P.1': data.get('addressEffectiveYear', ''),
        'fill_10_P.1': data.get('email', ''),
        'fill_18_P.1': data.get('signerName', ''),
        'fill_19_P.1': f"{data.get('signDateDay','')}/{data.get('signDateMonth','')}/{data.get('signDateYear','')}",
        'fill_20_P.1': data.get('presentorName', ''),
        'fill_21_P.1': data.get('presentorAddress', ''),
        'fill_22_P.1': data.get('presentorContact', ''),
    }

    # Email effective date — only if email is filled
    if (data.get('email', '') or '').strip():
        field_map['fill_11_P.1'] = data.get('emailEffectiveDay', '')
        field_map['fill_12_P.1'] = data.get('emailEffectiveMonth', '')
        field_map['fill_13_P.1'] = data.get('emailEffectiveYear', '')

    # Phone effective date — only if phone is filled
    if (data.get('phone', '') or '').strip():
        field_map['fill_14_P.1'] = data.get('phone', '')
        field_map['fill_15_P.1'] = data.get('phoneEffectiveDay', '')
        field_map['fill_16_P.1'] = data.get('phoneEffectiveMonth', '')
        field_map['fill_17_P.1'] = data.get('phoneEffectiveYear', '')
    else:
        field_map['fill_14_P.1'] = data.get('phone', '')

    for name, value in field_map.items():
        if name not in fmap:
            continue
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                try:
                    w.field_value = value if value else ''
                    w.update()
                except Exception:
                    pass
                break

    # ── Signer Capacity: use template dropdown to strike through ──
    # Dropdown1_P.1 (x≈156) = 董事旁, Dropdown2_P.1 (x≈226) = 公司秘书旁
    # 每个 dropdown 有 2 个选项: [0]=空白, [1]=横线(字型编码)
    # 用 xref 直接设 /I 索引 + /V 值 + draw_line 画线保底（对齐 NAR1 P.8 模式）
    signer_capacity = (data.get('signerCapacity', '') or '').strip()
    if signer_capacity in ('director', 'secretary'):
        # director → 划掉公司秘书(Dropdown2), secretary → 划掉董事(Dropdown1)
        cross_widget = 'Dropdown2_P.1' if signer_capacity == 'director' else 'Dropdown1_P.1'
        for widget_name in ('Dropdown1_P.1', 'Dropdown2_P.1'):
            if widget_name not in fmap:
                continue
            pi = fmap[widget_name]
            for w in doc[pi].widgets():
                if w.field_name == widget_name:
                    try:
                        use_dashes = (widget_name == cross_widget)
                        opt_idx = 1 if use_dashes else 0
                        # Set /I index and /V value via xref
                        doc.xref_set_key(w._annot.xref, 'I', f'[{opt_idx}]')
                        val = w.choice_values[opt_idx]
                        # NR1 dropdowns store (export, display) tuples; NAR1 stores plain strings
                        if isinstance(val, tuple):
                            val = val[0]  # export value e.g. 'Yes'
                        doc.xref_set_key(w._annot.xref, 'V', fitz.get_pdf_str(val))
                        doc.xref_set_key(w._annot.xref, 'F', '4')
                        # Draw visible line through the widget rect as fallback guarantee
                        if use_dashes:
                            doc[pi].draw_line(
                                fitz.Point(w.rect.x0 + 2, w.rect.y0 + w.rect.height / 2),
                                fitz.Point(w.rect.x1 - 2, w.rect.y0 + w.rect.height / 2),
                                color=(0, 0, 0), width=1.0
                            )
                    except Exception:
                        pass
                    break

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nr1-pdf', methods=['POST'])
def generate_nr1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nr1_pdf(data)
        # Auto-update company data from form (Phase 3.2)
        _apply_form_changes_to_company(data, 'nr1')
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NDR1 PDF 生成（撤銷註冊申請書） ───

# Global font cache for CJK widget AP rendering
_CJK_AP_FONT_CACHE = {}


def _get_cjk_font(fontfile):
    """Get or create a cached fitz.Font for glyph-index lookups."""
    if fontfile not in _CJK_AP_FONT_CACHE:
        _CJK_AP_FONT_CACHE[fontfile] = fitz.Font(fontfile=fontfile)
    return _CJK_AP_FONT_CACHE[fontfile]


def _set_widget_cjk_ap(doc, page, widget, text, fontsize, fontfile, font_pages, font_xref_map, align='left', valign='bottom'):
    """Build a custom widget AP stream with embedded CJK font.

    Unlike page-content overlay (which hides the widget with F=3),
    this writes the CJK text directly into the widget's /AP stream,
    so the widget stays visible as a blue editable box with readable CJK text.

    Automatically wraps text into multiple lines when it exceeds widget width
    at the given fontsize (instead of shrinking font below readability).

    align: 'left' (default) or 'center' — horizontal alignment per line.
    valign: 'bottom' (default) or 'center' — vertical alignment of the entire
            text block within the widget.
    """
    try:
        import re as _re
        pi = page.number
        fontname = 'CJK_TPL'

        # Embed CJK font on page if not already done (tracked per-page).
        if pi not in font_pages:
            page.clean_contents()
            font_xref = page.insert_font(fontname=fontname, fontfile=fontfile)
            font_pages.add(pi)
            font_xref_map[pi] = font_xref
        font_xref = font_xref_map.get(pi)
        if not font_xref:
            return  # font embedding failed, skip this widget

        rect = widget.rect
        w_width = rect.width
        h_height = rect.height

        cjk_font = _get_cjk_font(fontfile)

        def _char_width(ch):
            """Width of a single character in points at current fontsize."""
            return fontsize * (1.0 if ord(ch) > 127 else 0.5)

        def _build_hex(line):
            """Convert a text line to PDF hex string using glyph indices."""
            gids = []
            for ch in line:
                gid = cjk_font.has_glyph(ord(ch))
                gids.append(gid if gid else 0)
            return '<' + ''.join(f'{gid:04x}' for gid in gids) + '>'

        inset = 2.0
        usable_w = w_width - 2 * inset
        line_height = fontsize * 1.35  # line spacing

        # Build lines: word-aware wrapping with explicit \n support
        paragraphs = text.split('\n')

        def _wrap_paragraph(para):
            """Tokenize and wrap a single paragraph. Empty → [''], else wrapped lines."""
            if not para:
                return ['']
            tokens = []
            i = 0
            while i < len(para):
                ch = para[i]
                if ord(ch) > 127:
                    tokens.append(ch)
                    i += 1
                else:
                    j = i
                    while j < len(para) and ord(para[j]) <= 127:
                        j += 1
                    tokens.append(para[i:j])
                    i = j

            para_width = sum(_char_width(ch) for ch in para)
            if para_width <= usable_w:
                return [para]

            lines = []
            cur_tokens = []
            cur_w = 0.0
            for tok in tokens:
                tok_w = sum(_char_width(ch) for ch in tok)
                if tok_w > usable_w:
                    if cur_tokens:
                        lines.append(''.join(cur_tokens))
                        cur_tokens = []
                        cur_w = 0.0
                    for ch in tok:
                        ch_w = _char_width(ch)
                        if cur_w + ch_w > usable_w and cur_tokens:
                            lines.append(''.join(cur_tokens))
                            cur_tokens = [ch]
                            cur_w = ch_w
                        else:
                            cur_tokens.append(ch)
                            cur_w += ch_w
                elif cur_w + tok_w > usable_w:
                    lines.append(''.join(cur_tokens))
                    cur_tokens = [tok]
                    cur_w = tok_w
                else:
                    cur_tokens.append(tok)
                    cur_w += tok_w
            if cur_tokens:
                lines.append(''.join(cur_tokens))
            return lines

        lines = []
        for para in paragraphs:
            lines.extend(_wrap_paragraph(para))

        # Limit lines to what fits vertically in the widget
        max_lines = max(1, int((h_height - inset) / line_height))
        if len(lines) > max_lines:
            lines = lines[:max_lines]
            # Truncate last visible line with ellipsis
            if len(lines[-1]) > 1:
                lines[-1] = lines[-1][:-2] + '…'

        line_hexes = [_build_hex(ln) for ln in lines]

        # Calculate positions per line
        line_widths = [sum(_char_width(ch) for ch in ln) for ln in lines]
        total_block_h = len(lines) * line_height

        # Vertical start position for the entire text block
        if valign == 'center':
            block_top = (h_height - total_block_h) / 2.0
        else:
            block_top = 0.0

        # Build content stream — one Tm + Tj per line
        stream_parts = [f"/Tx BMC\nq\nBT\n/{fontname} {fontsize} Tf"]
        for i, (lh, lw) in enumerate(zip(line_hexes, line_widths)):
            if align == 'center':
                lx = max(inset, (w_width - lw) / 2.0)
            else:
                lx = inset
            ly = h_height - (block_top + i * line_height + fontsize * 1.05)
            stream_parts.append(f"1 0 0 1 {lx} {ly} Tm\n{lh} Tj")
        stream_parts.append("ET\nQ\nEMC")
        stream = '\n'.join(stream_parts)
    except Exception:
        return  # If font embedding or text processing fails, skip this widget

    try:
        # Create a new PDF object for the Form XObject
        form_xref = doc.get_new_xref()
        form_dict = (
            f'<<\n'
            f'/Type /XObject\n'
            f'/Subtype /Form\n'
            f'/BBox [0 0 {w_width} {h_height}]\n'
            f'/Resources << /Font << /{fontname} {font_xref} 0 R >> >>\n'
            f'/Length {len(stream.encode("utf-8"))}\n'
            f'>>'
        )
        doc.update_object(form_xref, form_dict)
        doc.update_stream(form_xref, stream.encode('utf-8'))

        # Set as widget's normal appearance
        annot_xref = widget._annot.xref
        ap_dict = f'<< /N {form_xref} 0 R >>'
        doc.xref_set_key(annot_xref, 'AP', ap_dict)

        # Set the field value (UTF-16BE is correct for AcroForm /V)
        doc.xref_set_key(annot_xref, 'V', fitz.get_pdf_str(text))
    except Exception:
        pass  # If anything fails, widget stays unfilled (better than crashing)


def _stamp_br_on_page(doc, page_index, br_text):
    """在指定頁面上用 insert_textbox 疊加 BR 號碼（數字字體，不依賴 AcroForm widget）。
    放在頁面頂部靠右位置，確保每一頁都有商業登記號碼。"""
    if not br_text:
        return
    page = doc[page_index]
    rect = page.rect
    text_rect = fitz.Rect(rect.width - 150, 15, rect.width - 10, 30)
    try:
        page.insert_textbox(
            text_rect,
            f"BR: {br_text}",
            fontname="helv",
            fontsize=7,
            color=(0.2, 0.2, 0.2),
            align=2,  # right-align
        )
    except Exception:
        pass  # 若頁面無法插入文字則靜默跳過


def _stamp_br_on_all_pages(doc, br_text):
    """在 PDF 每一頁疊加 BR 號碼"""
    if not br_text or not br_text.strip():
        return
    for pi in range(doc.page_count):
        _stamp_br_on_page(doc, pi, br_text)


def _fill_ndr1_pdf(data):
    """填充 NDR1 PDF 模板，返回 bytes

    模板實際佈局（千问 VL + PyMuPDF 驗證）：
    P.1: BR + 公司名 + 撤銷條件(6項文字) + A.申請人身份(cb_1~3) + 提交人資料(fill_3~11)
    P.2: B.申請人資料 — 自然人(中文名/英文姓/英文名) 或 法人團體(公司名) + 地址5欄 + 電郵/傳真
    P.3: C.獲提名自然人資料 — 僅當申請人為「上述公司」時填寫
    P.4: 簽署人 + 日期 + 聲明勾選 + Dropdown劃線
    P.5-P.8: 指引頁（無 widgets）
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NDR1-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate a system CJK font for blue widget AP rendering
    _tpl_cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _tpl_cjk_fontfile = _sf
            break

    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk_ap(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        """Fill field with blue editable CJK widget AP. Falls back to _set_text for ASCII."""
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _tpl_cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                field_h = w.rect.height
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                if field_h > 0:
                    fs = min(fs, max(min_fs, int(field_h - 3)))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _tpl_cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]
    app_capacity = data.get('applicantCapacity', '')  # 'company' | 'director' | 'member'

    # ═══ P.1: 公司資料 ═══
    _set_text(doc, fmap, 'fill_1_P.1', br8)
    _set_text(doc, fmap, 'fill_2_P.1', data.get('companyName', ''))

    # ═══ P.1: A.申請人身份 checkbox（cb_1~cb_3_P.1）═══
    # cb_1 = 上述公司 (the above named company)
    # cb_2 = 上述公司的一名董事 (a director of the above named company)
    # cb_3 = 上述公司的一名成員 (a member of the above named company)
    _check(doc, fmap, 'cb_1_P.1', app_capacity == 'company')
    _check(doc, fmap, 'cb_2_P.1', app_capacity == 'director')
    _check(doc, fmap, 'cb_3_P.1', app_capacity == 'member')

    # ═══ P.1 左下角：提交人資料 Presentor's Reference ═══
    _set_cjk_ap('fill_3_P.1', data.get('presenterNameCN', ''), align='left')
    _set_text(doc, fmap, 'fill_4_P.1', data.get('presenterNameEN', ''))
    _set_cjk_ap('fill_5_P.1', data.get('presenterAddress1', ''), min_fs=9, valign='bottom')
    _set_cjk_ap('fill_6_P.1', data.get('presenterAddress2', ''), min_fs=9, valign='bottom')
    _set_cjk_ap('fill_7_P.1', data.get('presenterAddress3', ''), min_fs=9, valign='bottom')
    _set_text_size(doc, fmap, 'fill_8_P.1', data.get('presenterTel', ''), 8)
    _set_text_size(doc, fmap, 'fill_9_P.1', data.get('presenterFax', ''), 8)
    _set_text_size(doc, fmap, 'fill_10_P.1', data.get('presenterEmail', ''), 8)
    _set_text_size(doc, fmap, 'fill_11_P.1', data.get('presenterReference', ''), 8)

    # ═══ P.2: B.申請人資料 ═══
    app_type = data.get('applicantType', 'natural')  # 'natural' | 'corporate'
    _set_text(doc, fmap, 'fill_1_P.2', br8)

    if app_type == 'corporate' or app_capacity == 'company':
        # 法人團體 — 只填法人名稱 (fill_5_P.2) + 地址 + 聯絡
        body_name = data.get('appBodyCorpName', '') or data.get('companyName', '')
        _set_cjk_ap('fill_5_P.2', body_name, align='center')
        # 自然人姓名三欄留空（不填）
    else:
        # 自然人 — 填中文名/英文姓/英文名 (fill_2~fill_4_P.2)
        cn = data.get('appChineseName', '')
        surname = data.get('appSurname', '')
        other = data.get('appOtherNames', '')
        # Fallback: parse from legacy/flat field
        if not surname and not other:
            en = data.get('appName', '') or data.get('applicantNameEN', '')
            if en:
                surname, other = _parse_english_name(en)
        if not cn:
            cn = data.get('applicantNameCN', '')
        _set_cjk_ap('fill_2_P.2', cn, align='center')
        _set_text(doc, fmap, 'fill_3_P.2', surname)
        _set_text(doc, fmap, 'fill_4_P.2', other)

    # P.2 地址（5 欄，fill_6~fill_10_P.2）
    _set_text(doc, fmap, 'fill_6_P.2', data.get('appAddrFlat', ''))
    _set_text(doc, fmap, 'fill_7_P.2', data.get('appAddrBuilding', ''))
    _set_text(doc, fmap, 'fill_8_P.2', data.get('appAddrStreet', ''))
    _set_text(doc, fmap, 'fill_9_P.2', data.get('appAddrDistrict', ''))
    _set_text(doc, fmap, 'fill_10_P.2', data.get('appAddrCountry', ''))

    # P.2 聯絡
    _set_text(doc, fmap, 'fill_11_P.2', data.get('appEmail', ''))
    fax_or_phone = data.get('appFax', '') or data.get('appTel', '')
    if fax_or_phone:
        _set_text(doc, fmap, 'fill_12_P.2', fax_or_phone)

    # ═══ P.3: C.獲提名自然人資料（僅當申請人是「上述公司」）═══
    if app_capacity == 'company':
        nom_cn = data.get('nomChineseName', '')
        nom_surname = data.get('nomSurname', '')
        nom_other = data.get('nomOtherNames', '')
        # Fallback from flat fields
        if not nom_surname and not nom_other:
            nom_en = data.get('nomName', '') or data.get('nomNameEnglish', '')
            if nom_en:
                nom_surname, nom_other = _parse_english_name(nom_en)

        _set_cjk_ap('fill_2_P.3', nom_cn, align='center')
        _set_text(doc, fmap, 'fill_3_P.3', nom_surname)
        _set_text(doc, fmap, 'fill_4_P.3', nom_other)

        _set_text(doc, fmap, 'fill_6_P.3', data.get('nomAddrFlat', ''))
        _set_text(doc, fmap, 'fill_7_P.3', data.get('nomAddrBuilding', ''))
        _set_text(doc, fmap, 'fill_8_P.3', data.get('nomAddrStreet', ''))
        _set_text(doc, fmap, 'fill_9_P.3', data.get('nomAddrDistrict', ''))
        _set_text(doc, fmap, 'fill_10_P.3', data.get('nomAddrCountry', ''))

        _set_text(doc, fmap, 'fill_11_P.3', data.get('nomEmail', ''))
        nom_fax = data.get('nomFax', '')
        if nom_fax:
            _set_text(doc, fmap, 'fill_12_P.3', nom_fax)
    # else: P.3 不填（自然人/董事/成員申請人不需要 Section 2C）

    # ═══ P.4: 簽署 ═══
    _set_text(doc, fmap, 'fill_1_P.4', br8)
    _set_text(doc, fmap, 'fill_2_P.4', data.get('signerName', ''))

    sign_date = data.get('signDate', '')
    if sign_date and '-' in sign_date:
        parts = sign_date.split('-')
        sign_date = f'{parts[2]}/{parts[1]}/{parts[0]}'
    elif not sign_date:
        dd = data.get('signDateDay', '')
        mm = data.get('signDateMonth', '')
        yy = data.get('signDateYear', '')
        if dd and mm and yy:
            sign_date = f'{dd}/{mm}/{yy}'
    _set_text(doc, fmap, 'fill_3_P.4', sign_date)

    # ═══ P.4: 簽署人身份 Capacity 勾選（二選一）═══
    # cb_1_P.4 = 本表格第2B項所述的申請人(自然人)
    # cb_2_P.4 = 本表格第2B項所述的申請人(法人團體)的董事/公司秘書/授權人
    is_corporate_applicant = (app_type == 'corporate' or app_capacity == 'company')
    _check(doc, fmap, 'cb_1_P.4', not is_corporate_applicant)   # 自然人
    _check(doc, fmap, 'cb_2_P.4', is_corporate_applicant)        # 法人團體代表

    # P.4 Dropdown 劃線（簽署人身份: director→劃Secretary, secretary→劃Director）
    signer_role = data.get('signerRole', 'director')
    # Dropdown 在P.4有6個（中英文雙行，每行三個）：前3個是Director行，後3個是Secretary行
    p4_widgets = []
    for w in doc[3].widgets():
        if w.field_name.startswith('Dropdown'):
            p4_widgets.append(w)
    dir_drops = sorted([w for w in p4_widgets if w.rect[0] > 380], key=lambda w: w.rect[1])
    sec_drops = sorted([w for w in p4_widgets if w.rect[0] < 380], key=lambda w: w.rect[1])
    # 選中劃線值（/I 1 = strike through line visible）
    for w in dir_drops:
        _select_dropdown(doc, fmap, w.field_name, 'Yes')
    for w in sec_drops:
        _select_dropdown(doc, fmap, w.field_name, 'Yes')
    # 然後畫線（draw_line）保底：director→劃secretary行, secretary→劃director行
    for w in (sec_drops if signer_role == 'director' else dir_drops):
        rect = w.rect
        doc[3].draw_line((rect.x0, rect.y0 + rect.height / 2), (rect.x1, rect.y0 + rect.height / 2), color=(0, 0, 0), width=1.5)

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    # Delete blank instruction pages (P.5-P.8)
    for pno in range(doc.page_count - 1, 3, -1):
        doc.delete_page(pno)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-ndr1-pdf', methods=['POST'])
def generate_ndr1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_ndr1_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── ND4 PDF 生成（董事/秘書辭任通知書） ───

def _fill_nd4_pdf(data):
    """填充 ND4 PDF 模板（公司秘書及董事辭任通知書），返回 bytes

    模板佈局：
    P.1: BR + 公司名 + 公司類型勾選 + 辭任人資料（中英文名/證件/地址） + 簽署 + 提交人
    P.2: 角色聲明勾選 + Dropdown 劃線 + 簽署人 + 日期
    P.3-P.6: 空白頁
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'ND4-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font for blue widget AP rendering
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break

    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        """Fill field with CJK widget AP. Falls back to _set_text for ASCII."""
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                field_h = w.rect.height
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                if field_h > 0:
                    fs = min(fs, max(min_fs, int(field_h - 3)))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]
    officer_type = data.get('officerType', 'director')  # director / secretary / alternate
    identity = data.get('identity', 'natural')

    # ── P.1: 公司資料 ──
    _set_text(doc, fmap, 'fill_1_P.1', br8)
    _set_text(doc, fmap, 'fill_2_P.1', data.get('companyName', ''))

    # 公司類型勾選 (cb_1=私人, cb_2=公眾, cb_3=擔保)
    company_type = (data.get('companyType') or '').lower()
    _check(doc, fmap, 'cb_1_P.1', '私人' in company_type or 'private' in company_type)
    _check(doc, fmap, 'cb_2_P.1', '公眾' in company_type or 'public' in company_type)
    _check(doc, fmap, 'cb_3_P.1', '擔保' in company_type)

    # ── P.1: 辭任人資料（按身份分區填寫）──
    # fill_3_P.1 = "代替 Alternate to" — only for alternate directors
    # fill_11/12/13_P.1 = 辭職日期 (Date of Resignation) — 模板標籤: "辭職日期 Date of Resignation" + "日DD 月MM 年YYYY"
    resign_day = data.get('resignationDay', '')
    resign_month = data.get('resignationMonth', '')
    resign_year = data.get('resignationYear', '')
    if resign_day and resign_month and resign_year:
        _set_text(doc, fmap, 'fill_11_P.1', resign_day)
        _set_text(doc, fmap, 'fill_12_P.1', resign_month)
        _set_text(doc, fmap, 'fill_13_P.1', resign_year)

    # fill_3_P.1 = "代替 Alternate to" — only for alternate directors
    if officer_type == 'alternate':
        _set_text(doc, fmap, 'fill_3_P.1', data.get('alternateTo', '') or data.get('officerNameEnglish', ''))

    if identity == 'natural':
        # 自然人：中文姓名 + 英文姓氏/名字 + HKID + 护照
        _set_cjk('fill_4_P.1', data.get('officerNameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_5_P.1', data.get('surname', ''))
        _set_text(doc, fmap, 'fill_6_P.1', data.get('otherNames', ''))
        _set_text(doc, fmap, 'fill_7_P.1', data.get('hkidPartial', ''), align='right')
        _set_text(doc, fmap, 'fill_8_P.1', data.get('passportCountry', ''))
        if 'fill_8b_P.1' in fmap:
            _set_text(doc, fmap, 'fill_8b_P.1', data.get('passportPartial', ''))
    else:
        # 法人團體：公司名稱 + 公司編號（fill_9/fill_10 是法人專區）
        _set_cjk('fill_9_P.1', data.get('corporateName', '') or data.get('officerNameEnglish', ''), align='left')
        _set_text(doc, fmap, 'fill_10_P.1', data.get('corporateNumber', '') or data.get('brNumber', ''))

    # ── P.1: 「是否仍然擔任」— 應為「否」(toggle_5_P.1) ──
    # toggle_4_P.1 = 是/Yes, toggle_5_P.1 = 否/No
    _check(doc, fmap, 'toggle_5_P.1', True)

    # ── P.1: 提交人 ──
    _set_cjk('fill_14_P.1', data.get('presentorName', ''), align='left')
    _set_cjk('fill_15_P.1', data.get('presentorAddress', ''), min_fs=8, valign='bottom')
    _set_text(doc, fmap, 'fill_16_P.1', data.get('presentorPhone', ''))
    _set_text(doc, fmap, 'fill_17_P.1', data.get('presentorFax', ''))
    _set_text(doc, fmap, 'fill_18_P.1', data.get('presentorEmail', ''))
    _set_text(doc, fmap, 'fill_19_P.1', data.get('presentorReference', ''))

    # ── P.2: 角色聲明 ──
    _set_text(doc, fmap, 'fill_1_P.2', br8)

    # 角色勾選：cb_1=董事辭任, cb_2=候補董事辭任, cb_3=秘書辭任
    _check(doc, fmap, 'cb_1_P.2', officer_type == 'director')
    _check(doc, fmap, 'cb_2_P.2', officer_type == 'alternate')
    _check(doc, fmap, 'cb_3_P.2', officer_type == 'secretary')

    # ── P.2: Dropdown 劃線聲明 ──
    # 模板有多個 Dropdown 實例（中英文各行），用於劃掉不適用的角色名稱。
    # Dropdown1/2 → "Director"/"Secretary" 或 "董事"/"公司秘書" 互斥劃線
    # Dropdown3/4 → 候補董事互斥
    # 按 NAR1 風格處理：設 /I + /V + F=4 讓被選中的選項顯示劃線
    _handle_nd4_dropdowns(doc, fmap, officer_type)

    # ── P.2: 簽署 ──
    _set_text(doc, fmap, 'fill_2_P.2', data.get('signerName', '') or data.get('presentorName', ''))
    sd = data.get('signDateDay', '') or str(date.today().day).zfill(2)
    sm = data.get('signDateMonth', '') or str(date.today().month).zfill(2)
    sy = data.get('signDateYear', '') or str(date.today().year)
    sign_date = f"{sd}/{sm}/{sy}"
    _set_text(doc, fmap, 'fill_3_P.2', sign_date)

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    # Remove blank instruction pages (P.3–P.6), only keep P.1 + P.2
    # Delete from the end to keep indices stable
    for pno in range(doc.page_count - 1, 1, -1):
        doc.delete_page(pno)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


def _handle_nd4_dropdowns(doc, fmap, officer_type):
    """處理 ND4 P.2 的 Dropdown 劃線聲明。

    Dropdown1/Dropdown2: 「董事／公司秘書」互斥（cb_1 / cb_3 對應）
    Dropdown3/Dropdown4: 候補董事互斥（cb_2 對應）
    Dropdown5/Dropdown6: 秘書辭任互斥
    Dropdown7/Dropdown8: 簽署人身份（董事/秘書）

    策略：設 widget /I 為對應選項索引，設 /V 為選中的選項值，設 F=4 顯示。
    """
    # 角色到 dropdown 選項值的映射（根據模板實際 choice_values）
    # Dropdown1/3/5/7 通常是第一選項（保留），Dropdown2/4/6/8 是第二選項（劃掉）
    _dropdown_role_map = {
        'director':   {'keep': 1, 'cross': 2},   # Dropdown1=保留董事, Dropdown2=劃掉秘書
        'secretary':  {'keep': 2, 'cross': 1},   # Dropdown1=劃掉董事, Dropdown2=保留秘書
        'alternate':  {'keep': 2, 'cross': 1},   # Dropdown3=劃掉董事, Dropdown4=保留候補
    }

    mapping = _dropdown_role_map.get(officer_type)
    if not mapping:
        return

    # 遍歷所有 P.2 dropdown 實例，按名稱前綴分組處理
    dropdown_groups = {}
    for name, pi in fmap.items():
        if not name.startswith('Dropdown') or not name.endswith('_P.2'):
            continue
        base = name.split('_P.2')[0]  # e.g. "Dropdown1"
        if base not in dropdown_groups:
            dropdown_groups[base] = []
        dropdown_groups[base].append(name)

    for base, names in dropdown_groups.items():
        # 提取數字後綴: Dropdown1 → 1
        m = re.match(r'Dropdown(\d+)', base)
        if not m:
            continue
        d_num = int(m.group(1))

        # Dropdown1/2: 董事 vs 秘書
        # Dropdown3/4: 候補董事互斥
        # Dropdown5/6: 秘書辭任
        # Dropdown7/8: 簽署人身份
        if d_num in (1, 2):
            # 董事辭任 → Dropdown1=keep 董事, Dropdown2=cross 秘書
            # 秘書辭任 → Dropdown1=cross 董事, Dropdown2=keep 秘書
            keep_dd = 'Dropdown1' if officer_type == 'director' else 'Dropdown2'
            cross_dd = 'Dropdown2' if officer_type == 'director' else 'Dropdown1'
        elif d_num in (3, 4):
            # 候補董事 → Dropdown4=keep
            keep_dd = 'Dropdown4' if officer_type == 'alternate' else 'Dropdown3'
            cross_dd = 'Dropdown3' if officer_type == 'alternate' else 'Dropdown4'
        elif d_num in (5, 6):
            # 秘書辭任 → Dropdown6=keep
            keep_dd = 'Dropdown6' if officer_type == 'secretary' else 'Dropdown5'
            cross_dd = 'Dropdown5' if officer_type == 'secretary' else 'Dropdown6'
        elif d_num in (7, 8):
            # 簽署人身份 → Dropdown7=董事, Dropdown8=秘書
            keep_dd = 'Dropdown7' if officer_type == 'director' else 'Dropdown8'
            cross_dd = 'Dropdown8' if officer_type == 'director' else 'Dropdown7'
        else:
            continue

        for name in names:
            pi = fmap[name]
            for w in doc[pi].widgets():
                if w.field_name != name:
                    continue
                try:
                    choices = w.choice_values
                    if not choices or len(choices) < 2:
                        continue
                    # 獲取選項值列表
                    choice_vals = [cv[0] for cv in choices]
                    # 決定此 widget 應選哪個
                    if name.startswith(keep_dd):
                        idx = 0 if officer_type == 'director' or (d_num in (5, 6) and officer_type == 'secretary') else 1
                        # 簡化：Dropdown1→選0, Dropdown2→選1 表示"保留此角色"
                        if d_num % 2 == 1:  # 奇數 dropdown → 保留第一選項
                            w.field_value = choice_vals[0]
                        else:
                            w.field_value = choice_vals[1] if len(choice_vals) > 1 else choice_vals[0]
                    else:
                        # 劃掉 → 選相反的
                        if d_num % 2 == 1:
                            w.field_value = choice_vals[1] if len(choice_vals) > 1 else choice_vals[0]
                        else:
                            w.field_value = choice_vals[0]
                    w.update()
                except Exception:
                    pass
                break


@app.route('/api/generate-nd4-pdf', methods=['POST'])
def generate_nd4_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nd4_pdf(data)
        _apply_form_changes_to_company(data, 'nd4')
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NSC1 PDF 生成（股份配發申報書） ───

def _fill_nsc1_pdf(data):
    """填充 NSC1 PDF 模板（股份配發申報書），返回 bytes

    Template: NSC1_fillable.pdf (14 pages, 226 widgets)
    Page structure (verified by Qwen VL 2026-08-04):
      P.1: BR, Company, Allotment Date FROM/TO, Section B (Currency|Amount),
           Section D (Class|Currency|Number|Paid|Unpaid), Presenter
      P.2: BR, Allottees 5-col table, 4 checkboxes, Particulars
      P.3: BR, Share Capital 6-col table (TOTAL post-allotment),
           Rights, Continuation counters, Signature
      P.7: Schedule 2 — Allottee personal details (name, address, shares)
      P.9-P.10: Share Capital continuation (6-col table)
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NSC1-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break
    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set(name, value, fontsize=10, align='left'):
        """Fill widget with CJK support."""
        if name not in fmap or not value:
            return False
        vstr = str(value)
        has_cjk = any(ord(c) > 127 for c in vstr)
        if not has_cjk:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                fs = fontsize
                field_w = w.rect.width
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * len(vstr) * 0.85
                    if est_w > usable_w:
                        fs = max(6, int(fs * usable_w / est_w * 0.95))
                if _cjk_fontfile:
                    _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                       _cjk_ap_font_pages, _cjk_ap_font_xref_map, align=align)
                    return True
                break
        return _set_text(doc, fmap, name, value)

    # ── Helper: set field if empty ──
    fields = data.get('fields', {})
    def _set_if_empty(name, value):
        if not value or not str(value).strip():
            return
        if name in fmap and not fields.get(name, '').strip():
            _set(name, str(value))

    # ── Auto-populate company data from DB ──
    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]
    company_name = ''
    company_id = data.get('company_id')
    if company_id:
        try:
            db = get_db()
            company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
            if company:
                c = dict(company)
                if not br8:
                    br8 = re.sub(r'[^0-9A-Za-z]', '', c.get('company_number', '') or '')[:8]
                company_name = c.get('name', '') or ''
        except Exception:
            pass

    # ── BR on every page via fill_1_P.X widgets ──
    if br8:
        for page_no in range(1, doc.page_count + 1):
            br_field = f'fill_1_P.{page_no}'
            if br_field in fmap:
                _set_text(doc, fmap, br_field, br8)

    # ── P.1: Company Name ──
    _set_if_empty('fill_2_P.1', company_name or data.get('companyName', ''))

    # ── P.1: Allotment Date Range ──
    # fill_3/4/5 = FROM (D/M/Y), fill_6/7/8 = TO (D/M/Y)
    today_str = datetime.now().strftime('%d/%m/%Y')
    allot_date = data.get('allotmentDate', today_str)
    parts = allot_date.split('/')
    dd, mm, yyyy = (parts + ['', '', ''])[:3]
    _set_if_empty('fill_3_P.1', dd)
    _set_if_empty('fill_4_P.1', mm)
    _set_if_empty('fill_5_P.1', yyyy)
    _set_if_empty('fill_6_P.1', dd)
    _set_if_empty('fill_7_P.1', mm)
    _set_if_empty('fill_8_P.1', yyyy)

    # ── P.1 Section B: Total consideration by currency ──
    currency = data.get('currency') or data.get('sectionB_currency') or 'HKD'
    total_consideration = data.get('totalConsideration') or data.get('sectionB_amount') or ''
    _set_if_empty('fill_9_P.1', currency)
    if total_consideration:
        _set_if_empty('fill_10_P.1', str(total_consideration))

    # ── P.1 Section D: New allotment details (Class|Currency|Number|Paid|Unpaid) ──
    share_class = data.get('shareClass') or data.get('allotteeClass') or 'Ordinary'
    shares = data.get('shares') or data.get('allotteeShares') or ''
    price_per_share = data.get('pricePerShare') or ''
    unpaid = data.get('unpaidPerShare') or '0.00'
    _set_if_empty('fill_15_P.1', share_class)
    _set_if_empty('fill_16_P.1', currency)
    if shares:
        _set_if_empty('fill_17_P.1', str(shares))
    if price_per_share:
        _set_if_empty('fill_18_P.1', str(price_per_share))
    _set_if_empty('fill_19_P.1', str(unpaid))

    # ── Presenter defaults (Twinsail) ──
    presenter_defaults = {
        'fill_30_P.1': 'Twinsail Consultants Limited',
        'fill_31_P.1': 'Room 1203, 12/F, Wing On Centre, 111 Connaught Road Central, Hong Kong',
        'fill_32_P.1': '+852 2521 3888',
        'fill_33_P.1': '+852 2521 3999',
        'fill_34_P.1': 'info@twinsail.com',
        'fill_35_P.1': 'TS-2026-001',
    }
    for k, v in presenter_defaults.items():
        _set_if_empty(k, v)

    # ── P.3: Share Capital Table (TOTAL post-allotment, Section 6) ──
    # 1. Query existing share_capital from DB
    # 2. Add the new allotment shares
    # 3. Group by class; single-currency → consolidate into one row
    if company_id:
        try:
            db = get_db()
            by_class = {}
            # Try share_capital table first (production), fall back to share_transactions
            try:
                sc_rows = db.execute(
                    "SELECT class_name, currency, COALESCE(total_number,0) as total_number, "
                    "COALESCE(total_amount,0) as total_amount, COALESCE(paid_up,0) as paid_up, "
                    "COALESCE(unpaid,0) as unpaid "
                    "FROM share_capital WHERE company_id = ?",
                    (company_id,)
                ).fetchall()
                for sc in sc_rows:
                    cls = (sc['class_name'] or 'Ordinary').strip()
                    cur = (sc['currency'] or 'HKD').strip()
                    if cls not in by_class:
                        by_class[cls] = {'currency': cur, 'shares': 0, 'paid': 0.0, 'unpaid': 0.0}
                    by_class[cls]['shares'] += int(sc['total_number'] or 0)
                    by_class[cls]['paid'] += float(sc['total_amount'] or 0)
                    by_class[cls]['unpaid'] += float(sc['unpaid'] or 0)
            except Exception:
                pass  # Table may not exist in local SQLite
            # If no share_capital rows, fall back to share_transactions
            if not by_class:
                try:
                    tx_rows = db.execute(
                        "SELECT share_type, currency, COALESCE(SUM(shares), 0) as total_shares "
                        "FROM share_transactions WHERE company_id = ? GROUP BY share_type",
                        (company_id,)
                    ).fetchall()
                    for tx in tx_rows:
                        cls = (tx['share_type'] or 'Ordinary').strip()
                        cur = (tx['currency'] or 'HKD').strip()
                        if cls not in by_class:
                            by_class[cls] = {'currency': cur, 'shares': 0, 'paid': 0.0, 'unpaid': 0.0}
                        by_class[cls]['shares'] += int(tx['total_shares'] or 0)
                except Exception:
                    pass
            # Add new allotment to totals
            new_class = (data.get('shareClass') or data.get('allotteeClass') or 'Ordinary').strip()
            new_shares = int(data.get('shares') or data.get('allotteeShares') or 0)
            new_currency = (data.get('currency') or 'HKD').strip()
            new_paid = float(data.get('pricePerShare') or 0)
            if new_class not in by_class:
                by_class[new_class] = {'currency': new_currency, 'shares': 0, 'paid': 0.0, 'unpaid': 0.0}
            by_class[new_class]['currency'] = by_class[new_class]['currency'] or new_currency
            by_class[new_class]['shares'] += new_shares
            by_class[new_class]['paid'] += new_shares * new_paid
            # Fill P.3 table (3 rows max, 6 cols: Class|Currency|Number|Paid|Unpaid|Total)
            currencies = set(v['currency'] for v in by_class.values() if v['currency'])
            if len(currencies) == 1 and by_class:
                # Single currency: consolidate all classes into one row
                total_shares = sum(v['shares'] for v in by_class.values())
                total_paid = sum(v['paid'] for v in by_class.values())
                total_unpaid = sum(v['unpaid'] for v in by_class.values())
                total_amount = total_paid + total_unpaid
                single_currency = list(currencies)[0] or 'HKD'
                _set_if_empty('fill_2_P.3', '普通股 Ordinary')
                _set_if_empty('fill_3_P.3', single_currency)
                _set_if_empty('fill_4_P.3', str(total_shares))
                if total_paid > 0:
                    _set_if_empty('fill_5_P.3', f'{total_paid:.2f}')
                _set_if_empty('fill_6_P.3', f'{total_unpaid:.2f}' if total_unpaid > 0 else '0.00')
                _set_if_empty('fill_7_P.3', f'{total_amount:.2f}')
            else:
                for i, (cls_name, v) in enumerate(by_class.items()):
                    if i >= 3:
                        break
                    base = 2 + i * 6  # fill_2, fill_8, fill_14
                    total = v['paid'] + v['unpaid']
                    _set_if_empty(f'fill_{base}_P.3', cls_name)
                    _set_if_empty(f'fill_{base+1}_P.3', v['currency'] or 'HKD')
                    _set_if_empty(f'fill_{base+2}_P.3', str(v['shares']))
                    if v['paid'] > 0:
                        _set_if_empty(f'fill_{base+3}_P.3', f"{v['paid']:.2f}")
                    _set_if_empty(f'fill_{base+4}_P.3', f"{v['unpaid']:.2f}" if v['unpaid'] > 0 else '0.00')
                    _set_if_empty(f'fill_{base+5}_P.3', f'{total:.2f}')
        except Exception:
            pass  # Non-critical; user can fill manually

    # ── Fill all provided fields (after auto-fill, so explicit fields take priority) ──
    for name, value in fields.items():
        if not value:
            continue
        vstr = str(value)
        if any(ord(c) > 127 for c in vstr):
            _set(name, vstr)
        else:
            _set_text(doc, fmap, name, vstr)

    # ── Align center fields ──
    for name in data.get('alignCenterFields', []):
        if name in fmap:
            pi = fmap[name]
            for w in doc[pi].widgets():
                if w.field_name == name:
                    try:
                        doc.xref_set_key(w._annot.xref, 'Q', '1')
                    except Exception:
                        pass
                    break

    # ── Checkboxes ──
    for name in (data.get('checkboxes', []) or []):
        _check(doc, fmap, name, True)
    # Default: share capital increased
    _check(doc, fmap, 'cb_1_P.1', True)

    # ── Allottee name (may come from either top-level or allottees[0]) ──
    allottee_name = (data.get('allotteeName') or '').strip()
    allottee_name_zh = (data.get('allotteeNameZh') or '').strip()
    allottees_list = data.get('allottees', [])
    if (allottee_name or allottee_name_zh) and not allottees_list:
        allottees_list = [{'nameEn': allottee_name, 'nameZh': allottee_name_zh,
                           'address': data.get('allotteeAddress', ''),
                           'shares': shares}]
    has_allottees = bool(allottees_list and len(allottees_list) > 0 and
                         any((a.get('nameEn', '') or a.get('nameZh', '') or '').strip() for a in allottees_list))

    # ── Non-cash consideration (P.1 + P.2 Section C) ──
    non_cash = data.get('nonCashConsideration', False)
    if non_cash:
        _check(doc, fmap, 'cb_2_P.1', True)  # P.1: non-cash consideration indicator
        non_cash_types = data.get('nonCashTypes', [])
        _TYPE_TO_CB = {
            'division2_part13': 'cb_1_P.2',
            'credited_fully_paid': 'cb_2_P.2',
            'written_contract_s142': 'cb_3_P.2',
        }
        for t in non_cash_types:
            cb_name = _TYPE_TO_CB.get(t)
            if cb_name:
                _check(doc, fmap, cb_name, True)
        # Fill Section C details text area
        details = data.get('nonCashDetails', '')
        if details:
            _set('fill_17_P.2', details)

    # P.3: cb_1 = allottee details in Schedule 2; P.2: cb_4_P.3 = same (template quark)
    if has_allottees:
        _check(doc, fmap, 'cb_1_P.3', True)
        _check(doc, fmap, 'cb_4_P.3', True)
        # P.3 continuation counters: fill_26 = Schedule 2 page count
        # (verified by Qwen VL: 5 boxes = A/B/C/Schedule1/Schedule2, fill_26 is rightmost)
        import math
        sched2_pages = max(1, math.ceil(len(allottees_list) / 2))
        _set_if_empty('fill_26_P.3', str(sched2_pages))

    # ── Overlays ──
    for ov in (data.get('overlays', []) or []):
        try:
            page = doc[ov.get('page', 0)]
            text = ov.get('text', '')
            if text:
                page.insert_textbox(
                    fitz.Rect(ov.get('x', 0), ov.get('y', 0),
                              ov.get('x', 0) + 200, ov.get('y', 0) + 20),
                    text, fontname="helv", fontsize=ov.get('fontsize', 10)
                )
        except Exception:
            pass

    # ── P.7: Schedule 2 — Allottee personal details ──
    if has_allottees:
        # Field specs for each allottee slot (widget name → object key)
        # P.7 widget layout (verified by Qwen VL + PyMuPDF text labels 2026-08-04):
        #   A1: fill_4=nameZh, fill_5=surname, fill_6=otherNames,
        #       fill_8=flat, fill_9=building, fill_10=street, fill_11=district, fill_12=country,
        #       fill_13=shares, cb_1=jointlyHeld
        #   A2: fill_15=nameZh, fill_16=surname, fill_17=otherNames,
        #       fill_19=flat, fill_20=building, fill_21=street, fill_22=district, fill_23=country,
        #       fill_24=shares, cb_2=jointlyHeld
        #   Note: fill_2/3 are section-level, fill_7/18 "英文名稱" removed per user request
        p7_specs_a1 = [
            ('fill_4_P.7', 'nameZh'), ('fill_5_P.7', 'surname'),
            ('fill_6_P.7', 'otherNames'),
            ('fill_8_P.7', 'flat'), ('fill_9_P.7', 'building'),
            ('fill_10_P.7', 'street'), ('fill_11_P.7', 'district'),
            ('fill_12_P.7', 'country'), ('fill_13_P.7', 'shares'),
        ]
        p7_specs_a2 = [
            ('fill_15_P.7', 'nameZh'), ('fill_16_P.7', 'surname'),
            ('fill_17_P.7', 'otherNames'),
            ('fill_19_P.7', 'flat'), ('fill_20_P.7', 'building'),
            ('fill_21_P.7', 'street'), ('fill_22_P.7', 'district'),
            ('fill_23_P.7', 'country'), ('fill_24_P.7', 'shares'),
        ]
        for idx, a in enumerate(allottees_list[:2]):  # P.7 fits 2 allottees
            specs = p7_specs_a1 if idx == 0 else p7_specs_a2
            name_en = (a.get('nameEn', '') or a.get('allotteeName', '')).strip()
            # Build normalized allottee dict with parsed name + structured address
            a_norm = {
                'nameEn': name_en,
                'nameZh': (a.get('nameZh', '') or a.get('allotteeNameZh', '')).strip(),
                'surname': a.get('surname', ''),
                'otherNames': a.get('otherNames', ''),
                'flat': a.get('flat', '') or a.get('allotteeFlat', ''),
                'building': a.get('building', '') or a.get('allotteeBuilding', ''),
                'street': a.get('street', '') or a.get('allotteeStreet', ''),
                'district': a.get('district', '') or a.get('allotteeDistrict', ''),
                'postal': a.get('postal', ''),
                'country': a.get('country', '') or a.get('allotteeCountry', '') or 'Hong Kong',
                'shares': str(a.get('shares', '') or a.get('allotteeShares', '') or shares or ''),
                'remarks': a.get('remarks', '') or a.get('allotteeRemarks', ''),
            }
            # Auto-parse surname/otherNames from nameEn if not explicitly provided
            # HK convention: first word = surname, rest = otherNames
            if name_en and not a_norm['surname']:
                parts = name_en.split()
                if len(parts) >= 2:
                    a_norm['surname'] = parts[0]
                    a_norm['otherNames'] = ' '.join(parts[1:])
                else:
                    a_norm['surname'] = name_en
            # Fallback: if no structured address but flat address exists, use it
            if not a_norm['flat']:
                addr = (a.get('address', '') or a.get('allotteeAddress', '')).strip()
                if addr:
                    a_norm['flat'] = addr
            # Fill all mapped fields
            for field_name, key in specs:
                val = a_norm.get(key, '')
                if val and str(val).strip():
                    _set_if_empty(field_name, str(val).strip())
            # Jointly held checkbox
            is_joint = a.get('jointlyHeld') or a.get('allotteeJointlyHeld')
            if is_joint:
                cb_name = 'cb_1_P.7' if idx == 0 else 'cb_2_P.7'
                try:
                    _check(doc, fmap, cb_name, True)
                except Exception:
                    pass
        # P.7 bottom page counter: "附表二第 _ 頁 Schedule 2 Page _"
        # Each P.7 fits 2 allottees → pages = ceil(count / 2)
        import math
        sched2_pages = max(1, math.ceil(len(allottees_list) / 2))
        _set_if_empty('fill_26_P.7', str(sched2_pages))
        _set_if_empty('fill_27_P.7', str(sched2_pages))
        # Only fallback to old top-level logic if allottees list is empty
    elif allottee_name or allottee_name_zh:
        # Allottee 1 name (backward compatibility)
        if allottee_name_zh:
            _set_if_empty('fill_4_P.7', allottee_name_zh)
        if allottee_name:
            name_parts = allottee_name.strip().split()
            if len(name_parts) >= 2:
                _set_if_empty('fill_5_P.7', name_parts[0])  # HK: first word = surname
                _set_if_empty('fill_6_P.7', ' '.join(name_parts[1:]))
            else:
                _set_if_empty('fill_5_P.7', allottee_name)
        for key, field in [('allotteeFlat', 'fill_8_P.7'), ('allotteeBuilding', 'fill_9_P.7'),
                           ('allotteeStreet', 'fill_10_P.7'), ('allotteeDistrict', 'fill_11_P.7'),
                           ('allotteeCountry', 'fill_12_P.7')]:
            val = data.get(key, '')
            if val:
                _set_if_empty(field, str(val))
        addr = data.get('allotteeAddress', '')
        if addr and not data.get('allotteeFlat', ''):
            _set_if_empty('fill_8_P.7', addr)
        if shares:
            _set_if_empty('fill_13_P.7', str(shares))

    # ── P.3: Signature — Director crossed out (index 1), Secretary kept (index 0) ──
    for dd_name in ('Dropdown1_P.3', 'Dropdown2_P.3'):
        if dd_name in fmap:
            pi = fmap[dd_name]
            for w in doc[pi].widgets():
                if w.field_name == dd_name:
                    try:
                        idx = 1 if dd_name == 'Dropdown1_P.3' else 0
                        w.field_value = 'Yes'
                        doc.xref_set_key(w._annot.xref, 'I', str(idx))
                    except Exception:
                        pass
                    if dd_name == 'Dropdown1_P.3':
                        try:
                            r = w.rect
                            page3 = doc[2]
                            line_y = (r.y0 + r.y1) / 2
                            page3.draw_line(
                                fitz.Point(r.x0 + 2, line_y),
                                fitz.Point(r.x1 - 2, line_y),
                                color=(0, 0, 0), width=1.0
                            )
                        except Exception:
                            pass
                    break

    # ── P.3: Signature date ──
    sign_date = data.get('signDate') or today_str
    if 'fill_28_P.3' in fmap:
        _set_text(doc, fmap, 'fill_28_P.3', sign_date)

    # ── Page management ──
    # Keep: P.1-P.3 always, P.7 (Schedule 2) if allottee data,
    #       any pages referenced in fields dict or overlays
    keep_indices = {0, 1, 2}  # P.1, P.2, P.3
    if has_allottees:
        keep_indices.add(6)  # P.7 = Schedule 2 allottee details
    # Keep pages referenced in fields dict
    for name in fields:
        m = re.match(r'_P\.?(\d+)$', name)
        if m:
            idx = int(m.group(1)) - 1
            if 0 <= idx < doc.page_count:
                keep_indices.add(idx)
    for ov in (data.get('overlays', []) or []):
        p = ov.get('page', -1)
        if 0 <= p < doc.page_count:
            keep_indices.add(p)
    for pno in range(doc.page_count - 1, -1, -1):
        if pno not in keep_indices:
            doc.delete_page(pno)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nsc1-pdf', methods=['POST'])
def generate_nsc1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nsc1_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NNC2 PDF 生成（更改公司名稱通知書） ───

def _fill_nnc2_pdf(data):
    """填充 NNC2 PDF 模板（更改公司名稱通知書），返回 bytes

    模板佈局（3 頁）：
    P.1: BR + 現有名稱(中/英) + 更改後名稱(中/英) + 生效日期 D/M/Y + 簽署 + 提交人
    P.2-P.3: 空白頁
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NNC2-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break
    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        """Fill field with CJK widget AP. Falls back to _set_text for ASCII."""
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]

    # ── P.1: Company Info ──
    _set_text(doc, fmap, 'fill_1_P.1', br8)                          # BR 號碼（右上角窄欄）
    _set_cjk('fill_2_P.1', data.get('oldName', '') or data.get('companyName', ''), align='center')  # 現有英文名稱
    _set_cjk('fill_3_P.1', data.get('newName', ''), align='center')  # 更改後英文名稱

    # 生效日期 D/M/Y（fill_4~6: x=393-571, y=329）
    _set_text(doc, fmap, 'fill_4_P.1', data.get('effectiveDay', ''))
    _set_text(doc, fmap, 'fill_5_P.1', data.get('effectiveMonth', ''))
    _set_text(doc, fmap, 'fill_6_P.1', data.get('effectiveYear', ''))

    # 中文名稱（fill_7-8: x=74, wide）
    _set_cjk('fill_7_P.1', data.get('oldChineseName', ''), align='center')  # 現有中文名稱
    _set_cjk('fill_8_P.1', data.get('newChineseName', ''), align='center')  # 更改後中文名稱

    # 簽署 + 日期（fill_9=簽署人姓名, fill_10=決議日期, y=554）
    _set_cjk('fill_9_P.1', data.get('signerName', '') or data.get('presentorName', ''))
    _set_text(doc, fmap, 'fill_10_P.1', data.get('resolutionDate', ''))

    # Dropdown 簽署人身份（Dropdown_1="Company Director 公司董事", Dropdown_2="Company Secretary 公司秘書"）
    signer_capacity = (data.get('signerCapacity') or 'director').lower()
    if 'Dropdown_1_P.1' in fmap:
        try:
            pi = fmap['Dropdown_1_P.1']
            for w in doc[pi].widgets():
                if w.field_name == 'Dropdown_1_P.1':
                    choices = w.choice_values
                    if choices and len(choices) >= 2:
                        if 'secretary' in signer_capacity:
                            w.field_value = choices[1][0]  # 選 Company Secretary
                        else:
                            w.field_value = choices[0][0]  # 選 Company Director
                        w.update()
                    break
        except Exception:
            pass

    # ── P.1: 提交人資料（fill_11-17）──
    _set_cjk('fill_11_P.1', data.get('presentorNameCn', '') or data.get('presentorNameChinese', ''))
    _set_text(doc, fmap, 'fill_12_P.1', data.get('presentorNameEn', '') or data.get('presentorNameEnglish', ''))
    _set_cjk('fill_13_P.1', data.get('presentorAddress', ''), min_fs=8, valign='bottom')
    _set_text(doc, fmap, 'fill_14_P.1', data.get('presentorPhone', ''))
    _set_text(doc, fmap, 'fill_15_P.1', data.get('presentorFax', ''))
    _set_text(doc, fmap, 'fill_16_P.1', data.get('presentorEmail', ''))
    _set_text(doc, fmap, 'fill_17_P.1', data.get('presentorRef', '') or data.get('presentorReference', ''))

    # ── Auto-populate from DB ──
    company_id = data.get('company_id')
    if company_id:
        try:
            db = get_db()
            company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
            if company:
                c = dict(company)
                if not br8:
                    br8 = re.sub(r'[^0-9A-Za-z]', '', c.get('company_number', '') or '')[:8]
                    _set_text(doc, fmap, 'fill_1_P.1', br8)
                # Auto-fill empty fields from DB
                if not data.get('oldName') and not data.get('companyName'):
                    _set_cjk('fill_2_P.1', c.get('name', ''), align='center')
                if not data.get('oldChineseName'):
                    _set_cjk('fill_7_P.1', c.get('chinese_name', ''), align='center')
        except Exception:
            pass

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nnc2-pdf', methods=['POST'])
def generate_nnc2_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        # Accept both formats: semantic keys OR {fields: {...}} from generic template
        if 'fields' in data and isinstance(data.get('fields'), dict):
            field_map = data['fields']
            # Merge common field mappings into top-level data keys
            data.setdefault('oldName', field_map.get('fill_2_P.1', ''))
            data.setdefault('newName', field_map.get('fill_3_P.1', ''))
            data.setdefault('effectiveDay', field_map.get('fill_4_P.1', ''))
            data.setdefault('effectiveMonth', field_map.get('fill_5_P.1', ''))
            data.setdefault('effectiveYear', field_map.get('fill_6_P.1', ''))
            data.setdefault('oldChineseName', field_map.get('fill_7_P.1', ''))
            data.setdefault('newChineseName', field_map.get('fill_8_P.1', ''))
            data.setdefault('signerName', field_map.get('fill_9_P.1', ''))
            data.setdefault('resolutionDate', field_map.get('fill_10_P.1', ''))
            data.setdefault('presentorNameCn', field_map.get('fill_11_P.1', ''))
            data.setdefault('presentorNameEn', field_map.get('fill_12_P.1', ''))
            data.setdefault('presentorAddress', field_map.get('fill_13_P.1', ''))
            data.setdefault('presentorPhone', field_map.get('fill_14_P.1', ''))
            data.setdefault('presentorFax', field_map.get('fill_15_P.1', ''))
            data.setdefault('presentorEmail', field_map.get('fill_16_P.1', ''))
            data.setdefault('presentorRef', field_map.get('fill_17_P.1', ''))
            if 'brNumber' not in data:
                data['brNumber'] = field_map.get('fill_1_P.1', '')
        pdf_bytes = _fill_nnc2_pdf(data)
        _apply_form_changes_to_company(data, 'nnc2')
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NNC1 PDF 生成（法團成立表格） ───

def _fill_nnc1_pdf(data):
    """填充 NNC1 PDF 模板（法團成立表格股份有限公司），返回 bytes

    模板佈局（24 頁）：
    P.1:  公司名稱(中/英) + 公司類別 checkbox + 業務性質 + 註冊地址 + 提交人
    P.2:  電郵/電話 + 股本結構表
    P.3:  創辦成員（股東）資料
    P.4:  公司秘書（自然人）
    P.5:  公司秘書（法人團體）
    P.6:  董事（自然人）
    P.7:  董事（法人團體）+ 簽署
    P.8:  創辦成員陳述書（簽署 + 續頁頁數）
    P.9-P.13: 續頁（同 P.3-P.7）
    P.14: 呈遞聲明書
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NNC1-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break
    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        """Fill field with CJK widget AP. Falls back to _set_text for ASCII."""
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                field_h = w.rect.height
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                if field_h > 0:
                    fs = min(fs, max(min_fs, int(field_h - 3)))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]

    # ── P.1: 公司名稱 ──
    _set_cjk('fill_1_P.1', data.get('companyName', '') or data.get('nameEnglish', ''), align='center')
    _set_cjk('fill_2_P.1', data.get('companyChinese', '') or data.get('nameChinese', ''), align='center')

    # 公司類別 checkbox
    company_type = (data.get('companyType') or '').lower()
    _check(doc, fmap, 'cb_1_P.1', '私人' in company_type or 'private' in company_type)
    _check(doc, fmap, 'cb_2_P.1', '公眾' in company_type or 'public' in company_type)

    # 業務性質
    _set_text(doc, fmap, 'fill_3_P.1', data.get('businessCode', ''))        # 業務編碼（窄欄 48px）
    _set_cjk('fill_4_P.1', data.get('businessNature', ''), align='center')  # 業務描述（寬欄 420px）

    # 註冊地址（fill_5-8: 室/大廈/街道/區）
    _set_cjk('fill_5_P.1', data.get('addrFlat', ''))
    _set_cjk('fill_6_P.1', data.get('addrBuilding', ''))
    _set_cjk('fill_7_P.1', data.get('addrStreet', ''))
    _set_cjk('fill_8_P.1', data.get('addrDistrict', ''))

    # 提交人（fill_9-15）
    _set_cjk('fill_9_P.1', data.get('presentorNameCn', '') or data.get('submitterNameCn', ''))
    _set_text(doc, fmap, 'fill_10_P.1', data.get('presentorNameEn', '') or data.get('submitterNameEn', ''))
    _set_cjk('fill_11_P.1', data.get('presentorAddress', '') or data.get('submitterAddress', ''), min_fs=7, valign='bottom')
    _set_text(doc, fmap, 'fill_12_P.1', data.get('presentorPhone', '') or data.get('submitterPhone', ''))
    _set_text(doc, fmap, 'fill_13_P.1', data.get('presentorFax', '') or data.get('submitterFax', ''))
    _set_text(doc, fmap, 'fill_14_P.1', data.get('presentorEmail', '') or data.get('submitterEmail', ''))
    _set_text(doc, fmap, 'fill_15_P.1', data.get('presentorRef', '') or data.get('submitterRef', ''))

    # ── P.2: 聯絡 + 股本 ──
    _set_text(doc, fmap, 'fill_1_P.2', data.get('companyEmail', ''))
    _set_text(doc, fmap, 'fill_2_P.2', data.get('companyPhone', ''))

    # 股本表 (Row 1: fill_3-8, Row 2: fill_9-14)
    shares_data = data.get('shares', [])
    # P.2 also accepts flat fields for single-share-class companies
    if not shares_data and data.get('shareClass'):
        shares_data = [{
            'class': data.get('shareClass', 'Ordinary'),
            'number': data.get('totalShares', ''),
            'currency': data.get('shareCurrency', 'HKD'),
            'totalAmount': data.get('shareCapital', ''),
            'paid': data.get('totalPaid', ''),
            'unpaid': data.get('totalUnpaid', ''),
        }]
    for idx, sh in enumerate(shares_data[:2]):  # P.2 最多兩行
        row_offset = idx * 8  # Row1=fill_3-8, Row2=fill_9-14 → offset diff = 6 per field
        # fill_3/9=類別, fill_4/10=數目, fill_5/11=貨幣, fill_6/12=總額, fill_7/13=已繳, fill_8/14=未繳
        _set_text(doc, fmap, f'fill_{3 + row_offset}_P.2', sh.get('class', ''))
        _set_text(doc, fmap, f'fill_{4 + row_offset}_P.2', str(sh.get('number', '')))
        _set_text(doc, fmap, f'fill_{5 + row_offset}_P.2', sh.get('currency', ''))
        _set_text(doc, fmap, f'fill_{6 + row_offset}_P.2', sh.get('totalAmount', ''))
        _set_text(doc, fmap, f'fill_{7 + row_offset}_P.2', sh.get('paid', ''))
        _set_text(doc, fmap, f'fill_{8 + row_offset}_P.2', sh.get('unpaid', ''))

    # 合計行 (fill_15=總股數, fill_16-19 = 貨幣/總額/已繳/未繳)
    total_shares_count = 0
    for sh in shares_data[:2]:
        try:
            total_shares_count += int(str(sh.get('number', '0')).replace(',', ''))
        except ValueError:
            pass
    _set_text(doc, fmap, 'fill_15_P.2', str(total_shares_count) if total_shares_count else '')
    _set_text(doc, fmap, 'fill_16_P.2', data.get('totalCurrency', 'HKD'))
    _set_text(doc, fmap, 'fill_17_P.2', data.get('totalCapital', ''))
    _set_text(doc, fmap, 'fill_18_P.2', data.get('totalPaid', ''))
    _set_text(doc, fmap, 'fill_19_P.2', data.get('totalUnpaid', ''))

    # P.2 備註（fill_24=窄列, fill_25=寬列）
    _set_text(doc, fmap, 'fill_24_P.2', data.get('notesLeft', '') or data.get('notesRef', ''))
    _set_cjk('fill_25_P.2', data.get('notes', '') or data.get('remarks', ''), min_fs=7, valign='bottom')

    # ── Helper: get first officer of each type ──
    officers = data.get('officers', [])
    first_sec_nat = next((o for o in officers if o.get('role') == 'secretary' and o.get('identity') != 'corporate'), None)
    first_sec_corp = next((o for o in officers if o.get('role') == 'secretary' and o.get('identity') == 'corporate'), None)
    first_dir_nat = next((o for o in officers if o.get('role') == 'director' and o.get('identity') != 'corporate'), None)
    first_dir_corp = next((o for o in officers if o.get('role') == 'director' and o.get('identity') == 'corporate'), None)

    shareholders = data.get('shareholders', [])
    first_sh = shareholders[0] if shareholders else None
    signer = data.get('signer', {}) or {}

    # Name parsing helpers
    def _parse_en_name(en):
        parts = (en or '').strip().split()
        return {'surname': parts[0] if parts else '', 'otherNames': ' '.join(parts[1:]) if len(parts) > 1 else ''}

    def _parse_addr(addr):
        parts = (addr or '').split(',')
        return {
            'flat': parts[0].strip() if len(parts) > 0 else '',
            'building': parts[1].strip() if len(parts) > 1 else '',
            'street': parts[2].strip() if len(parts) > 2 else '',
            'district': parts[3].strip() if len(parts) > 3 else '',
            'region': parts[4].strip() if len(parts) > 4 else '',
        }

    def _fmt_hkid(id_str):
        return (id_str or '').strip()[:4]

    # ── P.3: 創辦成員（股東）──
    if first_sh:
        _set_cjk('fill_1_P.3', first_sh.get('name', ''), align='center')       # 中文姓名
        en = _parse_en_name(first_sh.get('surname', '') + ' ' + first_sh.get('otherNames', ''))
        _set_text(doc, fmap, 'fill_2_P.3', first_sh.get('surname', '') or en['surname'])
        _set_text(doc, fmap, 'fill_3_P.3', first_sh.get('otherNames', '') or en['otherNames'])
        # fill_4=OR body-corporate alternative (法人備選), leave empty for natural person

        # 地址（fill_5-9: 室/大廈/街道/區/國家）
        addr = _parse_addr(first_sh.get('address', ''))
        _set_cjk('fill_5_P.3', addr['flat'])
        _set_cjk('fill_6_P.3', addr['building'])
        _set_cjk('fill_7_P.3', addr['street'])
        _set_cjk('fill_8_P.3', addr['district'])
        _set_cjk('fill_9_P.3', addr['region'])

        # 持股（fill_10-13: 類別/數目/貨幣/總額）
        _set_text(doc, fmap, 'fill_10_P.3', first_sh.get('shareType', 'Ordinary'))
        _set_text(doc, fmap, 'fill_11_P.3', str(first_sh.get('shares', '')))
        _set_text(doc, fmap, 'fill_12_P.3', 'HKD')
        _set_text(doc, fmap, 'fill_13_P.3', first_sh.get('amountPaid', ''))

        # 合計行（fill_18=總股數, fill_19=貨幣, fill_20=總款額）
        total_shares = sum(int(str(s.get('shares', '0')) or '0') for s in shareholders)
        total_paid = 0.0
        for s in shareholders:
            amt_str = str(s.get('amountPaid', '0'))
            amt_str = amt_str.replace('HKD', '').replace(',', '').strip()
            try:
                total_paid += float(amt_str)
            except ValueError:
                pass
        _set_text(doc, fmap, 'fill_18_P.3', str(total_shares) if total_shares else '')
        _set_text(doc, fmap, 'fill_19_P.3', 'HKD')
        _set_text(doc, fmap, 'fill_20_P.3', str(int(total_paid)) if total_paid else '')

    # ── P.4: 公司秘書（自然人）──
    if first_sec_nat:
        en = _parse_en_name(first_sec_nat.get('nameEnglish', ''))
        _set_cjk('fill_1_P.4', first_sec_nat.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.4', en['surname'])
        _set_text(doc, fmap, 'fill_3_P.4', en['otherNames'])
        # Previous Names & Alias
        _set_text(doc, fmap, 'fill_4_P.4', first_sec_nat.get('previousNameChinese', ''))
        _set_text(doc, fmap, 'fill_5_P.4', first_sec_nat.get('previousNameEnglish', ''))
        _set_text(doc, fmap, 'fill_6_P.4', first_sec_nat.get('aliasChinese', ''))
        _set_text(doc, fmap, 'fill_7_P.4', first_sec_nat.get('aliasEnglish', ''))
        # Address
        addr = _parse_addr(first_sec_nat.get('address', ''))
        _set_cjk('fill_8_P.4', addr['flat'])
        _set_cjk('fill_9_P.4', addr['building'])
        _set_cjk('fill_10_P.4', addr['street'])
        _set_cjk('fill_11_P.4', addr['district'])
        _set_text(doc, fmap, 'fill_12_P.4', first_sec_nat.get('email', ''))
        # HKID or Passport
        sec_id = (first_sec_nat.get('idNumber', '') or '').strip()
        sec_is_hkid = bool(re.match(r'^[A-Z]?\d', sec_id))
        _set_text(doc, fmap, 'fill_13_P.4', _fmt_hkid(sec_id) if sec_is_hkid else '', align='right')
        _set_text(doc, fmap, 'fill_14_P.4', first_sec_nat.get('passportCountry', '') if not sec_is_hkid and sec_id else '')
        _set_text(doc, fmap, 'fill_15_P.4', sec_id if not sec_is_hkid and sec_id else '')
        # TCSP licence
        tcsp_sec = first_sec_nat.get('tcspLicense', '') or ''
        _set_text(doc, fmap, 'fill_16_P.4', tcsp_sec)
        # TCSP checkbox — only check "not required" if no licence
        if not tcsp_sec:
            _check(doc, fmap, 'cb_1_P.4', True)  # 無須領有 TCSP 牌照

    # ── P.5: 公司秘書（法人團體）──
    if first_sec_corp:
        _set_cjk('fill_1_P.5', first_sec_corp.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.5', first_sec_corp.get('nameEnglish', ''))
        addr = _parse_addr(first_sec_corp.get('address', ''))
        _set_cjk('fill_3_P.5', addr['flat'])
        _set_cjk('fill_4_P.5', addr['building'])
        _set_cjk('fill_5_P.5', addr['street'])
        _set_cjk('fill_6_P.5', addr['district'])
        _set_text(doc, fmap, 'fill_7_P.5', first_sec_corp.get('email', ''))
        _set_text(doc, fmap, 'fill_8_P.5', first_sec_corp.get('companyNumberRef', '') or first_sec_corp.get('idNumber', ''))
        _set_text(doc, fmap, 'fill_9_P.5', first_sec_corp.get('tcspLicense', ''))
        # TCSP checkbox — only check "not required" if no licence
        if not first_sec_corp.get('tcspLicense', ''):
            _check(doc, fmap, 'cb_1_P.5', True)

    # ── P.6: 董事（自然人）──
    if first_dir_nat:
        en = _parse_en_name(first_dir_nat.get('nameEnglish', ''))
        _set_cjk('fill_1_P.6', first_dir_nat.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.6', en['surname'])
        _set_text(doc, fmap, 'fill_3_P.6', en['otherNames'])
        # Previous Names & Alias
        _set_text(doc, fmap, 'fill_4_P.6', first_dir_nat.get('previousNameChinese', ''))
        _set_text(doc, fmap, 'fill_5_P.6', first_dir_nat.get('previousNameEnglish', ''))
        _set_text(doc, fmap, 'fill_6_P.6', first_dir_nat.get('aliasChinese', ''))
        _set_text(doc, fmap, 'fill_7_P.6', first_dir_nat.get('aliasEnglish', ''))
        # Address
        addr = _parse_addr(first_dir_nat.get('address', ''))
        _set_cjk('fill_8_P.6', addr['flat'])
        _set_cjk('fill_9_P.6', addr['building'])
        _set_cjk('fill_10_P.6', addr['street'])
        _set_cjk('fill_11_P.6', addr['district'])
        _set_cjk('fill_12_P.6', addr['region'])
        _set_text(doc, fmap, 'fill_13_P.6', first_dir_nat.get('email', ''))
        # HKID or Passport
        dir_id = (first_dir_nat.get('idNumber', '') or '').strip()
        dir_is_hkid = bool(re.match(r'^[A-Z]?\d', dir_id))
        _set_text(doc, fmap, 'fill_14_P.6', _fmt_hkid(dir_id) if dir_is_hkid else '', align='right')
        _set_text(doc, fmap, 'fill_15_P.6', first_dir_nat.get('passportCountry', '') if not dir_is_hkid and dir_id else '')
        _set_text(doc, fmap, 'fill_16_P.6', dir_id if not dir_is_hkid and dir_id else '')
        _check(doc, fmap, 'cb_1_P.6', True)  # 同意擔任董事

    # ── P.7: 董事（法人團體）──
    if first_dir_corp:
        _set_cjk('fill_1_P.7', first_dir_corp.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.7', first_dir_corp.get('nameEnglish', ''))
        addr = _parse_addr(first_dir_corp.get('address', ''))
        _set_cjk('fill_3_P.7', addr['flat'])
        _set_cjk('fill_4_P.7', addr['building'])
        _set_cjk('fill_5_P.7', addr['street'])
        _set_cjk('fill_6_P.7', addr['district'])
        _set_cjk('fill_7_P.7', addr['region'])
        _set_text(doc, fmap, 'fill_8_P.7', first_dir_corp.get('email', ''))
        _set_text(doc, fmap, 'fill_9_P.7', first_dir_corp.get('companyNumberRef', '') or first_dir_corp.get('idNumber', ''))
        _check(doc, fmap, 'cb_1_P.7', True)
        # 簽署人（法人董事需要簽署人，用股東/創辦成員作為簽署人）
        signer_name = signer.get('nameEnglish', '') or signer.get('surname', '')
        if not signer_name and first_sh:
            signer_name = (first_sh.get('surname', '') + ' ' + first_sh.get('otherNames', '')).strip()
        _set_text(doc, fmap, 'fill_10_P.7', signer_name)

    # ── PI-NNC1 (P.14+): 首任公司秘書／董事(自然人) 受保護資料 ──
    # ⚠️ 每頁只填報一名自然人！需要多頁時自動複製P.14
    # cb_1/cb_2=身份（秘書/董事），非HKID/護照
    # 字段順序：fill_2-4=姓名 → fill_5-6=HKID → fill_7-8=護照 → fill_9-13=住址5欄
    # 優先使用前端送來的 piPersons（含所有自然人），fallback 從 officers 提取
    pi_raw = data.get('piPersons')
    if pi_raw:
        pi_nat_persons = []
        for p in pi_raw:
            en_surname = (p.get('surname', '') or '').strip()
            en_other = (p.get('otherNames', '') or '').strip()
            name_en = f"{en_surname} {en_other}".strip()
            # Reconstruct idNumber from hkidMain+hkidCheck or passportNumber
            if p.get('isHkid'):
                hk_main = (p.get('hkidMain', '') or '').strip()
                hk_check = (p.get('hkidCheck', '') or '').strip()
                id_number = f"{hk_main}({hk_check})" if hk_check else hk_main
            else:
                id_number = (p.get('passportNumber', '') or '').strip()
            # Reconstruct address string for _parse_addr
            addr_parts = [(p.get(k, '') or '').strip() for k in ('addrFlat', 'addrBuilding', 'addrStreet', 'addrDistrict', 'addrRegion')]
            address = ', '.join(addr_parts)
            pi_nat_persons.append({
                'nameChinese': (p.get('nameChinese', '') or '').strip(),
                'nameEnglish': name_en,
                'idNumber': id_number,
                'passportCountry': (p.get('passportCountry', '') or '').strip(),
                'address': address,
                'role': 'secretary' if p.get('isSecretary') else 'director',
                'identity': 'natural',
            })
    else:
        pi_nat_persons = []
        for o in officers:
            if o.get('identity') == 'corporate':
                continue
            if o.get('role') in ('secretary', 'director'):
                pi_nat_persons.append(o)

    def _fill_one_pi_nnc1(page, person, is_first_page=True):
        """Fill PI-NNC1 fields for one natural person on the given page."""
        pi_is_sec = person.get('role') == 'secretary'
        pi_en = _parse_en_name(person.get('nameEnglish', ''))
        pi_id = (person.get('idNumber', '') or '').strip()
        hkid_m = re.match(r'^([A-Z]?\d+)\s*\(?(\d)\)?$', pi_id)
        pi_is_hkid = bool(re.match(r'^[A-Z]?\d', pi_id))
        pi_addr = _parse_addr(person.get('address', ''))

        if is_first_page:
            # Use form fields on the original P.14
            _set_text(doc, fmap, 'fill_1_P.14', data.get('companyName', '') or data.get('nameEnglish', ''))
            _set_cjk('fill_2_P.14', person.get('nameChinese', ''), align='center')
            _set_text(doc, fmap, 'fill_3_P.14', pi_en['surname'])
            _set_text(doc, fmap, 'fill_4_P.14', pi_en['otherNames'])
            _set_text(doc, fmap, 'fill_5_P.14', (hkid_m.group(1) if hkid_m else pi_id) if pi_is_hkid else '')
            _set_text(doc, fmap, 'fill_6_P.14', (hkid_m.group(2) if hkid_m else '') if pi_is_hkid else '')
            _set_text(doc, fmap, 'fill_7_P.14', person.get('passportCountry', '') if not pi_is_hkid and pi_id else '')
            _set_text(doc, fmap, 'fill_8_P.14', pi_id if not pi_is_hkid and pi_id else '')
            _set_cjk('fill_9_P.14', pi_addr['flat'])
            _set_cjk('fill_10_P.14', pi_addr['building'])
            _set_cjk('fill_11_P.14', pi_addr['street'])
            _set_cjk('fill_12_P.14', pi_addr['district'])
            _set_cjk('fill_13_P.14', pi_addr['region'])
        else:
            # Use insert_textbox on copied pages (form fields share names with original)
            # White rectangles over field areas + text overlay
            pi_fields_rects = [
                ((207, 384, 562, 406), person.get('nameChinese', ''), 'cj'),      # fill_2
                ((207, 413, 562, 436), pi_en['surname'], ''),                     # fill_3
                ((207, 442, 562, 464), pi_en['otherNames'], ''),                  # fill_4
                ((257, 482, 513, 504), (hkid_m.group(1) if hkid_m else pi_id) if pi_is_hkid else '', ''),  # fill_5
                ((526, 482, 551, 504), (hkid_m.group(2) if hkid_m else '') if pi_is_hkid else '', ''),    # fill_6
                ((257, 510, 562, 532), person.get('passportCountry', '') if not pi_is_hkid and pi_id else '', ''),  # fill_7
                ((257, 539, 562, 560), pi_id if not pi_is_hkid and pi_id else '', ''),  # fill_8
                ((207, 584, 562, 612), pi_addr['flat'], 'cj'),                    # fill_9
                ((207, 618, 562, 646), pi_addr['building'], 'cj'),                # fill_10
                ((207, 652, 562, 680), pi_addr['street'], 'cj'),                  # fill_11
                ((207, 686, 562, 714), pi_addr['district'], 'cj'),                # fill_12
                ((207, 720, 562, 748), pi_addr['region'], 'cj'),                  # fill_13
            ]
            for fi, ((x0, y0, x1, y1), val, mode) in enumerate(pi_fields_rects):
                if val:
                    page.draw_rect(fitz.Rect(x0-2, y0-2, x1+2, y1+2), color=1, fill=1)  # white
                    # Center name fields (indices 0-2 = fill_2/fill_3/fill_4), left-align others
                    is_center = fi < 3  # first 3 fields are names
                    if mode == 'cj':
                        # Embed CJK font on the copied page if not already there
                        page.clean_contents()
                        try:
                            page.insert_font(fontname='TC', fontfile=_cjk_fontfile)
                        except Exception:
                            pass  # font already embedded
                        page.insert_textbox(fitz.Rect(x0+2, y0+2, x1-2, y1-2), val,
                            fontsize=10, fontname='TC', align=1 if is_center else 0)
                    else:
                        page.insert_textbox(fitz.Rect(x0+2, y0+2, x1-2, y1-2), val,
                            fontsize=10, fontname='Helv', align=1 if is_center else 0)

        # ── Checkbox: redact widget annotations + draw checkmark ──
        # ⚠️ ALL PI-NNC1 pages (original + copies) use the same approach:
        # Widgets on copied pages share field objects with the original P.14 —
        # setting field_value on one page affects ALL pages. And annotations
        # render above page content, so draw_rect cannot cover them.
        # Solution: 1) Redact both checkbox widget areas (removes widget
        # annotations from this page + fills area with white)
        # 2) Draw a checkmark using lines at the correct position.
        # cb_1_P.14 (Director/董事): Rect(207.3, 360.1, 221.8, 374.4)
        # cb_2_P.14 (Secretary/公司秘書):  Rect(313.5, 360.1, 328.0, 374.5)
        page.add_redact_annot(fitz.Rect(206, 359, 223, 376), fill=(1, 1, 1))
        page.add_redact_annot(fitz.Rect(312, 359, 329, 376), fill=(1, 1, 1))
        page.apply_redactions()

        # Redraw checkbox outlines (redaction removes the original widget borders)
        for box_cx, box_cy in [(214.5, 367.5), (320.5, 367.5)]:
            page.draw_rect(fitz.Rect(box_cx-5, box_cy-6, box_cx+5, box_cy+6), color=0, width=1)

        # Draw checkmark with lines at the correct position
        if pi_is_sec:
            cx, cy = 320.5, 367.5  # cb_2 center (Secretary)
        else:
            cx, cy = 214.5, 367.5  # cb_1 center (Director)
        page.draw_line(fitz.Point(cx-3, cy-1), fitz.Point(cx, cy+3), color=0, width=1.5)
        page.draw_line(fitz.Point(cx, cy+3), fitz.Point(cx+5, cy-4), color=0, width=1.5)

    if pi_nat_persons:
        pi_page_idx = 13  # 0-indexed P.14
        _fill_one_pi_nnc1(doc[pi_page_idx], pi_nat_persons[0], is_first_page=True)

        # Extra PI-NNC1 pages for additional natural persons
        # ⚠️ fullcopy_page inserts the copy at the END of the document, NOT after the source!
        for i, person in enumerate(pi_nat_persons[1:], start=1):
            doc.fullcopy_page(pi_page_idx)  # copy goes to end of document
            copy_idx = len(doc) - 1
            target_idx = pi_page_idx + i
            doc.move_page(copy_idx, target_idx)  # move to correct position after source
            _fill_one_pi_nnc1(doc[target_idx], person, is_first_page=False)

    # ── P.8: 創辦成員陳述書 ──
    # fill_1-5: 續頁頁數（A=創辦成員股東, B=秘書自然人, C=秘書法人, D=董事自然人, E=董事法人）
    sec_nat_count = sum(1 for o in officers if o.get('role') == 'secretary' and o.get('identity') != 'corporate')
    sec_corp_count = sum(1 for o in officers if o.get('role') == 'secretary' and o.get('identity') == 'corporate')
    dir_nat_count = sum(1 for o in officers if o.get('role') == 'director' and o.get('identity') != 'corporate')
    dir_corp_count = sum(1 for o in officers if o.get('role') == 'director' and o.get('identity') == 'corporate')

    # Only fill page counts for types that need continuation sheets (>1)
    # A=創辦成員(股東)  B=秘書自然人  C=秘書法人  D=董事自然人  E=董事法人
    sh_count = len(shareholders) if shareholders else 0
    _set_text(doc, fmap, 'fill_1_P.8', str(sh_count - 1) if sh_count > 1 else '')      # A: 股東
    _set_text(doc, fmap, 'fill_2_P.8', str(sec_nat_count - 1) if sec_nat_count > 1 else '')  # B: 秘書自然人
    _set_text(doc, fmap, 'fill_3_P.8', str(sec_corp_count - 1) if sec_corp_count > 1 else '') # C: 秘書法人
    _set_text(doc, fmap, 'fill_4_P.8', str(dir_nat_count - 1) if dir_nat_count > 1 else '')   # D: 董事自然人
    _set_text(doc, fmap, 'fill_5_P.8', str(dir_corp_count - 1) if dir_corp_count > 1 else '') # E: 董事法人

    # PI-NNC1 頁數（受保護資料頁）= 所有自然人秘書+董事
    pi_count = len(pi_nat_persons)
    _set_text(doc, fmap, 'fill_6_P.8', str(pi_count) if pi_count > 0 else '')

    # 創辦成員簽署（fill_7=姓名, fill_8=日期 DD/MM/YYYY）
    signer_fullname = signer.get('nameEnglish', '') or signer.get('fullName', '')
    if not signer_fullname and first_sh:
        signer_fullname = (first_sh.get('surname', '') + ' ' + first_sh.get('otherNames', '')).strip()
    _set_text(doc, fmap, 'fill_7_P.8', signer_fullname)
    sign_date = data.get('signerDate', '')
    if sign_date:
        _set_text(doc, fmap, 'fill_8_P.8', sign_date)

    # ── Auto-populate from DB ──
    company_id = data.get('company_id')
    if company_id:
        try:
            db = get_db()
            company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
            if company:
                c = dict(company)
                if not br8:
                    br8 = re.sub(r'[^0-9A-Za-z]', '', c.get('company_number', '') or '')[:8]
                if not data.get('companyName'):
                    _set_cjk('fill_1_P.1', c.get('name', ''), align='center')
                if not data.get('companyChinese'):
                    _set_cjk('fill_2_P.1', c.get('chinese_name', ''), align='center')
                # Address
                if not data.get('addrFlat'):
                    _set_cjk('fill_5_P.1', c.get('reg_flat', ''))
                    _set_cjk('fill_6_P.1', c.get('reg_building', ''))
                    _set_cjk('fill_7_P.1', c.get('reg_street', ''))
                    _set_cjk('fill_8_P.1', c.get('reg_district', ''))
                if not data.get('companyEmail'):
                    _set_text(doc, fmap, 'fill_1_P.2', c.get('email', ''))
                if not data.get('companyPhone'):
                    _set_text(doc, fmap, 'fill_2_P.2', c.get('phone', ''))
        except Exception:
            pass

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    # 刪除填表須知白頁 (P.15-24, 0-indexed: 14-23)
    # 若插入了額外PI-NNC1頁面，白頁位置會向後偏移
    extra_pi_pages = max(0, len(pi_nat_persons) - 1)
    white_start = 14 + extra_pi_pages
    white_end = 23 + extra_pi_pages
    for pno in range(white_end, white_start - 1, -1):
        if pno < len(doc):
            doc.delete_page(pno)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nnc1-pdf', methods=['POST'])
def generate_nnc1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        # Accept both formats: semantic keys OR {fields: {...}} from generic template
        if 'fields' in data and isinstance(data.get('fields'), dict):
            fm = data['fields']
            data.setdefault('companyName', fm.get('fill_1_P.1', ''))
            data.setdefault('companyChinese', fm.get('fill_2_P.1', ''))
            data.setdefault('businessCode', fm.get('fill_3_P.1', ''))
            data.setdefault('businessNature', fm.get('fill_4_P.1', ''))
            data.setdefault('addrFlat', fm.get('fill_5_P.1', ''))
            data.setdefault('addrBuilding', fm.get('fill_6_P.1', ''))
            data.setdefault('addrStreet', fm.get('fill_7_P.1', ''))
            data.setdefault('addrDistrict', fm.get('fill_8_P.1', ''))
            data.setdefault('presentorNameCn', fm.get('fill_9_P.1', ''))
            data.setdefault('presentorNameEn', fm.get('fill_10_P.1', ''))
            data.setdefault('presentorAddress', fm.get('fill_11_P.1', ''))
            data.setdefault('presentorPhone', fm.get('fill_12_P.1', ''))
            data.setdefault('presentorFax', fm.get('fill_13_P.1', ''))
            data.setdefault('presentorEmail', fm.get('fill_14_P.1', ''))
            data.setdefault('presentorRef', fm.get('fill_15_P.1', ''))
            data.setdefault('companyEmail', fm.get('fill_1_P.2', ''))
            data.setdefault('companyPhone', fm.get('fill_2_P.2', ''))
            # ── Extract P.2 share capital from fields dict ──
            if not data.get('shares') and not data.get('shareClass'):
                sc_class = fm.get('fill_3_P.2', 'Ordinary')
                sc_number = fm.get('fill_4_P.2', '')
                sc_currency = fm.get('fill_5_P.2', 'HKD')
                sc_total = fm.get('fill_6_P.2', '')
                sc_paid = fm.get('fill_7_P.2', '')
                sc_unpaid = fm.get('fill_8_P.2', '')
                if sc_number or sc_total:
                    data.setdefault('shareClass', sc_class)
                    data.setdefault('totalShares', sc_number)
                    data.setdefault('shareCurrency', sc_currency)
                    data.setdefault('shareCapital', sc_total)
                    data.setdefault('totalPaid', sc_paid)
                    data.setdefault('totalUnpaid', sc_unpaid)
                # Total row (fill_15-19)
                data.setdefault('totalCurrency', fm.get('fill_16_P.2', 'HKD'))
            # ── Extract shareholders (P.3 founder members) from fields dict ──
            if not data.get('shareholders'):
                sh_name = fm.get('fill_1_P.3', '')
                sh_surname = fm.get('fill_2_P.3', '')
                sh_other = fm.get('fill_3_P.3', '')
                sh_flat = fm.get('fill_5_P.3', '')
                sh_building = fm.get('fill_6_P.3', '')
                sh_street = fm.get('fill_7_P.3', '')
                sh_district = fm.get('fill_8_P.3', '')
                sh_region = fm.get('fill_9_P.3', '')
                sh_addr = ', '.join([sh_flat, sh_building, sh_street, sh_district, sh_region])
                sh_type = fm.get('fill_10_P.3', 'Ordinary')
                sh_shares = fm.get('fill_11_P.3', '0')
                sh_paid = fm.get('fill_13_P.3', '0')
                if sh_name or sh_surname or sh_other:
                    data['shareholders'] = [{
                        'name': sh_name, 'surname': sh_surname, 'otherNames': sh_other,
                        'address': sh_addr, 'shareType': sh_type,
                        'shares': sh_shares, 'amountPaid': sh_paid,
                    }]
                else:
                    data.setdefault('shareholders', [])
            # ── Extract officers (P.4-P.7) from fields dict ──
            if not data.get('officers'):
                officers = []
                # P.4: Secretary (Natural Person)
                sn_cn = fm.get('fill_1_P.4', '')
                sn_surname = fm.get('fill_2_P.4', '')
                sn_other = fm.get('fill_3_P.4', '')
                sn_flat = fm.get('fill_8_P.4', '')
                sn_building = fm.get('fill_9_P.4', '')
                sn_street = fm.get('fill_10_P.4', '')
                sn_district = fm.get('fill_11_P.4', '')
                sn_id = fm.get('fill_13_P.4', '') or fm.get('fill_15_P.4', '')
                if sn_cn or sn_surname or sn_other:
                    officers.append({
                        'role': 'secretary', 'identity': 'natural',
                        'nameChinese': sn_cn, 'nameEnglish': f'{sn_surname} {sn_other}'.strip(),
                        'idNumber': sn_id,
                        'address': ', '.join([sn_flat, sn_building, sn_street, sn_district, '']),
                    })
                # P.5: Secretary (Body Corporate)
                sc_cn = fm.get('fill_1_P.5', '')
                sc_en = fm.get('fill_2_P.5', '')
                sc_flat = fm.get('fill_3_P.5', '')
                sc_building = fm.get('fill_4_P.5', '')
                sc_street = fm.get('fill_5_P.5', '')
                sc_district = fm.get('fill_6_P.5', '')
                sc_br = fm.get('fill_8_P.5', '')
                if sc_cn or sc_en:
                    officers.append({
                        'role': 'secretary', 'identity': 'corporate',
                        'nameChinese': sc_cn, 'nameEnglish': sc_en,
                        'companyNumberRef': sc_br,
                        'address': ', '.join([sc_flat, sc_building, sc_street, sc_district, '']),
                    })
                # P.6: Director (Natural Person)
                dn_cn = fm.get('fill_1_P.6', '')
                dn_surname = fm.get('fill_2_P.6', '')
                dn_other = fm.get('fill_3_P.6', '')
                dn_flat = fm.get('fill_8_P.6', '')
                dn_building = fm.get('fill_9_P.6', '')
                dn_street = fm.get('fill_10_P.6', '')
                dn_district = fm.get('fill_11_P.6', '')
                dn_region = fm.get('fill_12_P.6', '')
                dn_id = fm.get('fill_14_P.6', '') or fm.get('fill_16_P.6', '')
                if dn_cn or dn_surname or dn_other:
                    officers.append({
                        'role': 'director', 'identity': 'natural',
                        'nameChinese': dn_cn, 'nameEnglish': f'{dn_surname} {dn_other}'.strip(),
                        'idNumber': dn_id,
                        'address': ', '.join([dn_flat, dn_building, dn_street, dn_district, dn_region]),
                    })
                # P.7: Director (Body Corporate)
                dc_cn = fm.get('fill_1_P.7', '')
                dc_en = fm.get('fill_2_P.7', '')
                dc_flat = fm.get('fill_3_P.7', '')
                dc_building = fm.get('fill_4_P.7', '')
                dc_street = fm.get('fill_5_P.7', '')
                dc_district = fm.get('fill_6_P.7', '')
                dc_region = fm.get('fill_7_P.7', '')
                dc_br = fm.get('fill_9_P.7', '')
                if dc_cn or dc_en:
                    officers.append({
                        'role': 'director', 'identity': 'corporate',
                        'nameChinese': dc_cn, 'nameEnglish': dc_en,
                        'companyNumberRef': dc_br,
                        'address': ', '.join([dc_flat, dc_building, dc_street, dc_district, dc_region]),
                    })
                if officers:
                    data['officers'] = officers
                else:
                    data.setdefault('officers', [])
            # ── Extract signer (P.8 founder member) from fields dict ──
            signer_en = fm.get('fill_7_P.8', '')
            if signer_en and not data.get('signer'):
                data['signer'] = {'nameEnglish': signer_en}
            # ── Checkboxes (P.1 company type) ──
            if 'checkboxes' in data and isinstance(data['checkboxes'], list):
                company_type = (data.get('companyType') or '').lower()
                if 'cb_1_P.1' in data['checkboxes'] and 'private' not in company_type:
                    data.setdefault('companyType', 'Private company limited by shares')
                if 'cb_2_P.1' in data['checkboxes'] and 'public' not in company_type:
                    data.setdefault('companyType', 'Public company limited by shares')
        pdf_bytes = _fill_nnc1_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── IRBR1 PDF 生成（申請公司註冊補充表格） ───

def _fill_irbr1_pdf(data):
    """Fill IRBR1 PDF template（致商業登記署通知書），returns bytes.

    1 page, 2 radio buttons (Yes=widget[0] / No=widget[1]) named:
    topmostSubform[0].Page1[0].RadioButtonList[0]

    Accepts: {irbr1_yes: bool, brNumber?: str}
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'IRBR1-template.pdf')
    doc = fitz.open(template_path)

    irbr1_yes = data.get('irbr1_yes', True)  # default Yes
    if isinstance(irbr1_yes, str):
        irbr1_yes = irbr1_yes.lower() in ('yes', 'true', '1', '是')

    # Radio buttons: widget[0] = Yes, widget[1] = No
    page = doc[0]
    widgets = list(page.widgets())
    if len(widgets) >= 2:
        if irbr1_yes:
            widgets[0].field_value = True
            widgets[0].update()
        else:
            widgets[1].field_value = True
            widgets[1].update()

    # Stamp BR number if available
    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]
    if br8:
        _stamp_br_on_all_pages(doc, br8)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-irbr1-pdf', methods=['POST'])
def generate_irbr1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_irbr1_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── IRBR2 PDF 生成（非香港公司致商業登記署通知書） ───

def _fill_irbr2_pdf(data):
    """Fill IRBR2 PDF template（致商業登記署通知書 — 非香港公司），returns bytes.

    1 page, 9 widgets:
    - TextField1[0]: Business Registration Number
    - TextField2[0]: Business name Chinese
    - TextField2[1]: Business name English
    - TextField2[2]: Description & nature of business
    - DateTimeField1[0]: Date of commencement of business
    - RadioButtonList[1] (top): Already registered under Cap.310? widget[0]=Yes, widget[1]=No
    - RadioButtonList[0] (bottom): Elect 3-year certificate validity? widget[0]=Yes, widget[1]=No

    Accepts: { brNumber?, businessNameChinese?, businessNameEnglish?, businessNature?,
               commencementDate?, irbr2_registered?, irbr2_elect3yr? }
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'IRBR2-template.pdf')
    doc = fitz.open(template_path)

    # ── Radio button helpers ──
    irbr2_registered = data.get('irbr2_registered', True)  # default Yes
    if isinstance(irbr2_registered, str):
        irbr2_registered = irbr2_registered.lower() in ('yes', 'true', '1', '是')

    irbr2_elect3yr = data.get('irbr2_elect3yr', True)  # default Yes
    if isinstance(irbr2_elect3yr, str):
        irbr2_elect3yr = irbr2_elect3yr.lower() in ('yes', 'true', '1', '是')

    page = doc[0]
    widgets = list(page.widgets())
    # Widget index layout (verified from template):
    # [0]=TextField1[0](BR), [1]=TextField2[0](CN), [2]=TextField2[1](EN),
    # [3]=TextField2[2](Nature), [4]=DateTimeField1[0](Date),
    # [5]=RadioButtonList[0]-Yes(bottom), [6]=RadioButtonList[0]-No(bottom),
    # [7]=RadioButtonList[1]-Yes(top),    [8]=RadioButtonList[1]-No(top)

    # ── Text fields ──
    br_number = str(data.get('brNumber', '') or '').strip()
    name_cn = str(data.get('businessNameChinese', '') or '').strip()
    name_en = str(data.get('businessNameEnglish', '') or '').strip()
    nature = str(data.get('businessNature', '') or '').strip()
    commencement = str(data.get('commencementDate', '') or '').strip()

    if br_number and len(widgets) > 0:
        widgets[0].field_value = br_number
        widgets[0].update()
    if name_cn and len(widgets) > 1:
        widgets[1].field_value = name_cn
        widgets[1].update()
    if name_en and len(widgets) > 2:
        widgets[2].field_value = name_en
        widgets[2].update()
    if nature and len(widgets) > 3:
        widgets[3].field_value = nature
        widgets[3].update()
    if commencement and len(widgets) > 4:
        widgets[4].field_value = commencement
        widgets[4].update()

    # ── RadioButtonList[1] (top, indices 7,8): Already registered under Cap.310? ──
    if len(widgets) > 8:
        if irbr2_registered:
            widgets[7].field_value = True   # Yes
            widgets[7].update()
        else:
            widgets[8].field_value = True   # No
            widgets[8].update()

    # ── RadioButtonList[0] (bottom, indices 5,6): Elect 3-year certificate? ──
    if len(widgets) > 6:
        if irbr2_elect3yr:
            widgets[5].field_value = True   # Yes
            widgets[5].update()
        else:
            widgets[6].field_value = True   # No
            widgets[6].update()

    # ── BR stamp ──
    if br_number:
        _stamp_br_on_all_pages(doc, br_number)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-irbr2-pdf', methods=['POST'])
def generate_irbr2_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_irbr2_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NN1 PDF 生成（註冊非香港公司申請書） ───

def _fill_nn1_pdf(data):
    """填充 NN1 PDF 模板（註冊非香港公司註冊申請書），返回 bytes

    模板佈局（27 頁）：
    P.1:  擬用公司名稱(中/英) + 成立地 + 成立日期 + 香港地址 + 提交人
    P.2:  電郵/電話 + 主要營業地
    P.3:  獲授權代表（自然人）
    P.4:  公司類別 checkbox + 獲授權代表（法人）
    P.5-P.6: 秘書/董事及其他詳細資料
    P.7-P.9: 續頁
    P.10: 股本結構
    P.11-P.16: 董事/秘書詳情（續）
    P.17: 呈遞聲明書
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NN1-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break
    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        """Fill field with CJK widget AP. Falls back to _set_text for ASCII."""
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                field_h = w.rect.height
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                if field_h > 0:
                    fs = min(fs, max(min_fs, int(field_h - 3)))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]

    # ── P.1: 公司名稱 + 基本資料 ──
    _set_cjk('fill_1_P.1', data.get('proposedNameEnglish', '') or data.get('companyName', ''), align='center')
    _set_cjk('fill_2_P.1', data.get('proposedNameChinese', '') or data.get('companyChinese', ''), align='center')

    # 成立日期 D/M/Y（fill_3-5）
    _set_text(doc, fmap, 'fill_3_P.1', data.get('incorpDay', '') or data.get('estDay', ''))
    _set_text(doc, fmap, 'fill_4_P.1', data.get('incorpMonth', '') or data.get('estMonth', ''))
    _set_text(doc, fmap, 'fill_5_P.1', data.get('incorpYear', '') or data.get('estYear', ''))

    # 成立地點 + 香港地址（fill_6-9: 室/大廈/街道/區）
    _set_cjk('fill_6_P.1', data.get('placeOfIncorporation', ''))
    _set_cjk('fill_7_P.1', data.get('flat', ''))
    _set_cjk('fill_8_P.1', data.get('building', ''))
    _set_cjk('fill_9_P.1', data.get('street', ''))

    # 提交人（fill_10-16）
    _set_cjk('fill_10_P.1', data.get('presentorNameCn', '') or data.get('presentorNameChinese', ''))
    _set_text(doc, fmap, 'fill_11_P.1', data.get('presentorNameEn', '') or data.get('presentorNameEnglish', ''))
    _set_cjk('fill_12_P.1', data.get('presentorAddress', ''), min_fs=7, valign='bottom')
    _set_text(doc, fmap, 'fill_13_P.1', data.get('presentorPhone', ''))
    _set_text(doc, fmap, 'fill_14_P.1', data.get('presentorFax', ''))
    _set_text(doc, fmap, 'fill_15_P.1', data.get('presentorEmail', ''))
    _set_text(doc, fmap, 'fill_16_P.1', data.get('presentorRef', '') or data.get('presentorReference', ''))

    # ── P.2: 聯絡 + 主要營業地 ──
    _set_text(doc, fmap, 'fill_1_P.2', data.get('companyEmail', ''))
    _set_text(doc, fmap, 'fill_2_P.2', data.get('companyPhone', ''))
    # Principal place of business (fill_3-12 = address rows)
    ppb_addr = data.get('principalPlaceOfBusiness', '') or data.get('ppbAddress', '')
    if ppb_addr:
        lines = ppb_addr.split(',')
        for i, line in enumerate(lines[:9]):  # fill_3_P.2 through fill_12_P.2
            field_name = f'fill_{i + 3}_P.2'
            _set_cjk(field_name, line.strip(), fontsize=9, min_fs=7)
    # 授權代表備註
    _set_cjk('fill_13_P.2', data.get('arNotes', '') or data.get('authRepNotes', ''))

    # ── Helper functions ──
    def _parse_en_name(en):
        parts = (en or '').strip().split()
        return {'surname': parts[0] if parts else '', 'otherNames': ' '.join(parts[1:]) if len(parts) > 1 else ''}

    def _parse_addr(addr_str):
        parts = (addr_str or '').split(',')
        return {
            'flat': parts[0].strip() if len(parts) > 0 else '',
            'building': parts[1].strip() if len(parts) > 1 else '',
            'street': parts[2].strip() if len(parts) > 2 else '',
            'district': parts[3].strip() if len(parts) > 3 else '',
        }

    def _fmt_hkid(id_str):
        return (id_str or '').strip()[:4]

    officers = data.get('officers', [])

    # ── P.3: 獲授權代表（自然人）──
    ar_nat = next((o for o in officers if o.get('isAuthRep') and o.get('identity') != 'corporate'), None)
    if ar_nat:
        en = _parse_en_name(ar_nat.get('nameEnglish', ''))
        _set_cjk('fill_1_P.3', ar_nat.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.3', en['surname'])
        _set_text(doc, fmap, 'fill_3_P.3', en['otherNames'])
        # Alias / former name (fill_4~7)
        addr = _parse_addr(ar_nat.get('address', ''))
        _set_cjk('fill_8_P.3', addr['flat'])
        _set_cjk('fill_9_P.3', addr['building'])
        _set_cjk('fill_10_P.3', addr['street'])
        _set_cjk('fill_11_P.3', addr['district'])
        _set_text(doc, fmap, 'fill_12_P.3', ar_nat.get('hkid', ''), align='right')
        _set_text(doc, fmap, 'fill_13_P.3', ar_nat.get('passportCountry', ''))
        _set_text(doc, fmap, 'fill_14_P.3', ar_nat.get('passportNumber', ''))

    # ── P.4: 公司類別 + 獲授權代表（法人）──
    company_type = (data.get('companyType') or '').lower()
    _check(doc, fmap, 'cb_1_P.4', '私人' in company_type or 'private' in company_type)
    _check(doc, fmap, 'cb_2_P.4', '公眾' in company_type or 'public' in company_type)

    ar_corp = next((o for o in officers if o.get('isAuthRep') and o.get('identity') == 'corporate'), None)
    if ar_corp:
        _set_cjk('fill_1_P.4', ar_corp.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.4', ar_corp.get('nameEnglish', ''))
        addr = _parse_addr(ar_corp.get('address', ''))
        _set_cjk('fill_3_P.4', addr['flat'])
        _set_cjk('fill_4_P.4', addr['building'])
        _set_cjk('fill_5_P.4', addr['street'])
        _set_cjk('fill_6_P.4', addr['district'])
        _set_text(doc, fmap, 'fill_7_P.4', ar_corp.get('placeIncorporated', ''))
        # Registration number D/M/Y
        _set_text(doc, fmap, 'fill_8_P.4', ar_corp.get('regDay', ''))
        _set_text(doc, fmap, 'fill_9_P.4', ar_corp.get('regMonth', ''))
        _set_text(doc, fmap, 'fill_10_P.4', ar_corp.get('regYear', ''))

    # ── P.5: 董事（自然人）──
    dir_nat = next((o for o in officers if o.get('role') == 'director' and o.get('identity') != 'corporate'), None)
    if dir_nat:
        en = _parse_en_name(dir_nat.get('nameEnglish', ''))
        _set_cjk('fill_1_P.5', dir_nat.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.5', en['surname'])
        _set_text(doc, fmap, 'fill_3_P.5', en['otherNames'])
        addr = _parse_addr(dir_nat.get('address', ''))
        _set_cjk('fill_8_P.5', addr['flat'])
        _set_cjk('fill_9_P.5', addr['building'])
        _set_cjk('fill_10_P.5', addr['street'])
        _set_cjk('fill_11_P.5', addr['district'])
        _set_cjk('fill_12_P.5', addr.get('region', ''))
        _set_text(doc, fmap, 'fill_13_P.5', dir_nat.get('email', ''))
        _set_text(doc, fmap, 'fill_14_P.5', _fmt_hkid(dir_nat.get('idNumber', '')), align='right')
        _set_text(doc, fmap, 'fill_15_P.5', dir_nat.get('passportCountry', ''))
        _set_text(doc, fmap, 'fill_16_P.5', dir_nat.get('passportNumber', ''))
        # D/M/Y date of appointment
        date_appt = (dir_nat.get('dateAppointed', '') or '').split('/')
        if len(date_appt) >= 3:
            _set_text(doc, fmap, 'fill_17_P.5', date_appt[0])
            _set_text(doc, fmap, 'fill_18_P.5', date_appt[1])
            _set_text(doc, fmap, 'fill_19_P.5', date_appt[2])

    # ── P.6: 董事（法人團體）──
    dir_corp = next((o for o in officers if o.get('role') == 'director' and o.get('identity') == 'corporate'), None)
    if dir_corp:
        _set_cjk('fill_1_P.6', dir_corp.get('nameChinese', ''), align='center')
        _set_text(doc, fmap, 'fill_2_P.6', dir_corp.get('nameEnglish', ''))
        addr = _parse_addr(dir_corp.get('address', ''))
        _set_cjk('fill_3_P.6', addr['flat'])
        _set_cjk('fill_4_P.6', addr['building'])
        _set_cjk('fill_5_P.6', addr['street'])
        _set_cjk('fill_6_P.6', addr['district'])
        _set_cjk('fill_7_P.6', addr.get('region', ''))
        _set_text(doc, fmap, 'fill_8_P.6', dir_corp.get('placeIncorporated', ''))
        _set_text(doc, fmap, 'fill_9_P.6', dir_corp.get('companyNumberRef', '') or dir_corp.get('idNumber', ''))
        # Registration date D/M/Y
        _set_text(doc, fmap, 'fill_10_P.6', dir_corp.get('regDay', ''))
        _set_text(doc, fmap, 'fill_11_P.6', dir_corp.get('regMonth', ''))
        _set_text(doc, fmap, 'fill_12_P.6', dir_corp.get('regYear', ''))

    # ── Auto-populate from DB ──
    company_id = data.get('company_id')
    if company_id:
        try:
            db = get_db()
            company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
            if company:
                c = dict(company)
                if not br8:
                    br8 = re.sub(r'[^0-9A-Za-z]', '', c.get('company_number', '') or '')[:8]
                if not data.get('proposedNameEnglish') and not data.get('companyName'):
                    _set_cjk('fill_1_P.1', c.get('name', ''), align='center')
                if not data.get('proposedNameChinese') and not data.get('companyChinese'):
                    _set_cjk('fill_2_P.1', c.get('chinese_name', ''), align='center')
                if not data.get('flat'):
                    _set_cjk('fill_7_P.1', c.get('reg_flat', ''))
                    _set_cjk('fill_8_P.1', c.get('reg_building', ''))
                    _set_cjk('fill_9_P.1', c.get('reg_street', ''))
                if not data.get('companyEmail'):
                    _set_text(doc, fmap, 'fill_1_P.2', c.get('email', ''))
                if not data.get('companyPhone'):
                    _set_text(doc, fmap, 'fill_2_P.2', c.get('phone', ''))
        except Exception:
            pass

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nn1-pdf', methods=['POST'])
def generate_nn1_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        # Accept both formats: semantic keys OR {fields: {...}} from generic template
        if 'fields' in data and isinstance(data.get('fields'), dict):
            fm = data['fields']
            data.setdefault('proposedNameEnglish', fm.get('fill_1_P.1', ''))
            data.setdefault('proposedNameChinese', fm.get('fill_2_P.1', ''))
            data.setdefault('incorpDay', fm.get('fill_3_P.1', ''))
            data.setdefault('incorpMonth', fm.get('fill_4_P.1', ''))
            data.setdefault('incorpYear', fm.get('fill_5_P.1', ''))
            data.setdefault('placeOfIncorporation', fm.get('fill_6_P.1', ''))
            data.setdefault('flat', fm.get('fill_7_P.1', ''))
            data.setdefault('building', fm.get('fill_8_P.1', ''))
            data.setdefault('street', fm.get('fill_9_P.1', ''))
            data.setdefault('presentorNameCn', fm.get('fill_10_P.1', ''))
            data.setdefault('presentorNameEn', fm.get('fill_11_P.1', ''))
            data.setdefault('presentorAddress', fm.get('fill_12_P.1', ''))
            data.setdefault('presentorPhone', fm.get('fill_13_P.1', ''))
            data.setdefault('presentorFax', fm.get('fill_14_P.1', ''))
            data.setdefault('presentorEmail', fm.get('fill_15_P.1', ''))
            data.setdefault('presentorRef', fm.get('fill_16_P.1', ''))
            data.setdefault('companyEmail', fm.get('fill_1_P.2', ''))
            data.setdefault('companyPhone', fm.get('fill_2_P.2', ''))
        pdf_bytes = _fill_nn1_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NN3 PDF 生成（註冊非香港公司周年申報表） ───

def _fill_nn3_pdf(data):
    """填充 NN3 PDF 模板（非香港公司周年申报表），返回 bytes

    模板佈局：20+ 頁
    P.1: BR + 公司名稱(中/英) + 註冊地 + 香港辦事處地址 + 申報日期 D/M/Y + 提交人
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NN3-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break
    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]

    # ── P.1: BR + 公司名稱 ──
    _set_text(doc, fmap, 'fill_1_P.1', br8)
    _set_cjk('fill_2_P.1', data.get('companyName', '') or data.get('nameEnglish', ''))
    _set_cjk('fill_3_P.1', data.get('companyChineseName', '') or data.get('nameChinese', ''))

    # ── P.1: 註冊地 ──
    _set_text(doc, fmap, 'fill_4_P.1', data.get('placeOfIncorporation', ''))

    # ── P.1: 香港主要辦事處地址（5 fields）──
    addr_src = data
    _set_cjk('fill_5_P.1',  addr_src.get('flat', '') or addr_src.get('regFlat', ''))
    _set_cjk('fill_6_P.1',  addr_src.get('building', '') or addr_src.get('regBuilding', ''))
    _set_cjk('fill_7_P.1',  addr_src.get('street', '') or addr_src.get('regStreet', ''))
    _set_cjk('fill_8_P.1',  addr_src.get('district', '') or addr_src.get('regDistrict', ''))
    _set_cjk('fill_9_P.1',  addr_src.get('region', '') or addr_src.get('regRegion', ''))

    # ── P.1: 申報日期 D/M/Y ──
    return_date = data.get('returnDate', '') or data.get('annualReturnDate', '')
    if return_date and '/' in return_date:
        parts = return_date.split('/')
        if len(parts) >= 3:
            _set_text(doc, fmap, 'fill_10_P.1', parts[0])  # D
            _set_text(doc, fmap, 'fill_11_P.1', parts[1])  # M
            _set_text(doc, fmap, 'fill_12_P.1', parts[2])  # Y

    # ── P.1: 提交人 ──
    _set_cjk('fill_13_P.1', data.get('presentorName', '') or data.get('presenterName', ''))
    _set_cjk('fill_14_P.1', data.get('presentorAddress', '') or data.get('presenterAddress', ''), min_fs=7)
    _set_text(doc, fmap, 'fill_15_P.1', data.get('presentorContact', '') or data.get('presenterContact', ''))

    # ── Auto-populate from DB ──
    company_id = data.get('company_id')
    if company_id:
        try:
            db = get_db()
            company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
            if company:
                c = dict(company)
                if not br8:
                    br8 = re.sub(r'[^0-9A-Za-z]', '', c.get('company_number', '') or '')[:8]
                    _set_text(doc, fmap, 'fill_1_P.1', br8)
                # 从 DB 补充空字段
                if not data.get('companyName') and not data.get('nameEnglish'):
                    _set_cjk('fill_2_P.1', c.get('name', ''))
                if not data.get('companyChineseName') and not data.get('nameChinese'):
                    _set_cjk('fill_3_P.1', c.get('chinese_name', ''))
                if not data.get('placeOfIncorporation'):
                    # Non-HK company may have jurisdiction/incorporation info
                    jurisdiction = c.get('jurisdiction', '') or c.get('place_of_incorporation', '')
                    if jurisdiction:
                        _set_text(doc, fmap, 'fill_4_P.1', jurisdiction)
                # 地址（从公司注册地址，如果表单没有指定）
                if not data.get('flat') and not data.get('regFlat'):
                    _set_cjk('fill_5_P.1', c.get('reg_flat', ''))
                    _set_cjk('fill_6_P.1', c.get('reg_building', ''))
                    _set_cjk('fill_7_P.1', c.get('reg_street', ''))
                    _set_cjk('fill_8_P.1', c.get('reg_district', ''))
                    _set_cjk('fill_9_P.1', c.get('reg_region', ''))
        except Exception:
            pass

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nn3-pdf', methods=['POST'])
def generate_nn3_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        # Accept both formats: semantic keys OR {fields: {...}} from generic template
        if 'fields' in data and isinstance(data.get('fields'), dict):
            fm = data['fields']
            data.setdefault('companyName', fm.get('fill_2_P.1', ''))
            data.setdefault('companyChineseName', fm.get('fill_3_P.1', ''))
            data.setdefault('placeOfIncorporation', fm.get('fill_4_P.1', ''))
            data.setdefault('flat', fm.get('fill_5_P.1', ''))
            data.setdefault('building', fm.get('fill_6_P.1', ''))
            data.setdefault('street', fm.get('fill_7_P.1', ''))
            data.setdefault('district', fm.get('fill_8_P.1', ''))
            data.setdefault('region', fm.get('fill_9_P.1', ''))
            data.setdefault('presentorName', fm.get('fill_13_P.1', ''))
            data.setdefault('presentorAddress', fm.get('fill_14_P.1', ''))
            data.setdefault('presentorContact', fm.get('fill_15_P.1', ''))
        pdf_bytes = _fill_nn3_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NN9 PDF 生成（非香港公司更改地址申報表） ───

def _fill_nn9_pdf(data):
    """填充 NN9 PDF 模板（非香港公司更改地址申报表），返回 bytes

    模板佈局：
    P.1: BR + 公司名稱(中/英) + 舊地址 + 新地址 + 更改日期 D/M/Y + 簽署 + 提交人
    """
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NN9-template.pdf')
    doc = fitz.open(template_path)
    fmap = _build_field_page_map(doc)

    # Locate CJK font
    _cjk_fontfile = None
    for _sf in ['C:/Windows/Fonts/simhei.ttf', 'C:/Windows/Fonts/simsun.ttc',
                 'C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf']:
        if os.path.exists(_sf):
            _cjk_fontfile = _sf
            break
    _cjk_ap_font_pages = set()
    _cjk_ap_font_xref_map = {}

    def _set_cjk(name, value, fontsize=10, align='left', valign='bottom', min_fs=6):
        if name not in fmap or not value:
            return False
        vstr = str(value)
        cjk_n = sum(1 for c in vstr if ord(c) > 127)
        if cjk_n == 0 or not _cjk_fontfile:
            return _set_text(doc, fmap, name, value)
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                asc_n = len(vstr) - cjk_n
                fs = fontsize
                field_w = w.rect.width
                if field_w > 0:
                    usable_w = field_w - 4.0
                    est_w = fs * (cjk_n * 1.0 + asc_n * 0.66)
                    if est_w > usable_w:
                        fs = max(min_fs, int(fs * usable_w / est_w * 0.95))
                _set_widget_cjk_ap(doc, doc[pi], w, vstr, fs, _cjk_fontfile,
                                   _cjk_ap_font_pages, _cjk_ap_font_xref_map,
                                   align=align, valign=valign)
                return True
        return False

    br8 = re.sub(r'[^0-9A-Za-z]', '', data.get('brNumber', '') or '')[:8]

    # ── P.1: BR + 公司名稱 ──
    _set_text(doc, fmap, 'fill_1_P.1', br8)
    _set_cjk('fill_2_P.1', data.get('companyName', '') or data.get('nameEnglish', ''))
    _set_cjk('fill_3_P.1', data.get('companyChineseName', '') or data.get('nameChinese', ''))

    # ── P.1: 舊地址（4 fields）──
    _set_cjk('fill_4_P.1', data.get('oldFlat', '') or data.get('flat', ''))
    _set_cjk('fill_5_P.1', data.get('oldBuilding', '') or data.get('building', ''))
    _set_cjk('fill_6_P.1', data.get('oldStreet', '') or data.get('street', ''))
    _set_cjk('fill_7_P.1', data.get('oldDistrict', '') or data.get('district', ''))

    # ── P.1: 新地址（4 fields）──
    _set_cjk('fill_8_P.1',  data.get('newFlat', '') or data.get('flat', ''))
    _set_cjk('fill_9_P.1',  data.get('newBuilding', '') or data.get('building', ''))
    _set_cjk('fill_10_P.1', data.get('newStreet', '') or data.get('street', ''))
    _set_cjk('fill_11_P.1', data.get('newDistrict', '') or data.get('district', ''))

    # ── P.1: 更改日期 D/M/Y ──
    change_date = data.get('changeDate', '') or data.get('effectiveDate', '')
    if change_date and '/' in change_date:
        parts = change_date.split('/')
        if len(parts) >= 3:
            _set_text(doc, fmap, 'fill_12_P.1', parts[0])  # D
            _set_text(doc, fmap, 'fill_13_P.1', parts[1])  # M
            _set_text(doc, fmap, 'fill_14_P.1', parts[2])  # Y

    # ── P.1: 簽署人 ──
    _set_cjk('fill_15_P.1', data.get('signerName', '') or data.get('presentorName', ''))
    sign_date = data.get('signDate', '') or data.get('resolutionDate', '')
    if sign_date and '/' in sign_date:
        parts = sign_date.split('/')
        if len(parts) >= 3:
            _set_text(doc, fmap, 'fill_16_P.1', parts[0])  # D
            _set_text(doc, fmap, 'fill_17_P.1', parts[1])  # M
            _set_text(doc, fmap, 'fill_18_P.1', parts[2])  # Y

    # ── P.1: 提交人 ──
    _set_cjk('fill_19_P.1', data.get('presentorName', '') or data.get('presenterName', ''))
    _set_cjk('fill_20_P.1', data.get('presentorAddress', '') or data.get('presenterAddress', ''), min_fs=7)
    _set_text(doc, fmap, 'fill_21_P.1', data.get('presentorContact', '') or data.get('presenterContact', ''))

    # ── Auto-populate from DB ──
    company_id = data.get('company_id')
    if company_id:
        try:
            db = get_db()
            company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
            if company:
                c = dict(company)
                if not br8:
                    br8 = re.sub(r'[^0-9A-Za-z]', '', c.get('company_number', '') or '')[:8]
                    _set_text(doc, fmap, 'fill_1_P.1', br8)
                # 从 DB 补充空字段
                if not data.get('companyName') and not data.get('nameEnglish'):
                    _set_cjk('fill_2_P.1', c.get('name', ''))
                if not data.get('companyChineseName') and not data.get('nameChinese'):
                    _set_cjk('fill_3_P.1', c.get('chinese_name', ''))
                # 自动取旧地址 = DB 中的当前注册地址
                if not data.get('oldFlat') and not data.get('flat'):
                    _set_cjk('fill_4_P.1', c.get('reg_flat', ''))
                    _set_cjk('fill_5_P.1', c.get('reg_building', ''))
                    _set_cjk('fill_6_P.1', c.get('reg_street', ''))
                    _set_cjk('fill_7_P.1', c.get('reg_district', ''))
        except Exception:
            pass

    # BR on all pages
    _stamp_br_on_all_pages(doc, br8)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nn9-pdf', methods=['POST'])
def generate_nn9_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        # Accept both formats: semantic keys OR {fields: {...}} from generic template
        if 'fields' in data and isinstance(data.get('fields'), dict):
            fm = data['fields']
            data.setdefault('companyName', fm.get('fill_2_P.1', ''))
            data.setdefault('companyChineseName', fm.get('fill_3_P.1', ''))
            data.setdefault('oldFlat', fm.get('fill_4_P.1', ''))
            data.setdefault('oldBuilding', fm.get('fill_5_P.1', ''))
            data.setdefault('oldStreet', fm.get('fill_6_P.1', ''))
            data.setdefault('oldDistrict', fm.get('fill_7_P.1', ''))
            data.setdefault('newFlat', fm.get('fill_8_P.1', ''))
            data.setdefault('newBuilding', fm.get('fill_9_P.1', ''))
            data.setdefault('newStreet', fm.get('fill_10_P.1', ''))
            data.setdefault('newDistrict', fm.get('fill_11_P.1', ''))
            data.setdefault('signerName', fm.get('fill_15_P.1', ''))
            data.setdefault('presentorName', fm.get('fill_19_P.1', ''))
            data.setdefault('presentorAddress', fm.get('fill_20_P.1', ''))
            data.setdefault('presentorContact', fm.get('fill_21_P.1', ''))
        pdf_bytes = _fill_nn9_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── ND2A PDF 生成（委任/停任董事秘書） ───

def _fill_nd2a_pdf(data, template='ND2A-template.pdf'):
    """Fill ND2A/NN6 PDF template using PyMuPDF"""
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', template)
    doc = fitz.open(template_path)
    fmap = {}
    for pi in range(doc.page_count):
        for w in doc[pi].widgets():
            if w.field_name:
                fmap[w.field_name] = pi

    def _set(name, value, align=None):
        if name not in fmap or not value:
            return
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                try:
                    w.field_value = str(value) if value else ''
                    if align == 'right':
                        doc.xref_set_key(w._annot.xref, 'Q', '2')
                    elif align == 'center':
                        doc.xref_set_key(w._annot.xref, 'Q', '1')
                    w.update()
                except Exception:
                    pass
                break

    def _check(name):
        if name not in fmap:
            return
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                try:
                    # PyMuPDF 1.28+: w._annot is an Annot object (not dict),
                    # so .get('AP') raises AttributeError. Use w.field_value = True
                    # which automatically selects the "On" appearance state.
                    w.field_value = True
                    w.update()
                except Exception:
                    pass
                break

    br8 = (data.get('brNumber', '') or '').replace(r'[^0-9A-Za-z]', '')[:8]

    # P.1: Company info
    _set('fill_1_P.1', br8)
    _set('fill_2_P.1', data.get('companyName', ''))

    # 每页右上角都填商业登记号码
    for pi in range(2, doc.page_count + 1):
        _set(f'fill_1_P.{pi}', br8)

    officers = data.get('officers') or []
    nat_appt_idx = 0   # 自然人委任計數
    nat_cess_idx = 0   # 自然人停任計數
    corp_idx = 0       # 法人計數
    for i, officer in enumerate(officers[:3]):
        is_natural = officer.get('identity') == 'natural'
        is_cessation = officer.get('type') == 'cessation'

        if is_natural:
            if is_cessation:
                # 自然人停任 → P.4（續頁A）
                page = 4
                nat_cess_idx += 1
            else:
                # 自然人委任：P.2（第1個詳細）→ P.6（PI-ND2A 續頁C，第2個）→ P.7（PI-ND2A 續頁，第3個）
                page = [2, 6, 7][nat_appt_idx] if nat_appt_idx < 3 else 7
                nat_appt_idx += 1
        else:
            # 法人：依序 P.3, P.5, P.7
            page = (corp_idx * 2) + 3
            corp_idx += 1
        p = page

        if is_natural:
            # Split English name: prefer explicit surname/otherNames, fallback to parse
            eng = officer.get('nameEnglish', '') or ''
            surname = officer.get('nameSurname', '') or ''
            other = officer.get('nameOtherNames', '') or officer.get('nameOther', '') or ''
            if not surname and eng:
                parts = eng.strip().split()
                surname = parts[-1] if len(parts) > 1 else (parts[0] if parts else '')
                other = ' '.join(parts[:-1]) if len(parts) > 1 else ''
            chinese = officer.get('nameChinese', '')

            # ── Page-specific field mapping (each ND2A page has different widget layout) ──
            if p == 2:
                # P.2: Natural person appointment — detailed info
                # fill_2 = 代替 Alternate to (only for alternate director)
                if officer.get('role') == 'alternate':
                    _set(f'fill_2_P.2', officer.get('alternateTo', ''))
                # Names
                _set('fill_3_P.2', chinese)
                _set('fill_4_P.2', surname)
                _set('fill_5_P.2', other)
                # Structured address: fill_10~14
                addr_fb = officer.get('addrFlatBlock', '')
                addr_bld = officer.get('addrBuilding', '')
                addr_se = officer.get('addrStreetEstate', '')
                addr_dist = officer.get('addrDistrict', '')
                addr_reg = officer.get('addrRegion', '')
                if any([addr_fb, addr_bld, addr_se, addr_dist, addr_reg]):
                    _set('fill_10_P.2', addr_fb)
                    _set('fill_11_P.2', addr_bld)
                    _set('fill_12_P.2', addr_se)
                    _set('fill_13_P.2', addr_dist)
                    _set('fill_14_P.2', addr_reg)
                else:
                    # Fallback: flat address → fill_10
                    _set('fill_10_P.2', officer.get('address', ''))
                # HKID + Passport
                _set('fill_16_P.2', (officer.get('idNumber', '') or '')[:4], align='right')
                if officer.get('passportCountry'):
                    _set('fill_17_P.2', officer.get('passportCountry', ''))
                if officer.get('passportNumber'):
                    _set('fill_18_P.2', _parse_passport_partial(officer['passportNumber']))
                # Date: fill_21/22/23 = D/M/Y（底部三窄栏）
                date_str = officer.get('dateAppointed') if officer.get('type') == 'appointment' else officer.get('dateCeased')
                if date_str:
                    parts = date_str.split('-')
                    if len(parts) >= 3:
                        _set('fill_21_P.2', parts[2])
                        _set('fill_22_P.2', parts[1])
                        _set('fill_23_P.2', parts[0])
                # Role checkboxes on P.2: cb_1=秘書, cb_2=董事, cb_3=候補
                # 每個角色只勾一個 checkbox（互斥）
                role = officer.get('role', 'director')
                if role == 'secretary':
                    _check('cb_1_P.2')
                elif role == 'alternate':
                    _check('cb_3_P.2')
                else:
                    _check('cb_2_P.2')

                # P.2 底部聲明：cross out "董事" or "候補董事" in consent statement
                # Dropdown_1_P.2 → "董事" (director), Dropdown_2_P.2 → "候補董事*" (alternate director)
                # role='director'  → cross out "候補董事" (Dropdown_2)
                # role='alternate' → cross out "董事" (Dropdown_1)
                # NOTE: ND2A has duplicate widget instances (Chinese + English rows), so no break
                if role in ('director', 'alternate'):
                    cross_widget = f'Dropdown_{2 if role == "director" else 1}_P.2'
                    for wn in (f'Dropdown_1_P.2', f'Dropdown_2_P.2'):
                        if wn not in fmap:
                            continue
                        pi = fmap[wn]
                        for w in doc[pi].widgets():
                            if w.field_name == wn:
                                try:
                                    use_dashes = (wn == cross_widget)
                                    opt_idx = 1 if use_dashes else 0
                                    doc.xref_set_key(w._annot.xref, 'I', f'[{opt_idx}]')
                                    val = w.choice_values[opt_idx]
                                    doc.xref_set_key(w._annot.xref, 'V', fitz.get_pdf_str(val))
                                    doc.xref_set_key(w._annot.xref, 'F', '4')
                                    if use_dashes:
                                        doc[pi].draw_line(
                                            fitz.Point(w.rect.x0 + 2, w.rect.y0 + w.rect.height / 2),
                                            fitz.Point(w.rect.x1 - 2, w.rect.y0 + w.rect.height / 2),
                                            color=(0, 0, 0), width=1.0
                                        )
                                except Exception:
                                    pass

                # Already director? cb_5=是, cb_6=否
                if officer.get('alreadyDirector') == 'yes':
                    _check('cb_5_P.2')
                elif officer.get('alreadyDirector') == 'no':
                    _check('cb_6_P.2')
            elif p == 4:
                # P.4: Natural person cessation continuation (續頁A 停任)
                _set('fill_3_P.4', chinese)
                _set('fill_4_P.4', surname)
                _set('fill_5_P.4', other)
                _set('fill_6_P.4', officer.get('idNumber', ''), align='right')
                _set('fill_7_P.4', officer.get('passportNumber', ''))
                # Address: two tall fields
                addr_parts = [officer.get('addrFlatBlock', ''), officer.get('addrBuilding', ''),
                              officer.get('addrStreetEstate', ''), officer.get('addrDistrict', ''),
                              officer.get('addrRegion', '')]
                addr_has = [x for x in addr_parts if x]
                if addr_has:
                    _set('fill_8_P.4', ', '.join(addr_has[:3]))
                    _set('fill_9_P.4', ', '.join(addr_has[3:]))
                else:
                    _set('fill_8_P.4', officer.get('address', ''))
                # Cessation date: fill_10/11/12 = D/M/Y
                date_str = officer.get('dateCeased') or officer.get('dateAppointed')
                if date_str:
                    parts = date_str.split('-')
                    if len(parts) >= 3:
                        _set('fill_10_P.4', parts[2])
                        _set('fill_11_P.4', parts[1])
                        _set('fill_12_P.4', parts[0])
                # Role
                role = officer.get('role', 'director')
                if role == 'secretary':
                    _check('cb_1_P.4')
                elif role == 'alternate':
                    _check('cb_3_P.4')
                else:
                    _check('cb_2_P.4')
                _check('cb_4_P.4')  # cessation
            elif p == 6:
                # P.6 (PI-ND2A 續頁C) — different layout from P.2
                # fill_2 = 代替 Alternate to (僅候補董事填)
                if officer.get('role') == 'alternate':
                    _set(f'fill_2_P.{p}', officer.get('alternateTo', ''))
                # fill_3 = 中文姓名, fill_4 = 姓氏, fill_5 = 名字
                _set(f'fill_3_P.{p}', chinese)
                _set(f'fill_4_P.{p}', surname)
                _set(f'fill_5_P.{p}', other)
                # fill_8 = 住址（五欄地址組合）, fill_9 = 國家／地區 (護照簽發國), fill_10 = 通訊地址
                addr_p6 = [officer.get('addrFlatBlock', ''), officer.get('addrBuilding', ''),
                           officer.get('addrStreetEstate', ''), officer.get('addrDistrict', ''),
                           officer.get('addrRegion', '')]
                addr_p6_has = [x for x in addr_p6 if x]
                _set(f'fill_8_P.{p}', ', '.join(addr_p6_has) if addr_p6_has else officer.get('address', ''))
                if officer.get('passportCountry'):
                    _set(f'fill_9_P.{p}', officer.get('passportCountry', ''))
                # fill_11 = 身份證號碼（完整號碼，不截斷）
                if officer.get('idNumber'):
                    _set(f'fill_11_P.{p}', officer['idNumber'], align='right')
                # 護照號碼：P.6 無獨立欄位，與身份證號碼共用 fill_11
                if officer.get('passportNumber'):
                    existing = officer.get('idNumber', '')
                    if existing:
                        _set(f'fill_11_P.{p}', f'{existing} / {officer["passportNumber"]}')
                    else:
                        _set(f'fill_11_P.{p}', officer['passportNumber'])
                # Dates: fill_14/15/16 = D/M/Y
                date_str = officer.get('dateAppointed') if officer.get('type') == 'appointment' else officer.get('dateCeased')
                if date_str:
                    parts = date_str.split('-')
                    if len(parts) >= 3:
                        _set(f'fill_14_P.{p}', parts[2])
                        _set(f'fill_15_P.{p}', parts[1])
                        _set(f'fill_16_P.{p}', parts[0])
                role = officer.get('role', 'director')
                if role == 'secretary':
                    _check(f'cb_1_P.{p}')
                elif role == 'alternate':
                    _check(f'cb_3_P.{p}')
                else:
                    _check(f'cb_2_P.{p}')
                if officer.get('type') == 'cessation':
                    _check(f'cb_4_P.{p}')

                # P.6 底部聲明：cross out "董事" or "候補董事" in consent statement
                # (same pattern as P.2 — only pages that have the dropdowns)
                if f'Dropdown_1_P.{p}' in fmap:
                    if role in ('director', 'alternate'):
                        cross_widget = f'Dropdown_{2 if role == "director" else 1}_P.{p}'
                        for wn in (f'Dropdown_1_P.{p}', f'Dropdown_2_P.{p}'):
                            if wn not in fmap:
                                continue
                            pi2 = fmap[wn]
                            for w in doc[pi2].widgets():
                                if w.field_name == wn:
                                    try:
                                        use_dashes = (wn == cross_widget)
                                        opt_idx = 1 if use_dashes else 0
                                        doc.xref_set_key(w._annot.xref, 'I', f'[{opt_idx}]')
                                        val = w.choice_values[opt_idx]
                                        doc.xref_set_key(w._annot.xref, 'V', fitz.get_pdf_str(val))
                                        doc.xref_set_key(w._annot.xref, 'F', '4')
                                        if use_dashes:
                                            doc[pi2].draw_line(
                                                fitz.Point(w.rect.x0 + 2, w.rect.y0 + w.rect.height / 2),
                                                fitz.Point(w.rect.x1 - 2, w.rect.y0 + w.rect.height / 2),
                                                color=(0, 0, 0), width=1.0
                                            )
                                    except Exception:
                                        pass
            elif p == 7:
                # ── P.7 (PI-ND2A) 受保護資料頁 ──
                # 公眾紀錄不會顯示此頁。完整 HKID 及護照號碼在此頁全部顯示。
                # 欄位佈局（由上至下）：
                #   fill_2 = 中文姓名, fill_3 = 英文姓氏, fill_4 = 英文名字
                #   fill_5 = 香港身份證(完整號碼), fill_6 = 括號校驗位 e.g. "(1)"
                #   fill_7 = 護照簽發國家/地區, fill_8 = 護照完整號碼
                #   fill_9 = 室/樓/座, fill_10 = 大廈, fill_11 = 街道/屋苑
                #   fill_12 = 區/市/省, fill_13 = 國家/地區
                _set(f'fill_2_P.{p}', chinese)
                _set(f'fill_3_P.{p}', surname)
                _set(f'fill_4_P.{p}', other)
                # HKID 完整號碼 + 括號校驗位
                id_full = officer.get('idNumber', '') or ''
                if id_full:
                    # 解析 HKID：主體部分 vs 括號校驗位
                    # 格式如 "Y231456(1)" → main="Y231456", check="(1)"
                    import re as _re
                    hkid_match = _re.match(r'^([A-Za-z]?\d+)\s*(\([^)]*\))?$', id_full.strip())
                    if hkid_match:
                        hkid_main = hkid_match.group(1)
                        hkid_check = hkid_match.group(2) or ''
                        _set(f'fill_5_P.{p}', hkid_main, align='right')
                        if hkid_check:
                            _set(f'fill_6_P.{p}', hkid_check)
                    else:
                        _set(f'fill_5_P.{p}', id_full, align='right')
                # 護照簽發國家 + 護照完整號碼
                if officer.get('passportCountry'):
                    _set(f'fill_7_P.{p}', officer.get('passportCountry', ''))
                if officer.get('passportNumber'):
                    _set(f'fill_8_P.{p}', officer['passportNumber'])
                # 通常住址（董事／候補董事）
                addr_fb = officer.get('addrFlatBlock', '')
                addr_bld = officer.get('addrBuilding', '')
                addr_se = officer.get('addrStreetEstate', '')
                addr_dist = officer.get('addrDistrict', '')
                addr_reg = officer.get('addrRegion', '')
                if any([addr_fb, addr_bld, addr_se, addr_dist, addr_reg]):
                    _set(f'fill_9_P.{p}', addr_fb)
                    _set(f'fill_10_P.{p}', addr_bld)
                    _set(f'fill_11_P.{p}', addr_se)
                    _set(f'fill_12_P.{p}', addr_dist)
                    _set(f'fill_13_P.{p}', addr_reg)
                else:
                    _set(f'fill_9_P.{p}', officer.get('address', ''))
                # Role checkboxes: cb_1=秘書, cb_2=董事, cb_3=候補
                role = officer.get('role', 'director')
                if role == 'secretary':
                    _check(f'cb_1_P.{p}')
                elif role == 'alternate':
                    _check(f'cb_3_P.{p}')
                else:
                    _check(f'cb_2_P.{p}')
                # Note: P.7 (PI-ND2A) has no D/M/Y date fields, no cessation checkbox,
                # no Dropdown cross-out fields, no "already director" checkboxes
            else:
                # Should not reach here for natural persons; fallback to P.6-style
                pass
        else:
            # ── 法人團體 (Body Corporate) ──
            # P.7 (PI-ND2A) layout is COMPLETELY different from P.3/P.5
            if p == 7:
                # ── P.7 (PI-ND2A) 法人團體 ──
                # PI-ND2A 是為自然人設計的頁面。法人團體無 HKID / 護照，
                # 只填姓名 + 地址 + 角色，fill_5/6/7/8（HKID/護照）留空。
                _set(f'fill_2_P.7', officer.get('nameChinese', ''))
                # English name: company name split into surname+otherNames for P.7
                eng_full = officer.get('companyName', officer.get('nameEnglish', ''))
                eng_parts = eng_full.strip().split()
                if len(eng_parts) > 1:
                    _set(f'fill_3_P.7', eng_parts[0])  # 首詞→姓氏位
                    _set(f'fill_4_P.7', ' '.join(eng_parts[1:]))  # 餘詞→名字位
                else:
                    _set(f'fill_3_P.7', eng_full)
                # 法人無 HKID / 護照 → fill_5/6/7/8 全部留空
                # 五欄地址
                addr_flat = officer.get('addrFlatBlock', '')
                addr_bld = officer.get('addrBuilding', '')
                addr_se = officer.get('addrStreetEstate', '')
                addr_dist = officer.get('addrDistrict', '')
                addr_reg = officer.get('addrRegion', '')
                if any([addr_flat, addr_bld, addr_se, addr_dist, addr_reg]):
                    _set(f'fill_9_P.7', addr_flat)
                    _set(f'fill_10_P.7', addr_bld)
                    _set(f'fill_11_P.7', addr_se)
                    _set(f'fill_12_P.7', addr_dist)
                    _set(f'fill_13_P.7', addr_reg)
                else:
                    _set(f'fill_9_P.7', officer.get('address', ''))
                # Role
                role = officer.get('role', 'director')
                if role == 'secretary':
                    _check(f'cb_1_P.7')
                elif role == 'alternate':
                    _check(f'cb_3_P.7')
                else:
                    _check(f'cb_2_P.7')
            else:
                # ── P.3/P.5 法人團體 (Body Corporate) ──
                # Field mapping verified by 千问 VL (2026-08-01):
                #   fill_3 = 中文名稱, fill_4 = 英文名稱
                #   fill_5 = Flat/Floor/Block, fill_6 = Building
                #   fill_7 = Street/Estate, fill_8 = District/City
                #   fill_9 = Country/Region, fill_10 = Email
                #   fill_11 = Business Registration Number (牌照) [right col]
                #   fill_12 = TCSP Licence No. [left col]
                #   fill_14/15/16 = Date of Appointment D/M/Y
                _set(f'fill_3_P.{p}', officer.get('nameChinese', ''))
                _set(f'fill_4_P.{p}', officer.get('companyName', officer.get('nameEnglish', '')))
                # 五欄地址：優先使用結構化地址，fallback 到 flat address
                addr_flat = officer.get('addrFlatBlock', '')
                addr_bld = officer.get('addrBuilding', '')
                addr_se = officer.get('addrStreetEstate', '')
                addr_dist = officer.get('addrDistrict', '')
                addr_reg = officer.get('addrRegion', '')
                if any([addr_flat, addr_bld, addr_se, addr_dist, addr_reg]):
                    _set(f'fill_5_P.{p}', addr_flat)
                    _set(f'fill_6_P.{p}', addr_bld)
                    _set(f'fill_7_P.{p}', addr_se)
                    _set(f'fill_8_P.{p}', addr_dist)
                    _set(f'fill_9_P.{p}', addr_reg)
                else:
                    # Fallback: parse flat address string
                    addr = officer.get('address', '')
                    if addr:
                        _set(f'fill_5_P.{p}', addr)
                # Email
                if officer.get('email'):
                    _set(f'fill_10_P.{p}', officer.get('email', ''))
                # Business Registration Number (商業登記號碼 = 牌照)
                if officer.get('companyNumber'):
                    _set(f'fill_11_P.{p}', officer.get('companyNumber', ''))
                # TCSP Licence No. (fill_12, left column) — only if applicable
                if officer.get('tcspLicence'):
                    _set(f'fill_12_P.{p}', officer.get('tcspLicence', ''))
                # Appointment/cessation date for corporate officers: fill_14/15/16 = D/M/Y
                date_str = None
                if officer.get('type') == 'appointment':
                    date_str = officer.get('dateAppointed')
                elif officer.get('type') == 'cessation':
                    date_str = officer.get('dateCeased')
                if date_str:
                    parts = date_str.split('-')
                    if len(parts) >= 3:
                        _set(f'fill_14_P.{p}', parts[2])
                        _set(f'fill_15_P.{p}', parts[1])
                        _set(f'fill_16_P.{p}', parts[0])
                # Role: cb_1=秘書, cb_2=董事, cb_3=候補董事
                role = officer.get('role', 'director')
                if role == 'secretary':
                    _check(f'cb_1_P.{p}')
                elif role == 'alternate':
                    _check(f'cb_3_P.{p}')
                else:
                    _check(f'cb_2_P.{p}')
                # Type: only cessation needs checkbox (appointment is implicit)
                if officer.get('type') == 'cessation':
                    _check(f'cb_4_P.{p}')

                # ── P.3/P.5 法人團體簽署橫線（千问 VL 驗證 2026-08-02）──
                # 第一簽署：董事(法人團體)的 董事／公司秘書／獲授權人士*
                #   Dropdown_3=董事, Dropdown_4=公司秘書, Dropdown_5=獲授權人士
                #   默認由法人團體的董事簽署 → KEEP 董事, CROSS OUT 公司秘書+獲授權人士
                # 第二簽署：董事Director／公司秘書Company Secretary*
                #   Dropdown_6=董事, Dropdown_7=公司秘書
                #   默認董事簽署 → KEEP 董事, CROSS OUT 公司秘書
                if p in (3, 5):
                    # First signature: Dropdown_3/4/5 (中英文雙行，不 break)
                    # KEEP Dropdown_3 (董事), CROSS OUT Dropdown_4+5
                    for dn in ('Dropdown_3', 'Dropdown_4', 'Dropdown_5'):
                        key = f'{dn}_P.{p}'
                        if key not in fmap:
                            continue
                        pi = fmap[key]
                        cross_out = dn in ('Dropdown_4', 'Dropdown_5')
                        for w in doc[pi].widgets():
                            if w.field_name == key:
                                try:
                                    opt_idx = 1 if cross_out else 0
                                    doc.xref_set_key(w._annot.xref, 'I', f'[{opt_idx}]')
                                    val = w.choice_values[opt_idx]
                                    doc.xref_set_key(w._annot.xref, 'V', fitz.get_pdf_str(val))
                                    doc.xref_set_key(w._annot.xref, 'F', '4')
                                    if cross_out:
                                        doc[pi].draw_line(
                                            fitz.Point(w.rect.x0 + 2, w.rect.y0 + w.rect.height / 2),
                                            fitz.Point(w.rect.x1 - 2, w.rect.y0 + w.rect.height / 2),
                                            color=(0, 0, 0), width=1.0
                                        )
                                except Exception:
                                    pass
                    # Second signature: Dropdown_6/7 (中英文雙行，不 break)
                    # KEEP Dropdown_6 (董事), CROSS OUT Dropdown_7 (公司秘書)
                    for dn in ('Dropdown_6', 'Dropdown_7'):
                        key = f'{dn}_P.{p}'
                        if key not in fmap:
                            continue
                        pi = fmap[key]
                        cross_out = (dn == 'Dropdown_7')
                        for w in doc[pi].widgets():
                            if w.field_name == key:
                                try:
                                    opt_idx = 1 if cross_out else 0
                                    doc.xref_set_key(w._annot.xref, 'I', f'[{opt_idx}]')
                                    val = w.choice_values[opt_idx]
                                    doc.xref_set_key(w._annot.xref, 'V', fitz.get_pdf_str(val))
                                    doc.xref_set_key(w._annot.xref, 'F', '4')
                                    if cross_out:
                                        doc[pi].draw_line(
                                            fitz.Point(w.rect.x0 + 2, w.rect.y0 + w.rect.height / 2),
                                            fitz.Point(w.rect.x1 - 2, w.rect.y0 + w.rect.height / 2),
                                            color=(0, 0, 0), width=1.0
                                        )
                                except Exception:
                                    pass

    # ── PI-ND2A 受保護資料頁（P.7）：始終填入第一個人的完整資料 ──
    # P.7 是獨立於主表格的受保護資料頁，公眾紀錄不會顯示。
    # 優先取第一個自然人（完整 HKID + 護照），若無自然人則取第一個法人團體（BR + 成立地）。
    first_nat = next((o for o in officers if o.get('identity') == 'natural'), None)
    first_corp = next((o for o in officers if o.get('identity') == 'corporate'), None)
    pi_subject = first_nat or first_corp  # 優先自然人，fallback 法人
    if pi_subject and fmap.get('fill_1_P.7') is not None:
        p = 7
        is_nat = pi_subject.get('identity') == 'natural'
        if is_nat:
            # ── 自然人 PI-ND2A：完整 HKID + 護照 ──
            eng = pi_subject.get('nameEnglish', '') or ''
            surname = pi_subject.get('nameSurname', '') or ''
            other = pi_subject.get('nameOtherNames', '') or pi_subject.get('nameOther', '') or ''
            if not surname and eng:
                parts = eng.strip().split()
                surname = parts[-1] if len(parts) > 1 else (parts[0] if parts else '')
                other = ' '.join(parts[:-1]) if len(parts) > 1 else ''
            chinese = pi_subject.get('nameChinese', '')
            _set(f'fill_2_P.{p}', chinese)
            _set(f'fill_3_P.{p}', surname)
            _set(f'fill_4_P.{p}', other)
            # HKID 完整號碼 + 括號校驗位
            id_full = pi_subject.get('idNumber', '') or ''
            if id_full:
                import re as _re
                hkid_match = _re.match(r'^([A-Za-z]?\d+)\s*(\([^)]*\))?$', id_full.strip())
                if hkid_match:
                    _set(f'fill_5_P.{p}', hkid_match.group(1), align='right')
                    if hkid_match.group(2):
                        _set(f'fill_6_P.{p}', hkid_match.group(2))
                else:
                    _set(f'fill_5_P.{p}', id_full, align='right')
            # 護照：簽發國家 + 完整號碼（不截斷）
            if pi_subject.get('passportCountry'):
                _set(f'fill_7_P.{p}', pi_subject.get('passportCountry', ''))
            if pi_subject.get('passportNumber'):
                _set(f'fill_8_P.{p}', pi_subject['passportNumber'])
            # 通常住址
            addr_fb = pi_subject.get('addrFlatBlock', '')
            addr_bld = pi_subject.get('addrBuilding', '')
            addr_se = pi_subject.get('addrStreetEstate', '')
            addr_dist = pi_subject.get('addrDistrict', '')
            addr_reg = pi_subject.get('addrRegion', '')
            if any([addr_fb, addr_bld, addr_se, addr_dist, addr_reg]):
                _set(f'fill_9_P.{p}', addr_fb)
                _set(f'fill_10_P.{p}', addr_bld)
                _set(f'fill_11_P.{p}', addr_se)
                _set(f'fill_12_P.{p}', addr_dist)
                _set(f'fill_13_P.{p}', addr_reg)
            else:
                _set(f'fill_9_P.{p}', pi_subject.get('address', ''))
            # Role checkbox
            role = pi_subject.get('role', 'director')
            if role == 'secretary':
                _check(f'cb_1_P.{p}')
            elif role == 'alternate':
                _check(f'cb_3_P.{p}')
            else:
                _check(f'cb_2_P.{p}')
        else:
            # ── 法人團體 PI-ND2A：只填姓名+地址+角色，不填HKID/護照（公司沒有）──
            _set(f'fill_2_P.7', pi_subject.get('nameChinese', ''))
            eng_full = pi_subject.get('companyName', pi_subject.get('nameEnglish', ''))
            eng_parts = eng_full.strip().split()
            if len(eng_parts) > 1:
                _set(f'fill_3_P.7', eng_parts[0])
                _set(f'fill_4_P.7', ' '.join(eng_parts[1:]))
            else:
                _set(f'fill_3_P.7', eng_full)
            # 法人團體無 HKID / 護照 → fill_5/6/7/8 留空
            # Address
            addr_fb = pi_subject.get('addrFlatBlock', '')
            addr_bld = pi_subject.get('addrBuilding', '')
            addr_se = pi_subject.get('addrStreetEstate', '')
            addr_dist = pi_subject.get('addrDistrict', '')
            addr_reg = pi_subject.get('addrRegion', '')
            if any([addr_fb, addr_bld, addr_se, addr_dist, addr_reg]):
                _set(f'fill_9_P.7', addr_fb)
                _set(f'fill_10_P.7', addr_bld)
                _set(f'fill_11_P.7', addr_se)
                _set(f'fill_12_P.7', addr_dist)
                _set(f'fill_13_P.7', addr_reg)
            else:
                _set(f'fill_9_P.7', pi_subject.get('address', ''))
            role = pi_subject.get('role', 'director')
            if role == 'secretary':
                _check(f'cb_1_P.7')
            elif role == 'alternate':
                _check(f'cb_3_P.7')
            else:
                _check(f'cb_2_P.7')

    # ── P.1 日期（僅停任，取自第一個 officer） ──
    # fill_11/12/13 = D/M/Y（三個並排窄框），委任不填 P.1 日期
    if officers and officers[0].get('type') == 'cessation':
        date_str = officers[0].get('dateCeased')
        if date_str:
            if '-' in date_str:
                parts = date_str.split('-')
                if len(parts) >= 3:
                    _set('fill_11_P.1', parts[2])  # day
                    _set('fill_12_P.1', parts[1])  # month
                    _set('fill_13_P.1', parts[0])  # year
            elif '/' in date_str:
                parts = date_str.split('/')
                if len(parts) >= 3:
                    _set('fill_11_P.1', parts[0])  # day
                    _set('fill_12_P.1', parts[1])  # month
                    _set('fill_13_P.1', parts[2])  # year

    # ── P.1 提交人信息 ──

    # 缩小提交人字段字号（12pt → 9pt），避免长文本溢出
    for field_name in ('fill_14_P.1', 'fill_16_P.1', 'fill_17_P.1', 'fill_18_P.1', 'fill_19_P.1'):
        if field_name in fmap:
            pi = fmap[field_name]
            for w in doc[pi].widgets():
                if w.field_name == field_name:
                    try:
                        xref = w._annot.xref
                        da = doc.xref_get_key(xref, 'DA')
                        if da[0] == 'string':
                            import re
                            new_da = re.sub(r'\d+(?:\.\d+)?\s+Tf', '9 Tf', da[1])
                            doc.xref_set_key(xref, 'DA', fitz.get_pdf_str(new_da))
                    except Exception:
                        pass
                    break

    # fill_14 = 提交人名稱, fill_15 = 提交人地址
    _set('fill_14_P.1', data.get('presentorName', ''))
    _set('fill_15_P.1', data.get('presentorAddress', ''))

    # fill_16-19 = 電話 / 傳真 / 電郵 / 檔號
    _set('fill_16_P.1', data.get('presentorPhone', data.get('presentorContact', '')))
    _set('fill_17_P.1', data.get('presentorFax', ''))
    _set('fill_18_P.1', data.get('presentorEmail', ''))
    _set('fill_19_P.1', data.get('presentorReference', ''))

    # Delete blank pages after P.7 (keep P.1~P.7, P.8+ are blank instruction pages)
    for pno in range(doc.page_count - 1, 6, -1):
        doc.delete_page(pno)

    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nd2a-pdf', methods=['POST'])
def generate_nd2a_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nd2a_pdf(data)
        _apply_form_changes_to_company(data, 'nd2a')
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NN6 PDF 生成（非香港公司更改秘書及董事） ───

@app.route('/api/generate-nn6-pdf', methods=['POST'])
def generate_nn6_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nd2a_pdf(data, template='NN6-template.pdf')
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── ND2B PDF 生成（更改董事秘書詳情） ───

def _fill_nd2b_pdf(data, template='ND2B-template.pdf'):
    """Fill ND2B/NN7 PDF template using PyMuPDF"""
    template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', template)
    doc = fitz.open(template_path)
    fmap = {}
    for pi in range(doc.page_count):
        for w in doc[pi].widgets():
            if w.field_name:
                fmap[w.field_name] = pi

    def _set(name, value, align=None):
        if name not in fmap or not value:
            return
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                try:
                    w.field_value = str(value) if value else ''
                    if align == 'right':
                        doc.xref_set_key(w._annot.xref, 'Q', '2')
                    elif align == 'center':
                        doc.xref_set_key(w._annot.xref, 'Q', '1')
                    w.update()
                except Exception:
                    pass
                break

    def _check(name):
        if name not in fmap:
            return
        pi = fmap[name]
        for w in doc[pi].widgets():
            if w.field_name == name:
                try:
                    # PyMuPDF 1.28+: w._annot is an Annot object (not dict),
                    # so .get('AP') raises AttributeError. Use w.field_value = True
                    # which automatically selects the "On" appearance state.
                    w.field_value = True
                    w.update()
                except Exception:
                    pass
                break

    br8 = (data.get('brNumber', '') or '').replace(r'[^0-9A-Za-z]', '')[:8]
    # Prefer explicit surname/otherNames from frontend, fall back to parsing nameEnglish
    surname = data.get('nameSurname', '')
    other = data.get('nameOtherNames', '')
    if not surname and not other:
        name_parts = (data.get('nameEnglish', '') or '').strip().split()
        surname = name_parts[0] if name_parts else ''  # First word = surname (Chinese convention)
        other = ' '.join(name_parts[1:]) if len(name_parts) > 1 else ''

    # ── Parse change types (support both new array and old string) ──
    change_types = data.get('changeTypes', [])
    if isinstance(change_types, str):
        change_types = [change_types] if change_types else []
    # backward compat: old singular changeType
    old_ct = data.get('changeType', '')
    if old_ct and old_ct not in change_types:
        change_types.append(old_ct)

    # ── Parse effective date for D/M/Y fields ──
    eff_date = data.get('effectiveDate', '')
    eff_day = eff_month = eff_year = ''
    if eff_date:
        # Try YYYY-MM-DD
        parts = eff_date.split('-')
        if len(parts) == 3:
            eff_day, eff_month, eff_year = parts[2], parts[1], parts[0]
        else:
            # Try DD/MM/YYYY
            parts = eff_date.split('/')
            if len(parts) == 3:
                eff_day, eff_month, eff_year = parts[0], parts[1], parts[2]

    # ── Build new address string ──
    new_addr_parts = [
        data.get('newFlat', ''),
        data.get('newBuilding', ''),
        data.get('newStreet', ''),
        data.get('newDistrict', ''),
        data.get('newRegion', '')
    ]
    new_address = ', '.join(p for p in new_addr_parts if p)
    # backward compat: old newAddress single string
    if not new_address:
        new_address = data.get('newAddress', '')

    # ── Build new English name ──
    new_surname = data.get('newNameSurname', '')
    new_other = data.get('newNameOtherNames', '')
    new_english = f"{new_surname} {new_other}".strip()
    new_chinese = data.get('newNameChinese', '')
    # backward compat: old newNameEnglish single string
    if not new_english:
        old_new_name = data.get('newNameEnglish', '')
        if old_new_name:
            nparts = old_new_name.strip().split()
            new_surname = nparts[0] if nparts else ''
            new_other = ' '.join(nparts[1:]) if len(nparts) > 1 else ''
            new_english = old_new_name

    # P.1: Company info
    _set('fill_1_P.1', br8)
    _set('fill_2_P.1', data.get('companyName', ''))

    is_natural = data.get('identity') == 'natural'
    role = data.get('role', '')

    if is_natural:
        if role == 'secretary':
            _check('cb_1_P.1')
        else:
            _check('cb_2_P.1')
        _set('fill_3_P.1', data.get('nameChinese', ''))
        _set('fill_4_P.1', surname)
        _set('fill_5_P.1', other)
        _set('fill_7_P.1', data.get('idNumber', ''), align='right')
        if data.get('passportCountry') or data.get('passportPlaceOfIssue'):
            _set('fill_7b_P.1', data.get('passportCountry') or data.get('passportPlaceOfIssue', ''))
        if data.get('passportNumber'):
            _set('fill_7c_P.1', _parse_passport_partial(data['passportNumber']))

        # ── P.2: Change details (multi-type support) ──
        # Effective date fills (shared across all change rows)
        if eff_day:
            _set('fill_5_P.2', eff_day)
            _set('fill_4_P.2', eff_month)
            _set('fill_3_P.2', eff_year)

        # (a) 姓名更改 Name Change
        if 'name' in change_types:
            # Old name (current) — fill_2 is description row for item 14
            current_name = data.get('nameEnglish', '') or f"{surname} {other}".strip()
            _set('fill_2_P.2', current_name)
            # New Chinese name
            if new_chinese:
                _set('fill_6_P.2', new_chinese)
            # New English name
            if new_english:
                _set('fill_7_P.2', new_english)
            # Name change effective date
            if eff_day:
                _set('fill_10_P.2', eff_day)
                _set('fill_9_P.2', eff_month)
                _set('fill_8_P.2', eff_year)

        # (b) 別名 Alias
        new_alias_eng = data.get('newAliasEnglish', '')
        new_alias_cn = data.get('newAliasChinese', '')
        if new_alias_eng or new_alias_cn:
            alias_text = f"{new_alias_eng} {new_alias_cn}".strip()
            _set('fill_12_P.2', alias_text)
            if eff_day:
                _set('fill_15_P.2', eff_day)
                _set('fill_14_P.2', eff_month)
                _set('fill_13_P.2', eff_year)

        # (d) 通訊地址更改 Address Change
        if 'address' in change_types:
            _set('fill_19_P.2', data.get('newFlat', ''))
            _set('fill_20_P.2', data.get('newBuilding', ''))
            _set('fill_21_P.2', data.get('newStreet', ''))
            _set('fill_22_P.2', data.get('newDistrict', ''))
            _set('fill_23_P.2', data.get('newRegion', ''))
            # backward compat: old newAddress single string → fill_19
            if not data.get('newFlat') and data.get('newAddress'):
                _set('fill_19_P.2', data.get('newAddress', ''))
            if eff_day:
                _set('fill_26_P.2', eff_day)
                _set('fill_25_P.2', eff_month)
                _set('fill_24_P.2', eff_year)

        # (f) 聯絡資料更改 Contact Change
        if 'contact' in change_types:
            new_email = data.get('newEmail', '')
            if new_email:
                _set('fill_27_P.2', new_email)
                if eff_day:
                    _set('fill_30_P.2', eff_day)
                    _set('fill_29_P.2', eff_month)
                    _set('fill_28_P.2', eff_year)

        # (g) 證件號碼更改 ID Change
        if 'id' in change_types:
            new_id = data.get('newIdNumber', '')
            if new_id:
                _set('fill_35_P.2', new_id, align='right')
                if eff_day:
                    _set('fill_34_P.2', eff_day)
                    _set('fill_33_P.2', eff_month)
                    _set('fill_32_P.2', eff_year)
            # Passport change
            new_passport = data.get('passportNumber', '')
            new_passport_country = data.get('passportPlaceOfIssue', '') or data.get('passportCountry', '')
            if new_passport:
                _set('fill_37_P.2', _parse_passport_partial(new_passport))
            if new_passport_country:
                _set('fill_36_P.2', new_passport_country)
            if (new_passport or new_passport_country) and eff_day:
                _set('fill_39_P.2', eff_day)
                _set('fill_38_P.2', eff_month)
                _set('fill_37_P.2', eff_year)  # year for passport row

        # ── P.6: Protected Information (PI-ND2B) ──
        if role == 'secretary':
            _check('cb_1_P.6')
        else:
            _check('cb_2_P.6')
        # Current values (mirror P.1)
        _set('fill_2_P.6', data.get('nameChinese', ''))
        _set('fill_3_P.6', surname)
        _set('fill_4_P.6', other)
        # New values (mirror changes)
        if 'address' in change_types and new_address:
            _set('fill_9_P.6', new_address)
        elif data.get('newAddress'):
            _set('fill_9_P.6', data.get('newAddress', ''))
        # HKID (new if changed, else current)
        if 'id' in change_types and data.get('newIdNumber'):
            _set('fill_5_P.6', data.get('newIdNumber', ''), align='right')
        else:
            _set('fill_5_P.6', data.get('idNumber', ''), align='right')
        # Passport
        ppoi = data.get('passportPlaceOfIssue', '') or data.get('passportCountry', '')
        if 'id' in change_types and ppoi:
            _set('fill_7_P.6', ppoi)
        elif ppoi:
            _set('fill_7_P.6', ppoi)

    # ── P.1 提交人信息 ──
    # fill_8 = 提交人名稱, fill_9 = 提交人地址
    _set('fill_8_P.1', data.get('presentorName', ''))
    _set('fill_9_P.1', data.get('presentorAddress', ''))
    # fill_10-13 = 電話 / 傳真 / 電郵 / 檔號
    _set('fill_10_P.1', data.get('presentorPhone', data.get('presentorContact', '')))
    _set('fill_11_P.1', data.get('presentorFax', ''))
    _set('fill_12_P.1', data.get('presentorEmail', ''))
    _set('fill_13_P.1', data.get('presentorReference', ''))

    # P.3: Signature
    _set('fill_30_P.3', data.get('signerName', ''))
    _set('fill_31_P.3', data.get('signDate', ''))

    # BR on all pages
    for pi in range(2, doc.page_count + 1):
        _set(f'fill_1_P.{pi}', br8)

    # ⚠️ 不删页：保留模板全部页面（仅 NAR1 可以删空页）
    pdf_bytes = doc.write(deflate=True)
    doc.close()
    return pdf_bytes


@app.route('/api/generate-nd2b-pdf', methods=['POST'])
def generate_nd2b_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nd2b_pdf(data)
        _apply_nd2b_changes_to_person(data)  # Apply changes to person record
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── NN7 PDF 生成（非香港公司更改秘書及董事詳情） ───

@app.route('/api/generate-nn7-pdf', methods=['POST'])
def generate_nn7_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _fill_nd2b_pdf(data, template='NN7-template.pdf')
        _apply_nd2b_changes_to_person(data)  # Apply changes to person record
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── Generic Form PDF 生成（Resolution / Rename / NewCompany） ───

@app.route('/api/generate-generic-form-pdf', methods=['POST'])
def generate_generic_form_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400

        pdf = FPDF()
        pdf.add_page()
        # Use system CJK font directly (font download in find_font() can hang)
        font_path = None
        for sys_font in ['C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf', 'C:/Windows/Fonts/simsun.ttc']:
            if os.path.exists(sys_font):
                font_path = sys_font
                break
        if font_path:
            pdf.add_font('TC', '', font_path)
            pdf.add_font('TC', 'B', font_path)
            font_name = 'TC'
        else:
            font_name = 'Helvetica'

        def safe_text(text):
            return (text or '').encode('utf-8', errors='replace').decode('utf-8')

        # Title
        pdf.set_font(font_name, 'B', 16)
        pdf.cell(0, 10, safe_text(data.get('title', '')), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
        if data.get('subtitle'):
            pdf.set_font(font_name, '', 10)
            pdf.cell(0, 7, safe_text(data['subtitle']), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
        pdf.ln(4)

        # Form code + company info
        pdf.set_font(font_name, '', 10)
        if data.get('formCode'):
            pdf.cell(30, 6, 'Form Code:', 0, 0)
            pdf.cell(0, 6, safe_text(data['formCode']), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if data.get('companyName'):
            pdf.cell(30, 6, 'Company:', 0, 0)
            pdf.cell(0, 6, safe_text(data['companyName']), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if data.get('brNumber'):
            pdf.cell(30, 6, 'BR No.:', 0, 0)
            pdf.cell(0, 6, safe_text(data['brNumber']), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

        # Sections
        sections = data.get('sections') or []
        for sec in sections:
            if sec.get('heading'):
                pdf.set_font(font_name, 'B', 11)
                pdf.cell(0, 7, safe_text(sec['heading']), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(1)

            if sec.get('rows'):
                pdf.set_font(font_name, '', 9)
                for row in sec['rows']:
                    label = safe_text(row[0]) if len(row) > 0 else ''
                    value = safe_text(row[1]) if len(row) > 1 else ''
                    pdf.cell(50, 5, label + ':', 0, 0)
                    pdf.cell(0, 5, value, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)

            if sec.get('paragraph'):
                pdf.set_font(font_name, '', 9)
                pdf.multi_cell(0, 5, safe_text(sec['paragraph']))
                pdf.ln(2)

            if sec.get('bullets'):
                pdf.set_font(font_name, '', 9)
                for b in sec['bullets']:
                    pdf.cell(5, 5, '', 0, 0)
                    pdf.cell(0, 5, chr(8226) + ' ' + safe_text(b), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)

        # Signature lines
        sig_lines = data.get('signatureLines') or []
        if sig_lines:
            pdf.ln(6)
            pdf.set_font(font_name, '', 10)
            for sl in sig_lines:
                pdf.cell(0, 8, safe_text(sl), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)

        pdf_bytes = bytes(pdf.output())
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── Auto-Fill CR Form PDF（自動填入公司資料的政府表格 PDF） ───

@app.route('/api/generate-cr-form-pdf', methods=['POST', 'OPTIONS'])
def generate_cr_form_pdf():
    """Auto-fill a CR form PDF from company database data.
    Request: {company_id, form_code (nar1|nd2a|...|nn9)}"""
    if request.method == 'OPTIONS':
        return ('', 204)
    try:
        u = get_user()
        if not u:
            return jsonify({'error': 'Not authenticated'}), 401

        data = request.get_json(force=True, silent=True) or {}
        company_id = data.get('company_id')
        form_code = (data.get('form_code') or '').lower()
        if not company_id or not form_code:
            return jsonify({'error': '缺少 company_id 或 form_code'}), 400

        meta = CR_FORM_META.get(form_code)
        if not meta:
            return jsonify({'error': f'不支援的表格代碼：{form_code}'}), 400

        db = get_db()
        bundle = _docx_company_bundle(db, company_id)
        if not bundle:
            return jsonify({'error': '找不到該公司'}), 404
        c = bundle['c']
        name_en = c.get('name') or ''
        name_cn = c.get('chinese_name') or ''
        br = c.get('company_number') or ''
        cr_num = c.get('ci_number') or ''

        # Build PDF with fpdf2
        pdf = FPDF()
        pdf.add_page()
        font_path = None
        for sys_font in ['C:/Windows/Fonts/msjh.ttc', 'C:/Windows/Fonts/Deng.ttf', 'C:/Windows/Fonts/simsun.ttc']:
            if os.path.exists(sys_font):
                font_path = sys_font
                break
        if font_path:
            pdf.add_font('TC', '', font_path)
            pdf.add_font('TC', 'B', font_path)
            ft = 'TC'
        else:
            ft = 'Helvetica'

        def _t(text):
            return (text or '').encode('utf-8', errors='replace').decode('utf-8')

        def _person_label(m):
            en = (m.get('name_english') or '').strip()
            cn = (m.get('name_chinese') or '').strip()
            if en and cn:
                return f"{en}（{cn}）"
            return en or cn or '—'

        # Header
        pdf.set_font(ft, 'B', 16)
        pdf.cell(0, 10, _t(meta['title']), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
        pdf.set_font(ft, '', 10)
        pdf.cell(0, 7, _t(f"{meta['code']} — {meta['title_en']}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
        pdf.cell(0, 7, _t(f"公司註冊處表格 · 由系統自動填入 {name_en or name_cn} 的資料生成草稿"),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')
        pdf.ln(4)

        # Company info section
        pdf.set_font(ft, 'B', 12)
        pdf.cell(0, 7, _t('公司基本資料'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font(ft, '', 9)
        info_rows = [
            ('英文名稱', name_en), ('中文名稱', name_cn),
            ('商業登記號碼 (BR)', br), ('公司註冊編號 (CR)', cr_num),
            ('公司類型', c.get('company_type')), ('成立日期', _docx_fmt_date(c.get('incorporation_date'))),
            ('狀態', c.get('status')), ('註冊辦事處地址', bundle['address']),
            ('電郵', c.get('email')), ('電話', c.get('phone')),
        ]
        for label, val in info_rows:
            if val:
                pdf.cell(50, 5, _t(label) + ':', 0, 0)
                pdf.cell(0, 5, _t(str(val)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

        # Directors & Secretaries (most forms)
        has_officers = form_code in ('nar1','nd2a','nd2b','nd4','nnc1','nn1','nn3','nn6','nn7')
        if has_officers:
            pdf.set_font(ft, 'B', 11)
            pdf.cell(0, 7, _t(f"董事（{len(bundle['directors'])} 人）"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(ft, '', 9)
            if bundle['directors']:
                for d in bundle['directors']:
                    pdf.cell(0, 5, _t(f"  {_person_label(d)}  |  {d.get('id_number') or d.get('passport_number') or ''}  |  委任: {_docx_fmt_date(d.get('date_appointed'))}  |  {d.get('address') or ''}"),
                             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                pdf.cell(0, 5, _t('  （無董事記錄）'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

            pdf.set_font(ft, 'B', 11)
            pdf.cell(0, 7, _t(f"公司秘書（{len(bundle['secretaries'])} 人）"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(ft, '', 9)
            if bundle['secretaries']:
                for s in bundle['secretaries']:
                    pdf.cell(0, 5, _t(f"  {_person_label(s)}  |  TCSP: {s.get('tcsp_number') or ''}  |  委任: {_docx_fmt_date(s.get('date_appointed'))}"),
                             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                pdf.cell(0, 5, _t('  （無秘書記錄）'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(3)

        # Shareholders
        has_shares = form_code in ('nar1','nsc1','nnc1','nn1','nn3')
        if has_shares:
            ts = bundle['total_shares']
            pdf.set_font(ft, 'B', 11)
            pdf.cell(0, 7, _t(f"股東／股本結構（總發行股數：{ts}）"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(ft, '', 9)
            if bundle['shareholders']:
                for sh in bundle['shareholders']:
                    pct = f"{round(int(sh.get('shares') or 0) * 100 / ts, 2)}%" if ts else '—'
                    pdf.cell(0, 5, _t(f"  {_person_label(sh)}  |  {sh.get('shares')} 股  |  {sh.get('share_type') or '普通股'}  |  {pct}"),
                             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                pdf.cell(0, 5, _t('  （無股東記錄）'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(3)

        # Form-specific blocks
        if form_code == 'nar1':
            pdf.set_font(ft, '', 9)
            pdf.cell(0, 5, _t('重要控制人登記冊 (SCR) 是否備存於公司註冊辦事處？  是 □  否 □'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.cell(0, 5, _t('截至申報日期之董事／秘書／股東資料以上表為準。'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if form_code in ('nr1', 'ndr1', 'nn9'):
            pdf.set_font(ft, '', 9)
            pdf.cell(0, 5, _t(f"現有註冊地址：{bundle['address'] or '（未填）'}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.cell(0, 5, _t('變更後註冊地址（請手動填寫）：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if form_code == 'nsc1':
            pdf.set_font(ft, '', 9)
            for line in ['配發日期：＿＿＿＿＿＿＿＿', '配發股份類別：＿＿＿＿＿＿＿＿', '每股發行價：＿＿＿＿＿＿＿＿', '配發總額：＿＿＿＿＿＿＿＿']:
                pdf.cell(0, 5, _t(line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if form_code in ('nnc1', 'nn1'):
            pdf.set_font(ft, '', 9)
            pdf.cell(0, 5, _t('擬成立公司之名稱／形式等詳情，請參考上方公司基本資料。'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if form_code in ('nnc2',):
            pdf.set_font(ft, '', 9)
            pdf.cell(0, 5, _t('更改後公司名稱（請手動填寫）：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        # Signature block
        pdf.ln(6)
        pdf.set_font(ft, 'B', 11)
        pdf.cell(0, 7, _t('簽署 / SIGNED:'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(5)
        pdf.set_font(ft, '', 10)
        pdf.cell(0, 8, _t('_______________________________'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 6, _t('董事 / Director       日期 Date：＿＿＿＿＿＿'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)
        pdf.cell(0, 8, _t('_______________________________'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.cell(0, 6, _t('公司秘書 / Company Secretary       日期 Date：＿＿＿＿＿＿'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        # Footer
        pdf.ln(6)
        pdf.set_font(ft, '', 7)
        pdf.cell(0, 5, _t(f"本文件由公司秘書管理系統自動生成 · {datetime.now().strftime('%Y-%m-%d %H:%M')}"),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

        pdf_bytes = bytes(pdf.output())
        safe_name = re.sub(r'[^\w一-鿿-]', '_', (name_en or name_cn or 'company'))[:30]
        filename = f"{meta['code']}_{meta['title']}_{safe_name}.pdf"
        return jsonify({
            'success': True,
            'pdf': base64.b64encode(pdf_bytes).decode('ascii'),
            'filename': filename,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── Generic Template PDF Filler（PyMuPDF 填充任何 AcroForm 模板） ───

@app.route('/api/generate-template-pdf', methods=['POST'])
def generate_template_pdf():
    """Fill any AcroForm PDF template with provided field values.
    Request: {template: 'ND4-template.pdf', fields: {'fill_1_P.1': 'val', ...}, checkboxes: ['cb_1_P.1', ...]}"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400

        template_name = data.get('template', '')
        if not template_name:
            return jsonify({'error': 'template name required'}), 400

        # Security: only allow .pdf files in templates dir
        safe_name = os.path.basename(template_name)
        if not safe_name.endswith('.pdf') or '..' in safe_name:
            return jsonify({'error': 'invalid template name'}), 400

        template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', safe_name)
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template not found: {safe_name}'}), 404

        doc = fitz.open(template_path)

        # Build field → page map
        fmap = {}
        for pi in range(doc.page_count):
            for w in doc[pi].widgets():
                if w.field_name:
                    fmap[w.field_name] = pi

        fields = data.get('fields') or {}
        checkboxes = data.get('checkboxes') or []

        # Fill text fields
        for name, value in fields.items():
            if name not in fmap:
                continue
            pi = fmap[name]
            for w in doc[pi].widgets():
                if w.field_name == name:
                    try:
                        w.field_value = str(value) if value else ''
                        w.update()
                    except Exception:
                        pass
                    break

        # Check checkboxes
        for name in checkboxes:
            if name not in fmap:
                continue
            pi = fmap[name]
            for w in doc[pi].widgets():
                if w.field_name == name:
                    try:
                        w.field_value = True
                        w.update()
                    except Exception:
                        pass
                    break

        pdf_bytes = doc.write(deflate=True)
        doc.close()

        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── Resolution AI 生成（DeepSeek） ───

@app.route('/api/generate-resolution', methods=['POST'])
def generate_resolution():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        if not data.get('companyName') or not data.get('resolutionType'):
            return jsonify({'error': 'companyName and resolutionType required'}), 400

        lang = data.get('language', 'bilingual')
        lang_instruction = (
            "Output in Traditional Chinese only."
            if lang == 'zh' else
            "Output in English only."
            if lang == 'en' else
            "Output bilingually: each major paragraph in English, followed by Traditional Chinese on a new line."
        )

        system_prompt = f"""You are a Hong Kong corporate secretarial assistant. Generate a formal company resolution.
Follow Hong Kong Companies Ordinance conventions. Use the WRITTEN RESOLUTION format unless the user asks for a meeting minute.
Structure:
1. Header: company name + Chinese name + BR number + "Written Resolution of Directors / Members"
2. Resolution number (e.g. "RESOLVED THAT:")
3. Body — clear, formal, numbered if multiple items
4. Effective date
5. Signature block (Director(s) / Members)
{lang_instruction}
Return ONLY the resolution body text, no markdown headers, no commentary, ready to paste into a PDF."""

        user_prompt = f"""Generate a {data['resolutionType']} resolution for:
Company: {data.get('companyName', '')}{' (' + data.get('companyChineseName', '') + ')' if data.get('companyChineseName') else ''}
BR Number: {data.get('brNumber', '—')}
Resolution Date: {data.get('resolutionDate', '')}

Context / Specific details from user:
{data.get('context', '(no extra context provided)')}"""

        api_key = os.environ.get('DEEPSEEK_API_KEY', '')
        if not api_key:
            return jsonify({'error': 'DEEPSEEK_API_KEY not configured'}), 500

        req_body = json.dumps({
            'model': 'deepseek-chat',
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt},
            ],
        }).encode('utf-8')

        ai_req = urllib.request.Request(
            'https://api.deepseek.com/v1/chat/completions',
            data=req_body,
            headers={
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json',
            }
        )
        ai_resp = urllib.request.urlopen(ai_req, timeout=120)
        result = json.loads(ai_resp.read())
        content = result.get('choices', [{}])[0].get('message', {}).get('content', '')

        return jsonify({'content': content})
    except urllib.error.HTTPError as e:
        return jsonify({'error': f'AI API error {e.code}: {e.reason}'}), 502
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─────────────────────────────────────────────
# ─────────────────────────────────────────────
# ROM PDF from RTF template (Word COM Find & Replace)
# Uses the Paul Tang RTF template directly — finds sample data in the
# template and replaces with actual company/shareholder/transaction data.
# ─────────────────────────────────────────────

def _rtf_rom_to_pdf(db, company_id):
    """Generate ROM PDF from Paul Tang .doc template via Word COM Table.Cell() API.

    Opens Testing ROM.doc (24-row x 15-col blank table), fills cells with
    real data from the database, then saves as PDF.

    Cell mapping:
      Header:     Cell(1,2)=Company Name, Cell(2,2)=Company Number
      SH1:        Rows 4-13  (labels at 4,1..8,10; data in col 2/4/6)
      SH2:        Rows 15-24 (same layout as SH1, offset +11 rows)
      Tx sub-table: Rows 9-13 (SH1), Rows 20-24 (SH2)

    For >2 shareholders, the empty SH2 block (rows 15-24) is copied and
    appended to the table for each additional shareholder.
    For 1 shareholder, the unused SH2 block is deleted.

    Returns bytes (PDF content) or None on failure.
    """
    if not _HAS_WORD_COM:
        return None
    if not os.path.exists(_RTF_ROM):
        return None

    # ── Load data ──
    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return None

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role = 'shareholder'",
        (company_id,)).fetchall()

    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        placeholders = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({placeholders})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    txs = db.execute(
        "SELECT * FROM share_transactions WHERE company_id = ? ORDER BY transaction_date",
        (company_id,)).fetchall()

    co_name = rget(company, 'name') or ''
    co_br = rget(company, 'company_number') or ''
    today = datetime.now()
    report_date = today.strftime('%d %B %Y').upper()

    # ── Prepare shareholders ──
    shareholders = []
    for role in roles:
        p = person_map.get(role['person_id'], {})
        name_en = (rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)')[:80]
        id_no = rget(p, 'id_number') or ''
        id_str = f"(HKID No: {id_no})" if id_no else ''
        full_name = f"{name_en} {id_str}".strip()

        addr, region = _get_person_address(db, p['id'])
        if not addr:
            addr = (rget(p, 'address') or '')[:100]
        if region and region not in (addr or ''):
            addr = f"{addr}, {region}".strip(', ')
        addr = (addr or '')[:100]

        occupation = rget(p, 'occupation') or ''
        date_app = rget(role, 'date_appointed') or '-'
        date_cea = rget(role, 'date_ceased') or '-'
        shares_held = int(rget(role, 'shares') or 0)
        cert_no = rget(role, 'certificate_number') or '-'
        currency = rget(role, 'currency') or 'HKD'
        issue_price = rget(role, 'issue_price') or '1.00'

        person_name_key = name_en.strip().upper()
        person_txs = [
            t for t in txs
            if (rget(t, 'from_name') or rget(t, 'to_name') or '').strip().upper() == person_name_key
        ]

        shareholders.append({
            'full_name': full_name,
            'name_en': name_en,
            'id_no': id_no,
            'occupation': occupation,
            'addr': addr,
            'date_app': date_app,
            'date_cea': date_cea,
            'shares_held': shares_held,
            'cert_no': cert_no,
            'currency': currency,
            'issue_price': issue_price,
            'txs': person_txs,
        })

    # ── Copy .doc to temp file (avoid Chinese path issues with Word COM) ──
    import shutil as _shutil
    tmp_doc = tempfile.mktemp(suffix='.doc')
    _shutil.copy2(_RTF_ROM, tmp_doc)

    # ── Compact currency prefix ──
    _CURRENCY_COMPACT = {
        'HKD': 'HK$', 'USD': 'US$', 'CNY': '\xa5', 'RMB': '\xa5',
        'GBP': '\xa3', 'EUR': '€', 'JPY': '\xa5',
        'AUD': 'A$', 'SGD': 'S$', 'CAD': 'C$', 'TWD': 'NT$',
    }

    # ── Word COM: open → fill cells → save as PDF ──
    with _word_lock:
        pythoncom.CoInitialize()
        word = None
        tmp_pdf_out = None
        try:
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0

            doc = word.Documents.Open(tmp_doc)
            table = doc.Tables(1)

            # Helper: safely set a cell's text
            def _set_cell(r, c, value):
                try:
                    cell = table.Cell(r, c)
                    cell.Range.Text = str(value or '')
                except Exception as e:
                    # Fallback: use Selection.Find for merged-cell edge cases
                    pass

            # ── Header ──
            _set_cell(1, 2, co_name)
            _set_cell(2, 2, co_br)
            # Update REGISTER OF MEMBERS title with report date
            try:
                title_cell = table.Cell(2, 4)
                title_cell.Range.Text = f'REGISTER OF MEMBERS AT {report_date}'
            except:
                pass

            # ── Duplicate SH2 block for shareholders beyond 2 ──
            # NOTE: table.Rows(index) fails on merged cells — use Cell() range instead
            SH2_START = 15   # first row of SH2 block
            SH2_END = 24     # last row of SH2 block
            SH_BLOCK_ROWS = SH2_END - SH2_START + 1  # 10 rows per SH block
            NUM_COLS = 15
            _sh2_start_pos = table.Cell(SH2_START, 1).Range.Start
            _sh2_end_pos = table.Cell(SH2_END, NUM_COLS).Range.End
            _cur_last_row = SH2_END  # track where to paste next block
            for _ in range(len(shareholders) - 2):
                # Copy SH2 block using cell-based range
                doc.Range(_sh2_start_pos, _sh2_end_pos).Copy()
                # Paste after current last row (cell-based, avoids Rows collection)
                table.Cell(_cur_last_row, NUM_COLS).Range.Select()
                word.Selection.Collapse(0)  # wdCollapseEnd
                word.Selection.Paste()
                _cur_last_row += SH_BLOCK_ROWS

            # ── Build dynamic sh_bases for all shareholders ──
            sh_bases = []
            for i in range(len(shareholders)):
                if i == 0:
                    base = 4
                else:
                    base = 15 + (i - 1) * SH_BLOCK_ROWS
                sh_bases.append((base, base + 5, base + 9))

            # ── Fill all shareholders ──
            for i, sh in enumerate(shareholders):
                base, tx_row_start, tx_row_end = sh_bases[i]

                # Basic info (label row base)
                _set_cell(base, 2, sh['full_name'])
                _set_cell(base, 4, sh['occupation'])
                _set_cell(base, 6, sh['date_app'])

                # Address + Date Ceased (label row base+1)
                _set_cell(base + 1, 2, sh['addr'])
                _set_cell(base + 1, 4, sh['date_cea'])

                # ── Transaction sub-table ──
                # Row layout within sub-table:
                #   base+2: Date | Shares Acquired | Shares Transferred | Total Shares Held | Remarks | Entry Made By
                #   base+3: Cert No | Distinctive Nos | No. of Shares | Consideration | Deed No | Cert No | ...
                #   base+4: From | To | From | To
                #   base+5+: Data rows

                txs_for_sh = sh.get('txs', [])
                shares_held = sh.get('shares_held', 0)
                cert_no = sh.get('cert_no', '-')
                ccy = _CURRENCY_COMPACT.get(sh.get('currency', 'HKD').upper(), sh.get('currency', 'HKD')[:2] + '$')

                # Row 1 (Subscription): use first tx or initial holding
                if txs_for_sh:
                    first_tx = txs_for_sh[0]
                    tx_date = rget(first_tx, 'transaction_date') or sh['date_app']
                    tx_shares = int(rget(first_tx, 'shares') or 0)
                    _set_cell(tx_row_start, 1, tx_date)
                    _set_cell(tx_row_start, 2, f"{tx_shares:,}")
                    _set_cell(tx_row_start, 3, '0')
                    _set_cell(tx_row_start, 4, f"{tx_shares:,}")
                    _set_cell(tx_row_start, 5, '-')
                    _set_cell(tx_row_start, 6, '-')
                else:
                    # No transactions: fill initial subscription from role data
                    _set_cell(tx_row_start, 1, sh['date_app'])
                    _set_cell(tx_row_start, 2, f"{shares_held:,}")
                    _set_cell(tx_row_start, 3, '0')
                    _set_cell(tx_row_start, 4, f"{shares_held:,}")
                    _set_cell(tx_row_start, 5, '-')
                    _set_cell(tx_row_start, 6, '-')

                # Certificate / Distinctive numbers row
                _set_cell(tx_row_start + 1, 2, str(cert_no))
                _set_cell(tx_row_start + 1, 3, f"({shares_held:,})")
                _set_cell(tx_row_start + 1, 4, f"{shares_held:,}")
                _set_cell(tx_row_start + 1, 5, f"{ccy}{sh.get('issue_price', '1.00')}")

                # Additional transaction rows (Allotment, Transfer In, etc.)
                if len(txs_for_sh) > 1:
                    for ti in range(1, min(len(txs_for_sh), 4)):  # max 4 more rows
                        tx = txs_for_sh[ti]
                        dr = tx_row_start + 2 + ti  # data row offset
                        if dr > tx_row_end:
                            break
                        tx_date = rget(tx, 'transaction_date') or '-'
                        tx_shares = int(rget(tx, 'shares') or 0)
                        tx_type = rget(tx, 'type') or ''
                        _set_cell(dr, 1, tx_date)
                        _set_cell(dr, 2, f"{tx_shares:,}")
                        _set_cell(dr, 4, f"{shares_held:,}")

            # Save as PDF
            tmp_pdf_out = tempfile.mktemp(suffix='.pdf')
            doc.SaveAs2(tmp_pdf_out, FileFormat=17)
            doc.Close(False)

            with open(tmp_pdf_out, 'rb') as f:
                pdf_bytes = f.read()
            return pdf_bytes

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[DOC ROM→PDF] Word COM error: {e}")
            return None
        finally:
            if word:
                try:
                    word.Quit()
                except:
                    pass
            if tmp_pdf_out and os.path.exists(tmp_pdf_out):
                try:
                    os.unlink(tmp_pdf_out)
                except:
                    pass
            if os.path.exists(tmp_doc):
                try:
                    os.unlink(tmp_doc)
                except:
                    pass
            pythoncom.CoUninitialize()


# Register DOCX generators (ROM + ROD)
# Build properly-structured Word documents with tables matching Paul Tang RTF format,
# then convert to PDF via Word COM.  Uses python-docx for native Word table support
# so any number of shareholders / officers is handled naturally.
# ─────────────────────────────────────────────

def _build_rom_register_docx(db, company_id):
    """Generate ROM (Register of Members) as a .docx matching Paul Tang RTF format.

    RTF reference: 秘书系统文件/rod rom/Testing ROM.rtf
    Layout (Landscape A4):
      - Page border + company name (blue bold) + Company Number + Quorum
      - "REGISTER OF MEMBERS AT <date>" right-aligned
      - Per-shareholder block:
        * Name (label) + full name with HKID/passport
        * Address (label) + full address
        * Separator lines
        * Security (label) + share class description + Date / Date Ceased
        * Grey-header transaction sub-table (10 cols)

    Returns path to temp .docx file, or None on failure.
    """
    if not _HAS_DOCX:
        return None

    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return None

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role = 'shareholder'",
        (company_id,)).fetchall()
    txs = db.execute(
        "SELECT * FROM share_transactions WHERE company_id = ? ORDER BY transaction_date",
        (company_id,)).fetchall()

    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        placeholders = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({placeholders})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    co_name = rget(company, 'name') or ''
    co_br = rget(company, 'company_number') or ''
    today_str = datetime.now().strftime('%d %B %Y')

    # ── Constants matching RTF reference ──
    BLUE = '003399'         # ~RTF cf1 blue for headers
    GREY = 'E3E3E3'         # RGB(227,227,227) table header background
    FONT_EA = 'DengXian'    # East-Asian fallback for CJK text
    LABEL_SIZE = Pt(9)
    BODY_SIZE = Pt(9)
    HEADER_SIZE = Pt(12)
    TITLE_SIZE = Pt(13)

    # ── Date formatting helper ──
    def _fmt_date(val):
        """Normalize a date value to DD/MM/YYYY string."""
        if not val or val == '-':
            return '-'
        val_str = str(val).strip()
        # Already formatted DD/MM/YYYY
        if re.match(r'^\d{1,2}/\d{1,2}/\d{4}$', val_str):
            return val_str
        # ISO format YYYY-MM-DD
        if re.match(r'^\d{4}-\d{2}-\d{2}$', val_str):
            try:
                dt = datetime.strptime(val_str, '%Y-%m-%d')
                return dt.strftime('%d/%m/%Y')
            except:
                pass
        # Raw number DDMMYYYY (e.g. 24062026)
        if re.match(r'^\d{8}$', val_str):
            try:
                dt = datetime.strptime(val_str, '%d%m%Y')
                return dt.strftime('%d/%m/%Y')
            except:
                pass
            try:
                dt = datetime.strptime(val_str, '%Y%m%d')
                return dt.strftime('%d/%m/%Y')
            except:
                pass
        return val_str

    # ── Helpers ──
    def _set_run(run, name='Arial', size=BODY_SIZE, bold=False, color=None):
        """Set font on a run with proper East-Asian fallback."""
        run.font.name = name
        run.font.size = size
        run.bold = bold
        if color:
            r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        # Set East-Asian font for CJK characters
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_EA)

    def _add_run(para, text, bold=False, size=BODY_SIZE, align=None, color=None):
        """Add a run to a paragraph with consistent formatting."""
        if align:
            para.alignment = {'left': 0, 'center': 1, 'right': 2}.get(align, 0)
        run = para.add_run(str(text or ''))
        _set_run(run, size=size, bold=bold, color=color)
        return run

    def _add_para(doc, text, bold=False, size=BODY_SIZE, align='left', color=None,
                  space_before=0, space_after=0):
        """Add a paragraph with one run."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)
        _add_run(para, text, bold=bold, size=size, align=align, color=color)
        return para

    def _set_cell_bg(cell, hex_color):
        """Set cell background shading."""
        tcPr = cell._element.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), hex_color)
        shading.set(qn('w:val'), 'clear')
        tcPr.insert(0, shading)

    def _add_horizontal_line(doc, width_pt=1.5, color='000000'):
        """Add a thin horizontal line paragraph (bottom border on empty para)."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        pPr = para._element.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)
        bottom = pBdr.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom')
            pBdr.append(bottom)
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(int(width_pt * 8)))  # eighths of a point
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        return para

    try:
        from docx.oxml import OxmlElement
        from docx.shared import RGBColor

        doc = DocxDocument()

        # ── Page setup: Landscape A4 ──
        section = doc.sections[0]
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

        # ── Page border (thin black, matching RTF reference) ──
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for edge in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '24')
            b.set(qn('w:color'), '000000')
            pgBorders.append(b)
        sectPr.append(pgBorders)

        # Default style
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = BODY_SIZE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        # Set East-Asian font on Normal style
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_EA)

        # ── Header block ──
        # Company name (blue, bold, left)
        _add_para(doc, co_name, bold=True, size=HEADER_SIZE, color=BLUE, space_after=2)

        # Company Number + Quorum (blue) — Quorum on the right
        hdr_table = doc.add_table(rows=1, cols=2)
        hdr_table.autofit = True
        c_br = hdr_table.rows[0].cells[0]
        c_br.text = ''
        _add_run(c_br.paragraphs[0], f"Company Number: {co_br}", bold=True, size=BODY_SIZE, color=BLUE)
        c_q = hdr_table.rows[0].cells[1]
        c_q.text = ''
        _add_run(c_q.paragraphs[0], "Quorum:  1", size=Pt(8), color=BLUE, align='right')
        # Remove table borders
        for cell in hdr_table.rows[0].cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'none')
                b.set(qn('w:sz'), '0')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'auto')
                tcBorders.append(b)
            tcPr.append(tcBorders)

        # Title: REGISTER OF MEMBERS AT <date> (blue, bold, right-aligned)
        _add_para(doc, f"REGISTER OF MEMBERS AT {today_str}",
                  bold=True, size=TITLE_SIZE, color=BLUE, align='right', space_before=6)

        # Thick separator line
        _add_horizontal_line(doc, width_pt=1.5)
        _add_para(doc, '', size=Pt(2), space_after=0)  # tiny spacer

        # ── Per-shareholder blocks ──
        for role in roles:
            p = person_map.get(role['person_id'], {})
            name_en = (rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)')[:80]
            name_zh = (rget(p, 'name_chinese') or '')[:40]
            # Address: use structured fields with fallback to address column
            addr, region = _get_person_address(db, p['id'])
            if not addr:
                addr = (rget(p, 'address') or '')[:100]
            if region and region not in (addr or ''):
                addr = f"{addr}, {region}".strip(', ')
            addr = (addr or '')[:100]
            id_type = rget(p, 'id_type') or 'HKID'
            id_no = rget(p, 'id_number') or ''
            date_app = rget(role, 'date_appointed') or '-'
            date_cea = rget(role, 'date_ceased') or '-'
            shares_held = rget(role, 'shares') or 0
            cert_no = rget(role, 'certificate_number') or '-'
            currency = rget(role, 'currency') or 'HKD'
            issue_price = rget(role, 'issue_price') or '1.00'
            # Build share class description
            share_class = f"ORD - {currency}${issue_price} ORDINARY FULLY PAID ({currency})"

            # Build ID string
            if id_no:
                id_str = f"({id_type} No: {id_no})"
            else:
                id_str = ''

            # Full shareholder name line
            full_name_line = f"{name_en} {id_str}".strip()
            if name_zh:
                full_name_line = f"{name_en} {name_zh} {id_str}".strip()

            # --- Name row ---
            name_table = doc.add_table(rows=1, cols=2)
            name_table.autofit = True
            nl = name_table.rows[0].cells[0]
            nl.width = Cm(1.5)
            nl.text = ''
            _add_run(nl.paragraphs[0], 'Name', bold=True, size=LABEL_SIZE, color=BLUE)
            nr = name_table.rows[0].cells[1]
            nr.text = ''
            _add_run(nr.paragraphs[0], full_name_line, bold=True, size=LABEL_SIZE, color=BLUE)
            # Remove borders
            for cell in name_table.rows[0].cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for edge in ['top', 'left', 'bottom', 'right']:
                    b = OxmlElement(f'w:{edge}')
                    b.set(qn('w:val'), 'none')
                    b.set(qn('w:sz'), '0')
                    b.set(qn('w:space'), '0')
                    b.set(qn('w:color'), 'auto')
                    tcBorders.append(b)
                tcPr.append(tcBorders)

            # --- Address row ---
            addr_table = doc.add_table(rows=1, cols=2)
            addr_table.autofit = True
            al = addr_table.rows[0].cells[0]
            al.width = Cm(1.5)
            al.text = ''
            _add_run(al.paragraphs[0], 'Address', size=LABEL_SIZE, color=BLUE)
            ar = addr_table.rows[0].cells[1]
            ar.text = ''
            _add_run(ar.paragraphs[0], addr, size=LABEL_SIZE, color=BLUE)
            for cell in addr_table.rows[0].cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for edge in ['top', 'left', 'bottom', 'right']:
                    b = OxmlElement(f'w:{edge}')
                    b.set(qn('w:val'), 'none')
                    b.set(qn('w:sz'), '0')
                    b.set(qn('w:space'), '0')
                    b.set(qn('w:color'), 'auto')
                    tcBorders.append(b)
                tcPr.append(tcBorders)

            # Separator (blue, matching Paul Tang Register of members template)
            _add_horizontal_line(doc, width_pt=1.0, color=BLUE)
            _add_para(doc, '', size=Pt(2), space_after=0)

            # --- Security row ---
            sec_table = doc.add_table(rows=1, cols=4)
            sec_table.autofit = True
            sl = sec_table.rows[0].cells[0]
            sl.width = Cm(1.5)
            sl.text = ''
            _add_run(sl.paragraphs[0], 'Security', bold=True, size=LABEL_SIZE, color=BLUE)
            sd = sec_table.rows[0].cells[1]
            sd.width = Cm(10)
            sd.text = ''
            _add_run(sd.paragraphs[0], share_class, bold=True, size=LABEL_SIZE, color=BLUE)
            sdt = sec_table.rows[0].cells[2]
            sdt.width = Cm(1.2)
            sdt.text = ''
            _add_run(sdt.paragraphs[0], 'Date', bold=True, size=LABEL_SIZE, color=BLUE, align='right')
            sdv = sec_table.rows[0].cells[3]
            sdv.width = Cm(2)
            sdv.text = ''
            _add_run(sdv.paragraphs[0], _fmt_date(date_app), size=LABEL_SIZE, color=BLUE)
            # Date Ceased
            row2 = sec_table.add_row()
            sc = row2.cells[2]
            sc.text = ''
            _add_run(sc.paragraphs[0], 'Date Ceased', bold=True, size=LABEL_SIZE, color=BLUE, align='right')
            scv = row2.cells[3]
            scv.text = ''
            if date_cea and date_cea != '-':
                _add_run(scv.paragraphs[0], _fmt_date(date_cea), size=LABEL_SIZE, color=BLUE)
            for row in sec_table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for edge in ['top', 'left', 'bottom', 'right']:
                        b = OxmlElement(f'w:{edge}')
                        b.set(qn('w:val'), 'none')
                        b.set(qn('w:sz'), '0')
                        b.set(qn('w:space'), '0')
                        b.set(qn('w:color'), 'auto')
                        tcBorders.append(b)
                    tcPr.append(tcBorders)

            # ── Transaction sub-table (grey header, 10 cols) ──
            person_name_key = name_en.strip().upper()
            person_txs = [t for t in txs
                          if (rget(t, 'from_name') or rget(t, 'to_name') or '').strip().upper() == person_name_key]

            # Always show initial subscription row + transaction rows
            total_data_rows = 1 + len(person_txs)
            tx_table = doc.add_table(rows=2 + total_data_rows, cols=9)
            tx_table.style = 'Table Grid'
            tx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tx_table.autofit = True

            # Column widths (approximate, matching RTF proportions)
            col_widths = [Cm(2.0), Cm(1.8), Cm(1.5), Cm(2.0), Cm(2.0),
                          Cm(1.5), Cm(1.5), Cm(3.0), Cm(2.5)]
            for ci, w in enumerate(col_widths):
                for row in tx_table.rows:
                    row.cells[ci].width = w

            # Grey header row 1
            hdr_texts = [
                "Date Entered\n/ Ceased",
                "Transaction\nType",
                "Units",
                "Par Value",
                "Paid Up Value",
                "Certificate\nNo",
                "Balance",
                "Transferred To/From,\nRedeemed, Reissued",
                "Distinctive\nNumbers",
            ]
            for ci, txt in enumerate(hdr_texts):
                cell = tx_table.rows[0].cells[ci]
                cell.text = ''
                _add_run(cell.paragraphs[0], txt, bold=True, size=Pt(7))
                _set_cell_bg(cell, GREY)

            # Grey header row 2 (sub-labels: Per Share under Par Value / Paid Up Value)
            sub_labels = {3: "Per Share", 4: "Per Share"}
            for ci, txt in sub_labels.items():
                cell = tx_table.rows[1].cells[ci]
                cell.text = ''
                _add_run(cell.paragraphs[0], txt, bold=True, size=Pt(7))
                _set_cell_bg(cell, GREY)
            # Fill remaining grey cells in row 2
            for ci in range(9):
                if ci not in sub_labels:
                    _set_cell_bg(tx_table.rows[1].cells[ci], GREY)

            # ── Row 1: Initial subscription/allotment ──
            initial_balance = int(shares_held) if shares_held else 0
            initial_data = [
                _fmt_date(date_app),
                'Subscription' if initial_balance > 0 else '-',
                str(initial_balance) if initial_balance > 0 else '-',
                f"{currency}${issue_price}" if initial_balance > 0 else '-',
                f"{currency}${issue_price}" if initial_balance > 0 else '-',
                str(cert_no),
                str(initial_balance),
                '', '']
            for ci, val in enumerate(initial_data):
                cell = tx_table.rows[2].cells[ci]
                cell.text = ''
                _add_run(cell.paragraphs[0], str(val)[:50], size=Pt(7))

            # ── Subsequent rows: transactions ──
            balance = initial_balance
            for ti, tx in enumerate(person_txs):
                tx_shares = int(rget(tx, 'shares') or 0)
                tx_date = rget(tx, 'transaction_date') or '-'
                tx_inst = rget(tx, 'instrument_number') or '-'
                tx_cert = rget(tx, 'certificate_number') or cert_no
                tx_price = rget(tx, 'price_per_share') or issue_price
                tx_currency = rget(tx, 'currency') or currency

                is_in = (rget(tx, 'to_name') or '').strip().upper() == person_name_key
                is_out = (rget(tx, 'from_name') or '').strip().upper() == person_name_key
                if is_in:
                    balance += tx_shares
                    tx_type = 'Transfer In'
                    counterparty = f"From: {rget(tx, 'from_name') or ''}"
                elif is_out:
                    balance -= tx_shares
                    tx_type = 'Transfer Out'
                    counterparty = f"To: {rget(tx, 'to_name') or ''}"
                else:
                    balance += tx_shares
                    tx_type = 'Allotment'
                    counterparty = ''

                row_data = [
                    _fmt_date(tx_date),
                    tx_type,
                    str(tx_shares),
                    f"{tx_currency}${tx_price}",
                    f"{tx_currency}${tx_price}",
                    str(tx_cert),
                    str(balance),
                    counterparty,
                    tx_inst,
                ]
                data_row_idx = 3 + ti  # after header rows (0,1) + initial row (2)
                for ci, val in enumerate(row_data):
                    cell = tx_table.rows[data_row_idx].cells[ci]
                    cell.text = ''
                    _add_run(cell.paragraphs[0], str(val)[:50], size=Pt(7))

            # Spacer between shareholders
            _add_para(doc, '', size=Pt(4), space_after=0)

        # ── Footer: page number ──
        _add_para(doc, '', size=Pt(6), space_before=12)
        _add_para(doc, '- 1 -', size=Pt(8), align='center')

        # ── Save ──
        tmp_docx = tempfile.mktemp(suffix='.docx')
        doc.save(tmp_docx)
        return tmp_docx

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[ROM DOCX] Error: {e}")
        return None


def _build_rod_register_docx(db, company_id):
    """Generate ROD (Register of Officers) as a .docx with Word tables.

    Matches the Paul Tang RTF reference format:
      - Portrait A4, 6-column table
      - Directors rendered first, then Secretaries
      - Columns: Name/Address | DOB/Place/Occupation | ID/Passport |
                 Position | Date Appointed | Reason/Date Ceased

    Returns path to temp .docx file, or None on failure.
    """
    if not _HAS_DOCX:
        return None

    company = db.execute("SELECT * FROM companies WHERE id = ?", (company_id,)).fetchone()
    if not company:
        return None

    roles = db.execute(
        "SELECT * FROM person_company_roles WHERE company_id = ? AND role IN ('director', 'secretary')",
        (company_id,)).fetchall()
    person_ids = [r['person_id'] for r in roles]
    person_map = {}
    if person_ids:
        placeholders = ','.join(['?'] * len(person_ids))
        persons = db.execute(
            f"SELECT * FROM persons WHERE id IN ({placeholders})", person_ids).fetchall()
        person_map = {p['id']: p for p in persons}

    directors = [r for r in roles if r['role'] == 'director']
    secretaries = [r for r in roles if r['role'] == 'secretary']

    co_name = rget(company, 'name') or ''
    co_br = rget(company, 'company_number') or ''
    today_str = datetime.now().strftime('%d %B %Y')
    quorum = len(directors) if directors else None

    # ── Helpers (same as ROM) ──
    def _set_run(run, name='Times New Roman', size=Pt(9), bold=False):
        run.font.name = name
        run.font.size = size
        run.bold = bold

    def _cell(cell, text, bold=False, size=Pt(7), align='left'):
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = {'left': 0, 'center': 1, 'right': 2}.get(align, 0)
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(str(text or '')[:120])
        _set_run(run, size=size, bold=bold)

    def _para_run(para, text, bold=False, size=Pt(9), align='left'):
        para.alignment = {'left': 0, 'center': 1, 'right': 2}.get(align, 0)
        run = para.add_run(str(text or ''))
        _set_run(run, size=size, bold=bold)

    try:
        doc = DocxDocument()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        # ── Header ──
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = True

        c_left = header_table.rows[0].cells[0]
        c_left.text = ''
        p = c_left.paragraphs[0]
        _para_run(p, co_name, bold=True, size=Pt(12))
        p2 = c_left.add_paragraph()
        _para_run(p2, f"Company Number: {co_br}", size=Pt(9))
        if quorum:
            p3 = c_left.add_paragraph()
            _para_run(p3, f"Quorum: {quorum}", size=Pt(9))

        c_right = header_table.rows[0].cells[1]
        c_right.text = ''
        pr = c_right.paragraphs[0]
        _para_run(pr, "REGISTER OF OFFICERS", bold=True, size=Pt(13), align='right')

        c_left.width = Cm(10)
        c_right.width = Cm(6)

        date_para = doc.add_paragraph()
        _para_run(date_para, f"AT {today_str}", bold=True, size=Pt(9), align='right')

        doc.add_paragraph()

        # ── 6-column table ──
        col_headers = [
            "Name / Service /\nResidential Address",
            "Date / Place Birth /\nPlace Incorporated /\nOccupation /",
            "ID No / Passport\nDetails",
            "Position",
            "Date(s) Appointed\n/Meeting",
            "Reason / Date(s)\nCeased",
        ]

        total_rows = 1 + len(directors) + len(secretaries)  # header + data
        total_rows = max(total_rows, 3)
        table = doc.add_table(rows=total_rows, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for ci, label in enumerate(col_headers):
            _cell(table.rows[0].cells[ci], label, bold=True, size=Pt(7), align='center')

        # ── Data rows ──
        data_row = 1

        def _render_section(items, is_secretary=False):
            nonlocal data_row
            for role in items:
                p = person_map.get(role['person_id'], {})
                name_en = rget(p, 'name_english') or rget(p, 'name_chinese') or '(unnamed)'
                name_ch = rget(p, 'name_chinese') if rget(p, 'name_english') else ''
                is_nat = rget(p, 'identity') != 'corporate'

                # Name / Address block
                addr = (rget(p, 'address') or '') if is_nat else (rget(p, 'registered_office') or rget(p, 'address') or '')
                name_block = name_en
                if name_ch:
                    name_block += f"\n{name_ch}"
                if addr:
                    name_block += f"\n{addr[:120]}"

                # DOB / Place / Nation / Occupation
                if is_nat:
                    dob = rget(p, 'date_of_birth') or '-'
                    pob = rget(p, 'place_of_birth') or '-'
                    nat = rget(p, 'nationality') or '-'
                    occupation = rget(p, 'occupation') or '-'
                    dob_block = f"{dob}\n{pob}\n{nat}\n{occupation}"
                else:
                    poi = rget(p, 'place_incorporated') or '-'
                    dob_block = f"{poi}\n-\n-\n-"

                # ID block
                id_info = (rget(p, 'id_number') or rget(p, 'passport_number') or '-') if is_nat else (rget(p, 'company_number_ref') or '-')

                # Position
                if is_secretary:
                    position = "Secretary"
                else:
                    position = "Reserve Director" if rget(role, 'is_reserve') else "Director"

                # Date Appointed
                date_app = rget(role, 'date_appointed') or '-'

                # Date Ceased / Reason
                date_cea = rget(role, 'date_ceased')
                if date_cea:
                    reason_block = f"Resigned\n{date_cea}"
                else:
                    reason_block = "Current\n現任"

                row_data = [name_block, dob_block, id_info, position, date_app, reason_block]
                for ci, val in enumerate(row_data):
                    _cell(table.rows[data_row].cells[ci], val, size=Pt(7))
                data_row += 1

        # Directors first
        if directors:
            _render_section(directors)
        else:
            _cell(table.rows[data_row].cells[0], "(No directors / 尚無董事記錄)", size=Pt(8))
            for ci in range(1, 6):
                _cell(table.rows[data_row].cells[ci], '', size=Pt(7))
            data_row += 1

        # Secretaries below
        if secretaries:
            _render_section(secretaries, is_secretary=True)

        # ── Save ──
        tmp_docx = tempfile.mktemp(suffix='.docx')
        doc.save(tmp_docx)
        return tmp_docx

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[ROD DOCX] Error: {e}")
        return None


# ─── DOCX Document Generation (功能說明書 7.1) ───
# Generate Microsoft Word (.docx) company secretarial documents from system data.

DOCX_TYPES = {
    'company_profile':   '公司資料摘要',
    'directors_register': '董事名冊',
    'members_register':  '成員（股東）名冊',
    'board_resolution':  '董事會書面決議',
    'meeting_minutes':   '董事會會議記錄',
    'cr_form':           '政府表格 (Word)',
}
CR_FORM_META = {
    'nar1':  {'code': 'NAR1',  'title': '周年申報表',           'title_en': 'Annual Return'},
    'nd2a':  {'code': 'ND2A',  'title': '更改公司秘書及董事通知書（委任／停任）', 'title_en': 'Notice of Change of Company Secretary and Director (Appointment/Cessation)'},
    'nd2b':  {'code': 'ND2B',  'title': '更改公司秘書及董事詳情通知書',       'title_en': 'Notice of Change in Particulars of Company Secretary and Director'},
    'nd4':   {'code': 'ND4',   'title': '公司秘書及董事辭任通知書',           'title_en': 'Notice of Resignation of Company Secretary and Director'},
    'ndr1':  {'code': 'NDR1',  'title': '撤銷註冊申請書',                    'title_en': 'Application for Deregistration'},
    'nr1':   {'code': 'NR1',   'title': '註冊辦事處地址變更通知書',           'title_en': 'Notice of Change of Registered Office Address'},
    'nsc1':  {'code': 'NSC1',  'title': '股份配發申報書',                    'title_en': 'Return of Allotment'},
    'nnc1':  {'code': 'NNC1',  'title': '法團成立表格（股份有限公司）',        'title_en': 'Incorporation Form (Company Limited by Shares)'},
    'nnc2':  {'code': 'NNC2',  'title': '更改公司名稱通知書',                 'title_en': 'Notice of Change of Company Name'},
    'nn1':   {'code': 'NN1',   'title': '註冊非香港公司註冊申請書',            'title_en': 'Application for Registration as Registered Non-Hong Kong Company'},
    'nn3':   {'code': 'NN3',   'title': '註冊非香港公司周年申報表',            'title_en': 'Annual Return of Registered Non-Hong Kong Company'},
    'nn6':   {'code': 'NN6',   'title': '非香港公司更改秘書及董事（委任／停任）', 'title_en': 'Change of Company Secretary and Director of Non-Hong Kong Company'},
    'nn7':   {'code': 'NN7',   'title': '非香港公司更改秘書及董事詳情',         'title_en': 'Change in Particulars of Company Secretary and Director of Non-Hong Kong Company'},
    'nn9':   {'code': 'NN9',   'title': '非香港公司更改地址申報表',            'title_en': 'Notice of Change of Address of Non-Hong Kong Company'},
}


def _docx_fmt_date(s):
    """Normalise stored dates. DDMMYYYY -> DD/MM/YYYY; pass through otherwise."""
    if not s:
        return ''
    s = str(s).strip()
    if len(s) == 8 and s.isdigit():
        return f"{s[0:2]}/{s[2:4]}/{s[4:8]}"
    return s


def _docx_company_bundle(db, company_id):
    """Assemble a company + its members (directors / secretaries / shareholders)."""
    row = db.execute("SELECT * FROM companies WHERE id=?", (company_id,)).fetchone()
    if not row:
        return None
    c = dict(row)
    members = [dict(r) for r in db.execute(
        """SELECT pcr.role, pcr.shares, pcr.share_type, pcr.currency, pcr.paid_up,
                  pcr.date_appointed, pcr.date_ceased, pcr.is_reserve,
                  p.name_english, p.name_chinese, p.id_number, p.passport_number,
                  p.address, p.service_address, p.email, p.phone, p.identity, p.tcsp_number
           FROM person_company_roles pcr JOIN persons p ON p.id = pcr.person_id
           WHERE pcr.company_id=? AND (pcr.date_ceased IS NULL OR pcr.date_ceased='')
           ORDER BY pcr.role, p.name_english""", (company_id,)).fetchall()]
    directors = [m for m in members if m['role'] == 'director']
    secretaries = [m for m in members if m['role'] == 'secretary']
    shareholders = [m for m in members if m['role'] == 'shareholder']
    total_shares = sum(int(m['shares'] or 0) for m in shareholders)
    addr = ', '.join(filter(None, [
        c.get('reg_flat'), c.get('reg_building'), c.get('reg_street'),
        c.get('reg_district'), c.get('reg_region')]))
    return {
        'c': c, 'address': addr,
        'directors': directors, 'secretaries': secretaries,
        'shareholders': shareholders, 'total_shares': total_shares,
    }


def _docx_set_cjk_font(doc, font='Microsoft JhengHei'):
    """Ensure East-Asian glyphs render (Traditional Chinese) in the Normal style."""
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:eastAsia'), font)


def _docx_person_label(m):
    en = (m.get('name_english') or '').strip()
    cn = (m.get('name_chinese') or '').strip()
    if en and cn:
        return f"{en}（{cn}）"
    return en or cn or '—'


def _build_docx(db, company_id, doc_type, extra=None):
    """Return (docx_bytes, filename) for the requested document type."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    extra = extra or {}
    bundle = _docx_company_bundle(db, company_id)
    if not bundle:
        return None, None
    c = bundle['c']
    name_en = c.get('name') or ''
    name_cn = c.get('chinese_name') or ''
    br = c.get('company_number') or ''
    cr = c.get('ci_number') or ''

    doc = Document()
    _docx_set_cjk_font(doc)

    def h(text, size=16, bold=True, center=True, space_after=6):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def para(text, size=11, bold=False, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        return p

    def kv_table(rows):
        t = doc.add_table(rows=0, cols=2)
        t.style = 'Light Grid Accent 1'
        for k, v in rows:
            cells = t.add_row().cells
            cells[0].text = k
            cells[1].text = v or '—'
            cells[0].paragraphs[0].runs[0].bold = True
        return t

    def members_table(headers, data_rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Light Grid Accent 1'
        for i, hd in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = hd
            cell.paragraphs[0].runs[0].bold = True
        for dr in data_rows:
            cells = t.add_row().cells
            for i, val in enumerate(dr):
                cells[i].text = str(val) if val not in (None, '') else '—'
        return t

    # Company header block (shared)
    def company_header():
        h(name_en or name_cn or '公司', size=18)
        if name_cn and name_en:
            h(name_cn, size=14, bold=False)
        sub = []
        if br:
            sub.append(f"商業登記號碼 (BR)：{br}")
        if cr:
            sub.append(f"公司註冊編號 (CR)：{cr}")
        if sub:
            para('　｜　'.join(sub), size=10, center=True)
        doc.add_paragraph()

    label = DOCX_TYPES.get(doc_type, '公司文件')

    if doc_type == 'company_profile':
        company_header()
        h('公司資料摘要', size=15)
        kv_table([
            ('英文名稱', name_en),
            ('中文名稱', name_cn),
            ('商業登記號碼 (BR)', br),
            ('公司註冊編號 (CR)', cr),
            ('公司類型', c.get('company_type')),
            ('成立日期', _docx_fmt_date(c.get('incorporation_date'))),
            ('狀態', c.get('status')),
            ('註冊辦事處地址', bundle['address']),
            ('電郵', c.get('email')),
            ('電話', c.get('phone')),
        ])
        doc.add_paragraph()
        h(f"董事（{len(bundle['directors'])}）", size=13, center=False)
        members_table(['姓名', '身份證／護照', '委任日期', '地址'],
                      [[_docx_person_label(m), m.get('id_number') or m.get('passport_number'),
                        _docx_fmt_date(m.get('date_appointed')), m.get('address')]
                       for m in bundle['directors']] or [['（無）', '', '', '']])
        doc.add_paragraph()
        h(f"公司秘書（{len(bundle['secretaries'])}）", size=13, center=False)
        members_table(['姓名', 'TCSP 號碼', '委任日期', '地址'],
                      [[_docx_person_label(m), m.get('tcsp_number'),
                        _docx_fmt_date(m.get('date_appointed')), m.get('address')]
                       for m in bundle['secretaries']] or [['（無）', '', '', '']])
        doc.add_paragraph()
        ts = bundle['total_shares']
        h(f"股東 / 股本結構（總發行股數：{ts}）", size=13, center=False)
        members_table(['股東', '持股', '股份類別', '佔比'],
                      [[_docx_person_label(m), m.get('shares'), m.get('share_type') or '普通股',
                        (f"{round(int(m.get('shares') or 0) * 100 / ts, 2)}%" if ts else '—')]
                       for m in bundle['shareholders']] or [['（無）', '', '', '']])

    elif doc_type == 'directors_register':
        company_header()
        h('董事名冊 / Register of Directors', size=15)
        para('依據《公司條例》(第622章) 第 641 條備存。', size=10)
        doc.add_paragraph()
        members_table(['姓名', '身份證／護照', '委任日期', '住址', '電郵'],
                      [[_docx_person_label(m), m.get('id_number') or m.get('passport_number'),
                        _docx_fmt_date(m.get('date_appointed')), m.get('address'), m.get('email')]
                       for m in bundle['directors']] or [['（無董事記錄）', '', '', '', '']])

    elif doc_type == 'members_register':
        company_header()
        h('成員（股東）名冊 / Register of Members', size=15)
        para('依據《公司條例》(第622章) 第 627 條備存。', size=10)
        doc.add_paragraph()
        ts = bundle['total_shares']
        members_table(['股東', '持股數', '股份類別', '已繳股款', '佔比'],
                      [[_docx_person_label(m), m.get('shares'), m.get('share_type') or '普通股',
                        m.get('paid_up'),
                        (f"{round(int(m.get('shares') or 0) * 100 / ts, 2)}%" if ts else '—')]
                       for m in bundle['shareholders']] or [['（無股東記錄）', '', '', '', '']])
        doc.add_paragraph()
        para(f"總發行股數：{ts}", bold=True)

    elif doc_type in ('board_resolution', 'meeting_minutes'):
        is_min = doc_type == 'meeting_minutes'
        company_header()
        m_date = extra.get('meeting_date') or ''
        location = extra.get('location') or '公司註冊辦事處'
        if is_min:
            h('董事會會議記錄', size=15)
            h('MINUTES OF MEETING OF THE BOARD OF DIRECTORS', size=11, bold=False)
            doc.add_paragraph()
            kv_table([
                ('會議日期 Date', m_date),
                ('會議地點 Venue', location),
                ('出席董事 Present', '；'.join(_docx_person_label(m) for m in bundle['directors']) or '—'),
                ('主席 Chairman', _docx_person_label(bundle['directors'][0]) if bundle['directors'] else '—'),
            ])
        else:
            h('董事會書面決議', size=15)
            h('WRITTEN RESOLUTION OF THE DIRECTORS', size=11, bold=False)
            doc.add_paragraph()
            kv_table([
                ('決議日期 Date', m_date),
                ('簽署董事 Directors', '；'.join(_docx_person_label(m) for m in bundle['directors']) or '—'),
            ])
        doc.add_paragraph()
        para('議決事項 / RESOLVED THAT:', bold=True, size=12)
        content = (extra.get('content') or '').strip()
        if content:
            for line in content.split('\n'):
                para(line, size=11)
        else:
            para('1. ＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿', size=11)
            para('2. ＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿', size=11)
        doc.add_paragraph()
        doc.add_paragraph()
        para('簽署 / SIGNED:', bold=True)
        for m in (bundle['directors'] or [{'name_english': '＿＿＿＿＿＿'}]):
            doc.add_paragraph()
            para('_______________________________', size=11)
            para(f"{_docx_person_label(m)}　董事 / Director", size=10)
    elif doc_type == 'cr_form':
        form_code = (extra.get('form_code') or '').lower()
        meta = CR_FORM_META.get(form_code, {})
        if not meta:
            return None, None
        label = f"{meta['code']}_{meta['title']}"  # 檔名含表格編號+中文名
        company_header()
        h(f"{meta['title']} {meta['code']}", size=15)
        h(meta['title_en'], size=11, bold=False)
        para(f'公司註冊處表格 {meta["code"]} — 由系統自動填入公司資料生成草稿', size=9)
        doc.add_paragraph()

        # ── 共用欄位：公司基本資料 ──
        kv_table([
            ('公司英文名稱', name_en),
            ('公司中文名稱', name_cn),
            ('商業登記號碼 (BR)', br),
            ('公司註冊編號 (CR)', cr),
            ('公司類型', c.get('company_type')),
            ('註冊辦事處地址', bundle['address']),
            ('電郵', c.get('email')),
            ('電話', c.get('phone')),
            ('成立日期', _docx_fmt_date(c.get('incorporation_date'))),
            ('公司狀態', c.get('status')),
        ])
        doc.add_paragraph()

        # ── 董事 / 秘書（大部分表格都需要）──
        has_officers = form_code in ('nar1','nd2a','nd2b','nd4','nnc1','nn1','nn3','nn6','nn7')
        if has_officers:
            h(f"董事（{len(bundle['directors'])}）", size=13, center=False)
            members_table(
                ['姓名', '身份', '身份證／護照／公司編號', '委任日期', '辭任日期', '地址', '電郵'],
                [[_docx_person_label(m),
                  '法人' if m.get('identity') == 'corporate' else '自然人',
                  m.get('id_number') or m.get('passport_number') or m.get('tcsp_number') or '',
                  _docx_fmt_date(m.get('date_appointed')),
                  _docx_fmt_date(m.get('date_ceased')),
                  m.get('address') or '',
                  m.get('email') or '']
                 for m in bundle['directors']] or [['（無）', '', '', '', '', '', '']])
            doc.add_paragraph()
            h(f"公司秘書（{len(bundle['secretaries'])}）", size=13, center=False)
            members_table(
                ['姓名', '身份', 'TCSP／公司編號', '委任日期', '地址', '電郵'],
                [[_docx_person_label(m),
                  '法人' if m.get('identity') == 'corporate' else '自然人',
                  m.get('tcsp_number') or '',
                  _docx_fmt_date(m.get('date_appointed')),
                  m.get('address') or '',
                  m.get('email') or '']
                 for m in bundle['secretaries']] or [['（無）', '', '', '', '', '']])
            doc.add_paragraph()

        # ── 股東 / 股本 ──
        has_shares = form_code in ('nar1','nsc1','nnc1','nn1','nn3')
        if has_shares:
            ts = bundle['total_shares']
            h(f"股東／股本結構（總發行股數：{ts}）", size=13, center=False)
            members_table(
                ['股東', '持股數', '股份類別', '已繳股款', '佔比'],
                [[_docx_person_label(m),
                  m.get('shares') or '0',
                  m.get('share_type') or '普通股',
                  m.get('paid_up') or '',
                  (f"{round(int(m.get('shares') or 0) * 100 / ts, 2)}%" if ts else '—')]
                 for m in bundle['shareholders']] or [['（無）', '', '', '', '']])
            doc.add_paragraph()

        # ── NAR1 額外欄位 ──
        if form_code == 'nar1':
            para('重要控制人登記冊 (SCR) 是否備存於公司註冊辦事處？　是 □　否 □', size=11)
            para(f"截至申報日期之董事／秘書／股東資料以上表為準。", size=10)
            doc.add_paragraph()

        # ── 地址變更類 (NR1 / NDR1 / NN9) ──
        if form_code in ('nr1', 'ndr1', 'nn9'):
            para('現有註冊地址：', bold=True, size=11)
            para(bundle['address'] or '（未填）', size=11)
            para('變更後註冊地址（請手動填寫）：', bold=True, size=11)
            para('＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿', size=11)
            doc.add_paragraph()

        # ── 股份配發 (NSC1) 額外欄位 ──
        if form_code == 'nsc1':
            para('配發日期：＿＿＿＿＿＿＿＿', size=11)
            para('配發股份類別：＿＿＿＿＿＿＿＿', size=11)
            para('每股發行價：＿＿＿＿＿＿＿＿', size=11)
            para('配發總額：＿＿＿＿＿＿＿＿', size=11)
            doc.add_paragraph()

        # ── 簽署區 ──
        doc.add_paragraph()
        para('簽署 / SIGNED:', bold=True, size=11)
        doc.add_paragraph()
        para('_______________________________', size=11)
        para(f"董事 / Director　　日期 Date：＿＿＿＿＿＿＿＿", size=10)
        doc.add_paragraph()
        para('_______________________________', size=11)
        para(f"公司秘書 / Company Secretary　　日期 Date：＿＿＿＿＿＿＿＿", size=10)

    else:
        return None, None

    # Footer note
    doc.add_paragraph()
    footer = para(f"本文件由公司秘書管理系統自動生成 · {datetime.now().strftime('%Y-%m-%d %H:%M')}", size=8)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    import io as _io
    buf = _io.BytesIO()
    doc.save(buf)
    safe_name = re.sub(r'[^\w一-鿿-]', '_', (name_en or name_cn or 'company'))[:40]
    filename = f"{safe_name}_{label}.docx"
    return buf.getvalue(), filename


@app.route('/api/generate-docx', methods=['POST', 'OPTIONS'])
def generate_docx():
    if request.method == 'OPTIONS':
        return ('', 204)
    try:
        u = get_user()
        if not u:
            return jsonify({'error': 'Not authenticated'}), 401

        data = request.get_json(force=True, silent=True) or {}
        company_id = data.get('company_id')
        doc_type = data.get('doc_type')
        if not company_id:
            return jsonify({'error': '缺少 company_id'}), 400
        if doc_type not in DOCX_TYPES:
            return jsonify({'error': f'不支援的文件類型：{doc_type}',
                            'supported': list(DOCX_TYPES.keys())}), 400
        db = get_db()
        docx_bytes, filename = _build_docx(
            db, company_id, doc_type,
            extra={
                'content': data.get('content'),
                'meeting_date': data.get('meeting_date'),
                'location': data.get('location'),
                'form_code': data.get('form_code'),
            })
        if docx_bytes is None:
            return jsonify({'error': '找不到該公司'}), 404
        return jsonify({
            'success': True,
            'docx': base64.b64encode(docx_bytes).decode('ascii'),
            'filename': filename,
            'doc_type': doc_type,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/docx-types', methods=['GET'])
def docx_types():
    return jsonify([{'key': k, 'label': v} for k, v in DOCX_TYPES.items()])


# ─── 檢索服務 (功能說明書 6.1–6.6) ───

@app.route('/api/company-registers', methods=['GET'])
def company_registers():
    """6.2–6.6 公司登記冊明細：當前/歷史董事、股東、股份轉讓、SCR。"""
    company_id = request.args.get('company_id')
    if not company_id:
        return jsonify({'error': '缺少 company_id'}), 400
    db = get_db()
    company = db.execute("SELECT id, name, chinese_name, company_number FROM companies WHERE id=?",
                         (company_id,)).fetchone()
    if not company:
        return jsonify({'error': '找不到該公司'}), 404

    def roles(role, historical):
        cond = ("(pcr.date_ceased IS NOT NULL AND pcr.date_ceased != '')" if historical
                else "(pcr.date_ceased IS NULL OR pcr.date_ceased = '')")
        return [dict(r) for r in db.execute(
            f"""SELECT pcr.role, pcr.shares, pcr.share_type, pcr.currency, pcr.paid_up, pcr.unpaid,
                       pcr.date_appointed, pcr.date_ceased, pcr.is_reserve,
                       p.id AS person_id, p.identity, p.name_english, p.name_chinese,
                       p.id_number, p.passport_number, p.address, p.email, p.phone
                FROM person_company_roles pcr JOIN persons p ON p.id = pcr.person_id
                WHERE pcr.company_id = ? AND pcr.role = ? AND {cond}
                ORDER BY p.name_english""", (company_id, role)).fetchall()]

    share_tx = [dict(r) for r in db.execute(
        """SELECT * FROM share_transactions WHERE company_id = ?
           ORDER BY transaction_date DESC, created_at DESC""", (company_id,)).fetchall()]
    scr = [dict(r) for r in db.execute(
        "SELECT * FROM significant_controllers WHERE company_id = ? ORDER BY name_english",
        (company_id,)).fetchall()]

    return jsonify({
        'company': dict(company),
        'current_directors': roles('director', False),
        'historical_directors': roles('director', True),
        'current_shareholders': roles('shareholder', False),
        'historical_shareholders': roles('shareholder', True),
        'secretaries': roles('secretary', False),
        'share_transactions': share_tx,
        'scr': scr,
    })


# ─── NAR1 Field Listing & Debug PDF ───

@app.route('/api/nar1-fields', methods=['GET'])
def nar1_fields():
    """List all AcroForm fields in the NAR1 template (for FieldMapping page)."""
    try:
        template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NAR1-template-new.pdf')
        if not os.path.exists(template_path):
            return jsonify({'error': 'NAR1 template not found'}), 404
        doc = fitz.open(template_path)
        fields = []
        for pi in range(doc.page_count):
            for w in doc[pi].widgets():
                if w.field_name:
                    fields.append({
                        'name': w.field_name,
                        'type': 'checkbox' if w.field_name.startswith('cb_') else
                                'dropdown' if w.field_name.startswith('Dropdown') else 'text',
                        'page': pi + 1,
                    })
        doc.close()
        return jsonify({'fields': fields})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/nar1-debug-pdf', methods=['POST'])
def nar1_debug_pdf():
    """Generate a debug PDF with all field names filled into their widgets."""
    try:
        template_path = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates', 'NAR1-template-new.pdf')
        if not os.path.exists(template_path):
            return jsonify({'error': 'NAR1 template not found'}), 404
        doc = fitz.open(template_path)
        for pi in range(doc.page_count):
            for w in doc[pi].widgets():
                if not w.field_name:
                    continue
                try:
                    if w.field_name.startswith('cb_'):
                        # Check all checkboxes
                        w.field_value = True
                        w.update()
                    else:
                        w.field_value = w.field_name
                        w.update()
                except Exception:
                    pass
        pdf_bytes = doc.write(deflate=True)
        doc.close()
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── Full Export ───

@app.route('/api/export-all', methods=['POST'])
def export_all():
    """Export all tables as a JSON dump. Returns JSON (zip requires extra deps)."""
    try:
        db = get_db()
        export = {}
        for table in TABLES:
            try:
                rows = db.execute(f"SELECT * FROM {table}").fetchall()
                export[table] = [dict(r) for r in rows]
            except sqlite3.OperationalError:
                export[table] = []
        return jsonify({'success': True, 'data': export, 'exported_at': datetime.now().isoformat()})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# Handle OPTIONS preflight
@app.route('/api/<path:path>', methods=['OPTIONS'])
def handle_options(path=''):
    return '', 204

# ─── Form History endpoints (mirrors Cloudflare Functions /api/form-history/*) ───

@app.route('/api/form-history/list', methods=['GET'])
def form_history_list():
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401
    form_type = request.args.get('formType', '')
    if not form_type:
        return jsonify({'entries': []})
    db = get_db()
    rows = db.execute(
        'SELECT id, label, form_type, submission_index, created_at FROM form_history WHERE user_id = ? AND form_type = ? ORDER BY submission_index DESC',
        (u['id'], form_type)
    ).fetchall()
    entries = [{'id': r['id'], 'label': r['label'], 'form_type': r['form_type'],
                'submission_index': r['submission_index'], 'created_at': r['created_at']} for r in rows]
    return jsonify({'entries': entries})


@app.route('/api/form-history/load', methods=['GET'])
def form_history_load():
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401
    entry_id = request.args.get('id', '')
    if not entry_id:
        return jsonify({'error': 'Missing id'}), 400
    db = get_db()
    row = db.execute(
        'SELECT id, form_data FROM form_history WHERE id = ? AND user_id = ?',
        (entry_id, u['id'])
    ).fetchone()
    if not row:
        return jsonify({'error': 'Not found'}), 404
    form_data = json.loads(row['form_data']) if row['form_data'] else {}
    return jsonify({'entry': {'id': row['id'], 'form_data': form_data}})


@app.route('/api/form-history/save', methods=['POST'])
def form_history_save():
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401
    data = request.get_json(force=True, silent=True) or {}
    form_type = data.get('formType', '')
    form_data = data.get('formData')
    if not form_type or form_data is None:
        return jsonify({'error': 'formType and formData required'}), 400
    db = get_db()
    max_row = db.execute(
        'SELECT COALESCE(MAX(submission_index), 0) as max_idx FROM form_history WHERE user_id = ? AND form_type = ?',
        (u['id'], form_type)
    ).fetchone()
    next_idx = (max_row['max_idx'] or 0) + 1
    today = datetime.now().date().isoformat()
    label = f'{today}_{form_type}_{next_idx}'
    cursor = db.execute(
        'INSERT INTO form_history (user_id, user_email, form_type, submission_index, label, form_data) VALUES (?, ?, ?, ?, ?, ?)',
        (u['id'], u.get('email', ''), form_type, next_idx, label, json.dumps(form_data, ensure_ascii=False))
    )
    db.commit()
    new_id = cursor.lastrowid
    # ── Write-through: sync to cloud ──
    _sync_to_cloud("POST", "/api/form_history", {
        "id": new_id,
        "user_id": u['id'],
        "user_email": u.get('email', ''),
        "form_type": form_type,
        "submission_index": next_idx,
        "label": label,
        "form_data": json.dumps(form_data, ensure_ascii=False),
    })
    return jsonify({'id': new_id, 'label': label, 'submission_index': next_idx}), 201


@app.route('/api/form-history/<entry_id>', methods=['DELETE'])
def form_history_delete(entry_id):
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401
    if not entry_id:
        return jsonify({'error': 'Missing id'}), 400
    db = get_db()
    row = db.execute(
        'SELECT form_type, submission_index FROM form_history WHERE id = ? AND user_id = ?',
        (entry_id, u['id'])
    ).fetchone()
    if not row:
        return jsonify({'error': 'Not found'}), 404
    form_type = row['form_type']
    deleted_idx = row['submission_index']
    db.execute('DELETE FROM form_history WHERE id = ? AND user_id = ?', (entry_id, u['id']))
    db.execute(
        'UPDATE form_history SET submission_index = submission_index - 1 WHERE user_id = ? AND form_type = ? AND submission_index > ?',
        (u['id'], form_type, deleted_idx)
    )
    rows = db.execute(
        'SELECT id, created_at, submission_index FROM form_history WHERE user_id = ? AND form_type = ? AND submission_index >= ? ORDER BY submission_index',
        (u['id'], form_type, deleted_idx)
    ).fetchall()
    for r in rows:
        date_part = (r['created_at'] or '')[:10]
        new_label = f'{date_part}_{form_type}_{r["submission_index"]}'
        db.execute('UPDATE form_history SET label = ? WHERE id = ?', (new_label, r['id']))
    db.commit()
    # ── Write-through: sync to cloud ──
    _sync_to_cloud("DELETE", f"/api/form_history/{entry_id}")
    return jsonify({'ok': True})


# ─── Form Linkages (Phase 5.1: 表单关联查询) ───

@app.route('/api/form-linkages', methods=['GET'])
def form_linkages():
    """Return active form linkage rules, optionally filtered by primary form."""
    primary = request.args.get('primary', '')
    db = get_db()
    if primary:
        rows = db.execute(
            'SELECT id, primary_form, linked_form, linkage_type, description FROM form_linkages WHERE primary_form = ? AND is_active = 1',
            (primary,)
        ).fetchall()
    else:
        rows = db.execute(
            'SELECT id, primary_form, linked_form, linkage_type, description FROM form_linkages WHERE is_active = 1'
        ).fetchall()
    linkages = [{
        'id': r['id'],
        'primary_form': r['primary_form'],
        'linked_form': r['linked_form'],
        'linkage_type': r['linkage_type'],
        'description': r['description'],
    } for r in rows]
    return jsonify({'linkages': linkages})


# ─── IRC3111A PDF 生成（Phase 5.4: 税務局更改業務地址通知）───

def _build_irc3111a_pdf(data):
    """Build IRC 3111A — Notification of Change of Business Address (税務局).
    Free-form layout, Portrait A4, fpdf2. Matches IRD form style."""
    pdf = create_pdf(landscape=False)
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    M = 50  # margin
    PW, PH = 595, 842
    label_x = M
    value_x = M + 170  # label ~170pt wide

    co_name = data.get('companyName', data.get('name', ''))
    br = data.get('brNumber', data.get('br_number', ''))
    old_addr = data.get('oldAddress', '')
    new_addr = data.get('newAddress', '')
    effective_date = data.get('changeDate', data.get('effectiveDate', ''))
    signer = data.get('signerName', '')
    sign_date = data.get('signDate', '')

    # Build address strings if structured fields provided
    if not new_addr:
        parts = [
            data.get('newFlat', data.get('flat', '')),
            data.get('newBuilding', data.get('building', '')),
            data.get('newStreet', data.get('street', '')),
            data.get('newDistrict', data.get('district', '')),
        ]
        new_addr = ', '.join(p for p in parts if p)
    if not old_addr:
        parts_old = [
            data.get('oldFlat', ''),
            data.get('oldBuilding', ''),
            data.get('oldStreet', ''),
            data.get('oldDistrict', ''),
        ]
        old_addr = ', '.join(p for p in parts_old if p)
    if not effective_date:
        d = data.get('addressEffectiveDay', '')
        m = data.get('addressEffectiveMonth', '')
        y = data.get('addressEffectiveYear', '')
        if d or m or y:
            effective_date = f'{d}/{m}/{y}'
    if not sign_date:
        d = data.get('signDateDay', '')
        m = data.get('signDateMonth', '')
        y = data.get('signDateYear', '')
        if d or m or y:
            sign_date = f'{d}/{m}/{y}'

    def tnr(text, x, y, size=11, bold=False, align='L'):
        style = 'B' if bold else ''
        pdf.set_font('TNR', style, size)
        pdf.set_text_color(0, 0, 0)
        if align == 'C':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw / 2
        pdf.set_xy(x, y)
        pdf.cell(0, size + 4, str(text or ''))

    def tc(text, x, y, size=11, bold=False):
        style = 'B' if bold else ''
        pdf.set_font('TC', style, size)
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(x, y)
        pdf.cell(0, size + 4, str(text or ''))

    def draw_label(label_cn, label_en, y_pos, size=11):
        tc(label_cn, label_x, y_pos, size=size, bold=True)
        tnr(label_en, label_x, y_pos + size + 4, size=size - 2)

    def draw_value(value, y_pos, size=11):
        y_line = y_pos + 6
        tnr(str(value or ''), value_x, y_line, size=size)
        # Underline
        if value:
            tw = pdf.get_string_width(str(value or '')) + 4
        else:
            tw = PW - value_x - M - 10
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.5)
        pdf.line(value_x, y_line + size + 3, min(value_x + tw, PW - M), y_line + size + 3)
        pdf.set_line_width(0.3)

    # ── Header ──
    y = PH - 55

    # IRD reference
    tnr('Inland Revenue Department', PW / 2, y, size=9, align='C')
    tc('税 務 局', PW / 2, y + 12, size=9)
    y -= 35

    # Form title
    tnr('IR 3111A', PW / 2, y, size=16, bold=True, align='C')
    y -= 22
    tnr('Notification of Change of Business Address', PW / 2, y, size=12, bold=True, align='C')
    tc('通知更改業務地址', PW / 2, y + 16, size=10)
    y -= 50

    # Thick separator line
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(1.2)
    pdf.line(M, y, PW - M, y)
    pdf.set_line_width(0.3)
    y -= 22

    # ── Section 1: Business Registration Number ──
    draw_label('商業登記號碼', 'Business Registration No.', y)
    draw_value(br, y, size=12)
    y -= 36

    # ── Section 2: Name of Business ──
    draw_label('商業名稱', 'Name of Business', y)
    draw_value(co_name, y, size=12)
    y -= 42

    # ── Section 3: Old Business Address ──
    draw_label('舊業務地址', 'Old Business Address', y)
    draw_value(old_addr, y, size=11)
    y -= 42

    # ── Section 4: New Business Address ──
    draw_label('新業務地址', 'New Business Address', y)
    draw_value(new_addr, y, size=11)
    y -= 42

    # ── Section 5: Effective Date of Change ──
    draw_label('更改生效日期', 'Effective Date of Change', y)
    draw_value(effective_date, y, size=12)
    y -= 50

    # ── Section 6: Declaration ──
    tnr('Declaration', label_x, y, size=11, bold=True)
    tc('聲明', label_x + 70, y, size=11, bold=True)
    y -= 20

    declaration_text = (
        'I hereby declare that the above particulars are true and correct. '
        '本人謹此聲明，以上填報的詳情均屬真實和正確。'
    )
    tc(declaration_text, label_x + 5, y, size=9)
    y -= 28

    # Signer
    draw_label('簽署人姓名', 'Name of Signatory', y)
    draw_value(signer, y, size=12)
    y -= 42

    # Date
    draw_label('日期', 'Date', y)
    draw_value(sign_date, y, size=12)
    y -= 50

    # ── Footer note ──
    tnr('Notes:', label_x, y, size=8, bold=True)
    tc('註：', label_x + 35, y, size=8, bold=True)
    y -= 14
    note = (
        '1. This form should be completed and submitted to the Inland Revenue Department '
        'within 1 month of the change.\n'
        '2. 本表格須於更改生效後1個月內提交税務局。'
    )
    tc(note, label_x + 5, y, size=7.5)

    return pdf.output()


@app.route('/api/generate-irc3111a-pdf', methods=['POST'])
def generate_irc3111a_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _build_irc3111a_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── IR1263 PDF 生成（Phase 5.5: 稅務局結束營業通知書）───

def _build_ir1263_pdf(data):
    """Build IR1263 — Notice of Cessation of Business (稅務局).
    Free-form layout, Portrait A4, fpdf2. Ported from Cloud TS version."""
    pdf = create_pdf(landscape=False)
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    M = 50
    PW, PH = 595, 842
    label_x = M
    value_x = M + 180
    BLUE = (0, 51, 153)

    co_name = data.get('companyName', data.get('name', ''))
    br = data.get('brNumber', data.get('br_number', ''))
    app_date = data.get('applicationDate', data.get('cessationDate', ''))

    def tnr(text, x, y, size=11, bold=False, align='L', color=(0, 0, 0)):
        style = 'B' if bold else ''
        pdf.set_font('TNR', style, size)
        pdf.set_text_color(*color)
        if align == 'C':
            tw = pdf.get_string_width(str(text or ''))
            x = x - tw / 2
        pdf.set_xy(x, y)
        pdf.cell(0, size + 4, str(text or ''))

    def tc(text, x, y, size=11, bold=False, color=(0, 0, 0)):
        style = 'B' if bold else ''
        pdf.set_font('TC', style, size)
        pdf.set_text_color(*color)
        pdf.set_xy(x, y)
        pdf.cell(0, size + 4, str(text or ''))

    def draw_label(label_cn, label_en, y_pos, size=11):
        tc(label_cn, label_x, y_pos, size=size, bold=True)
        tnr(label_en, label_x, y_pos + size + 4, size=size - 2)

    def draw_value(value, y_pos, size=11):
        y_line = y_pos + 6
        tnr(str(value or ''), value_x, y_line, size=size)
        if value:
            tw = pdf.get_string_width(str(value)) + 4
        else:
            tw = PW - value_x - M - 10
        pdf.set_draw_color(0, 0, 0)
        pdf.set_line_width(0.5)
        pdf.line(value_x, y_line + size + 3, min(value_x + tw, PW - M), y_line + size + 3)
        pdf.set_line_width(0.3)

    # ── Header ──
    y = PH - 55

    # IRD reference
    tnr('Inland Revenue Department', PW / 2, y, size=9, align='C', color=BLUE)
    tc('税 務 局', PW / 2, y + 12, size=9, color=BLUE)
    y -= 35

    # Form title (centered, blue)
    tnr('IR1263 — Notice of Cessation of Business', PW / 2, y, size=14, bold=True, align='C', color=BLUE)
    y -= 20
    tc('IR1263 — 結束營業通知書', PW / 2, y, size=12, bold=True, color=BLUE)
    y -= 38

    # Separator line
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(1.2)
    pdf.line(M, y, PW - M, y)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.3)
    y -= 22

    # ── Company Name ──
    draw_label('公司名稱', 'Company Name', y)
    draw_value(co_name, y, size=12)
    y -= 42

    # ── Business Registration Number ──
    draw_label('商業登記號碼', 'Business Registration No.', y)
    draw_value(br, y, size=12)
    y -= 42

    # ── Date of Cessation ──
    draw_label('結束營業日期', 'Date of Cessation', y)
    draw_value(app_date, y, size=12)
    y -= 50

    # ── Declaration ──
    tnr('Declaration', label_x, y, size=11, bold=True)
    tc('聲明', label_x + 70, y, size=11, bold=True)
    y -= 20
    declaration = (
        'I hereby declare that the above-mentioned business has ceased operation. '
        '本人特此聲明上述業務已結束營業。'
    )
    tc(declaration, label_x + 5, y, size=9)
    y -= 36

    # ── Signature ──
    draw_label('簽署', 'Signature', y)
    draw_value('', y, size=11)
    y -= 42

    # ── Date ──
    draw_label('日期', 'Date', y)
    draw_value('', y, size=11)
    y -= 50

    # ── Footer ──
    from datetime import date
    today = date.today().isoformat()
    tnr(f'Generated by Company Secretary Management System · {today}', PW / 2, y, size=7, align='C')

    return pdf.output()


@app.route('/api/generate-ir1263-pdf', methods=['POST'])
def generate_ir1263_pdf():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        pdf_bytes = _build_ir1263_pdf(data)
        import base64 as b64
        return jsonify({'pdf': b64.b64encode(pdf_bytes).decode('ascii')})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── Batch Generation (Phase 5.3: 批量生成關聯表格) ───

def _translate_form_data(primary_data, primary_form, linked_form, db=None):
    """Translate primary form data into linked form fill data."""
    result = dict(primary_data)

    if primary_form == 'NR1' and linked_form == 'IRC3111A':
        new_addr = ', '.join(filter(None, [
            primary_data.get('flat', ''), primary_data.get('building', ''),
            primary_data.get('street', ''), primary_data.get('district', ''),
            primary_data.get('region', '')
        ]))
        eff_date = '/'.join(filter(None, [
            primary_data.get('addressEffectiveDay', ''),
            primary_data.get('addressEffectiveMonth', ''),
            primary_data.get('addressEffectiveYear', ''),
        ]))
        sign_date = '/'.join(filter(None, [
            primary_data.get('signDateDay', ''),
            primary_data.get('signDateMonth', ''),
            primary_data.get('signDateYear', ''),
        ]))
        result = {
            'companyName': primary_data.get('companyName', ''),
            'brNumber': primary_data.get('brNumber', ''),
            'oldAddress': '',
            'newAddress': new_addr,
            'changeDate': eff_date,
            'signerName': primary_data.get('signerName', ''),
            'signDate': sign_date,
        }
        # Fetch old address from DB
        company_id = primary_data.get('company_id', '')
        if company_id and db:
            row = db.execute('SELECT reg_flat, reg_building, reg_street, reg_district, reg_region FROM companies WHERE id=?', (company_id,)).fetchone()
            if row:
                old_addr = ', '.join(filter(None, [row['reg_flat'], row['reg_building'], row['reg_street'], row['reg_district'], row['reg_region']]))
                result['oldAddress'] = old_addr

    elif primary_form == 'NN9' and linked_form == 'IRC3111A':
        new_addr = ', '.join(filter(None, [
            primary_data.get('newFlat', primary_data.get('flat', '')),
            primary_data.get('newBuilding', primary_data.get('building', '')),
            primary_data.get('newStreet', primary_data.get('street', '')),
            primary_data.get('newDistrict', primary_data.get('district', '')),
        ]))
        result = {
            'companyName': primary_data.get('companyName', ''),
            'brNumber': primary_data.get('brNumber', ''),
            'oldAddress': '',
            'newAddress': new_addr,
            'changeDate': primary_data.get('changeDate', ''),
            'signerName': primary_data.get('signerName', ''),
            'signDate': primary_data.get('signDate', ''),
        }

    elif primary_form == 'NDR1' and linked_form == 'IR1263':
        result = {
            'companyName': primary_data.get('companyName', ''),
            'brNumber': primary_data.get('brNumber', ''),
            'applicationDate': primary_data.get('applicationDate', ''),
        }

    elif primary_form == 'ND2A' and linked_form == 'ND4':
        # Transfer cessation-related data
        result = primary_data

    elif primary_form == 'NNC1' and linked_form == 'IRBR1':
        result = {
            'irbr1_yes': True,
            'brNumber': primary_data.get('brNumber', ''),
        }

    elif primary_form == 'NN1' and linked_form == 'IRBR2':
        fields = primary_data.get('fields', primary_data)
        result = {
            'brNumber': primary_data.get('brNumber', fields.get('br_number', '')),
            'businessNameChinese': primary_data.get('companyNameChinese', fields.get('nameChinese', '')),
            'businessNameEnglish': primary_data.get('companyName', fields.get('nameEnglish', '')),
            'businessNature': primary_data.get('businessNature', fields.get('businessNature', '')),
            'commencementDate': primary_data.get('commencementDate', fields.get('commencementDate', '')),
            'irbr2_registered': True,
            'irbr2_elect3yr': True,
        }

    return result


@app.route('/api/generate-related-forms', methods=['POST'])
def generate_related_forms():
    """Batch generate primary form + linked forms."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'Empty request body'}), 400
        primary_form = data.get('primary_form', '')
        form_data = data.get('form_data', {})
        linked_forms = data.get('linked_forms', [])
        company_id = data.get('company_id', '')
        db = get_db()
        import base64 as b64

        results = []
        for form_code in linked_forms:
            try:
                linked_data = _translate_form_data(form_data, primary_form, form_code, db)
                if form_code == 'IRC3111A':
                    pdf_bytes = _build_irc3111a_pdf(linked_data)
                elif form_code == 'ND4':
                    # Use existing ND4 fill function
                    pdf_bytes = _fill_nd4_pdf(linked_data) if '_fill_nd4_pdf' in dir() else None
                    if pdf_bytes is None:
                        results.append({'form_code': form_code, 'error': 'ND4 function not available'})
                        continue
                elif form_code == 'IR1263':
                    pdf_bytes = _build_ir1263_pdf(linked_data)
                elif form_code == 'IRBR1':
                    pdf_bytes = _fill_irbr1_pdf(linked_data)
                elif form_code == 'IRBR2':
                    pdf_bytes = _fill_irbr2_pdf(linked_data)
                else:
                    results.append({'form_code': form_code, 'error': f'Unsupported form: {form_code}'})
                    continue
                results.append({
                    'form_code': form_code,
                    'pdf': b64.b64encode(pdf_bytes).decode('ascii'),
                    'filename': f'{form_code}_{linked_data.get("companyName", "form")}.pdf'
                })
            except Exception as e:
                import traceback
                traceback.print_exc()
                results.append({'form_code': form_code, 'error': str(e)})

        return jsonify({'success': True, 'forms': results})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════
# AI / OCR Mock 端点（对标云端 Functions，AI_MOCK 默认 true）
# ═══════════════════════════════════════════════════════════════

# ─── Mock 数据 ───
_MOCK_BR_DATA = {
    "companyName": "ABC TESTING COMPANY LIMITED",
    "chineseName": "ABC 測試有限公司",
    "brNumber": "12345678",
    "tradingName": "ABC Testing",
    "businessNature": "Import and Export",
    "businessCode": "46210",
    "companyType": "私人公司 Private company",
    "registerDate": "01/01/2020",
}

_MOCK_CI_DATA = {
    "companyName": "ABC TESTING COMPANY LIMITED",
    "chineseName": "ABC 測試有限公司",
    "companyNumber": "1234567",
    "incorporationDate": "01/01/2020",
    "companyType": "私人公司 Private company",
    "jurisdiction": "Hong Kong",
}

_MOCK_RESOLUTION_DATA = {
    "companyName": "ABC TESTING COMPANY LIMITED",
    "chineseName": "ABC 測試有限公司",
    "brNumber": "12345678",
    "companyNumber": "1234567",
    "incorporationDate": "01/01/2020",
    "companyType": "私人公司 Private company",
    "jurisdiction": "Hong Kong",
    "businessNature": "Import and Export",
    "tradingName": "ABC Testing",
    "regFlat": "Room 1201, 12/F",
    "regBuilding": "Tower A, Regent Centre",
    "regStreet": "63 Wo Yi Hop Road",
    "regDistrict": "Kwai Chung",
    "regRegion": "新界 New Territories",
    "contactPhone": "2123 4567",
    "contactEmail": "info@abctesting.com.hk",
    "directors": [
        {"nameEnglish": "CHAN Tai Man", "nameChinese": "陳大文", "idNumber": "A123456(3)", "address": "Flat A, 1/F, Block 1, Tai Po Garden, Tai Po, N.T.", "identity": "natural"},
        {"nameEnglish": "LEE Siu Ming", "nameChinese": "李小明", "idNumber": "B654321(2)", "address": "Room 5, 10/F, Kornhill Apartments, Quarry Bay, Hong Kong", "identity": "natural"},
    ],
    "secretaries": [
        {"nameEnglish": "Twinsail Consultants Limited", "nameChinese": "", "idNumber": "", "address": "Room 1201, 12/F, Tower A, Regent Centre, 63 Wo Yi Hop Road, Kwai Chung, N.T.", "identity": "corporate"},
    ],
    "shareholders": [
        {"nameEnglish": "CHAN Tai Man", "nameChinese": "陳大文", "idNumber": "A123456(3)", "address": "Flat A, 1/F, Block 1, Tai Po Garden, Tai Po, N.T.", "shares": 5000, "shareType": "Ordinary", "identity": "natural"},
        {"nameEnglish": "LEE Siu Ming", "nameChinese": "李小明", "idNumber": "B654321(2)", "address": "Room 5, 10/F, Kornhill Apartments, Quarry Bay, Hong Kong", "shares": 5000, "shareType": "Ordinary", "identity": "natural"},
    ],
}

_MOCK_CHAT_RESPONSES = [
    "根據資料庫記錄，目前共有 3 間公司、5 位董事和 2 位秘書。請問您需要查詢哪一間公司的詳細資料？",
    "已為您找到相關公司資料。如需生成 NAR1 年報表格，請在公司管理頁面點擊「生成 NAR1」按鈕。",
    "已成功新增董事記錄。您可以在公司詳情頁面查看和編輯相關資料。",
    "根據《公司條例》第 622 章，每間私人公司必須在成立周年日後 42 天內提交周年申報表（NAR1）。逾期提交可能會被罰款。",
    "已為您更新公司資料。如有其他需要，請隨時告訴我。",
]

# ─── Mock 响应辅助函数 ───
def _mock_json(data, status=200):
    return jsonify(data), status


def _ai_file_to_mock(form_data, mock_data):
    """提取上传文件信息并返回 mock 数据（模拟 AI OCR）"""
    file = form_data.get('file') if form_data else None
    filename = file.filename if file else '(no file)'
    print(f"[AI_MOCK] OCR simulated for: {filename}")
    return mock_data


# ─── POST /api/chat-assistant ───
@app.route('/api/chat-assistant', methods=['POST', 'OPTIONS'])
def chat_assistant():
    if request.method == 'OPTIONS':
        return ('', 204)
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        data = request.json or {}
        messages = data.get('messages', [])

        if AI_MOCK:
            # 返回预设 Mock 响应，根据用户最后一条消息选择
            last_msg = messages[-1]['content'] if messages else ''
            import random
            idx = hash(last_msg) % len(_MOCK_CHAT_RESPONSES) if last_msg else 0
            content = _MOCK_CHAT_RESPONSES[idx]
            print(f"[AI_MOCK] chat-assistant: returning mock response (AI_MOCK=true)")
            return jsonify({'content': content})

        # AI_MOCK=false: 调用真实 AI API
        LOVABLE_API_KEY = os.environ.get('LOVABLE_API_KEY', '')
        if not LOVABLE_API_KEY:
            return jsonify({'error': 'LOVABLE_API_KEY not configured'}), 500

        resp = requests.post(
            'https://ai.gateway.lovable.dev/v1/chat/completions',
            headers={
                'Authorization': f'Bearer {LOVABLE_API_KEY}',
                'Content-Type': 'application/json',
            },
            json={
                'model': 'google/gemini-3-flash-preview',
                'messages': [
                    {'role': 'system', 'content': '你是一個公司秘書管理系統的 AI 助手。回覆時使用繁體中文。'},
                    *messages,
                ],
            },
            timeout=60,
        )
        if not resp.ok:
            return jsonify({'error': f'AI gateway error: {resp.status_code}'}), resp.status_code

        result = resp.json()
        content = result.get('choices', [{}])[0].get('message', {}).get('content', '抱歉，無法生成回覆。')
        return jsonify({'content': content})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── POST /api/extract-br-info ───
@app.route('/api/extract-br-info', methods=['POST', 'OPTIONS'])
def extract_br_info():
    if request.method == 'OPTIONS':
        return ('', 204)
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        form_data = request.form if request.form else None
        extracted = _ai_file_to_mock(form_data, dict(_MOCK_BR_DATA)) if AI_MOCK else None

        if not AI_MOCK:
            LOVABLE_API_KEY = os.environ.get('LOVABLE_API_KEY', '')
            if not LOVABLE_API_KEY:
                return jsonify({'error': 'LOVABLE_API_KEY not configured'}), 500

            file = request.files.get('file')
            if not file:
                return jsonify({'error': 'No file uploaded'}), 400

            import base64 as b64
            file_bytes = file.read()
            file_b64 = b64.b64encode(file_bytes).decode()
            mime = file.mimetype or 'application/pdf'

            resp = requests.post(
                'https://ai.gateway.lovable.dev/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {LOVABLE_API_KEY}',
                    'Content-Type': 'application/json',
                },
                json={
                    'model': 'google/gemini-2.5-flash',
                    'messages': [
                        {'role': 'system', 'content': 'You are a Hong Kong Business Registration certificate data extractor. Extract: companyName, chineseName, brNumber, tradingName, businessNature, businessCode, companyType, registerDate. Return ONLY valid JSON.'},
                        {'role': 'user', 'content': [
                            {'type': 'text', 'text': 'Extract company info from this BR certificate.'},
                            {'type': 'image_url', 'image_url': {'url': f'data:{mime};base64,{file_b64}'}},
                        ]},
                    ],
                },
                timeout=60,
            )
            if not resp.ok:
                return jsonify({'error': f'AI gateway error: {resp.status_code}'}), resp.status_code

            result = resp.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            import re as _re
            cleaned = _re.sub(r'```json\n?|```\n?', '', content).strip()
            try:
                extracted = json.loads(cleaned)
            except json.JSONDecodeError:
                return jsonify({'error': '無法解析 AI 回應'}), 500

        return jsonify({'data': extracted})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── POST /api/extract-ci-info ───
@app.route('/api/extract-ci-info', methods=['POST', 'OPTIONS'])
def extract_ci_info():
    if request.method == 'OPTIONS':
        return ('', 204)
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        form_data = request.form if request.form else None
        extracted = _ai_file_to_mock(form_data, dict(_MOCK_CI_DATA)) if AI_MOCK else None

        if not AI_MOCK:
            LOVABLE_API_KEY = os.environ.get('LOVABLE_API_KEY', '')
            if not LOVABLE_API_KEY:
                return jsonify({'error': 'LOVABLE_API_KEY not configured'}), 500

            file = request.files.get('file')
            if not file:
                return jsonify({'error': 'No file uploaded'}), 400

            import base64 as b64
            file_bytes = file.read()
            file_b64 = b64.b64encode(file_bytes).decode()
            mime = file.mimetype or 'application/pdf'

            resp = requests.post(
                'https://ai.gateway.lovable.dev/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {LOVABLE_API_KEY}',
                    'Content-Type': 'application/json',
                },
                json={
                    'model': 'google/gemini-2.5-flash',
                    'messages': [
                        {'role': 'system', 'content': 'You are a Hong Kong Certificate of Incorporation data extractor. Extract: companyName, chineseName, companyNumber, incorporationDate, companyType, jurisdiction. Return ONLY valid JSON.'},
                        {'role': 'user', 'content': [
                            {'type': 'text', 'text': 'Extract info from this CI certificate.'},
                            {'type': 'image_url', 'image_url': {'url': f'data:{mime};base64,{file_b64}'}},
                        ]},
                    ],
                },
                timeout=60,
            )
            if not resp.ok:
                return jsonify({'error': f'AI gateway error: {resp.status_code}'}), resp.status_code

            result = resp.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            import re as _re
            cleaned = _re.sub(r'```json\n?|```\n?', '', content).strip()
            try:
                extracted = json.loads(cleaned)
            except json.JSONDecodeError:
                return jsonify({'error': '無法解析 AI 回應'}), 500

        return jsonify({'data': extracted})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ─── POST /api/extract-resolution-info ───
@app.route('/api/extract-resolution-info', methods=['POST', 'OPTIONS'])
def extract_resolution_info():
    if request.method == 'OPTIONS':
        return ('', 204)
    u = get_user()
    if not u:
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        form_data = request.form if request.form else None
        extracted = _ai_file_to_mock(form_data, dict(_MOCK_RESOLUTION_DATA)) if AI_MOCK else None

        if not AI_MOCK:
            LOVABLE_API_KEY = os.environ.get('LOVABLE_API_KEY', '')
            if not LOVABLE_API_KEY:
                return jsonify({'error': 'LOVABLE_API_KEY not configured'}), 500

            file = request.files.get('file')
            if not file:
                return jsonify({'error': 'No file uploaded'}), 400

            import base64 as b64
            file_bytes = file.read()
            file_b64 = b64.b64encode(file_bytes).decode()
            mime = file.mimetype or 'application/pdf'

            resp = requests.post(
                'https://ai.gateway.lovable.dev/v1/chat/completions',
                headers={
                    'Authorization': f'Bearer {LOVABLE_API_KEY}',
                    'Content-Type': 'application/json',
                },
                json={
                    'model': 'google/gemini-2.5-flash',
                    'messages': [
                        {'role': 'system', 'content': 'You are a Hong Kong company document extractor. Extract company info + directors/secretaries/shareholders arrays from uploaded document. Return ONLY valid JSON.'},
                        {'role': 'user', 'content': [
                            {'type': 'text', 'text': 'Extract all company information from this document.'},
                            {'type': 'image_url', 'image_url': {'url': f'data:{mime};base64,{file_b64}'}},
                        ]},
                    ],
                },
                timeout=60,
            )
            if not resp.ok:
                return jsonify({'error': f'AI gateway error: {resp.status_code}'}), resp.status_code

            result = resp.json()
            content = result.get('choices', [{}])[0].get('message', {}).get('content', '')
            import re as _re
            cleaned = _re.sub(r'```json\n?|```\n?', '', content).strip()
            try:
                extracted = json.loads(cleaned)
            except json.JSONDecodeError:
                return jsonify({'error': '無法解析 AI 回應'}), 500

        return jsonify({'data': extracted})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    init_db()
    auto_migrate()
    # Start the scheduled-email background thread once (avoid the Flask reloader
    # child spawning a second copy).
    if os.environ.get('WERKZEUG_RUN_MAIN') == 'true' or not app.debug:
        threading.Thread(target=scheduler_loop, daemon=True).start()
        print("[SERVER] Scheduled-email worker started (checks every 60s)")
    print("[SERVER] Local API running at http://localhost:5000")
    print("[SERVER] Admin account: admin@localhost / admin123")
    print("[SERVER] Register new accounts at /api/auth/register (no admin required)")
    ai_mock_status = "MOCK (AI_MOCK=true)" if AI_MOCK else "LIVE (AI_MOCK=false)"
    print(f"[SERVER] AI/OCR endpoints: {ai_mock_status}")
    if RESEND_API_KEY:
        print(f"[SERVER] Email: Resend API configured (onboarding@resend.dev)")
    elif SMTP_HOST:
        print(f"[SERVER] SMTP: configured ({SMTP_HOST})")
    else:
        print(f"[SERVER] Email: NOT configured — emails are SIMULATED & logged")
    app.run(host='0.0.0.0', port=5000, debug=True)
