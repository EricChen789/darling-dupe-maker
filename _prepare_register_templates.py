"""
Prepare Paul Tang DOCX register templates with {{PLACEHOLDER}} markers.
This script opens each template, replaces sample data with placeholders,
and saves cleaned versions ready for ZIP extraction.
"""
import os, sys
from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy

BASE = r"D:\myproject\秘书系统文件"
OUT = r"D:\myproject\darling-dupe-maker\_templates_cleaned"
os.makedirs(OUT, exist_ok=True)

def set_cell_text(cell, text):
    """Replace all text in a cell with a single placeholder."""
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run(text)
    return run

def clean_run_text(paragraph, old_text, new_text):
    """Replace text within paragraph runs."""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

# ═══════════════════════════ ROM ═══════════════════════════
print("=== Cleaning ROM template ===")
rom_path = os.path.join(BASE, r"登记册\股東登記冊_PaulTang格式.docx")
if os.path.exists(rom_path):
    doc = Document(rom_path)

    # P[0]: "Name of Company" → {{CO_NAME}}
    clean_run_text(doc.paragraphs[0], "Name of Company", "{{CO_NAME}}")

    # P[2]: "Company Number  ...  REGISTER OF MEMBERS" → placeholders
    p2 = doc.paragraphs[2]
    for run in p2.runs:
        if "Company Number" in run.text:
            run.text = "Company Number: {{CO_BR}}"
        if "REGISTER OF MEMBERS" in run.text:
            run.text = "REGISTER OF MEMBERS as at {{REPORT_DATE}}"

    # Table[3] = 18-column data table. Rows 0,1 = headers, Rows 2+ = sample data
    table = doc.tables[3]
    print(f"  Data table: {len(table.rows)} rows x {len(table.columns)} cols")

    # Clean sample data rows (rows 2+)
    # Keep row 2 as template row with placeholders
    if len(table.rows) > 2:
        template_row = table.rows[2]
        placeholders = [
            "{{SH_FULL_NAME}}",      # 0: Full Name
            "{{SH_ADDR}}",           # 1: Address
            "{{SH_OCCUPATION}}",     # 2: Occupation
            "{{SH_MERCHANT}}",       # 3: Merchant
            "{{SH_DATE_APP}}",       # 4: Date Entered as Member
            "{{SH_DATE_CEA}}",       # 5: Date Ceasing to be Member
            "{{SH_CERT_NO}}",        # 6: Cert No (Acq)
            "{{SH_DISTINCTIVE}}",    # 7: Distinctive Nos (Acq)
            "{{SH_SHARES_ACQ}}",     # 8: No. of Shares (Acq)
            "{{SH_CONSIDERATION}}",  # 9: Consideration Paid (Acq)
            "{{SH_TRANSFER_DEED}}",  # 10: Transfer Deed No (Acq)
            "{{SH_XFER_CERT}}",      # 11: Cert No (Xfer)
            "{{SH_XFER_DIST}}",      # 12: Distinctive Nos (Xfer)
            "{{SH_XFER_SHARES}}",    # 13: No. of Shares (Xfer)
            "{{SH_XFER_CONS}}",      # 14: Consideration Paid (Xfer)
            "{{SH_TOTAL_SHARES}}",   # 15: Total Shares Held
            "{{SH_REMARKS}}",        # 16: Remarks
            "{{SH_ENTRY_BY}}",       # 17: Entry Made By
        ]
        for ci, ph in enumerate(placeholders):
            if ci < len(template_row.cells):
                set_cell_text(template_row.cells[ci], ph)
        print(f"  Set {len(placeholders)} placeholders in template row")

        # Remove extra sample rows (rows 3+)
        for i in range(len(table.rows) - 1, 2, -1):
            table._tbl.remove(table.rows[i]._tr)
        print(f"  Removed extra sample rows, kept 1 template row")

    # Remove dummy tables 0,1,2 (empty placeholder tables)
    # Actually keep them — they're part of the template layout

    out_path = os.path.join(OUT, "ROM_register_template.docx")
    doc.save(out_path)
    print(f"  Saved: {out_path}")
else:
    print(f"  NOT FOUND: {rom_path}")

# ═══════════════════════════ SCR ═══════════════════════════
print("\n=== Cleaning SCR template ===")
scr_path = os.path.join(BASE, r"登记册\重要控制人登記冊_PaulTang格式.docx")
if os.path.exists(scr_path):
    doc = Document(scr_path)

    # Table[0]: Company name + title
    # Row[0], Cell[0]: "NAME OF COMPANY:  PAUL TANG AND COMPANY LIMITED\n公司名稱:  ..."
    t0 = doc.tables[0]
    cell0_text = t0.rows[0].cells[0].text
    print(f"  Header cell text: {cell0_text[:200]}")

    # Replace company name in Table[0], Row[0], Cell[0]
    for p in t0.rows[0].cells[0].paragraphs:
        for run in p.runs:
            if "PAUL TANG AND COMPANY LIMITED" in run.text:
                run.text = run.text.replace("PAUL TANG AND COMPANY LIMITED", "{{CO_NAME_EN}}")
            # Also replace any Chinese company name
            if "PAUL TANG" in run.text:
                run.text = run.text.replace("PAUL TANG", "{{CO_NAME_EN}}")

    # P[1], P[2]: JURISDICTION
    clean_run_text(doc.paragraphs[1], "HONG KONG", "{{JURISDICTION}}")
    clean_run_text(doc.paragraphs[2], "HONG KONG", "{{JURISDICTION}}")

    # Table[1] = Data table (8 cols)
    # Row[0] = headers, Row[1] = sample data
    t1 = doc.tables[1]
    print(f"  Data table: {len(t1.rows)} rows x {len(t1.columns)} cols")
    if len(t1.rows) > 1:
        sample_row = t1.rows[1]
        scr_placeholders = [
            "{{SCR_ENTRY_DATE}}",    # 0: Entry Date
            "{{SCR_NAME}}",          # 1: Name (merged with 0?)
            "{{SCR_ADDR}}",          # 2: Address
            "{{SCR_ID}}",            # 3: ID/PPT No
            "{{SCR_NATURE}}",        # 4: Nature of Control
            "{{SCR_DATES}}",         # 5: Becoming/Cessation Dates
            "{{SCR_REMARKS}}",       # 6: Remarks
        ]
        # Note: Table has 8 columns but only 7 data fields
        # Col 0 and Col 1 might be the same (Entry Date)
        all_phs = [
            "{{SCR_ENTRY_DATE}}",    # 0
            "{{SCR_NAME}}",          # 1
            "{{SCR_ADDR}}",          # 2
            "{{SCR_ID}}",            # 3
            "{{SCR_NATURE}}",        # 4
            "{{SCR_DATES}}",         # 5
            "{{SCR_REMARKS}}",       # 6
            "",                       # 7 (blank)
        ]
        for ci, ph in enumerate(all_phs):
            if ph and ci < len(sample_row.cells):
                set_cell_text(sample_row.cells[ci], ph)
        print(f"  Set placeholders in sample row")

    # Table[2] = Additional Matters (keep as-is with empty cells)

    out_path = os.path.join(OUT, "SCR_register_template.docx")
    doc.save(out_path)
    print(f"  Saved: {out_path}")
else:
    print(f"  NOT FOUND: {scr_path}")

# ═══════════════════════════ Transfer Resolutions ═══════════════════════════
print("\n=== Cleaning Transfer Resolutions template ===")
tr_path = os.path.join(BASE, r"Transfer\Transfer\Testing - Transfer resolutions.docx")
if os.path.exists(tr_path):
    doc = Document(tr_path)
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs[:20]):
        if p.text.strip():
            print(f'  P[{i}]: "{p.text[:150]}"')
    print(f"  Tables: {len(doc.tables)}")
    for ti, table in enumerate(doc.tables):
        print(f"  Table[{ti}]: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows[:3]):
            cells_text = [c.text[:60] for c in row.cells]
            print(f"    Row[{ri}]: {cells_text}")

    # For now, save a copy for inspection
    out_path = os.path.join(OUT, "Transfer_resolutions_template.docx")
    doc.save(out_path)
    print(f"  Saved copy: {out_path}")
else:
    print(f"  NOT FOUND: {tr_path}")

print("\n=== Done! ===")
print(f"Cleaned templates saved to: {OUT}")
for f in os.listdir(OUT):
    print(f"  {f}")
